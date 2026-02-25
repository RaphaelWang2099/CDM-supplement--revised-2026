import os
import re
import math
import sys
import time
import json
import pickle
import shutil
import hashlib
import threading
import tempfile
import pathlib
import warnings
from datetime import datetime
from collections import defaultdict
from dataclasses import dataclass, field
from typing import Any, Dict, List, Optional, Tuple
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
from tkinter.scrolledtext import ScrolledText

# Ensure matplotlib/fontconfig cache directories are writable on macOS environments
# where HOME cache paths may be read-only.
_tmp_cache_root = os.path.join(tempfile.gettempdir(), "lex_small_pipe_cache")
_mpl_cache_dir = os.path.join(_tmp_cache_root, "mpl")
_xdg_cache_dir = os.path.join(_tmp_cache_root, "xdg")
os.makedirs(_mpl_cache_dir, exist_ok=True)
os.makedirs(_xdg_cache_dir, exist_ok=True)
os.environ.setdefault("MPLCONFIGDIR", _mpl_cache_dir)
os.environ.setdefault("XDG_CACHE_HOME", _xdg_cache_dir)
os.environ.setdefault("TOKENIZERS_PARALLELISM", "false")
os.environ.setdefault("TRANSFORMERS_VERBOSITY", "error")
os.environ.setdefault("JOBLIB_MULTIPROCESSING", "0")
warnings.filterwarnings(
    "ignore",
    message=r"resource_tracker: There appear to be \d+ leaked semaphore objects to clean up at shutdown:.*",
    category=UserWarning,
)


def _import_matplotlib_safe():
    try:
        import matplotlib as _mpl
        try:
            _mpl.use("Agg", force=True)
        except Exception:
            pass

        if sys.platform == "darwin":
            # macOS font discovery may shell out to `system_profiler`, which can stall.
            import subprocess
            import plistlib

            _orig_check_output = subprocess.check_output
            _empty_fonts_plist = plistlib.dumps([])

            def _check_output_with_timeout(*args, **kwargs):
                cmd = args[0] if args else kwargs.get("args")
                if isinstance(cmd, (list, tuple)) and cmd and cmd[0] == "system_profiler":
                    kwargs.setdefault("timeout", 8)
                    try:
                        return _orig_check_output(*args, **kwargs)
                    except Exception:
                        return _empty_fonts_plist
                return _orig_check_output(*args, **kwargs)

            subprocess.check_output = _check_output_with_timeout
            try:
                import matplotlib.font_manager as _fm
                import matplotlib.pyplot as _plt
            finally:
                subprocess.check_output = _orig_check_output
        else:
            import matplotlib.font_manager as _fm
            import matplotlib.pyplot as _plt
        return _mpl, _fm, _plt
    except Exception:
        return None, None, None


matplotlib, fm, plt = _import_matplotlib_safe()
if matplotlib is None or fm is None or plt is None:
    raise RuntimeError("matplotlib import failed.")
_CJK_RC_FALLBACK_FAMILIES = [
    "STHeiti",
    "Heiti TC",
    "PingFang HK",
    "PingFang SC",
    "Songti SC",
    "Songti TC",
    "Hiragino Sans GB",
    "Arial Unicode MS",
    "Arial Unicode",
    "Noto Sans CJK SC",
    "Noto Sans CJK TC",
    "SimHei",
    "Microsoft YaHei",
    "DejaVu Sans",
]
plt.rcParams["font.family"] = "sans-serif"
plt.rcParams["font.sans-serif"] = list(_CJK_RC_FALLBACK_FAMILIES)
plt.rcParams["axes.unicode_minus"] = False  # 解决保存图像是负号'-'显示为方块的问题
import numpy as np
import pandas as pd
import openpyxl
from docx import Document as DocxWriter
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

# Heavy ML/scientific imports are loaded lazily to keep UI startup responsive.
CountVectorizer = None
TfidfVectorizer = None
TfidfTransformer = None
cosine_similarity = None
PCA = None
linkage = None
scipy_dendrogram = None
fcluster = None
squareform = None
MDS = None
silhouette_score = None
nx = None
_ML_STACK_READY = False


def _ensure_ml_stack():
    global CountVectorizer, TfidfVectorizer, TfidfTransformer
    global cosine_similarity, PCA, linkage, scipy_dendrogram, fcluster
    global squareform, MDS, silhouette_score, nx, _ML_STACK_READY
    if _ML_STACK_READY:
        return
    try:
        from sklearn.feature_extraction.text import (
            CountVectorizer as _CountVectorizer,
            TfidfVectorizer as _TfidfVectorizer,
            TfidfTransformer as _TfidfTransformer,
        )
        from sklearn.metrics.pairwise import cosine_similarity as _cosine_similarity
        from sklearn.decomposition import PCA as _PCA
        from sklearn.manifold import MDS as _MDS
        from sklearn.metrics import silhouette_score as _silhouette_score
        from scipy.cluster.hierarchy import (
            linkage as _linkage,
            dendrogram as _scipy_dendrogram,
            fcluster as _fcluster,
        )
        from scipy.spatial.distance import squareform as _squareform
        import networkx as _nx
    except Exception as exc:
        raise RuntimeError(
            f"Failed to import sklearn/scipy stack on Python {sys.version.split()[0]}. "
            "Please reinstall scientific dependencies in a supported environment."
        ) from exc

    CountVectorizer = _CountVectorizer
    TfidfVectorizer = _TfidfVectorizer
    TfidfTransformer = _TfidfTransformer
    cosine_similarity = _cosine_similarity
    PCA = _PCA
    linkage = _linkage
    scipy_dendrogram = _scipy_dendrogram
    fcluster = _fcluster
    squareform = _squareform
    MDS = _MDS
    silhouette_score = _silhouette_score
    nx = _nx
    _ML_STACK_READY = True

def _unique_nonempty(items: List[str]) -> List[str]:
    seen: set = set()
    out: List[str] = []
    for x in items:
        s = str(x or "").strip()
        if not s or s in seen:
            continue
        seen.add(s)
        out.append(s)
    return out


def _split_env_paths(raw: str) -> List[str]:
    text = str(raw or "")
    if not text:
        return []
    parts = [text]
    for sep in (",", ";", "\n", os.pathsep):
        next_parts: List[str] = []
        for p in parts:
            next_parts.extend(p.split(sep))
        parts = next_parts
    return [x.strip() for x in parts if x.strip()]


def _script_dir() -> str:
    try:
        return str(pathlib.Path(__file__).resolve().parent)
    except Exception:
        return os.getcwd()


def _detect_mpl_ttf_dir() -> str:
    try:
        import matplotlib as _mpl

        data_dir = str(_mpl.get_data_path() or "").strip()
        if not data_dir:
            return ""
        ttf_dir = os.path.join(data_dir, "fonts", "ttf")
        if os.path.isdir(ttf_dir):
            return ttf_dir
    except Exception:
        pass
    return ""


EXCEL_EXTS = {".xlsx", ".xlsm", ".xls"}
LEGACY_TTF_DIR = _detect_mpl_ttf_dir()
LEGACY_TTF_CANDIDATES = ["simhei.ttf", "simkai.ttf", "simsun.ttf"]
_env_default_font_path = str(os.environ.get("NGRAM_FONT_PATH", "")).strip()
_default_font_file_candidates = _unique_nonempty(
    [
        _env_default_font_path,
        "/System/Library/Fonts/STHeiti Medium.ttc",
        "/System/Library/Fonts/STHeiti Light.ttc",
        "/System/Library/Fonts/Supplemental/PingFang.ttc",
        "/System/Library/Fonts/Supplemental/Songti.ttc",
        "/Library/Fonts/Arial Unicode.ttf",
        "/Users/lab/Desktop/字体合集/公文字体/公文字体合集/simhei.ttf",
    ]
)


def _first_existing_file(paths: List[str]) -> str:
    for p in paths:
        pp = os.path.expanduser(str(p or "").strip())
        if pp and os.path.isfile(pp):
            return pp
    return ""


DEFAULT_HEATMAP_FONT_PATH = _first_existing_file(_default_font_file_candidates)
DEFAULT_HEATMAP_FONT_CANDIDATES = [
    "STHeiti",
    "Heiti TC",
    "PingFang HK",
    "PingFang SC",
    "Songti SC",
    "Songti TC",
    "STSong",
    "Heiti SC",
    "Hiragino Sans GB",
    "Noto Sans CJK SC",
    "Noto Sans CJK TC",
    "SimHei",
    "Microsoft YaHei",
    "WenQuanYi Zen Hei",
    "Arial Unicode MS",
]
_script_root = _script_dir()
_desktop_font_bundle_dir = os.path.join(os.path.expanduser("~"), "Desktop", "字体合集")
_font_env_dirs = _split_env_paths(os.environ.get("NGRAM_USER_FONT_DIRS", ""))
_font_env_single = str(os.environ.get("NGRAM_USER_FONT_DIR", "")).strip()
if _font_env_single:
    _font_env_dirs.append(_font_env_single)
USER_DEFAULT_FONT_DIRS = _unique_nonempty(
    _font_env_dirs
    + [
        _desktop_font_bundle_dir,
        os.path.join(_script_root, "fonts"),
        os.path.join(_script_root, "font"),
        os.path.join(_script_root, "assets", "fonts"),
        os.path.join(os.path.expanduser("~"), "Library", "Fonts"),
        os.path.join(os.path.expanduser("~"), ".fonts"),
        os.path.join(os.path.expanduser("~"), ".local", "share", "fonts"),
    ]
)
MAX_COLLATION_DP_CELLS = 2_000_000  # safeguard for display-only auto-collation
GRANULARITY_CHOICES = ["h1", "h2", "normal", "all"]
ANALYSIS_GRANULARITIES = ["h1", "h2", "normal"]

# Semantic / reranker defaults (pipeline13-style, raw scores only)
MODEL_SIZE_CHOICES = ["0.6B", "4B", "8B"]


_HF_HUB_DIR_CANDIDATES = _unique_nonempty(
    [
        str(os.environ.get("NGRAM_HF_HUB_DIR", "")).strip(),
        str(pathlib.Path.home() / ".cache" / "huggingface" / "hub"),
    ]
)


def _first_existing_dir(paths: List[str]) -> str:
    for p in paths:
        if os.path.isdir(p):
            return p
    return ""


_HF_HUB_DIR = (
    _first_existing_dir(_HF_HUB_DIR_CANDIDATES)
    or (_HF_HUB_DIR_CANDIDATES[-1] if _HF_HUB_DIR_CANDIDATES else str(pathlib.Path.home() / ".cache" / "huggingface" / "hub"))
)


def _pick_model_cache_dir(model_folder: str) -> str:
    for root in _HF_HUB_DIR_CANDIDATES:
        cand = os.path.join(root, model_folder)
        if os.path.isdir(cand):
            return cand
    return os.path.join(_HF_HUB_DIR, model_folder)


QWEN_EMBED_MODEL_DIRS: Dict[str, str] = {
    size: _pick_model_cache_dir(f"models--Qwen--Qwen3-Embedding-{size}")
    for size in MODEL_SIZE_CHOICES
}
QWEN_RERANK_MODEL_DIRS: Dict[str, str] = {
    size: _pick_model_cache_dir(f"models--Qwen--Qwen3-Reranker-{size}")
    for size in MODEL_SIZE_CHOICES
}
DEFAULT_EMBED_MODEL_SIZE = str(os.environ.get("NGRAM_EMBED_MODEL_SIZE", "0.6B")).strip() or "0.6B"
DEFAULT_RERANK_MODEL_SIZE = str(os.environ.get("NGRAM_RERANK_MODEL_SIZE", "0.6B")).strip() or "0.6B"
RERANK_TOP_PERCENT_NON_H1 = 3.0
SEMANTIC_MAXSIM_CHUNK_SIZE = 500
SEMANTIC_MAXSIM_CHUNK_STRIDE = 400
DEFAULT_TOPK_CANDIDATES = 20
H2_FULL_PAIR_UNITS_THRESHOLD = 41
RERANK_AUTO_SCALE_NON_H1 = True
RERANK_AUTO_NON_H1_MAX_CANDIDATES = 1200
RERANK_AUTO_NON_H1_MIN_PERCENT = 0.2
PIPELINE_CODE_VERSION = "small_pipe4_result_cache_v1"
RESULT_CACHE_SCHEMA_VERSION = "run_cache_v1"
RESULT_CACHE_ANALYSIS_VERSION = "small_pipe4_analysis_signature_v1"
_SCRIPT_SHA1_CACHE = ""

_SEMANTIC_BACKEND_CACHE: Dict[str, dict] = {}
_RERANK_BACKEND_CACHE: Dict[str, dict] = {}


def _script_sha1() -> str:
    global _SCRIPT_SHA1_CACHE
    if _SCRIPT_SHA1_CACHE:
        return str(_SCRIPT_SHA1_CACHE)
    try:
        with open(__file__, "rb") as f:
            _SCRIPT_SHA1_CACHE = hashlib.sha1(f.read()).hexdigest()
    except Exception:
        _SCRIPT_SHA1_CACHE = ""
    return str(_SCRIPT_SHA1_CACHE)


def _file_sha1(path: str, chunk_size: int = 1024 * 1024) -> str:
    p = str(path or "").strip()
    if not p or not os.path.isfile(p):
        return ""
    h = hashlib.sha1()
    try:
        with open(p, "rb") as f:
            while True:
                b = f.read(int(chunk_size))
                if not b:
                    break
                h.update(b)
        return h.hexdigest()
    except Exception:
        return ""


def _json_sha1(obj: Any) -> str:
    try:
        raw = json.dumps(obj, ensure_ascii=False, sort_keys=True, separators=(",", ":"))
    except Exception:
        raw = repr(obj)
    return hashlib.sha1(raw.encode("utf-8", errors="ignore")).hexdigest()


@dataclass
class Para:
    pid: int
    text: str
    style: str


@dataclass
class Unit:
    uid: int
    level: str
    h1: str
    h2: str
    text: str
    para_ids: List[int]
    normal_chapter: str = ""
    # For lexical n-gram generation: generate within each segment, then merge counts.
    # This prevents cross-chapter boundary n-grams.
    ngram_segments: List[str] = field(default_factory=list)


# ------------------- 文本處理 -------------------
# ------------------- Excel輸入與切分（同 pipeline12 思路） -------------------
def normalize_style(style: str) -> str:
    s = str(style or "").strip().lower()
    if s in ("h1", "title", "heading1", "heading 1"):
        return "h1"
    if s in ("h2", "heading2", "heading 2"):
        return "h2"
    if s in ("h3", "heading3", "heading 3"):
        return "h3"
    return "normal"


def normalize_granularity_choice(value: str) -> str:
    s = str(value or "").strip().lower()
    if s in ("all", "全部", "全量"):
        return "all"
    if s in ("h1", "h2", "normal"):
        return s
    return "normal"


def _aggregate_units_from_normal(normal_units: List[Unit], granularity: str) -> List[Unit]:
    g = normalize_style(granularity)
    if g == "normal":
        units: List[Unit] = []
        uid = 0
        for u in normal_units:
            # Strict body-under-heading policy: only keep正文 under h1/h2 context.
            if not (str(u.h1 or "").strip() or str(u.h2 or "").strip()):
                continue
            units.append(
                Unit(
                    uid=uid,
                    level="normal",
                    h1=u.h1,
                    h2=u.h2,
                    text=u.text,
                    para_ids=list(u.para_ids or []),
                    normal_chapter=str(u.normal_chapter or ""),
                    ngram_segments=(
                        [str(s) for s in list(u.ngram_segments or []) if str(s or "").strip()]
                        or ([str(u.text or "")] if str(u.text or "").strip() else [])
                    ),
                )
            )
            uid += 1
        return units

    grouped: Dict[Tuple[str, ...], List[Unit]] = {}
    for u in normal_units:
        h1 = str(u.h1 or "").strip()
        h2 = str(u.h2 or "").strip()
        # Strict body-under-heading policy for aggregated levels.
        if g == "h1":
            if not h1:
                continue
            key = (h1,)
        else:
            if not h2:
                continue
            key = (h1 or "(empty)", h2)
        grouped.setdefault(key, []).append(u)

    out: List[Unit] = []
    uid = 0
    for items in grouped.values():
        texts = [str(x.text or "").strip() for x in items if str(x.text or "").strip()]
        merged = "\n".join(texts).strip()
        if not merged:
            continue
        para_ids: List[int] = []
        ngram_segments: List[str] = []
        for x in items:
            para_ids.extend(list(x.para_ids or []))
            segs = [str(s) for s in list(x.ngram_segments or []) if str(s or "").strip()]
            if not segs and str(x.text or "").strip():
                segs = [str(x.text or "")]
            ngram_segments.extend(segs)
        out.append(
            Unit(
                uid=uid,
                level=g,
                h1=str(items[0].h1 or ""),
                h2=str(items[0].h2 or ""),
                text=merged,
                para_ids=para_ids,
                normal_chapter=str(items[0].normal_chapter or ""),
                ngram_segments=ngram_segments,
            )
        )
        uid += 1
    return out


# ---- Text preprocessing for similarity: keep ONLY Chinese Han characters ----
# Includes CJK Unified Ideographs + Extensions commonly used in classical texts.
_HAN_RE = re.compile(
    r"[\u3400-\u4DBF\u4E00-\u9FFF\U00020000-\U0002A6DF\U0002A700-\U0002B73F\U0002B740-\U0002B81F\U0002B820-\U0002CEAF\U0002CEB0-\U0002EBEF]"
)


def preprocess_keep_han(text: str) -> str:
    """For similarity only: remove punctuation/symbols and keep ONLY Han characters."""
    if not text:
        return ""
    return "".join(_HAN_RE.findall(str(text)))


def _normalize_chapter_value(v: Any) -> str:
    if v is None:
        return ""
    if isinstance(v, float):
        if math.isnan(v) or math.isinf(v):
            return ""
        if abs(v - round(v)) < 1e-9:
            return str(int(round(v)))
        return str(v).strip()
    if isinstance(v, int):
        return str(v)
    s = str(v).strip()
    return s


def read_excel_as_units_table(path: str) -> Optional[List[Unit]]:
    """Units 表：需包含 level + text/正文；可选 h1/h2。"""
    wb = openpyxl.load_workbook(path, data_only=True)
    ws = wb.active
    headers = [str(c.value).strip().lower() if c.value is not None else "" for c in ws[1]]

    if "level" not in headers:
        return None
    if "text" in headers:
        idx_text = headers.index("text")
    elif "正文" in headers:
        idx_text = headers.index("正文")
    else:
        return None

    idx_level = headers.index("level")
    idx_h1 = headers.index("h1") if "h1" in headers else None
    idx_h2 = headers.index("h2") if "h2" in headers else None
    idx_chapter = None
    for cand in (
        "normal",
        "chapter",
        "chapter_no",
        "chapter no",
        "章节号",
        "章節號",
        "章节編號",
        "章節編號",
        "章节",
        "章節",
        "序号",
        "序號",
    ):
        if cand in headers:
            idx_chapter = headers.index(cand)
            break

    units: List[Unit] = []
    uid = 0
    for row in ws.iter_rows(min_row=2, values_only=True):
        if row is None:
            continue
        txt = row[idx_text] if idx_text < len(row) else ""
        txt = "" if txt is None else str(txt)
        if not txt.strip():
            continue

        lvl = row[idx_level] if idx_level < len(row) else "normal"
        h1 = row[idx_h1] if (idx_h1 is not None and idx_h1 < len(row)) else ""
        h2 = row[idx_h2] if (idx_h2 is not None and idx_h2 < len(row)) else ""
        chapter = row[idx_chapter] if (idx_chapter is not None and idx_chapter < len(row)) else ""
        units.append(
            Unit(
                uid=uid,
                level=normalize_style(lvl),
                h1="" if h1 is None else str(h1).strip(),
                h2="" if h2 is None else str(h2).strip(),
                text=txt,
                para_ids=[uid],
                normal_chapter=_normalize_chapter_value(chapter),
                ngram_segments=[txt],
            )
        )
        uid += 1
    return units


def read_excel_as_source_table(path: str) -> Optional[List[Dict[str, str]]]:
    """
    Source 表（同 pipeline12）：
    支持标题如：书名/h1、篇名/h2、正文/text。
    """
    wb = openpyxl.load_workbook(path, data_only=True)
    ws = wb.active
    raw_headers = [str(c.value).strip() if c.value is not None else "" for c in ws[1]]
    low_headers = [h.lower() for h in raw_headers]

    def find_idx(names: List[str]) -> Optional[int]:
        for n in names:
            if n in raw_headers:
                return raw_headers.index(n)
            low = n.lower()
            if low in low_headers:
                return low_headers.index(low)
        return None

    idx_text = find_idx(["正文", "text"])
    idx_h1 = find_idx(["书名", "h1"])
    idx_h2 = find_idx(["篇名", "h2"])
    idx_chapter = find_idx([
        "章节号",
        "章節號",
        "章节編號",
        "章節編號",
        "章节",
        "章節",
        "normal",
        "chapter",
        "chapter_no",
        "chapter no",
        "序号",
        "序號",
    ])
    if idx_text is None or idx_h1 is None or idx_h2 is None:
        return None

    rows: List[Dict[str, str]] = []
    for ridx, row in enumerate(ws.iter_rows(min_row=2, values_only=True), start=2):
        if row is None:
            continue
        txt = row[idx_text] if idx_text < len(row) else ""
        txt = "" if txt is None else str(txt)
        if not txt.strip():
            continue
        h1 = row[idx_h1] if idx_h1 < len(row) else ""
        h2 = row[idx_h2] if idx_h2 < len(row) else ""
        chapter = row[idx_chapter] if (idx_chapter is not None and idx_chapter < len(row)) else ""
        rows.append(
            {
                "h1": "" if h1 is None else str(h1).strip(),
                "h2": "" if h2 is None else str(h2).strip(),
                "text": txt,
                "normal_chapter": _normalize_chapter_value(chapter),
                "row_id": str(ridx),
            }
        )
    return rows


def read_excel_as_paragraphs_table(path: str) -> List[Para]:
    """
    Paragraph 表（同 pipeline12）：
    A列=text，B列=style(h1/h2/normal)。
    """
    wb = openpyxl.load_workbook(path, data_only=True)
    ws = wb.active

    rows = list(ws.iter_rows(min_row=1, values_only=True))
    start_idx = 0
    if rows:
        c1 = str(rows[0][0]).strip().lower() if len(rows[0]) >= 1 and rows[0][0] is not None else ""
        c2 = str(rows[0][1]).strip().lower() if len(rows[0]) >= 2 and rows[0][1] is not None else ""
        if c1 in ("text", "正文", "content") and c2 in ("style", "level", "类型", "樣式", "样式"):
            start_idx = 1

    paras: List[Para] = []
    pid = 0
    for row in rows[start_idx:]:
        if row is None:
            continue
        txt = row[0] if len(row) >= 1 else ""
        sty = row[1] if len(row) >= 2 else "normal"
        txt = "" if txt is None else str(txt)
        if not txt.strip():
            continue
        paras.append(Para(pid=pid, text=txt, style=normalize_style(sty)))
        pid += 1
    return paras


def build_units_from_source_rows(rows: List[Dict[str, str]], granularity: str) -> List[Unit]:
    normal_units: List[Unit] = []
    for i, row in enumerate(rows):
        normal_units.append(
            Unit(
                uid=i,
                level="normal",
                h1=row.get("h1", ""),
                h2=row.get("h2", ""),
                text=row.get("text", ""),
                para_ids=[i],
                normal_chapter=str(row.get("normal_chapter", "") or ""),
                ngram_segments=[str(row.get("text", "") or "")],
            )
        )
    return _aggregate_units_from_normal(normal_units, granularity)


def segment_units_from_paragraphs(paras: List[Para], granularity: str) -> List[Unit]:
    g = normalize_style(granularity)
    units: List[Unit] = []
    cur_h1 = ""
    cur_h2 = ""

    if g == "normal":
        uid = 0
        for p in paras:
            if p.style == "h1":
                cur_h1 = p.text.strip()
                cur_h2 = ""
                continue
            if p.style == "h2":
                cur_h2 = p.text.strip()
                continue
            if p.style == "h3":
                # h3 is heading-only and must not enter normal-body similarity.
                continue
            txt = p.text.strip()
            if not txt:
                continue
            # Strict policy: only正文 under an existing heading context enters computation.
            if not (cur_h1 or cur_h2):
                continue
            units.append(
                Unit(
                    uid=uid,
                    level="normal",
                    h1=cur_h1,
                    h2=cur_h2,
                    text=txt,
                    para_ids=[p.pid],
                    normal_chapter="",
                    ngram_segments=[txt],
                )
            )
            uid += 1
        return units

    uid = 0
    current: Optional[Unit] = None

    def ensure_current() -> Unit:
        nonlocal uid, current
        if current is None:
            current = Unit(
                uid=uid,
                level=g,
                h1=cur_h1,
                h2=cur_h2,
                text="",
                para_ids=[],
                normal_chapter="",
                ngram_segments=[],
            )
            uid += 1
        return current

    def finalize_current():
        nonlocal current
        if current is not None and current.text.strip():
            if not current.ngram_segments:
                current.ngram_segments = [str(current.text or "")]
            units.append(current)
        current = None

    for p in paras:
        if p.style == "h1":
            finalize_current()
            cur_h1 = p.text.strip()
            cur_h2 = ""
            continue

        if p.style == "h2":
            if g == "h2":
                finalize_current()
            cur_h2 = p.text.strip()
            continue

        if p.style == "h3":
            # h3 is heading-only; excluded from body content.
            continue

        txt = p.text.strip()
        if not txt:
            continue
        # Strict policy: h1/h2 granularity requires corresponding heading context.
        if g == "h1" and not cur_h1:
            continue
        if g == "h2" and not cur_h2:
            continue
        u = ensure_current()
        u.text += (txt + "\n")
        u.para_ids.append(p.pid)

    finalize_current()
    return units


def select_units_from_units_table(units: List[Unit], granularity: str) -> List[Unit]:
    g = normalize_style(granularity)
    normal_units = [u for u in units if normalize_style(u.level) == "normal"]
    if g in ("h1", "h2"):
        # h1/h2 计算口径：仅由 normal 正文聚合，标题行本身不参与计算。
        return _aggregate_units_from_normal(normal_units, g)
    if g == "normal":
        return _aggregate_units_from_normal(normal_units, "normal")
    return []


def build_units_from_excel(path: str, granularity: str) -> Tuple[List[Unit], str]:
    units_tbl = read_excel_as_units_table(path)
    if units_tbl is not None:
        return select_units_from_units_table(units_tbl, granularity), "Units table"

    source_rows = read_excel_as_source_table(path)
    if source_rows is not None:
        return build_units_from_source_rows(source_rows, granularity), "Source table(书名/篇名/正文)"

    paras = read_excel_as_paragraphs_table(path)
    return segment_units_from_paragraphs(paras, granularity), "Paragraphs table(A=text,B=style)"


def _split_cn_title_and_tail(s: str) -> Tuple[Optional[str], str]:
    t = str(s or "").strip()
    if not t:
        return None, ""
    m = re.match(r"^\s*《([^》]+)》\s*(.*?)\s*$", t)
    if not m:
        return None, t
    return (m.group(1) or "").strip(), (m.group(2) or "").strip()


def _merge_book_article_path(h1: str, h2: str) -> str:
    b = str(h1 or "").strip()
    p = str(h2 or "").strip()
    if not b and not p:
        return ""
    if not b:
        return p
    if not p:
        return b

    bt, bs = _split_cn_title_and_tail(b)
    pt, ps = _split_cn_title_and_tail(p)

    if bt and pt:
        merged = f"《{bt}·{pt}》"
        tail = f"{bs}{ps}".strip()
        return f"{merged}{tail}" if tail else merged

    return f"{b}·{p}"


def _normal_chapter_no(u: Unit) -> str:
    user_chapter = str(getattr(u, "normal_chapter", "") or "").strip()
    if user_chapter:
        return user_chapter
    pos = (u.para_ids[0] + 1) if (u.para_ids and len(u.para_ids) > 0) else (u.uid + 1)
    return f"#{int(pos)}"


def unit_path(u: Unit) -> str:
    g = normalize_style(u.level)
    h1 = str(u.h1 or "").strip()
    h2 = str(u.h2 or "").strip()
    merged_h1_h2 = _merge_book_article_path(h1, h2)

    if g == "h1":
        path = h1 or h2
    elif g == "h2":
        path = merged_h1_h2 or h1 or h2
    else:
        # normal path follows pipeline14-like style and keeps chapter index.
        chapter = _normal_chapter_no(u)
        if merged_h1_h2:
            path = f"{merged_h1_h2} / {chapter}"
        elif h1 or h2:
            path = f"{(h1 or h2)} / {chapter}"
        else:
            path = chapter
    return path if path else f"unit_{int(u.uid)}"


def unit_display_name(u: Unit) -> str:
    return unit_path(u)


def generate_ngrams_from_tokens(tokens: List[str], n: int) -> List[str]:
    if len(tokens) < n:
        return []
    return [''.join(tokens[i:i+n]) for i in range(len(tokens)-n+1)]


def _unit_text_segments_for_ngrams(u: Unit) -> List[str]:
    segs = [str(s) for s in list(getattr(u, "ngram_segments", []) or []) if str(s or "").strip()]
    if segs:
        return segs
    txt = str(getattr(u, "text", "") or "")
    return [txt] if txt.strip() else []


def _build_lexical_inputs_by_segments(units: List[Unit], n: int) -> Tuple[List[str], List[List[str]], List[List[str]]]:
    """
    Scheme-1 boundary policy:
    - generate n-grams inside each segment only
    - merge segment n-gram counts as one unit vector
    """
    doc_texts_clean: List[str] = []
    doc_tokens_list: List[List[str]] = []
    all_ngram_tokens: List[List[str]] = []

    for u in units:
        segs_raw = _unit_text_segments_for_ngrams(u)
        segs_clean = [preprocess_keep_han(seg) for seg in segs_raw]
        segs_clean = [x for x in segs_clean if x]

        merged_clean = "".join(segs_clean)
        doc_texts_clean.append(merged_clean)
        doc_tokens_list.append([ch for ch in merged_clean])

        ngram_tokens: List[str] = []
        for seg in segs_clean:
            seg_tokens = [ch for ch in seg]
            ngram_tokens.extend(generate_ngrams_from_tokens(seg_tokens, n))
        all_ngram_tokens.append(ngram_tokens)

    return doc_texts_clean, doc_tokens_list, all_ngram_tokens


def _compute_jaccard_matrix(all_ngram_tokens: List[List[str]]) -> np.ndarray:
    n = len(all_ngram_tokens)
    mats = np.eye(n, dtype=float)
    token_sets = [set(ts) for ts in all_ngram_tokens]
    for i in range(n):
        si = token_sets[i]
        for j in range(i + 1, n):
            sj = token_sets[j]
            union = len(si | sj)
            score = (len(si & sj) / union) if union > 0 else 0.0
            mats[i, j] = score
            mats[j, i] = score
    return mats


def _normalize_embed_size_label(model_size: str) -> str:
    s = str(model_size or "").strip().upper().replace("Ｂ", "B")
    alias = {
        "0.6": "0.6B",
        "0.6B": "0.6B",
        "4": "4B",
        "4B": "4B",
        "8": "8B",
        "8B": "8B",
        "0.8B": "8B",
    }
    return alias.get(s, s if s else "0.6B")


def _normalize_reranker_size_label(model_size: str) -> str:
    s = str(model_size or "").strip().upper().replace("Ｂ", "B")
    alias = {
        "0.6": "0.6B",
        "0.6B": "0.6B",
        "4": "4B",
        "4B": "4B",
        "8": "8B",
        "8B": "8B",
        "0.8B": "8B",
    }
    return alias.get(s, s if s else "0.6B")


def _fmt_elapsed(seconds: float) -> str:
    s = max(0.0, float(seconds))
    if s < 60.0:
        return f"{s:.1f}s"
    total = int(round(s))
    m, sec = divmod(total, 60)
    h, m = divmod(m, 60)
    if h > 0:
        return f"{h}h{m:02d}m{sec:02d}s"
    return f"{m}m{sec:02d}s"


def _embed_model_label(model_size: str) -> str:
    return f"Qwen3-Embedding-{_normalize_embed_size_label(model_size)}"


def _reranker_model_label(model_size: str) -> str:
    return f"Qwen3-Reranker-{_normalize_reranker_size_label(model_size)}"


def _quiet_transformers_logging() -> None:
    try:
        from transformers.utils import logging as hf_logging
        hf_logging.set_verbosity_error()
    except Exception:
        pass


def _ensure_padding_token(tokenizer, model=None) -> bool:
    pad_id = getattr(tokenizer, "pad_token_id", None)
    if pad_id is None:
        eos_tok = getattr(tokenizer, "eos_token", None)
        if eos_tok is not None:
            try:
                tokenizer.pad_token = eos_tok
            except Exception:
                pass
        pad_id = getattr(tokenizer, "pad_token_id", None)
    if pad_id is None:
        eos_id = getattr(tokenizer, "eos_token_id", None)
        if eos_id is not None:
            try:
                tokenizer.pad_token_id = int(eos_id)
            except Exception:
                pass
        pad_id = getattr(tokenizer, "pad_token_id", None)
    if pad_id is not None and model is not None:
        try:
            if getattr(model.config, "pad_token_id", None) is None:
                model.config.pad_token_id = int(pad_id)
        except Exception:
            pass
    return pad_id is not None


def _resolve_hf_snapshot_dir(model_cache_dir: str) -> str:
    p = pathlib.Path(model_cache_dir)
    if not p.exists() or not p.is_dir():
        raise RuntimeError(f"Model dir not found: {model_cache_dir}")
    snaps = p / "snapshots"
    if not snaps.exists() or not snaps.is_dir():
        return str(p)
    ref_main = p / "refs" / "main"
    if ref_main.exists():
        try:
            commit = ref_main.read_text(encoding="utf-8").strip()
            cand = snaps / commit
            if cand.exists() and cand.is_dir():
                return str(cand)
        except Exception:
            pass
    subdirs = sorted([x for x in snaps.iterdir() if x.is_dir()], key=lambda x: x.stat().st_mtime, reverse=True)
    if subdirs:
        return str(subdirs[0])
    return str(p)


def _rank_order(values: List[Tuple[int, float]], descending: bool = True) -> Dict[int, int]:
    if not values:
        return {}
    vals = sorted(values, key=lambda x: x[1], reverse=descending)
    out: Dict[int, int] = {}
    rank = 0
    prev = None
    for pos, (idx, v) in enumerate(vals, start=1):
        if prev is None or v != prev:
            rank = pos
        out[int(idx)] = int(rank)
        prev = v
    return out


def _annotate_metric_rank(rows: List[Dict[str, Any]], score_key: str, rank_key: str, descending: bool = True) -> None:
    vals: List[Tuple[int, float]] = []
    for i, r in enumerate(rows):
        v = r.get(score_key)
        if v is None:
            continue
        try:
            x = float(v)
        except Exception:
            continue
        if math.isnan(x) or math.isinf(x):
            continue
        vals.append((i, x))
    rank_map = _rank_order(vals, descending=descending)
    for i, r in enumerate(rows):
        r[rank_key] = int(rank_map[i]) if i in rank_map else None


def _embed_text_for_semantic(text: str) -> str:
    # Keep preprocessing consistent across lexical/semantic/reranker:
    # remove punctuation/symbols/digits/Latin and keep Han only.
    return preprocess_keep_han(text or "")


def _load_semantic_backend(model_size: str) -> dict:
    model_size = _normalize_embed_size_label(model_size or DEFAULT_EMBED_MODEL_SIZE)
    if model_size not in QWEN_EMBED_MODEL_DIRS:
        raise RuntimeError(f"Unsupported embedding model size: {model_size}")
    model_dir = _resolve_hf_snapshot_dir(QWEN_EMBED_MODEL_DIRS[model_size])
    cache_key = f"{model_size}:{model_dir}"
    if cache_key in _SEMANTIC_BACKEND_CACHE:
        return _SEMANTIC_BACKEND_CACHE[cache_key]

    _quiet_transformers_logging()
    try:
        import torch
        from transformers import AutoTokenizer, AutoModel
    except Exception as e:
        raise RuntimeError(f"Semantic backend requires torch+transformers: {e}")

    if torch.cuda.is_available():
        device = "cuda"
    elif getattr(torch.backends, "mps", None) is not None and torch.backends.mps.is_available():
        device = "mps"
    else:
        device = "cpu"

    tokenizer = AutoTokenizer.from_pretrained(model_dir, trust_remote_code=True, local_files_only=True)
    _ensure_padding_token(tokenizer)
    load_kwargs = {"trust_remote_code": True, "local_files_only": True}
    if device == "cuda":
        load_kwargs["torch_dtype"] = torch.float16

    try:
        model = AutoModel.from_pretrained(model_dir, **load_kwargs)
    except Exception:
        model = AutoModel.from_pretrained(model_dir, trust_remote_code=True, local_files_only=True)
    _ensure_padding_token(tokenizer, model=model)
    model.to(device)
    model.eval()

    backend = {
        "torch": torch,
        "tokenizer": tokenizer,
        "model": model,
        "device": device,
        "model_size": model_size,
    }
    _SEMANTIC_BACKEND_CACHE[cache_key] = backend
    return backend


def _encode_semantic_embeddings(
    texts: List[str],
    model_size: str,
    batch_size: int = 8,
    max_length: int = 512,
) -> np.ndarray:
    if not texts:
        return np.zeros((0, 1), dtype=np.float32)
    backend = _load_semantic_backend(model_size)
    torch = backend["torch"]
    tokenizer = backend["tokenizer"]
    model = backend["model"]
    device = backend["device"]

    vecs: List[np.ndarray] = []
    bs = int(max(1, batch_size))
    if getattr(tokenizer, "pad_token_id", None) is None or getattr(getattr(model, "config", None), "pad_token_id", None) is None:
        bs = 1
    ml = int(max(64, max_length))

    for i in range(0, len(texts), bs):
        batch = texts[i: i + bs]
        with torch.no_grad():
            toks = tokenizer(
                batch,
                padding=(len(batch) > 1),
                truncation=True,
                max_length=ml,
                return_tensors="pt",
            )
            toks = {k: v.to(device) for k, v in toks.items()}
            out = model(**toks)
            pooler = None
            last_hidden = None
            if isinstance(out, dict):
                pooler = out.get("pooler_output")
                last_hidden = out.get("last_hidden_state")
            else:
                pooler = getattr(out, "pooler_output", None)
                last_hidden = getattr(out, "last_hidden_state", None)
                if last_hidden is None and isinstance(out, (tuple, list)) and len(out) > 0:
                    last_hidden = out[0]

            if pooler is not None:
                emb = pooler
            else:
                if last_hidden is None:
                    raise RuntimeError("Embedding model output missing hidden states.")
                attn = toks.get("attention_mask")
                if attn is None:
                    emb = last_hidden[:, 0]
                else:
                    mask = attn.unsqueeze(-1).expand_as(last_hidden).float()
                    emb = (last_hidden * mask).sum(dim=1) / torch.clamp(mask.sum(dim=1), min=1e-9)

            emb = torch.nn.functional.normalize(emb, p=2, dim=1)
            vecs.append(emb.detach().float().cpu().numpy().astype(np.float32, copy=False))

    return np.vstack(vecs) if vecs else np.zeros((0, 1), dtype=np.float32)


def _split_semantic_chunks(
    text: str,
    chunk_size: int = SEMANTIC_MAXSIM_CHUNK_SIZE,
    stride: int = SEMANTIC_MAXSIM_CHUNK_STRIDE,
) -> List[str]:
    # `text` is expected to be preprocessed (Han-only) by caller.
    t = str(text or "")
    if not t:
        return []
    c = int(max(80, chunk_size))
    s = int(max(1, stride))
    if s > c:
        s = c
    if len(t) <= c:
        return [t]
    out: List[str] = []
    i = 0
    n = len(t)
    while i < n:
        seg = t[i: i + c]
        if seg:
            out.append(seg)
        if i + c >= n:
            break
        i += s
    return out


def _l2_normalize_vec(v: np.ndarray) -> np.ndarray:
    arr = np.asarray(v, dtype=np.float32)
    den = float(np.linalg.norm(arr))
    if den <= 1e-12:
        return arr.astype(np.float32, copy=False)
    return (arr / den).astype(np.float32, copy=False)


def _maxsim_bidirectional(a: np.ndarray, b: np.ndarray) -> float:
    if a is None or b is None:
        return 0.0
    A = np.asarray(a, dtype=np.float32)
    B = np.asarray(b, dtype=np.float32)
    if A.ndim != 2 or B.ndim != 2 or A.shape[0] <= 0 or B.shape[0] <= 0:
        return 0.0
    sim = np.matmul(A, B.T)
    if sim.size <= 0:
        return 0.0
    row_max = np.max(sim, axis=1)
    col_max = np.max(sim, axis=0)
    score = 0.5 * (float(np.mean(row_max)) + float(np.mean(col_max)))
    return float(max(0.0, min(1.0, score)))


def _compute_embed_similarity_matrix(
    units: List[Unit],
    granularity: str,
    model_size: str,
    unit_texts_clean: Optional[List[str]] = None,
) -> Tuple[np.ndarray, str]:
    n = len(units)
    if n <= 0:
        return np.zeros((0, 0), dtype=float), ("maxsim" if normalize_style(granularity) == "h1" else "avgpool")
    g = normalize_style(granularity)
    clean_texts = list(unit_texts_clean or [])
    if len(clean_texts) != n:
        clean_texts = [_embed_text_for_semantic(u.text or "") for u in units]

    if g == "h1":
        all_chunks: List[str] = []
        spans: List[Tuple[int, int]] = []
        for txt_clean in clean_texts:
            chunks = _split_semantic_chunks(txt_clean or "")
            if not chunks:
                fb = str(txt_clean or "")
                chunks = [fb] if fb else ["空"]
            st = len(all_chunks)
            all_chunks.extend(chunks)
            spans.append((st, len(chunks)))
        chunk_emb = _encode_semantic_embeddings(all_chunks, model_size=model_size)
        d = int(chunk_emb.shape[1]) if chunk_emb.ndim == 2 else 1
        if d <= 0:
            d = 1
        chunk_map: Dict[int, np.ndarray] = {}
        for uid, (st, cnt) in enumerate(spans):
            arr = chunk_emb[st: st + cnt] if chunk_emb.ndim == 2 and chunk_emb.shape[0] >= st + cnt else np.zeros((1, d), dtype=np.float32)
            if arr.ndim != 2 or arr.shape[0] <= 0:
                arr = np.zeros((1, d), dtype=np.float32)
            chunk_map[uid] = arr.astype(np.float32, copy=False)

        mat = np.eye(n, dtype=float)
        for i in range(n):
            for j in range(i + 1, n):
                v = _maxsim_bidirectional(chunk_map.get(i), chunk_map.get(j))
                mat[i, j] = v
                mat[j, i] = v
        return mat, "maxsim"

    texts = [str(t or "") for t in clean_texts]
    embs = _encode_semantic_embeddings(texts, model_size=model_size)
    if embs.ndim != 2 or embs.shape[0] != n:
        return np.eye(n, dtype=float), "avgpool"
    mat = np.matmul(embs, embs.T).astype(float, copy=False)
    mat = np.clip(mat, 0.0, 1.0)
    np.fill_diagonal(mat, 1.0)
    return mat, "avgpool"


def _load_reranker_backend(model_size: str) -> dict:
    model_size = _normalize_reranker_size_label(model_size or DEFAULT_RERANK_MODEL_SIZE)
    if model_size not in QWEN_RERANK_MODEL_DIRS:
        raise RuntimeError(f"Unsupported reranker model size: {model_size}")
    model_key = model_size
    try:
        model_dir = _resolve_hf_snapshot_dir(QWEN_RERANK_MODEL_DIRS[model_key])
    except Exception:
        fallback_key = "0.6B"
        if model_key != fallback_key:
            model_key = fallback_key
            model_dir = _resolve_hf_snapshot_dir(QWEN_RERANK_MODEL_DIRS[model_key])
        else:
            raise
    cache_key = f"{model_size}:{model_dir}"
    if cache_key in _RERANK_BACKEND_CACHE:
        return _RERANK_BACKEND_CACHE[cache_key]

    _quiet_transformers_logging()
    try:
        import torch
        from transformers import AutoTokenizer, AutoModelForSequenceClassification
    except Exception as e:
        raise RuntimeError(f"Reranker backend requires torch+transformers: {e}")

    if torch.cuda.is_available():
        device = "cuda"
    elif getattr(torch.backends, "mps", None) is not None and torch.backends.mps.is_available():
        device = "mps"
    else:
        device = "cpu"

    tokenizer = AutoTokenizer.from_pretrained(model_dir, trust_remote_code=True, local_files_only=True)
    load_kwargs = {"trust_remote_code": True, "local_files_only": True}
    if device == "cuda":
        load_kwargs["torch_dtype"] = torch.float16
    try:
        model = AutoModelForSequenceClassification.from_pretrained(model_dir, **load_kwargs)
    except Exception:
        model = AutoModelForSequenceClassification.from_pretrained(model_dir, trust_remote_code=True, local_files_only=True)
    _ensure_padding_token(tokenizer, model=model)
    model.to(device)
    model.eval()

    backend = {
        "torch": torch,
        "tokenizer": tokenizer,
        "model": model,
        "device": device,
        "model_size": model_key,
    }
    _RERANK_BACKEND_CACHE[cache_key] = backend
    return backend


def _truncate_for_rerank(text: str, max_chars: int = 1200) -> str:
    # `text` is expected to be preprocessed (Han-only) by caller.
    t = str(text or "")
    m = int(max(64, max_chars))
    return t if len(t) <= m else t[:m]


def _score_reranker_pairs(
    pairs: List[Tuple[str, str]],
    model_size: str,
    batch_size: int = 8,
    max_length: int = 512,
) -> Tuple[List[float], str]:
    backend = _load_reranker_backend(model_size)
    torch = backend["torch"]
    tokenizer = backend["tokenizer"]
    model = backend["model"]
    device = backend["device"]
    model_used = str(backend.get("model_size") or model_size)
    if not pairs:
        return [], model_used

    out_scores: List[float] = []
    bs = int(max(1, batch_size))
    if getattr(tokenizer, "pad_token_id", None) is None or getattr(model.config, "pad_token_id", None) is None:
        bs = 1
    ml = int(max(64, max_length))

    for i in range(0, len(pairs), bs):
        chunk = pairs[i: i + bs]
        a_texts = [x[0] for x in chunk]
        b_texts = [x[1] for x in chunk]
        with torch.no_grad():
            toks = tokenizer(
                a_texts,
                b_texts,
                padding=(len(chunk) > 1),
                truncation=True,
                max_length=ml,
                return_tensors="pt",
            )
            toks = {k: v.to(device) for k, v in toks.items()}
            out = model(**toks)
            logits = None
            if isinstance(out, dict):
                logits = out.get("logits")
            else:
                logits = getattr(out, "logits", None)
                if logits is None and isinstance(out, (tuple, list)) and len(out) > 0:
                    logits = out[0]
            if logits is None:
                raise RuntimeError("Reranker output missing logits.")

            if logits.dim() == 1:
                scores = logits
            elif logits.size(-1) == 1:
                scores = logits.squeeze(-1)
            else:
                scores = logits[:, 0]
            out_scores.extend([float(x) for x in scores.detach().float().cpu().tolist()])
    return out_scores, model_used


def _estimate_reranker_candidate_count(total_rows: int, granularity: str, top_percent_non_h1: float = RERANK_TOP_PERCENT_NON_H1) -> int:
    n_rows = int(max(0, int(total_rows)))
    if n_rows <= 0:
        return 0
    g = normalize_style(granularity)
    # h1/h2 always use full reranker coverage.
    if g in ("h1", "h2"):
        return n_rows
    pct = float(top_percent_non_h1)
    if pct <= 0:
        pct = RERANK_TOP_PERCENT_NON_H1
    n_sel = int(max(1, round(n_rows * (pct / 100.0))))
    return int(max(1, min(n_sel, n_rows)))


def _effective_reranker_percent_non_h1(
    total_rows: int,
    granularity: str,
    top_percent_non_h1: float,
) -> Tuple[float, Dict[str, Any]]:
    n_rows = int(max(0, int(total_rows)))
    g = normalize_style(granularity)
    try:
        cfg_pct = float(top_percent_non_h1)
    except Exception:
        cfg_pct = float(RERANK_TOP_PERCENT_NON_H1)
    if cfg_pct <= 0:
        cfg_pct = float(RERANK_TOP_PERCENT_NON_H1)
    cfg_pct = max(0.001, min(100.0, cfg_pct))

    out = {
        "configured_percent": float(cfg_pct),
        "effective_percent": float(cfg_pct),
        "auto_scaled": False,
        "max_candidates_cap": int(max(1, int(RERANK_AUTO_NON_H1_MAX_CANDIDATES))),
        "total_rows": int(n_rows),
    }

    # h1/h2 keep full reranker by design.
    if g in ("h1", "h2"):
        out["effective_percent"] = 100.0
        return 100.0, out
    if n_rows <= 0:
        return float(cfg_pct), out

    if g != "normal":
        return float(cfg_pct), out

    if not bool(RERANK_AUTO_SCALE_NON_H1):
        return float(cfg_pct), out

    cap = int(max(1, int(RERANK_AUTO_NON_H1_MAX_CANDIDATES)))
    if n_rows <= cap:
        return float(cfg_pct), out

    pct_cap = 100.0 * float(cap) / float(n_rows)
    pct_cap = max(float(RERANK_AUTO_NON_H1_MIN_PERCENT), min(100.0, pct_cap))
    if pct_cap < cfg_pct:
        out["effective_percent"] = float(pct_cap)
        out["auto_scaled"] = True
        return float(pct_cap), out

    return float(cfg_pct), out


def _annotate_reranker_for_pairs(
    rows: List[Dict[str, Any]],
    units: List[Unit],
    granularity: str,
    model_size: str,
    top_percent_non_h1: float = RERANK_TOP_PERCENT_NON_H1,
    unit_texts_clean: Optional[List[str]] = None,
) -> Optional[str]:
    for r in rows:
        r["reranker_raw"] = None
        r["rank_reranker_raw"] = None
        r["rerank_applied"] = 0
        r["reranker_model"] = None

    if not rows or not units:
        return None
    clean_texts = list(unit_texts_clean or [])
    if len(clean_texts) != len(units):
        clean_texts = [_embed_text_for_semantic(u.text or "") for u in units]

    n_sel = _estimate_reranker_candidate_count(len(rows), granularity=granularity, top_percent_non_h1=top_percent_non_h1)
    order = sorted(range(len(rows)), key=lambda idx: float(rows[idx].get("cos_raw") or 0.0), reverse=True)
    sel = order[:n_sel]
    pairs: List[Tuple[str, str]] = []
    for idx in sel:
        r = rows[idx]
        id1 = int(r.get("id1", 0)) - 1
        id2 = int(r.get("id2", 0)) - 1
        if id1 < 0 or id2 < 0 or id1 >= len(units) or id2 >= len(units):
            pairs.append(("", ""))
            continue
        pairs.append(
            (
                _truncate_for_rerank(clean_texts[id1] if id1 < len(clean_texts) else "", max_chars=1200),
                _truncate_for_rerank(clean_texts[id2] if id2 < len(clean_texts) else "", max_chars=1200),
            )
        )

    scores, model_used = _score_reranker_pairs(pairs, model_size=model_size)
    if len(scores) != len(sel):
        raise RuntimeError(f"reranker score size mismatch: got={len(scores)} expect={len(sel)}")
    for i, ridx in enumerate(sel):
        rows[ridx]["reranker_raw"] = float(scores[i])
        rows[ridx]["rerank_applied"] = 1
        rows[ridx]["reranker_model"] = str(model_used)

    _annotate_metric_rank(rows, "reranker_raw", "rank_reranker_raw", descending=True)
    return str(model_used)


def _candidate_policy_for_granularity(
    granularity: str,
    n_units: int,
    normal_topk: int = DEFAULT_TOPK_CANDIDATES,
) -> Tuple[str, Optional[int]]:
    g = normalize_style(granularity)
    if g == "h1":
        return "all", None
    if g == "h2":
        return "all", None
    # normal / fallback
    return "topk", int(max(1, int(normal_topk)))


def build_candidate_pairs(
    doc_names: List[str],
    sim_cos_raw: np.ndarray,
    sim_tfidf: np.ndarray,
    sim_jaccard: np.ndarray,
    sim_embed: Optional[np.ndarray] = None,
    candidate_mode: str = "all",
    topk: int = DEFAULT_TOPK_CANDIDATES,
) -> List[Dict[str, float]]:
    rows: List[Dict[str, float]] = []
    n = len(doc_names)
    mode = str(candidate_mode or "all").strip().lower()
    if mode not in ("all", "topk"):
        mode = "all"

    def _append_pair(i: int, j: int):
        emb_v: Optional[float] = None
        if isinstance(sim_embed, np.ndarray) and sim_embed.ndim == 2 and sim_embed.shape[0] > i and sim_embed.shape[1] > j:
            emb_v = float(sim_embed[i, j])
        rows.append(
            {
                "id1": i + 1,
                "id2": j + 1,
                "path1": str(doc_names[i]),
                "path2": str(doc_names[j]),
                # legacy aliases kept for backward compatibility
                "name1": str(doc_names[i]),
                "name2": str(doc_names[j]),
                "cos_raw": float(sim_cos_raw[i, j]),
                "tfidf": float(sim_tfidf[i, j]),
                "jaccard": float(sim_jaccard[i, j]),
                "embed": emb_v,
                "reranker_raw": None,
            }
        )

    if mode == "all" or n <= 2:
        for i in range(n):
            for j in range(i + 1, n):
                _append_pair(i, j)
    else:
        k = int(max(1, topk))
        seen: set = set()
        for i in range(n):
            row = np.asarray(sim_cos_raw[i], dtype=float).copy()
            if row.size != n:
                continue
            row[i] = -np.inf
            k_eff = int(min(k, max(0, n - 1)))
            if k_eff <= 0:
                continue
            if k_eff >= n - 1:
                idxs = np.argsort(-row)
            else:
                idxs = np.argpartition(-row, k_eff)[:k_eff]
                idxs = idxs[np.argsort(-row[idxs])]
            for j in idxs:
                jj = int(j)
                if jj == i:
                    continue
                a, b = (i, jj) if i < jj else (jj, i)
                key = (a, b)
                if key in seen:
                    continue
                seen.add(key)
                _append_pair(a, b)

    rows.sort(key=lambda r: (r["cos_raw"], r["tfidf"], r["jaccard"]), reverse=True)
    for rank, r in enumerate(rows, start=1):
        r["rank"] = rank
        r["rank_cos_raw"] = rank
    _annotate_metric_rank(rows, "tfidf", "rank_tfidf", descending=True)
    _annotate_metric_rank(rows, "jaccard", "rank_jaccard", descending=True)
    _annotate_metric_rank(rows, "embed", "rank_embed", descending=True)
    _annotate_metric_rank(rows, "reranker_raw", "rank_reranker_raw", descending=True)
    return rows


def compute_similarity_from_ngrams(all_ngram_tokens: List[List[str]]):
    texts_for_vectorizer = [' '.join(tokens) for tokens in all_ngram_tokens]
    if not any(t.strip() for t in texts_for_vectorizer):
        raise ValueError("No n-gram tokens available for similarity computation.")
    vectorizer = CountVectorizer(tokenizer=lambda x: x.split(), token_pattern=None, lowercase=False)
    X_raw = vectorizer.fit_transform(texts_for_vectorizer)
    feature_names = vectorizer.get_feature_names_out()
    doc_vectors = X_raw.toarray()
    sim_cos_raw = cosine_similarity(X_raw)

    tfidf_vectorizer = TfidfVectorizer(tokenizer=lambda x: x.split(), token_pattern=None, lowercase=False)
    X_tfidf = tfidf_vectorizer.fit_transform(texts_for_vectorizer)
    sim_tfidf = cosine_similarity(X_tfidf)

    sim_jaccard = _compute_jaccard_matrix(all_ngram_tokens)

    return {
        "cos_raw": sim_cos_raw,
        "tfidf": sim_tfidf,
        "jaccard": sim_jaccard,
        "feature_names": feature_names,
        "doc_vectors": doc_vectors,
        "vectorizer": vectorizer,
    }


def _split_csv_values(text: str) -> List[str]:
    return [x.strip() for x in str(text or "").split(",") if str(x).strip()]


def _set_cjk_rcparams(primary_family: str = "") -> List[str]:
    families = _unique_nonempty([str(primary_family or "").strip()] + list(_CJK_RC_FALLBACK_FAMILIES))
    try:
        installed = {str(f.name or "").strip() for f in fm.fontManager.ttflist}
    except Exception:
        installed = set()
    if installed:
        filtered = [x for x in families if x in installed]
    else:
        filtered = list(families)
    if not filtered:
        filtered = ["DejaVu Sans"]
    matplotlib.rcParams["axes.unicode_minus"] = False
    matplotlib.rcParams["font.family"] = "sans-serif"
    matplotlib.rcParams["font.sans-serif"] = filtered
    plt.rcParams["axes.unicode_minus"] = False
    plt.rcParams["font.family"] = "sans-serif"
    plt.rcParams["font.sans-serif"] = filtered
    return filtered


def _font_from_file(path: str) -> Optional[fm.FontProperties]:
    p = os.path.expanduser(str(path or "").strip())
    if not p or not os.path.isfile(p):
        return None
    try:
        return fm.FontProperties(fname=p)
    except Exception:
        return None


def _iter_user_font_files() -> List[str]:
    exts = (".ttf", ".ttc", ".otf")
    out: List[str] = []
    for d in USER_DEFAULT_FONT_DIRS:
        dd = os.path.expanduser(str(d or "").strip())
        if not dd or not os.path.isdir(dd):
            continue
        try:
            for name in os.listdir(dd):
                p = os.path.join(dd, name)
                if os.path.isfile(p) and name.lower().endswith(exts):
                    out.append(p)
        except Exception:
            continue
    return out


def _find_user_default_font() -> Optional[fm.FontProperties]:
    files = _iter_user_font_files()
    if not files:
        return None
    # 优先中文常见字体文件名，再退回目录中第一个可用字体。
    keywords = [
        "simhei",
        "simsun",
        "simkai",
        "yahei",
        "pingfang",
        "song",
        "heiti",
        "kaiti",
        "noto",
        "sourcehan",
        "source han",
        "fandol",
    ]
    ranked = sorted(
        files,
        key=lambda p: (
            0 if any(k in os.path.basename(p).lower() for k in keywords) else 1,
            os.path.basename(p).lower(),
        ),
    )
    for p in ranked:
        fp = _font_from_file(p)
        if fp is not None:
            return fp
    return None


def resolve_heatmap_font(font_setting: str = "") -> Optional[fm.FontProperties]:
    # 优先级：
    # 默认指定字体文件 > 环境变量字体文件 > UI输入(可逗号分隔) >
    # （默认空设置时）旧固定目录 > 用户默认字体目录 >
    # （有自定义设置时）旧固定目录 >
    # 环境变量候选/默认候选
    fp = _font_from_file(DEFAULT_HEATMAP_FONT_PATH)
    if fp is not None:
        return fp

    env_font_path = str(os.environ.get("NGRAM_FONT_PATH", "")).strip()
    fp = _font_from_file(env_font_path)
    if fp is not None:
        return fp

    setting_candidates = _split_csv_values(font_setting)
    visited = set()
    for cand in setting_candidates:
        c = str(cand).strip()
        if not c or c in visited:
            continue
        visited.add(c)

        # cand 既可为字体文件路径，也可为字体家族名
        fp = _font_from_file(c)
        if fp is not None:
            return fp
        try:
            path = fm.findfont(fm.FontProperties(family=[c]), fallback_to_default=False)
        except Exception:
            continue
        fp = _font_from_file(path)
        if fp is not None:
            return fp

    if not setting_candidates:
        # 默认行为：优先保留原版图的字体风格
        for fname in LEGACY_TTF_CANDIDATES:
            legacy_path = os.path.join(LEGACY_TTF_DIR, fname)
            fp = _font_from_file(legacy_path)
            if fp is not None:
                return fp

    # 次优先：用户默认字体目录（统一目录管理）
    fp = _find_user_default_font()
    if fp is not None:
        return fp

    if setting_candidates:
        # 用户有自定义设置但不可用时，再尝试旧固定目录
        for fname in LEGACY_TTF_CANDIDATES:
            legacy_path = os.path.join(LEGACY_TTF_DIR, fname)
            fp = _font_from_file(legacy_path)
            if fp is not None:
                return fp

    env_candidates = _split_csv_values(str(os.environ.get("NGRAM_FONT_CANDIDATES", "")))
    fallback_candidates = env_candidates if env_candidates else list(DEFAULT_HEATMAP_FONT_CANDIDATES)
    visited2 = set(visited)
    for cand in fallback_candidates:
        c = str(cand).strip()
        if not c or c in visited2:
            continue
        visited2.add(c)
        fp = _font_from_file(c)
        if fp is not None:
            return fp
        try:
            path = fm.findfont(fm.FontProperties(family=[c]), fallback_to_default=False)
        except Exception:
            continue
        fp = _font_from_file(path)
        if fp is not None:
            return fp
    return None


def build_heatmap_labels(doc_names: List[str], strip_digits: bool = False) -> List[str]:
    raw_labels: List[str] = []
    for name in doc_names:
        label = _image_label_short(name, ascii_dot=True)
        if strip_digits:
            label = re.sub(r"\d+", "", label)
        label = re.sub(r"\s+", " ", str(label or "")).strip()
        if not label:
            label = "untitled"
        raw_labels.append(label)

    # 去重：避免“去数字”后重名导致标签歧义
    seen: Dict[str, int] = {}
    out: List[str] = []
    for lb in raw_labels:
        cnt = seen.get(lb, 0) + 1
        seen[lb] = cnt
        out.append(lb if cnt == 1 else f"{lb} ({cnt})")
    return out


_VERTICAL_PUNCT_MAP = {
    # 避免罕见竖排标点字形（如 ︽︾）在部分字体缺失而显示方块。
    "(": "（",
    ")": "）",
}


def _verticalize_heatmap_label(label: str) -> str:
    s = str(label or "")
    if not s:
        return ""
    # 竖排标签中，将书名号/括号替换为竖排形态；并去掉空格避免出现空行。
    s = "".join(_VERTICAL_PUNCT_MAP.get(ch, ch) for ch in s).replace(" ", "")
    return "\n".join(s)


# =================== 聚類與中心性分析 ===================

_ANALECTS_CHAPTERS = [
    "學而", "爲政", "八佾", "里仁", "公冶長", "雍也", "述而", "泰伯", "子罕", "鄉党",
    "先進", "颜渊", "子路", "憲问", "衛靈公", "季氏", "陽貨", "微子", "子張", "尧曰",
]
_ANALECTS_CHAPTER_SET = set(_ANALECTS_CHAPTERS)

_TEXT_NORMALIZE_CHAR_MAP = str.maketrans(
    {
        "论": "論",
        "语": "語",
        "记": "記",
        "传": "傳",
        "镜": "鏡",
        "选": "選",
        "乡": "鄉",
        "党": "黨",
        "颜": "顏",
        "渊": "淵",
        "宪": "憲",
        "问": "問",
        "卫": "衛",
        "阳": "陽",
        "尧": "堯",
        "为": "爲",
        "為": "爲",
        "﹒": "·",
        "．": "·",
        "•": "·",
        ".": "·",
    }
)

_ANALECTS_CHAPTER_ALIAS = {
    "學而": "學而",
    "爲政": "爲政",
    "八佾": "八佾",
    "里仁": "里仁",
    "公冶長": "公冶長",
    "雍也": "雍也",
    "述而": "述而",
    "泰伯": "泰伯",
    "子罕": "子罕",
    "鄉黨": "鄉党",
    "鄉党": "鄉党",
    "先進": "先進",
    "顏淵": "颜渊",
    "顏渊": "颜渊",
    "颜渊": "颜渊",
    "子路": "子路",
    "憲問": "憲问",
    "憲问": "憲问",
    "衛靈公": "衛靈公",
    "衛灵公": "衛靈公",
    "季氏": "季氏",
    "陽貨": "陽貨",
    "陽货": "陽貨",
    "微子": "微子",
    "子張": "子張",
    "堯曰": "尧曰",
    "尧曰": "尧曰",
}

_NODE_COLOR_MAP: Dict[str, str] = {
    "衣鏡·孔子傳": "#FFC0CB",
    "衣鏡·弟子傳": "#7FFFAA",
    "史記·孔子世家·甲": "#FF4500",
    "史記·孔子世家·乙": "#FF4500",
    "史記·孔子世家·丙": "#FF4500",
    "史記·孔子世家·丁": "#FF4500",
    "史記·孔子世家·戊": "#FF4500",
    "史記·孔子世家·己": "#FF4500",
    "史記·孔子世家·庚": "#FF4500",
    "史記·弟子列傳·選": "#98FB98",
    "家語·本姓解": "#FF00FF",
    "家語·七十二弟子解·選": "#ADFF2F",
}
for _ch in _ANALECTS_CHAPTERS:
    _NODE_COLOR_MAP[_ch] = "#97C2FC"

_BOOK_GROUP_COLORS = {
    "論語": "#97C2FC",
    "史記·孔子世家": "#FF4500",
    "史記·弟子列傳": "#98FB98",
    "家語·本姓解": "#FF00FF",
    "家語·七十二弟子解": "#ADFF2F",
    "衣鏡·孔子傳": "#FFC0CB",
    "衣鏡·弟子傳": "#7FFFAA",
}

_BOOK_GROUP_MARKERS = {
    "論語": "o",
    "史記·孔子世家": "s",
    "史記·弟子列傳": "^",
    "家語·本姓解": "D",
    "家語·七十二弟子解": "P",
    "衣鏡·孔子傳": "*",
    "衣鏡·弟子傳": "X",
}


def _normalize_color_key(name: str) -> str:
    t = str(name or "").strip()
    if not t:
        return ""
    t = t.replace("《", "").replace("》", "")
    t = re.sub(r"^[0-9０-９]+", "", t).strip()
    t = t.translate(_TEXT_NORMALIZE_CHAR_MAP)
    t = t.replace("(", "（").replace(")", "）")
    t = re.sub(r"\s+", "", t)
    t = re.sub(r"\s*·\s*", "·", t)
    t = re.sub(r"·+", "·", t).strip("·")
    t = re.sub(r"第[一二三四五六七八九十百零〇0-9０-９]+$", "", t)

    t = t.replace("孔子家語", "家語")
    t = t.replace("孔子衣鏡", "衣鏡")
    t = t.replace("（選）", "選")

    if t.startswith("論語·"):
        chap = t.split("·", 1)[1]
        chap = _ANALECTS_CHAPTER_ALIAS.get(chap, chap)
        if chap in _ANALECTS_CHAPTER_SET:
            return chap
    chap2 = _ANALECTS_CHAPTER_ALIAS.get(t, t)
    if chap2 in _ANALECTS_CHAPTER_SET:
        return chap2

    m = re.match(r"^史記·孔子世家·?([甲乙丙丁戊己庚])$", t)
    if m:
        return f"史記·孔子世家·{m.group(1)}"

    if t.startswith("史記·弟子列傳"):
        return "史記·弟子列傳·選" if "選" in t else "史記·弟子列傳"

    if t.startswith("家語·本姓解"):
        return "家語·本姓解"
    if t.startswith("家語·七十二弟子解"):
        return "家語·七十二弟子解·選" if "選" in t else "家語·七十二弟子解"

    if t.startswith("衣鏡·孔子傳"):
        return "衣鏡·孔子傳"
    if t.startswith("衣鏡·弟子傳"):
        return "衣鏡·弟子傳"

    return t


def _detect_book_group(name: str) -> str:
    key = _normalize_color_key(name)
    if key in _ANALECTS_CHAPTER_SET:
        return "論語"
    if key.startswith("史記·孔子世家·"):
        return "史記·孔子世家"
    if key.startswith("史記·弟子列傳"):
        return "史記·弟子列傳"
    if key.startswith("家語·本姓解"):
        return "家語·本姓解"
    if key.startswith("家語·七十二弟子解"):
        return "家語·七十二弟子解"
    if key.startswith("衣鏡·孔子傳"):
        return "衣鏡·孔子傳"
    if key.startswith("衣鏡·弟子傳"):
        return "衣鏡·弟子傳"
    return "論語"


def _node_color(name: str) -> str:
    key = _normalize_color_key(name)
    if key in _NODE_COLOR_MAP:
        return _NODE_COLOR_MAP[key]
    grp = _detect_book_group(name)
    return _BOOK_GROUP_COLORS.get(grp, "#97C2FC")


def _image_label_short(name: str, ascii_dot: bool = True) -> str:
    key = _normalize_color_key(name)
    out = str(key or "").strip()
    if not out:
        out = _pairwise_doc_label(name)
    # 使用中点「·」保持垂直居中显示，避免句点「.」下沉。
    out = out.replace(".", "·")
    return out


def compute_centrality_metrics(sim_matrix, doc_names):
    """计算 Strength / Eigenvector / Betweenness 中心性。"""
    n = len(doc_names)
    G = nx.Graph()
    for i in range(n):
        G.add_node(i, label=doc_names[i])
    for i in range(n):
        for j in range(i + 1, n):
            w = float(sim_matrix[i, j])
            if w > 0:
                G.add_edge(i, j, weight=w)

    strength = dict(G.degree(weight="weight"))
    eigenvector = nx.eigenvector_centrality_numpy(G, weight="weight")
    G_dist = G.copy()
    for u, v, d in G_dist.edges(data=True):
        d["weight"] = 1.0 / max(d["weight"], 1e-9)
    betweenness = nx.betweenness_centrality(G_dist, weight="weight")

    rows = []
    for i in range(n):
        rows.append({
            "name": doc_names[i],
            "book": _detect_book_group(doc_names[i]),
            "strength": strength[i],
            "eigenvector": eigenvector[i],
            "betweenness": betweenness[i],
        })
    df = pd.DataFrame(rows)
    df["rank_strength"] = df["strength"].rank(ascending=False, method="min").astype(int)
    df["rank_eigenvector"] = df["eigenvector"].rank(ascending=False, method="min").astype(int)
    df["rank_betweenness"] = df["betweenness"].rank(ascending=False, method="min").astype(int)
    return df.sort_values("rank_strength")


_CLUSTER_DEFAULT_FONT = DEFAULT_HEATMAP_FONT_PATH
_CLUSTER_PLOT_RATIO_W = 9.0
_CLUSTER_PLOT_RATIO_H = 16.0
_CLUSTER_BODY_FONTSIZE = 13
_CLUSTER_TITLE_FONTSIZE = _CLUSTER_BODY_FONTSIZE * 2
_CLUSTER_EXPORT_DPI = 600
_CLUSTER_GRIDSPEC_HSPACE = 0.11
_CLUSTER_NOTE_BOX_HEIGHT_SCALE = 0.40
# 固定四张图的版式配置，保证本地重复运行得到同一结构。
_CLUSTER_FIXED_STYLE: Dict[str, Dict[str, float]] = {
    "ward": {
        "width_ratio": 12.0,
        "height_ratio": 16.0,
        "fig_height": 14.4,
        "note_box_width": 0.92,
        "note_min_ratio": 0.20,
    },
    "mds": {
        "width_ratio": 12.0,
        "height_ratio": 12.0,
        "fig_height": 12.4,
        "note_box_width": 0.92,
        "note_min_ratio": 0.22,
    },
    "pca": {
        "width_ratio": 16.0,
        "height_ratio": 15.0,
        "fig_height": 16.0,
        "note_box_width": 0.92,
        "note_min_ratio": 0.22,
        "note_box_width_scale": 0.70,
    },
    "strength": {
        "width_ratio": 13.0,
        "height_ratio": 16.0,
        "fig_height": 13.6,
        "note_box_width": 0.92,
        "note_min_ratio": 0.24,
        "hspace": 0.09,
    },
}


def _cluster_figsize_by_ratio(width_ratio: float, height_ratio: float, height: float = 12.8) -> Tuple[float, float]:
    """按给定宽高比生成画布尺寸。"""
    h = max(8.0, float(height))
    hr = max(1e-6, float(height_ratio))
    wr = max(1e-6, float(width_ratio))
    w = h * (wr / hr)
    return (w, h)


def _cluster_figsize_9_16(height: float = 12.8) -> Tuple[float, float]:
    """固定 9:16（宽:高）画布比例；默认给足高度避免长标签拥挤。"""
    return _cluster_figsize_by_ratio(_CLUSTER_PLOT_RATIO_W, _CLUSTER_PLOT_RATIO_H, height)


def _cluster_fixed_layout(style_key: str, info_text: str) -> Tuple[Tuple[float, float], Dict[str, float]]:
    """按固定样式键返回画布与说明框布局，避免比例反复漂移。"""
    spec = _CLUSTER_FIXED_STYLE.get(str(style_key), _CLUSTER_FIXED_STYLE["ward"])
    fig_size = _cluster_figsize_by_ratio(
        float(spec.get("width_ratio", 12.0)),
        float(spec.get("height_ratio", 16.0)),
        float(spec.get("fig_height", 12.8)),
    )
    layout = _center_note_box_layout(
        _compute_endnote_layout(fig_size, info_text),
        box_width=float(spec.get("note_box_width", 0.92)),
        min_note_ratio=float(spec.get("note_min_ratio", 0.22)),
    )
    note_box_width_scale = float(spec.get("note_box_width_scale", 1.0))
    if abs(note_box_width_scale - 1.0) > 1e-9:
        layout["note_box_width"] = float(np.clip(float(layout.get("note_box_width", 0.92)) * note_box_width_scale, 0.30, 0.97))
        layout["note_box_x"] = float((1.0 - float(layout["note_box_width"])) / 2.0)
    layout["note_box_height_scale"] = float(
        np.clip(float(spec.get("note_box_height_scale", _CLUSTER_NOTE_BOX_HEIGHT_SCALE)), 0.20, 1.00)
    )
    layout["hspace"] = float(spec.get("hspace", _CLUSTER_GRIDSPEC_HSPACE))
    return fig_size, layout


def _center_note_box_layout(layout: Dict[str, float], box_width: float = 0.92, min_note_ratio: float = 0.22) -> Dict[str, float]:
    """尾注框居中并保留足够尾注高度，避免与主图纵轴区域重合。"""
    out = dict(layout or {})
    note_ratio = float(np.clip(max(float(out.get("note_ratio", 0.18)), float(min_note_ratio)), 0.16, 0.40))
    out["note_ratio"] = note_ratio
    out["main_ratio"] = float(max(0.50, 1.0 - note_ratio))
    bw = float(np.clip(float(box_width), 0.80, 0.97))
    out["note_box_width"] = bw
    out["note_box_x"] = float((1.0 - bw) / 2.0)
    return out


def _apply_font_to_rcparams(fp):
    """复用 plot_heatmap 的字体设置方式；如 fp 为 None 则尝试加载默认 SimHei。"""
    import matplotlib
    matplotlib.rcParams["axes.unicode_minus"] = False
    if fp is None and os.path.isfile(_CLUSTER_DEFAULT_FONT):
        fp = fm.FontProperties(fname=_CLUSTER_DEFAULT_FONT)
    primary_family = ""
    if fp is not None:
        # 注册字体文件到 fontManager，使 rcParams 能找到
        fpath = fp.get_file()
        if fpath and os.path.isfile(fpath):
            try:
                fm.fontManager.addfont(fpath)
            except Exception:
                pass
        primary_family = str(fp.get_name() or "").strip()
    _set_cjk_rcparams(primary_family)
    return fp


def _chart_font_props(font_setting: str = ""):
    """图表字体策略：中文统一黑体优先，整体统一粗体。"""
    import matplotlib

    matplotlib.rcParams["axes.unicode_minus"] = False
    _set_cjk_rcparams("")

    # 尝试注册外部字体文件（如用户传入），但不强制覆盖英文字体。
    p = str(font_setting or "").strip()
    primary_family = ""
    if p and os.path.isfile(p):
        try:
            fm.fontManager.addfont(p)
            primary_family = str(fm.FontProperties(fname=p).get_name() or "").strip()
        except Exception:
            pass

    family_list = _set_cjk_rcparams(primary_family)
    title_fp = fm.FontProperties(
        family=family_list,
        weight="bold",
        size=float(_CLUSTER_TITLE_FONTSIZE),
    )
    body_fp = fm.FontProperties(
        family=family_list,
        weight="bold",
        size=float(_CLUSTER_BODY_FONTSIZE),
    )
    return title_fp, body_fp


def _compute_endnote_layout(figsize: Tuple[float, float], info_text: str) -> Dict[str, float]:
    """按图像长宽比与说明文字行数，计算「上图下尾注」布局参数。"""
    try:
        fig_w = float(figsize[0])
        fig_h = float(figsize[1])
    except Exception:
        fig_w, fig_h = 14.0, 10.0
    aspect = fig_w / max(fig_h, 1e-6)
    n_lines = max(1, len(str(info_text or "").splitlines()))

    # 说明区高度：随行数增加；图越宽（aspect 越大）说明区可略矮。
    note_ratio = 0.13 + (0.0048 * n_lines * (1.15 / max(0.9, aspect)))
    note_ratio = float(np.clip(note_ratio, 0.14, 0.34))
    main_ratio = float(max(0.55, 1.0 - note_ratio))

    # 尾注文本框宽度：接近整图但保留左右留白，避免出格。
    note_box_width = 0.86 + min(0.10, max(0.0, 0.04 * (aspect - 1.0)))
    note_box_width = float(np.clip(note_box_width, 0.86, 0.94))
    note_box_x = float((1.0 - note_box_width) / 2.0)

    return {
        "main_ratio": main_ratio,
        "note_ratio": note_ratio,
        "note_box_width": note_box_width,
        "note_box_x": note_box_x,
    }


def _draw_endnote_box(
    ax_note,
    info_text: str,
    fp,
    layout: Dict[str, float],
    fontsize: float = 9.0,
    linespacing: float = 1.45,
) -> None:
    """在底部绘制横向长方形尾注文本框。"""
    from matplotlib.patches import FancyBboxPatch

    def _sanitize_endnote_text(text: str) -> str:
        # 避免 SimHei 缺失数学/下标字符导致 "□" 占位。
        s = str(text or "")
        replacements = {
            "√": "sqrt",
            "−": "-",
            "–": "-",
            "—": "-",
            "×": "*",
            "÷": "/",
            "∑": "sum",
            "‖": "||",
            "ₐ": "a",
            "ᵢ": "i",
            "ⱼ": "j",
            "′": "'",
            "·": "*",
            "≤": "<=",
            "≥": ">=",
            "→": "->",
        }
        for src, dst in replacements.items():
            s = s.replace(src, dst)
        s = re.sub(r"[ \t]+", " ", s)
        s = re.sub(r" *\n *", "\n", s)
        s = re.sub(r"\n{2,}", "\n", s)
        return s

    ax_note.axis("off")
    text_to_draw = _sanitize_endnote_text(info_text)
    # 恢复 v10 的横向长度（宽度），纵向高度单独缩放。
    box_w_base = float(layout.get("note_box_width", 0.94))
    box_w_base = float(np.clip(box_w_base, 0.80, 0.97))
    box_x_base = float((1.0 - box_w_base) / 2.0)
    box_h_full = 0.94
    box_h_scale = float(np.clip(float(layout.get("note_box_height_scale", _CLUSTER_NOTE_BOX_HEIGHT_SCALE)), 0.20, 1.00))
    box_h_base = box_h_full * box_h_scale
    box_top = 0.96

    # 依据当前画布比例与文本内容，尝试多个宽度并选取最居中的可用位置。
    cand_lo = max(0.80, box_w_base - 0.06)
    cand_hi = min(0.97, box_w_base + 0.03)
    candidates = np.linspace(cand_lo, cand_hi, 9)
    best_score = float("inf")
    best_w = box_w_base
    best_x = box_x_base
    for w in candidates:
        x = float((1.0 - float(w)) / 2.0)
        y = float(box_top - box_h_base)
        t = ax_note.text(
            x + 0.014,
            box_top - 0.02,
            text_to_draw,
            transform=ax_note.transAxes,
            fontsize=float(fontsize),
            va="top",
            ha="left",
            fontproperties=fp,
            linespacing=float(linespacing),
        )
        try:
            ax_note.figure.canvas.draw()
            renderer = ax_note.figure.canvas.get_renderer()
            bb = t.get_window_extent(renderer=renderer).transformed(ax_note.transAxes.inverted())
            overflow = (
                max(0.0, (x + 0.008) - float(bb.x0))
                + max(0.0, float(bb.x1) - (x + float(w) - 0.008))
                + max(0.0, (y + 0.008) - float(bb.y0))
                + max(0.0, float(bb.y1) - (box_top - 0.008))
            )
            center_penalty = abs(((float(bb.x0) + float(bb.x1)) / 2.0) - 0.5)
            width_penalty = abs(float(w) - box_w_base) * 0.02
            score = overflow * 100.0 + center_penalty + width_penalty
        except Exception:
            score = 1e9
        t.remove()
        if score < best_score:
            best_score = float(score)
            best_w = float(w)
            best_x = x

    box_w = best_w
    box_x = best_x
    # 用户要求：纵向高度缩到 40%；若文字仍溢出则仅做最小高度补偿，确保文本被框覆盖。
    box_h = box_h_base
    box_y = box_top - box_h
    _probe = ax_note.text(
        box_x + 0.014,
        box_top - 0.02,
        text_to_draw,
        transform=ax_note.transAxes,
        fontsize=float(fontsize),
        va="top",
        ha="left",
        fontproperties=fp,
        linespacing=float(linespacing),
    )
    try:
        ax_note.figure.canvas.draw()
        renderer = ax_note.figure.canvas.get_renderer()
        bbp = _probe.get_window_extent(renderer=renderer).transformed(ax_note.transAxes.inverted())
        needed_h = float((bbp.y1 - bbp.y0) + 0.03)
        if needed_h > box_h:
            box_h = min(0.94, needed_h)
            box_y = box_top - box_h
    except Exception:
        pass
    _probe.remove()
    # 固定长度的尾注框（可视），避免 bbox 仅包裹文字导致框过短。
    ax_note.add_patch(
        FancyBboxPatch(
            (box_x, box_y),
            box_w,
            box_h,
            boxstyle="round,pad=0.015",
            transform=ax_note.transAxes,
            facecolor="#F8F9FA",
            edgecolor="#DEE2E6",
            linewidth=1.0,
            alpha=0.95,
            zorder=0,
        )
    )
    # 透明占位框，辅助保持渲染边界一致。
    ax_note.add_patch(
        plt.Rectangle(
            (box_x, box_y),
            box_w,
            box_h,
            transform=ax_note.transAxes,
            facecolor="none",
            edgecolor="none",
            linewidth=0.0,
        )
    )
    ax_note.text(
        box_x + 0.014,
        box_top - 0.02,
        text_to_draw,
        transform=ax_note.transAxes,
        fontsize=float(fontsize),
        va="top",
        ha="left",
        fontproperties=fp,
        linespacing=float(linespacing),
    )


def _plot_dendrogram_terminal_colored(
    ax,
    Z: np.ndarray,
    doc_names: List[str],
    labels: List[str],
    colored_line_width: float = 2.0,
    gray_line_width: float = 1.1,
) -> None:
    """只给「叶子到第一次合并」末梢线着色，其余分支统一灰色。"""
    gray = "#777777"
    n = int(len(doc_names))
    if n <= 0:
        return

    d = scipy_dendrogram(
        Z,
        orientation="left",
        labels=list(labels),
        no_plot=True,
    )
    leaves = list(d.get("leaves") or [])
    if not leaves:
        return

    leaf_y: Dict[int, float] = {leaf: 5.0 + (10.0 * i) for i, leaf in enumerate(leaves)}
    node_x: Dict[int, float] = {}
    node_y: Dict[int, float] = {}

    rows = np.array(Z, dtype=float)
    for row_idx, row in enumerate(rows):
        node_id = n + row_idx
        left = int(row[0])
        right = int(row[1])
        dist = float(row[2])

        xl = 0.0 if left < n else float(node_x[left])
        yl = float(leaf_y[left] if left < n else node_y[left])
        xr = 0.0 if right < n else float(node_x[right])
        yr = float(leaf_y[right] if right < n else node_y[right])

        cl = _BOOK_GROUP_COLORS.get(_detect_book_group(str(doc_names[left])), gray) if left < n else gray
        cr = _BOOK_GROUP_COLORS.get(_detect_book_group(str(doc_names[right])), gray) if right < n else gray
        lw_l = float(colored_line_width if left < n else gray_line_width)
        lw_r = float(colored_line_width if right < n else gray_line_width)

        # 两条水平线：若 child 为 leaf，则这是末梢线，着分组色；否则灰色。
        ax.plot([xl, dist], [yl, yl], color=cl, linewidth=lw_l, solid_capstyle="round", zorder=2)
        ax.plot([xr, dist], [yr, yr], color=cr, linewidth=lw_r, solid_capstyle="round", zorder=2)
        # 竖线（合并主干）统一灰色。
        ax.plot(
            [dist, dist],
            [min(yl, yr), max(yl, yr)],
            color=gray,
            linewidth=float(gray_line_width),
            solid_capstyle="round",
            zorder=1,
        )

        node_x[node_id] = dist
        node_y[node_id] = (yl + yr) / 2.0

    y_ticks = [leaf_y[leaf] for leaf in leaves]
    y_labels = [str(labels[leaf]) if 0 <= leaf < len(labels) else str(leaf) for leaf in leaves]
    ax.set_yticks(y_ticks)
    ax.set_yticklabels(y_labels)

    y_max = 10.0 * len(leaves) + 5.0
    x_max = float(np.max(rows[:, 2])) if rows.size else 1.0
    ax.set_ylim(0.0, y_max)
    ax.set_xlim(max(1e-6, x_max * 1.05), 0.0)


def plot_cluster_dendrogram(sim_matrix, doc_names, outpath, font_setting=""):
    """Ward 层次聚类树状图，底部附算法说明。"""
    title_fp, body_fp = _chart_font_props(font_setting)

    dist = 1.0 - np.array(sim_matrix, dtype=float)
    np.fill_diagonal(dist, 0)
    dist = np.clip(dist, 0, None)
    condensed = squareform(dist, checks=False)
    Z = linkage(condensed, method="ward")

    info_text = (
        "Ward 層次聚類（Ward Hierarchical Clustering）\n"
        "━━━━━━━━━━━━━━━━━━\n"
        "合併距離公式：d(A,B)=sqrt(2|A||B|/(|A|+|B|))*||c_a-c_b||\n"
        "算法原理：自底向上逐步合併距離最近的簇，每次合併選擇使組內方差增量最小的一對。\n"
        "距離矩陣：dist(i,j)=1-cos_sim(i,j)\n"
        "應用場景：無需預設分組數，適合小規模語料分類，樹形結構可視化親疏關係。"
    )
    fig_size, layout = _cluster_fixed_layout("ward", info_text)
    fig = plt.figure(figsize=fig_size)
    gs = fig.add_gridspec(
        2, 1,
        height_ratios=[layout["main_ratio"], layout["note_ratio"]],
        hspace=float(layout.get("hspace", _CLUSTER_GRIDSPEC_HSPACE)),
    )
    ax = fig.add_subplot(gs[0, 0])
    ax_note = fig.add_subplot(gs[1, 0])
    dendro_labels = build_heatmap_labels(doc_names, strip_digits=False)
    _plot_dendrogram_terminal_colored(ax, Z, doc_names, dendro_labels, colored_line_width=2.2, gray_line_width=2.2)
    ax.yaxis.tick_right()
    ax.tick_params(axis="y", which="both", right=True, labelright=True, left=False, labelleft=False, pad=4, labelsize=9.6)
    leaf_labels = ax.get_ymajorticklabels()
    for lbl in leaf_labels:
        lbl.set_color("black")
        lbl.set_fontweight("bold")
        lbl.set_horizontalalignment("left")
        lbl.set_fontsize(9.6)
        lbl.set_fontproperties(body_fp)

    # k=2, k=4, k=6 截断线
    y_top = max(ax.get_ylim())
    y_bot = min(ax.get_ylim())
    _k_cut_configs = [
        (2, "#E74C3C", 0.98),   # 紅色，頂部
        (4, "#2E86C1", 0.60),   # 藍色，中部
        (6, "#27AE60", 0.22),   # 綠色，底部
    ]
    for _k, _color, _y_frac in _k_cut_configs:
        if len(Z) >= _k:
            # 截斷線放在產生 k 簇的合併距離與前一步之間的中點
            _d_merge = Z[-(_k - 1), 2]
            _d_prev = Z[-_k, 2] if len(Z) >= _k else _d_merge * 0.95
            _d_cut = (_d_merge + _d_prev) / 2
            ax.axvline(x=_d_cut, color=_color, linestyle="--", linewidth=1.2, alpha=0.7)
            _y_pos = y_bot + (y_top - y_bot) * _y_frac
            ax.text(_d_cut + 0.008, _y_pos, f"k={_k}",
                    fontsize=10, color=_color, va="center", fontproperties=body_fp, fontweight="bold")

    # k=2/4/6 顶部居中彩色小方框（同一水平线）
    from matplotlib.patches import Patch
    k_square_handles = [
        Patch(facecolor=_k_color, edgecolor=_k_color, label=f"k={_k}")
        for _k, _k_color, _ in _k_cut_configs
    ]
    k_square_legend = ax.legend(
        handles=k_square_handles,
        loc="upper center",
        bbox_to_anchor=(0.5, 1.01),
        ncol=3,
        frameon=False,
        fontsize=9.0,
        prop=body_fp,
        handlelength=0.9,
        handletextpad=0.35,
        columnspacing=1.2,
    )
    ax.add_artist(k_square_legend)

    ax.set_title(
        "Ward 層次聚類樹狀圖（餘弦距離）",
        fontproperties=title_fp,
        fontweight="bold",
        pad=10,
    )
    ax.set_xlabel("Ward 合併距離", fontsize=14, fontproperties=body_fp, fontweight="bold")
    for lbl in ax.get_xticklabels():
        lbl.set_fontproperties(body_fp)
        lbl.set_fontsize(11)
        lbl.set_fontweight("bold")

    _draw_endnote_box(ax_note, info_text, body_fp, layout, fontsize=float(_CLUSTER_BODY_FONTSIZE), linespacing=1.34)

    # 图例
    legend_handles = [Patch(facecolor=c, label=str(g))
                      for g, c in _BOOK_GROUP_COLORS.items()]
    leg = ax.legend(handles=legend_handles, loc="lower right", fontsize=8.8,
                    prop=body_fp, framealpha=0.8)
    for t in leg.get_texts():
        t.set_fontweight("bold")

    fig.subplots_adjust(
        left=0.07,
        right=0.84,
        top=0.96,
        bottom=0.05,
        hspace=float(layout.get("hspace", _CLUSTER_GRIDSPEC_HSPACE)),
    )
    fig.savefig(outpath, dpi=_CLUSTER_EXPORT_DPI)
    plt.close(fig)
    return Z


def plot_mds_scatter(sim_matrix, doc_names, outpath, font_setting=""):
    """MDS 二维降维散点图，底部附算法说明。"""
    title_fp, body_fp = _chart_font_props(font_setting)

    dist = 1.0 - np.array(sim_matrix, dtype=float)
    np.fill_diagonal(dist, 0)
    dist = np.clip(dist, 0, None)
    mds = MDS(n_components=2, dissimilarity="precomputed", random_state=42,
              normalized_stress="auto")
    coords = mds.fit_transform(dist)

    info_text = (
        "MDS 多維尺度分析（Multidimensional Scaling）\n"
        "━━━━━━━━━━━━━━━━━━\n"
        "目標函數（Stress）：s = sqrt(Sum(d_ij - d'_ij)^2 / Sum(d_ij)^2)\n"
        "算法原理：將高維距離矩陣投影到二維平面，盡量保留文本間的相對距離關係。\n"
        "距離定義：dist(i,j) = 1 - cos_sim(i,j)\n"
        "解讀：距離近=用詞相似度高；聚集區域=同類文本群；stress < 0.1 為良好擬合"
    )
    fig_size, layout = _cluster_fixed_layout("mds", info_text)
    fig = plt.figure(figsize=fig_size)
    gs = fig.add_gridspec(
        2, 1,
        height_ratios=[layout["main_ratio"], layout["note_ratio"]],
        hspace=float(layout.get("hspace", _CLUSTER_GRIDSPEC_HSPACE)),
    )
    ax = fig.add_subplot(gs[0, 0])
    ax_note = fig.add_subplot(gs[1, 0])

    for i, name in enumerate(doc_names):
        grp = _detect_book_group(name)
        marker = _BOOK_GROUP_MARKERS.get(grp, "o")
        color = _node_color(name)
        ax.scatter(coords[i, 0], coords[i, 1], c=color, marker=marker,
                   s=80, edgecolors="black", linewidths=0.5, zorder=3)

    # 标签偏移
    for i, name in enumerate(doc_names):
        short = _image_label_short(name, ascii_dot=True)
        ax.annotate(short, (coords[i, 0], coords[i, 1]),
                    textcoords="offset points", xytext=(6, 4),
                    fontsize=7.4, fontproperties=body_fp, alpha=0.9, fontweight="bold")

    ax.set_title(
        "MDS 二維空間可視化（餘弦距離）",
        fontproperties=title_fp,
        fontweight="bold",
        pad=12,
    )
    ax.set_xlabel("MDS 維度 1", fontsize=14, fontproperties=body_fp, fontweight="bold")
    ax.set_ylabel("MDS 維度 2", fontsize=14, fontproperties=body_fp, fontweight="bold")
    for lbl in ax.get_xticklabels() + ax.get_yticklabels():
        lbl.set_fontproperties(body_fp)
        lbl.set_fontsize(11)
        lbl.set_fontweight("bold")
    ax.set_aspect("equal", adjustable="box")
    ax.grid(True, alpha=0.3)

    # 计算 normalized stress-1：sqrt(raw_stress / sum(dist_sq))
    from scipy.spatial.distance import pdist, squareform as _sqf
    raw_stress = mds.stress_
    dist_upper = pdist(dist, metric="euclidean")  # 原始距离矩阵的上三角
    # 但 MDS 的 stress_ 是对 dissimilarity 的，直接用 condensed dist
    diss_upper = squareform(dist, checks=False)
    sum_diss_sq = np.sum(diss_upper ** 2)
    stress_norm = np.sqrt(raw_stress / sum_diss_sq) if sum_diss_sq > 0 else raw_stress
    ax.text(0.02, 0.02, f"Stress-1 = {stress_norm:.4f}",
            transform=ax.transAxes, fontsize=10, fontproperties=body_fp,
            bbox=dict(facecolor="white", alpha=0.8), fontweight="bold")

    from matplotlib.lines import Line2D
    legend_handles = []
    seen_groups: List[str] = []
    for name in doc_names:
        g = _detect_book_group(name)
        if g not in seen_groups:
            seen_groups.append(g)
    for grp in seen_groups:
        mk = _BOOK_GROUP_MARKERS.get(grp, "o")
        cl = _BOOK_GROUP_COLORS.get(grp, "#97C2FC")
        legend_handles.append(Line2D([0], [0], marker=mk, color="w",
                                     markerfacecolor=cl, markersize=8,
                                     markeredgecolor="black",
                                     markeredgewidth=0.5, label=str(grp)))
    leg = ax.legend(handles=legend_handles, loc="upper right", fontsize=8.8,
                    prop=body_fp, framealpha=0.8)
    for t in leg.get_texts():
        t.set_fontweight("bold")

    _draw_endnote_box(ax_note, info_text, body_fp, layout, fontsize=float(_CLUSTER_BODY_FONTSIZE), linespacing=1.36)

    fig.subplots_adjust(
        left=0.07,
        right=0.93,
        top=0.95,
        bottom=0.05,
        hspace=float(layout.get("hspace", _CLUSTER_GRIDSPEC_HSPACE)),
    )
    fig.savefig(outpath, dpi=_CLUSTER_EXPORT_DPI)
    plt.close(fig)
    return coords, stress_norm


def plot_pca_biplot(doc_vectors, feature_names, doc_names, outpath,
                    font_setting="", n_top_loadings=10, use_tfidf=True):
    """PCA 二維散點圖（biplot），載荷箭頭顯示驅動分離的 n-gram，底部附算法說明。"""
    import scipy.sparse as sp
    title_fp, body_fp = _chart_font_props(font_setting)

    # ---- 1. TF-IDF 轉換（從 raw count 矩陣恢復） ----
    if sp.issparse(doc_vectors):
        X_input = doc_vectors
    else:
        X_input = np.array(doc_vectors, dtype=float)
    if use_tfidf:
        X = TfidfTransformer().fit_transform(X_input)
        if sp.issparse(X):
            X = X.toarray()
    else:
        X = X_input if isinstance(X_input, np.ndarray) else X_input.toarray()

    n_docs, n_features = X.shape
    n_comp = min(2, n_docs, n_features)
    if n_comp < 2:
        print(f"[PCA] skip: n_docs={n_docs}, n_features={n_features}, need >=2")
        return None, None

    # ---- 2. PCA 降維 ----
    pca = PCA(n_components=n_comp)
    coords = pca.fit_transform(X)
    ev_ratio = pca.explained_variance_ratio_

    # ---- 3. 散點圖 ----
    info_text = (
        "PCA 主成分分析（Principal Component Analysis）\n"
        "━━━━━━━━━━━━━━━━━━\n"
        "投影方式：對 TF-IDF 特徵矩陣進行特徵值分解，提取方差最大的兩個方向作為座標軸。\n"
        "載荷向量（箭頭）：箭頭方向和長度表示該 n-gram 對主成分的貢獻大小，指向某群文本代表其高頻特徵。\n"
        "解釋方差比：PC1+PC2 合計越高，二維投影越能代表原始高維特徵空間。\n"
        "解讀：距離近=TF-IDF 特徵相似；箭頭指向=區分性 n-gram；合計 > 50% 為良好投影。"
    )
    fig_size, layout = _cluster_fixed_layout("pca", info_text)

    # 固定比例画布下按 MDS（12.4 高度）等比放大视觉元素，避免“图标和文字偏小”。
    pca_vis_scale = float(fig_size[0]) / 12.4
    marker_size = float(80.0 * (pca_vis_scale ** 2))
    marker_edge_lw = float(0.5 * pca_vis_scale)
    text_fs = float(7.4 * pca_vis_scale)
    ev_fs = float(10.0 * pca_vis_scale)
    axis_fs = float(14.0 * pca_vis_scale)
    tick_fs = float(11.0 * pca_vis_scale)
    legend_fs = float(8.8 * pca_vis_scale)
    legend_marker = float(8.0 * pca_vis_scale)

    title_fp_pca = fm.FontProperties(
        family=title_fp.get_family(),
        style=title_fp.get_style(),
        variant=title_fp.get_variant(),
        weight=title_fp.get_weight(),
        stretch=title_fp.get_stretch(),
        size=float(title_fp.get_size_in_points() * pca_vis_scale),
    )
    body_fp_pca = fm.FontProperties(
        family=body_fp.get_family(),
        style=body_fp.get_style(),
        variant=body_fp.get_variant(),
        weight=body_fp.get_weight(),
        stretch=body_fp.get_stretch(),
        size=float(body_fp.get_size_in_points() * pca_vis_scale),
    )

    fig = plt.figure(figsize=fig_size)
    gs = fig.add_gridspec(
        2, 1,
        height_ratios=[layout["main_ratio"], layout["note_ratio"]],
        hspace=float(layout.get("hspace", _CLUSTER_GRIDSPEC_HSPACE)),
    )
    ax = fig.add_subplot(gs[0, 0])
    ax_note = fig.add_subplot(gs[1, 0])

    for i, name in enumerate(doc_names):
        grp = _detect_book_group(name)
        marker = _BOOK_GROUP_MARKERS.get(grp, "o")
        color = _node_color(name)
        ax.scatter(coords[i, 0], coords[i, 1], c=color, marker=marker,
                   s=marker_size, edgecolors="black", linewidths=marker_edge_lw, zorder=3)

    # 標籤偏移
    for i, name in enumerate(doc_names):
        short = _image_label_short(name, ascii_dot=True)
        ax.annotate(short, (coords[i, 0], coords[i, 1]),
                    textcoords="offset points", xytext=(6, 4),
                    fontsize=text_fs, fontproperties=body_fp_pca, alpha=0.9, fontweight="bold")

    # ---- 4. 載荷箭頭（Biplot） ----
    loadings = pca.components_.T  # shape: (V, 2)
    magnitude = np.sqrt(loadings[:, 0] ** 2 + loadings[:, 1] ** 2)
    n_show = min(n_top_loadings, len(magnitude))
    top_idx = np.argsort(magnitude)[-n_show:]

    # 縮放箭頭使其與散點圖範圍匹配
    coord_range = max(np.abs(coords).max(), 1e-9)
    loading_range = max(magnitude[top_idx].max(), 1e-9)
    arrow_scale = coord_range / loading_range * 0.65

    feature_names_arr = np.array(feature_names) if not isinstance(feature_names, np.ndarray) else feature_names
    for idx in top_idx:
        lx = loadings[idx, 0] * arrow_scale
        ly = loadings[idx, 1] * arrow_scale
        ax.annotate(
            "", xy=(lx, ly), xytext=(0, 0),
            arrowprops=dict(arrowstyle="->", color="#888888", lw=(1.0 * pca_vis_scale), alpha=0.6),
        )
        ax.text(lx * 1.08, ly * 1.08, str(feature_names_arr[idx]),
                fontsize=text_fs, color="#555555", fontproperties=body_fp_pca, alpha=0.85, fontweight="bold",
                ha="center", va="center")

    # ---- 5. 解釋方差標註 ----
    ev_text = (f"PC1: {ev_ratio[0]*100:.1f}%   PC2: {ev_ratio[1]*100:.1f}%   "
               f"合計: {sum(ev_ratio[:2])*100:.1f}%")
    ax.text(0.02, 0.02, ev_text,
            transform=ax.transAxes, fontsize=ev_fs, fontproperties=body_fp_pca,
            bbox=dict(facecolor="white", alpha=0.8), fontweight="bold")

    ax.set_title(
        "PCA 二維空間可視化（TF-IDF 特徵）",
        fontproperties=title_fp_pca,
        fontweight="bold",
        pad=12,
    )
    ax.set_xlabel(f"PC1（{ev_ratio[0]*100:.1f}%）", fontsize=axis_fs, fontproperties=body_fp_pca, fontweight="bold")
    ax.set_ylabel(f"PC2（{ev_ratio[1]*100:.1f}%）", fontsize=axis_fs, fontproperties=body_fp_pca, fontweight="bold")
    from matplotlib.ticker import MultipleLocator
    # PCA 横纵轴主刻度统一为 0.1，避免读图比例不一致。
    ax.xaxis.set_major_locator(MultipleLocator(0.1))
    ax.yaxis.set_major_locator(MultipleLocator(0.1))
    for lbl in ax.get_xticklabels() + ax.get_yticklabels():
        lbl.set_fontproperties(body_fp_pca)
        lbl.set_fontsize(tick_fs)
        lbl.set_fontweight("bold")
    # 正方形绘图区，避免比例改动扭曲坐标关系。
    ax.set_aspect("equal", adjustable="box")
    ax.grid(True, alpha=0.3)
    ax.axhline(y=0, color="gray", linewidth=0.5, alpha=0.3)
    ax.axvline(x=0, color="gray", linewidth=0.5, alpha=0.3)

    # 圖例
    from matplotlib.lines import Line2D
    legend_handles = []
    seen_groups: List[str] = []
    for name in doc_names:
        g = _detect_book_group(name)
        if g not in seen_groups:
            seen_groups.append(g)
    for grp in seen_groups:
        mk = _BOOK_GROUP_MARKERS.get(grp, "o")
        cl = _BOOK_GROUP_COLORS.get(grp, "#97C2FC")
        legend_handles.append(Line2D([0], [0], marker=mk, color="w",
                                     markerfacecolor=cl, markersize=legend_marker,
                                     markeredgecolor="black",
                                     markeredgewidth=marker_edge_lw, label=str(grp)))
    leg = ax.legend(handles=legend_handles, loc="upper right", fontsize=legend_fs,
                    prop=body_fp_pca, framealpha=0.8)
    for t in leg.get_texts():
        t.set_fontweight("bold")

    _draw_endnote_box(ax_note, info_text, body_fp_pca, layout, fontsize=float(body_fp_pca.get_size_in_points()), linespacing=1.36)

    fig.subplots_adjust(
        left=0.07,
        right=0.93,
        top=0.96,
        bottom=0.05,
        hspace=float(layout.get("hspace", _CLUSTER_GRIDSPEC_HSPACE)),
    )
    fig.savefig(outpath, dpi=_CLUSTER_EXPORT_DPI)
    plt.close(fig)
    return coords, ev_ratio


def plot_network_graph(sim_matrix, doc_names, outpath,
                       font_setting="", score_key="cos_raw",
                       min_weight=0.0, layout_k=None, layout_seed=42,
                       figsize=(16, 10)):
    """ctext 風格力導向網絡關係圖，底部附算法說明。

    節點為圓角矩形填色塊（按書類著色），邊粗細 ∝ 相似度權重，
    同組邊用該組顏色，跨組邊用灰色。
    """
    from matplotlib.patches import FancyBboxPatch
    from matplotlib.lines import Line2D

    fp = resolve_heatmap_font(font_setting)
    fp = _apply_font_to_rcparams(fp)

    sim = np.array(sim_matrix, dtype=float)
    n = len(doc_names)
    if n < 2:
        return

    # ---- 1. 構建加權無向圖 ----
    G = nx.Graph()
    for i in range(n):
        G.add_node(i, name=doc_names[i])
    for i in range(n):
        for j in range(i + 1, n):
            w = float(sim[i, j])
            if w > min_weight:
                G.add_edge(i, j, weight=w)

    if G.number_of_edges() == 0:
        print(f"[Network] skip: no edges above min_weight={min_weight}")
        return

    # ---- 2. 力導向佈局 ----
    pos = nx.spring_layout(G, k=layout_k, seed=layout_seed,
                           weight="weight", iterations=100)

    # ---- 3. 邊的權重範圍（用於線寬/alpha 映射） ----
    weights = [d["weight"] for _, _, d in G.edges(data=True)]
    w_min, w_max = min(weights), max(weights)
    w_span = w_max - w_min if w_max > w_min else 1e-9

    def _edge_lw(w):
        return 0.3 + (w - w_min) / w_span * 5.7  # [0.3, 6.0]

    def _edge_alpha(w):
        return 0.15 + (w - w_min) / w_span * 0.65  # [0.15, 0.8]

    # 書類分組顏色（用於邊著色）
    grp_list = [_detect_book_group(name) for name in doc_names]

    def _edge_color(i, j):
        if grp_list[i] == grp_list[j]:
            return _BOOK_GROUP_COLORS.get(grp_list[i], "#888888")
        return "#888888"

    metric_label = {"cos_raw": "餘弦相似度", "tfidf": "TF-IDF 餘弦相似度",
                    "jaccard": "Jaccard 係數"}.get(score_key, score_key)
    threshold_line = (f"weight < {min_weight:.3f} 的邊已隱藏"
                      if min_weight > 0 else "顯示所有邊（無閾值）")
    info_text = (
        "力導向網絡圖（Force-Directed Network）\n"
        "━━━━━━━━━━━━━━━━━━\n"
        "佈局算法：Spring Layout（Fruchterman-Reingold），節點間斥力與邊引力達到平衡後的穩定佈局。\n"
        f"邊寬度：線寬 = {metric_label}；粗邊=高相似度，細邊=低相似度。\n"
        "邊顏色：同組邊=該組顏色；跨組邊=灰色。節點顏色：按文獻來源分組著色。\n"
        f"閾值：{threshold_line}；邊數：{G.number_of_edges()}；節點數：{G.number_of_nodes()}。"
    )
    fig_size = tuple(figsize)
    layout = _compute_endnote_layout(fig_size, info_text)
    fig = plt.figure(figsize=fig_size)
    gs = fig.add_gridspec(2, 1, height_ratios=[layout["main_ratio"], layout["note_ratio"]], hspace=0.03)
    ax = fig.add_subplot(gs[0, 0])
    ax_note = fig.add_subplot(gs[1, 0])

    # 繪製邊
    for u, v, d in G.edges(data=True):
        w = d["weight"]
        x0, y0 = pos[u]
        x1, y1 = pos[v]
        ax.plot([x0, x1], [y0, y1],
                color=_edge_color(u, v),
                linewidth=_edge_lw(w),
                alpha=_edge_alpha(w),
                solid_capstyle="round", zorder=1)

    # 繪製節點（圓角矩形 + 中文標籤）
    # 先 draw 一次以獲取 renderer（Agg backend 必須）
    fig.canvas.draw()
    renderer = fig.canvas.get_renderer()
    for i, name in enumerate(doc_names):
        x, y = pos[i]
        short = _image_label_short(name, ascii_dot=True)
        color = _node_color(name)

        # 用 text 來測量尺寸
        txt = ax.text(x, y, short, fontsize=7, fontproperties=fp,
                      ha="center", va="center", zorder=4,
                      color="black")

        # 畫背景圓角矩形
        bbox = txt.get_window_extent(renderer=renderer)
        bbox_data = bbox.transformed(ax.transData.inverted())
        pad_x = (bbox_data.width) * 0.35
        pad_y = (bbox_data.height) * 0.45
        rect = FancyBboxPatch(
            (bbox_data.x0 - pad_x, bbox_data.y0 - pad_y),
            bbox_data.width + 2 * pad_x,
            bbox_data.height + 2 * pad_y,
            boxstyle="round,pad=0.02",
            facecolor=color, edgecolor="black", linewidth=0.5,
            alpha=0.85, zorder=3,
            transform=ax.transData,
        )
        ax.add_patch(rect)

    ax.set_title(f"文本關係網絡圖（{metric_label}）", fontsize=13, fontproperties=fp)
    ax.set_aspect("equal")
    ax.axis("off")

    # 圖例（用 Patch 而非 FancyBboxPatch，後者無標準 legend handler）
    import matplotlib.patches as mpatches
    seen_grps = []
    for g in grp_list:
        if g not in seen_grps:
            seen_grps.append(g)
    legend_handles = []
    for grp in seen_grps:
        c = _BOOK_GROUP_COLORS.get(grp, "#97C2FC")
        legend_handles.append(
            mpatches.Patch(facecolor=c, edgecolor="black", linewidth=0.5, label=str(grp)))
    if legend_handles:
        ax.legend(handles=legend_handles, loc="lower left", fontsize=8,
                  prop=fp, framealpha=0.8, handlelength=1.5, handleheight=1.0)

    _draw_endnote_box(ax_note, info_text, fp, layout, fontsize=float(_CLUSTER_BODY_FONTSIZE), linespacing=1.4)

    fig.subplots_adjust(left=0.04, right=0.98, top=0.95, bottom=0.04, hspace=0.03)
    fig.savefig(outpath, dpi=600)
    plt.close(fig)


def plot_centrality_bars(centrality_df, outpath, font_setting=""):
    """Strength 中心性水平柱状图，底部附算法说明。"""
    title_fp, body_fp = _chart_font_props(font_setting)

    df = centrality_df.sort_values("strength", ascending=True)

    info_text = (
        "加權度中心性（Strength Centrality）\n"
        "━━━━━━━━━━━━━━━━━━\n"
        "公式：Strength(i) = Sum_j sim(i,j)。即節點 i 與所有其他節點的餘弦相似度之和。\n"
        "解讀：值越高表示與其他文本的總體相似度越高，排名靠前節點可視為語料庫中的核心文本。\n"
        "互補指標：Eigenvector（遞迴權重）與 Betweenness（最短路徑橋接作用）。"
    )
    fig_size, layout = _cluster_fixed_layout("strength", info_text)
    fig = plt.figure(figsize=fig_size)
    gs = fig.add_gridspec(
        2, 1,
        height_ratios=[layout["main_ratio"], layout["note_ratio"]],
        hspace=float(layout.get("hspace", _CLUSTER_GRIDSPEC_HSPACE)),
    )
    ax = fig.add_subplot(gs[0, 0])
    ax_note = fig.add_subplot(gs[1, 0])
    colors = [_node_color(n) for n in df["name"]]
    ax.barh(range(len(df)), df["strength"], color=colors,
            edgecolor="black", linewidth=0.3)
    ax.set_yticks(range(len(df)))
    ax.set_yticklabels([_image_label_short(x, ascii_dot=True) for x in df["name"]], fontsize=9.6, fontproperties=body_fp)
    for lbl in ax.get_yticklabels():
        lbl.set_fontproperties(body_fp)
        lbl.set_fontweight("bold")
    ax.set_xlabel("Strength（加權度中心性）", fontsize=14, fontproperties=body_fp, fontweight="bold")
    for lbl in ax.get_xticklabels():
        lbl.set_fontproperties(body_fp)
        lbl.set_fontsize(11)
        lbl.set_fontweight("bold")
    ax.set_title(
        "加權度中心性排名（Strength）",
        fontproperties=title_fp,
        fontweight="bold",
        pad=10,
    )
    ax.grid(axis="x", alpha=0.3)

    from matplotlib.patches import Patch
    legend_handles = [Patch(facecolor=c, label=str(g))
                      for g, c in _BOOK_GROUP_COLORS.items()]
    leg = ax.legend(handles=legend_handles, loc="lower right", fontsize=8.8,
                    prop=body_fp, framealpha=0.8)
    for t in leg.get_texts():
        t.set_fontweight("bold")

    _draw_endnote_box(ax_note, info_text, body_fp, layout, fontsize=float(_CLUSTER_BODY_FONTSIZE), linespacing=1.36)

    fig.subplots_adjust(
        left=0.20,
        right=0.97,
        top=0.96,
        bottom=0.05,
        hspace=float(layout.get("hspace", _CLUSTER_GRIDSPEC_HSPACE)),
    )
    fig.savefig(outpath, dpi=_CLUSTER_EXPORT_DPI)
    plt.close(fig)


def _set_docx_cell_font(
    cell,
    text,
    cn_font: str = "宋体",
    en_font: str = "Times New Roman",
    font_size: float = 10.5,
    bold: bool = False,
    align: str = "left",
):
    """设置 Word 表格单元格字体（中文宋体、英文 Times New Roman、黑色）。"""
    from docx.shared import Pt

    cell.text = ""
    para = cell.paragraphs[0]
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    para.paragraph_format.line_spacing = 1.0
    para.paragraph_format.first_line_indent = Pt(0)
    para.paragraph_format.left_indent = Pt(0)
    para.paragraph_format.right_indent = Pt(0)
    al = str(align or "left").strip().lower()
    if al == "center":
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif al == "right":
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run(str(text))
    run.font.name = str(en_font)
    run.font.size = Pt(font_size)
    run.font.bold = bool(bold)
    run.font.color.rgb = RGBColor(0, 0, 0)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:eastAsia"), str(cn_font))
    rFonts.set(qn("w:ascii"), str(en_font))
    rFonts.set(qn("w:hAnsi"), str(en_font))


def _set_docx_table_fixed_layout(table) -> None:
    try:
        table.autofit = False
    except Exception:
        pass
    try:
        tbl = table._tbl
        tblPr = tbl.tblPr
        if tblPr is None:
            return
        tbl_layout = tblPr.find(qn("w:tblLayout"))
        if tbl_layout is None:
            tbl_layout = OxmlElement("w:tblLayout")
            tblPr.append(tbl_layout)
        tbl_layout.set(qn("w:type"), "fixed")
    except Exception:
        pass


def _set_docx_table_col_widths(table, widths_cm: List[float]) -> None:
    widths = [float(w) for w in list(widths_cm or []) if float(w) > 0]
    if not widths:
        return
    for row in table.rows:
        for i, wcm in enumerate(widths):
            if i >= len(row.cells):
                break
            cell = row.cells[i]
            try:
                cell.width = Cm(wcm)
            except Exception:
                pass
            try:
                tcPr = cell._tc.get_or_add_tcPr()
                tcW = tcPr.find(qn("w:tcW"))
                if tcW is None:
                    tcW = OxmlElement("w:tcW")
                    tcPr.append(tcW)
                tcW.set(qn("w:type"), "dxa")
                tcW.set(qn("w:w"), str(int(round(wcm * 567.0))))
            except Exception:
                pass


def append_clustering_section(doc, sim_matrix, doc_names, plot_paths,
                              centrality_df, font_setting=""):
    """在 Word 文档中追加聚类分析章节。"""
    from docx.shared import Inches

    _configure_word_builtin_styles(doc)

    doc.add_heading("聚類與中心性分析", level=1)

    # --- 中心性指标 ---
    doc.add_heading("中心性指標", level=2)
    p = doc.add_paragraph(style="Normal")
    p.add_run("Strength（加權度中心性）").bold = True
    p.add_run("：節點所有邊權之和，反映文本與語料庫整體的相似程度。")
    p = doc.add_paragraph(style="Normal")
    p.add_run("Eigenvector（特徵向量中心性）").bold = True
    p.add_run("：遞迴計算，與重要節點相連的節點獲得更高權重。")
    p = doc.add_paragraph(style="Normal")
    p.add_run("Betweenness（中介中心性）").bold = True
    p.add_run("：衡量節點在最短路徑上的橋接作用，值越高表示節點越處於連接不同群組的關鍵位置。")

    has_centrality = centrality_df is not None and len(centrality_df) > 0
    if not has_centrality:
        doc.add_paragraph("中心性資料不可用（本輪未生成或文本數量不足）。", style="Normal")

    # 中心性排名表
    if has_centrality:
        df = centrality_df.sort_values("rank_strength")
        table = doc.add_table(rows=1 + len(df), cols=7, style="Table Grid")
        headers = ["排名", "文本", "書籍", "Strength", "Eigenvector", "Betweenness", "Betw.排名"]
        for j, h in enumerate(headers):
            _set_docx_cell_font(table.rows[0].cells[j], h, bold=True)
        for idx, (_, row) in enumerate(df.iterrows()):
            cells = table.rows[idx + 1].cells
            _set_docx_cell_font(cells[0], row["rank_strength"])
            _set_docx_cell_font(cells[1], row["name"])
            _set_docx_cell_font(cells[2], row["book"])
            _set_docx_cell_font(cells[3], f'{row["strength"]:.4f}')
            _set_docx_cell_font(cells[4], f'{row["eigenvector"]:.4f}')
            _set_docx_cell_font(cells[5], f'{row["betweenness"]:.4f}')
            _set_docx_cell_font(cells[6], row["rank_betweenness"])

        # 按书籍分组统计
        doc.add_paragraph("按書籍分組平均中心性：", style="Normal")
        grp_df = centrality_df.groupby("book")[["strength", "eigenvector", "betweenness"]].mean()
        grp_df = grp_df.sort_values("strength", ascending=False)
        t2 = doc.add_table(rows=1 + len(grp_df), cols=4, style="Table Grid")
        for j, h in enumerate(["書籍", "Avg Strength", "Avg Eigenvector", "Avg Betweenness"]):
            _set_docx_cell_font(t2.rows[0].cells[j], h, bold=True)
        for idx, (book, row) in enumerate(grp_df.iterrows()):
            cells = t2.rows[idx + 1].cells
            _set_docx_cell_font(cells[0], book)
            _set_docx_cell_font(cells[1], f'{row["strength"]:.4f}')
            _set_docx_cell_font(cells[2], f'{row["eigenvector"]:.4f}')
            _set_docx_cell_font(cells[3], f'{row["betweenness"]:.4f}')

    # --- Ward 聚类 ---
    doc.add_heading("Ward 層次聚類", level=2)
    p = doc.add_paragraph(style="Normal")
    p.add_run("Ward 法通過最小化合併後的組內方差來逐步構建層次聚類結構。"
              "距離矩陣定義為 dist(i,j) = 1 − cos_sim(i,j)。")

    dist = 1.0 - np.array(sim_matrix, dtype=float)
    np.fill_diagonal(dist, 0)
    dist = np.clip(dist, 0, None)
    condensed = squareform(dist, checks=False)
    Z = linkage(condensed, method="ward")

    for k in [2, 4]:
        labels = fcluster(Z, t=k, criterion="maxclust")
        doc.add_paragraph(f"k={k} 分群結果：", style="Normal")
        for c in sorted(set(labels)):
            members = [doc_names[i] for i in range(len(doc_names)) if labels[i] == c]
            doc.add_paragraph(f"群 {c}：{', '.join(members)}", style="Normal")

    sil = silhouette_score(dist, fcluster(Z, t=2, criterion="maxclust"),
                           metric="precomputed")
    doc.add_paragraph(f"Silhouette Score (k=2)：{sil:.4f}", style="Normal")

    if "dendrogram" in plot_paths and os.path.isfile(plot_paths["dendrogram"]):
        doc.add_picture(plot_paths["dendrogram"], width=Inches(6.0))

    # --- MDS ---
    doc.add_heading("MDS 降維可視化", level=2)
    p = doc.add_paragraph(style="Normal")
    p.add_run("MDS（多維尺度分析）將高維距離矩陣投影到二維平面，"
              "盡量保留文本間的相對距離關係。")
    if "mds" in plot_paths and os.path.isfile(plot_paths["mds"]):
        doc.add_picture(plot_paths["mds"], width=Inches(6.0))

    # --- PCA ---
    if "pca" in plot_paths and os.path.isfile(plot_paths.get("pca", "")):
        doc.add_heading("PCA 主成分分析", level=2)
        p = doc.add_paragraph(style="Normal")
        p.add_run("PCA（主成分分析）對 TF-IDF 特徵矩陣進行特徵值分解，"
                  "提取方差最大的兩個方向作為座標軸。載荷向量（箭頭）"
                  "表示各 n-gram 對主成分的貢獻，可揭示驅動文本群分離的"
                  "關鍵詞彙特徵。")
        doc.add_picture(plot_paths["pca"], width=Inches(6.0))

    # --- 網絡圖 ---
    if "network" in plot_paths and os.path.isfile(plot_paths.get("network", "")):
        doc.add_heading("網絡關係圖", level=2)
        p = doc.add_paragraph(style="Normal")
        p.add_run("力導向網絡圖（Spring Layout）以節點代表文本，"
                  "邊寬度反映餘弦相似度大小。同組文本間的邊以該組"
                  "顏色著色，跨組邊以灰色表示。節點按文獻來源分組著色。")
        doc.add_picture(plot_paths["network"], width=Inches(6.0))

    # --- 中心性图 ---
    doc.add_heading("中心性排名圖", level=2)
    if "centrality" in plot_paths and os.path.isfile(plot_paths["centrality"]):
        doc.add_picture(plot_paths["centrality"], width=Inches(6.0))


def run_clustering_standalone(csv_path, out_dir, font_setting=_CLUSTER_DEFAULT_FONT):
    """从 CSV 余弦相似度矩阵运行聚类+中心性分析，生成图表和专属 Word 报告。"""
    _ensure_ml_stack()
    os.makedirs(out_dir, exist_ok=True)
    ts_prefix = datetime.now().strftime("%Y%m%d%H%M")

    df = pd.read_csv(csv_path, index_col=0)
    sim_matrix = df.values.astype(float)
    np.fill_diagonal(sim_matrix, 1.0)
    doc_names = list(df.index)

    centrality_df = compute_centrality_metrics(sim_matrix, doc_names)

    path_dendrogram = os.path.join(out_dir, f"{ts_prefix}_圖4.1 Ward層次聚類樹狀圖（餘弦距離）.png")
    path_mds = os.path.join(out_dir, f"{ts_prefix}_圖4.2 MDS二維空間可視化（餘弦距離）.png")
    path_centrality = os.path.join(out_dir, f"{ts_prefix}_圖4.3 加權度中心性排名（Strength）.png")

    print("生成树状图 ...")
    plot_cluster_dendrogram(sim_matrix, doc_names, path_dendrogram, font_setting)
    print("生成 MDS 散点图 ...")
    plot_mds_scatter(sim_matrix, doc_names, path_mds, font_setting)
    print("生成中心性柱状图 ...")
    plot_centrality_bars(centrality_df, path_centrality, font_setting)

    path_network = os.path.join(out_dir, f"{ts_prefix}_圖4.5 文本關係網絡圖（餘弦相似度）.png")
    print("生成網絡關係圖 ...")
    try:
        plot_network_graph(sim_matrix, doc_names, path_network, font_setting)
    except Exception as e_net:
        path_network = ""
        print(f"[Network] Warning: {e_net}")

    plot_paths = {
        "dendrogram": path_dendrogram,
        "mds": path_mds,
        "centrality": path_centrality,
    }
    if path_network and os.path.isfile(path_network):
        plot_paths["network"] = path_network

    from docx import Document as DocxDoc
    doc = DocxDoc()
    append_clustering_section(doc, sim_matrix, doc_names, plot_paths,
                              centrality_df, font_setting)
    report_path = os.path.join(out_dir, f"{ts_prefix}_聚類與中心性分析報告.docx")
    doc.save(report_path)
    print(f"完成。報告：{report_path}")
    print(f"圖表：{path_dendrogram}")
    print(f"      {path_mds}")
    print(f"      {path_centrality}")


# ------------------- 矩陣熱力圖 -------------------
def plot_heatmap(
    sim_matrix,
    doc_names,
    outpath,
    title=None,
    strip_label_digits: bool = False,
    font_setting: str = "",
):
    # 固定画布与字体大小（用户指定）
    tick_fs = 12
    ann_fs = 12
    plt.figure(figsize=(21, 12))
    masked_matrix = sim_matrix.copy()
    np.fill_diagonal(masked_matrix, 1.0)

    font_prop = resolve_heatmap_font(font_setting)
    font_prop = _apply_font_to_rcparams(font_prop)

    # 紅白色對比色階: 0為白色，1為紅色
    from matplotlib.colors import LinearSegmentedColormap
    reds_white = LinearSegmentedColormap.from_list('reds_white', ['#ffffff', '#ff0000'])
    cmap = reds_white

    im = plt.imshow(masked_matrix, interpolation='nearest', aspect='auto', cmap=cmap, vmin=0, vmax=1)

    labels = build_heatmap_labels(doc_names, strip_digits=strip_label_digits)
    # X轴竖排显示（从上至下），避免使用稀有竖排标点字形导致缺字。
    xlabels = [_verticalize_heatmap_label(label) for label in labels]
    if font_prop:
        plt.xticks(
            range(len(doc_names)), xlabels,
            rotation=0, fontproperties=font_prop, va='top', fontsize=tick_fs
        )
        plt.yticks(range(len(doc_names)), labels, fontproperties=font_prop, fontsize=tick_fs)
    else:
        plt.xticks(range(len(doc_names)), xlabels, rotation=0, va='top', fontsize=tick_fs)
        plt.yticks(range(len(doc_names)), labels, fontsize=tick_fs)

    # 竖排X轴标签首字符紧贴x轴线，通过set_y或set_position调整
    ax = plt.gca()
    for tick in ax.get_xticklabels():
        # 默认y=0，向上偏移一点。我们要让首字符更靠近x轴，适当上移
        tick.set_y(0)

    ax = plt.gca()
    # 對角線顏色設為正紅色
    for i in range(masked_matrix.shape[0]):
        ax.add_patch(
            plt.Rectangle((i-0.5, i-0.5), 1, 1, fill=True, color=(1,0,0), linewidth=0)
        )

    # 顯示格子內數值（保留三位小數，顯示為數值，不帶百分號）
    for i in range(masked_matrix.shape[0]):
        for j in range(masked_matrix.shape[1]):
            val = masked_matrix[i, j]
            color = 'black'
            if i == j:
                color = 'white'
            # 以百分比顯示，保留两位小數，不帶百分號
            txt = f"{val*100:.2f}"
            if font_prop:
                plt.text(
                    j, i, txt, ha='center', va='center',
                    color=color,
                    fontsize=ann_fs,
                    fontproperties=font_prop
                )
            else:
                plt.text(
                    j, i, txt, ha='center', va='center',
                    color=color,
                    fontsize=ann_fs
                )
    # 移除額外的右側大標註，只保留 colorbar 的百分號顯示

    if title is None:
        title = '圖1.1《論語》《史記》《家語》《衣鏡》（長度歸一化）餘弦相似度矩陣'
    if font_prop:
        plt.title(title, fontproperties=font_prop, fontsize=24)
    else:
        plt.title(title, fontsize=24)
    cbar = plt.colorbar(im, fraction=0.046, pad=0.04)
    cbar.ax.tick_params(labelsize=tick_fs)
    # colorbar 百分比顯示，右側顯示百分號
    cbar_label = "相似度 (%)"
    if font_prop:
        cbar.set_label(cbar_label, fontproperties=font_prop)
    else:
        cbar.set_label(cbar_label)
    cbar.ax.yaxis.set_major_formatter(matplotlib.ticker.FuncFormatter(lambda x, _: f"{x*100:.0f}%"))
    if font_prop:
        for t in cbar.ax.get_yticklabels():
            t.set_fontproperties(font_prop)
    plt.tight_layout()
    plt.savefig(outpath, dpi=600)
    plt.close()


_TOP_CANDIDATES_RECALL_RATIO = 0.10
_TOP_CANDIDATES_MAX_ROWS_SINGLE_PAGE = 30
_TOP_CANDIDATES_TITLE = "Top30 TF-COS 候选对摘要（辅助指标：TF-IDF / Jaccard / Embed / Reranker）"


def plot_top_candidates_image(
    candidate_rows: List[Dict[str, Any]],
    outpath: str,
    recall_ratio: float = _TOP_CANDIDATES_RECALL_RATIO,
    title: str = _TOP_CANDIDATES_TITLE,
    font_setting: str = "",
    max_rows_single_page: int = _TOP_CANDIDATES_MAX_ROWS_SINGLE_PAGE,
):
    """
    生成候選對摘要圖（A4 直式表格）：
    - 以 cos raw 為主指標排序（輸入 rows 已按 cos 排序）
    - 召回 top 10%（可調 recall_ratio），並限制為單頁可讀行數
    - 指標欄位格式：rank | score（rank 在前）
    - 表頭居中、首列「序號」稍加寬，整體緊湊排版
    """
    if not candidate_rows:
        raise ValueError("No candidate rows for plotting.")
    font_prop = resolve_heatmap_font(font_setting)
    _apply_font_to_rcparams(font_prop)

    # 固定候選對圖片版式（對齊用戶確認版）
    style = {
        "fig_w": 11.20,
        "fig_h": 11.69,
        "title_y": 0.995,
        "subtitle_y": 0.955,
        "algo_y": 0.942,
        # 标题放大 1 倍（相对当前版本）。
        "title_fs": 21.0,
        # 注释字号与正文保持一致。
        "subtitle_fs": 8.2,
        "algo_fs": 8.2,
        "header_fs": 8.4,
        "body_fs": 8.2,
        # 上移表格，压缩标题区与表头之间的留白。
        "table_top": 0.928,
        "table_bottom": 0.030,
        "table_bbox_x": 0.020,
        "table_bbox_w": 0.960,
        "edge_color": "#E7CBA6",
        "edge_lw": 0.35,
        "header_bg": "#D96B2B",
        "header_fg": "white",
        "row_odd_bg": "#FFF1D9",
        "row_even_bg": "#FFF8E8",
        "body_fg": "#6A3F1D",
        "title_color": "#A54A1F",
        "subtitle_color": "#BF5A24",
        "algo_color": "#B8511F",
        "header_h_scale": 1.05,
        "body_h_scale": 0.98,
        "save_dpi": _CLUSTER_EXPORT_DPI,
        "save_pad_inches": 0.03,
        "margin_l": 0.02,
        "margin_r": 0.98,
        "margin_t": 0.99,
        "margin_b": 0.02,
        "path_max_len": 12,
    }

    def _fmt_path_short(s: str, max_len: int = 12) -> str:
        t = _image_label_short(s, ascii_dot=False).replace("\n", " ").strip()
        # 僅在本圖中：論語篇章補全為「論語·篇名」格式
        if t in _ANALECTS_CHAPTER_SET:
            t = f"論語·{t}"
        if len(t) > max_len:
            t = t[: max(1, max_len - 1)] + "…"
        return f"《{t}》"

    def _fmt_rank_score(rank_val, score_val) -> str:
        rk = "-"
        try:
            if rank_val is not None:
                rk = str(int(rank_val))
        except Exception:
            rk = "-"
        return f"{rk} | {_fmt_metric(score_val)}"

    total = len(candidate_rows)
    n_recall = max(1, int(math.ceil(total * float(recall_ratio))))
    max_rows = int(max(1, max_rows_single_page))
    n_keep = min(n_recall, max_rows)
    top_rows = candidate_rows[:n_keep]

    # ---- 構建表格數據 ----
    col_labels = [
        "序號",
        "文本A",
        "文本B",
        "TF COS\n(rank|score)",
        "TF-IDF\n(rank|score)",
        "Jaccard\n(rank|score)",
        "Embed\n(rank|score)",
        "Reranker\n(rank|score)",
    ]
    cell_text = []
    for idx, r in enumerate(top_rows):
        p1 = _fmt_path_short(r.get("path1", r.get("name1", "")), max_len=int(style["path_max_len"]))
        p2 = _fmt_path_short(r.get("path2", r.get("name2", "")), max_len=int(style["path_max_len"]))
        cell_text.append([
            str(idx + 1),
            p1,
            p2,
            _fmt_rank_score(r.get("rank_cos_raw", r.get("rank")), r.get("cos_raw")),
            _fmt_rank_score(r.get("rank_tfidf"), r.get("tfidf")),
            _fmt_rank_score(r.get("rank_jaccard"), r.get("jaccard")),
            _fmt_rank_score(r.get("rank_embed"), r.get("embed")),
            _fmt_rank_score(r.get("rank_reranker_raw"), r.get("reranker_raw")),
        ])

    n_rows = len(cell_text)
    # 拉寬橫軸：文本A/B雙列，首列略寬
    col_widths = [0.08, 0.23, 0.23, 0.092, 0.092, 0.092, 0.092, 0.092]
    # 確保總和歸一
    cw_sum = sum(col_widths)
    col_widths = [w / cw_sum for w in col_widths]

    # 固定縱向，橫向加寬（便於放下 top30）
    fig_w = float(style["fig_w"])
    fig_h = float(style["fig_h"])

    fig, ax = plt.subplots(figsize=(fig_w, fig_h))
    ax.axis("off")
    fig.patch.set_facecolor("white")

    # ---- 標題 ----
    ax.text(
        0.50, float(style["title_y"]), title,
        transform=ax.transAxes, ha="center", va="top",
        fontsize=float(style["title_fs"]), fontweight="bold", color=str(style["title_color"]),
        fontproperties=font_prop,
    )
    trunc_note = ""
    if n_keep < n_recall:
        trunc_note = f"；單頁顯示上限={n_keep}（原召回={n_recall}）"
    ax.text(
        0.50, float(style["subtitle_y"]),
        f"召回規則：TF COS 前 {int(round(recall_ratio * 100))}%　|　顯示：{n_keep} / {total}{trunc_note}",
        transform=ax.transAxes, ha="center", va="top",
        fontsize=float(style["subtitle_fs"]), color=str(style["subtitle_color"]),
        fontproperties=font_prop,
    )
    ax.text(
        0.02, float(style["algo_y"]),
        "算法：TF COS=cos(tf)；TF-IDF=cos(tfidf)；Jaccard=|A∩B|/|A∪B|；"
        "Embed=cos(emb)；Reranker=Cross-Encoder 分數（僅重排）",
        transform=ax.transAxes, ha="left", va="top",
        fontsize=float(style["algo_fs"]), color=str(style["algo_color"]),
        fontproperties=font_prop,
    )

    # ---- 表格 ----
    table_top = float(style["table_top"])
    table_bottom = float(style["table_bottom"])
    table_h = table_top - table_bottom
    table = ax.table(
        cellText=cell_text,
        colLabels=col_labels,
        colWidths=col_widths,
        bbox=[float(style["table_bbox_x"]), table_bottom, float(style["table_bbox_w"]), table_h],
        cellLoc="center",
    )

    # 固定字號，避免不同輪次因行數變化導致版式漂移。
    body_fontsize = float(style["body_fs"])
    header_fontsize = float(style["header_fs"])
    row_h = table_h / max(1, (n_rows + 1))

    table.auto_set_font_size(False)

    for (row, col), cell in table.get_celld().items():
        cell.set_edgecolor(str(style["edge_color"]))
        cell.set_linewidth(float(style["edge_lw"]))
        if row == 0:
            # 暖色表頭（偏紅黃）
            cell.set_facecolor(str(style["header_bg"]))
            cell.set_text_props(
                color=str(style["header_fg"]), fontsize=header_fontsize,
                fontproperties=font_prop, fontweight="bold",
            )
            cell.set_height(row_h * float(style["header_h_scale"]))
        else:
            # 暖色斑馬紋（偏黃、柔和）
            if row % 2 == 0:
                cell.set_facecolor(str(style["row_even_bg"]))
            else:
                cell.set_facecolor(str(style["row_odd_bg"]))
            cell.set_text_props(
                color=str(style["body_fg"]), fontsize=body_fontsize,
                fontproperties=font_prop,
            )
            cell.set_height(row_h * float(style["body_h_scale"]))
            # 文本A/文本B列左對齊
            if col in (1, 2):
                cell.set_text_props(ha="left")
                cell._loc = "left"

    plt.subplots_adjust(
        left=float(style["margin_l"]),
        right=float(style["margin_r"]),
        top=float(style["margin_t"]),
        bottom=float(style["margin_b"]),
    )
    plt.savefig(
        outpath,
        dpi=int(style["save_dpi"]),
        bbox_inches="tight",
        pad_inches=float(style["save_pad_inches"]),
    )
    plt.close(fig)
    return outpath


# ------------------- 報告生成 -------------------
def _to_float_or_default(v, default: float = 0.0) -> float:
    try:
        x = float(v)
    except Exception:
        return float(default)
    if math.isnan(x) or math.isinf(x):
        return float(default)
    return x


def _fmt_metric(v) -> str:
    try:
        x = float(v)
    except Exception:
        return "n/a"
    if math.isnan(x) or math.isinf(x):
        return "n/a"
    return f"{x:.4f}"


def _resolve_pair_indices(raw_i, raw_j, n_docs: int) -> Optional[Tuple[int, int]]:
    try:
        i = int(raw_i)
        j = int(raw_j)
    except Exception:
        return None

    # Support both 1-based and 0-based ids.
    if 1 <= i <= n_docs and 1 <= j <= n_docs:
        i -= 1
        j -= 1

    if not (0 <= i < n_docs and 0 <= j < n_docs):
        return None
    if i == j:
        return None
    if i > j:
        i, j = j, i
    return i, j


def _build_report_pair_rows(
    doc_names: List[str],
    sim_matrix: np.ndarray,
    candidate_pairs: Optional[List[Dict[str, Any]]] = None,
) -> List[Dict[str, Any]]:
    n_docs = len(doc_names)
    rows: List[Dict[str, Any]] = []
    seen: set = set()

    for r in candidate_pairs or []:
        if not isinstance(r, dict):
            continue
        idx = _resolve_pair_indices(r.get("id1"), r.get("id2"), n_docs)
        if idx is None:
            continue
        i, j = idx
        if (i, j) in seen:
            continue
        seen.add((i, j))
        rows.append(
            {
                "idx1": i,
                "idx2": j,
                "id1": i + 1,
                "id2": j + 1,
                "path1": str(r.get("path1") or r.get("name1") or doc_names[i]),
                "path2": str(r.get("path2") or r.get("name2") or doc_names[j]),
                # legacy aliases kept for backward compatibility
                "name1": str(r.get("path1") or r.get("name1") or doc_names[i]),
                "name2": str(r.get("path2") or r.get("name2") or doc_names[j]),
                "cos_raw": _to_float_or_default(r.get("cos_raw"), _to_float_or_default(sim_matrix[i, j], 0.0)),
                "tfidf": (_to_float_or_default(r.get("tfidf"), 0.0) if ("tfidf" in r) else None),
                "jaccard": (_to_float_or_default(r.get("jaccard"), 0.0) if ("jaccard" in r) else None),
                "embed": (_to_float_or_default(r.get("embed"), 0.0) if ("embed" in r and r.get("embed") is not None) else None),
                "reranker_raw": (_to_float_or_default(r.get("reranker_raw"), 0.0) if ("reranker_raw" in r and r.get("reranker_raw") is not None) else None),
                "rank_cos_raw": (int(r.get("rank_cos_raw")) if r.get("rank_cos_raw") is not None else None),
                "rank_tfidf": (int(r.get("rank_tfidf")) if r.get("rank_tfidf") is not None else None),
                "rank_jaccard": (int(r.get("rank_jaccard")) if r.get("rank_jaccard") is not None else None),
                "rank_embed": (int(r.get("rank_embed")) if r.get("rank_embed") is not None else None),
                "rank_reranker_raw": (int(r.get("rank_reranker_raw")) if r.get("rank_reranker_raw") is not None else None),
            }
        )

    if not rows:
        for i in range(n_docs):
            for j in range(i + 1, n_docs):
                rows.append(
                    {
                        "idx1": i,
                        "idx2": j,
                        "id1": i + 1,
                        "id2": j + 1,
                        "path1": str(doc_names[i]),
                        "path2": str(doc_names[j]),
                        # legacy aliases kept for backward compatibility
                        "name1": str(doc_names[i]),
                        "name2": str(doc_names[j]),
                        "cos_raw": _to_float_or_default(sim_matrix[i, j], 0.0),
                        "tfidf": None,
                        "jaccard": None,
                        "embed": None,
                        "reranker_raw": None,
                        "rank_cos_raw": None,
                        "rank_tfidf": None,
                        "rank_jaccard": None,
                        "rank_embed": None,
                        "rank_reranker_raw": None,
                    }
                )

    rows.sort(
        key=lambda x: (
            _to_float_or_default(x.get("cos_raw"), 0.0),
            _to_float_or_default(x.get("tfidf"), 0.0),
            _to_float_or_default(x.get("jaccard"), 0.0),
        ),
        reverse=True,
    )
    for rk, r in enumerate(rows, start=1):
        r["rank"] = rk
        if r.get("rank_cos_raw") is None:
            r["rank_cos_raw"] = rk
    _annotate_metric_rank(rows, "tfidf", "rank_tfidf", descending=True)
    _annotate_metric_rank(rows, "jaccard", "rank_jaccard", descending=True)
    _annotate_metric_rank(rows, "embed", "rank_embed", descending=True)
    _annotate_metric_rank(rows, "reranker_raw", "rank_reranker_raw", descending=True)
    return rows


def _clip_report_text(text: str, max_chars: int = 120) -> str:
    t = str(text or "")
    m = int(max(1, max_chars))
    if len(t) <= m:
        return t
    return t[:m] + "..."


def _report_doc_label(name: str) -> str:
    t = str(name or "").strip()
    return t if t else "untitled"


def _pairwise_doc_label(name: str) -> str:
    """Pairwise 文档标签：去掉书名号，保留现有命名体系。"""
    t = _report_doc_label(name)
    t = t.replace("《", "").replace("》", "")
    t = re.sub(r"^[0-9０-９]+", "", t).strip()
    return t if t else "untitled"


def _pairwise_color_label(name: str) -> str:
    """Pairwise 颜色清单显示名：统一到颜色映射规范名称。"""
    key = _normalize_color_key(name)
    if key == "衣鏡·孔子傳":
        return "衣鏡·孔子傳"
    if key == "衣鏡·弟子傳":
        return "衣鏡·弟子傳"
    if key.startswith("史記·孔子世家·"):
        return key
    if key == "史記·弟子列傳·選":
        return "史記·弟子列傳·選"
    if key == "家語·本姓解":
        return "家語·本姓解"
    if key == "家語·七十二弟子解·選":
        return "家語·七十二弟子解·選"
    if key in _ANALECTS_CHAPTER_SET:
        return key
    return _pairwise_doc_label(name)


def generate_pairwise_weights_doc(
    out_docx_path: str,
    doc_names: List[str],
    sim_matrix: np.ndarray,
    candidate_pairs: Optional[List[Dict[str, Any]]] = None,
    heading_title: str = "Pairwise Similarity Weights (cos raw tf)",
) -> None:
    """
    单独输出 Pairwise Similarity Weights 文档。
    格式示例：學而 -- 爲政 [weight=0.0271]
    """
    doc = DocxWriter()
    _configure_word_builtin_styles(doc)
    try:
        sec = doc.sections[0]
        sec.page_width = Cm(21.0)
        sec.page_height = Cm(29.7)
        sec.left_margin = Cm(2.0)
        sec.right_margin = Cm(2.0)
        sec.top_margin = Cm(2.0)
        sec.bottom_margin = Cm(2.0)
    except Exception:
        pass

    doc.add_heading(str(heading_title or "Pairwise Similarity Weights"), level=1)
    doc.add_paragraph(
        "表格格式：各相似度欄位均為 rank | score（rank 在前，score 在後）；"
        "首列為序號。為適配 A4 直式排版，採用緊湊表格。",
        style="Normal",
    )
    doc.add_paragraph(
        "算法說明：Cos Raw = cos(tf)（字面詞頻相似）；"
        "TF-IDF = cos(tfidf)（抑制高頻詞）；"
        "Jaccard = |A∩B|/|A∪B|（集合重疊率）；"
        "Embed = cos(embA, embB)（語義向量相似）；"
        "Reranker = Cross-Encoder 分數（候選重排）。",
        style="Normal",
    )

    pair_rows = _build_report_pair_rows(doc_names, sim_matrix, candidate_pairs=candidate_pairs)
    pair_rows.sort(
        key=lambda r: (
            int(r.get("rank_cos_raw") or 10**9),
            -_to_float_or_default(r.get("cos_raw"), 0.0),
        )
    )

    def _rank_score(rank_val, score_val) -> str:
        rk = "-"
        try:
            if rank_val is not None:
                rk = str(int(rank_val))
        except Exception:
            rk = "-"
        return f"{rk} | {_fmt_metric(score_val)}"

    headers = [
        "序號",
        "文本對",
        "Cos Raw\n(rank|score)",
        "TF-IDF\n(rank|score)",
        "Jaccard\n(rank|score)",
        "Embed\n(rank|score)",
        "Reranker\n(rank|score)",
    ]
    table = doc.add_table(rows=1, cols=len(headers), style="Table Grid")
    _set_docx_table_fixed_layout(table)
    _set_docx_table_col_widths(table, [1.2, 5.8, 1.8, 1.8, 1.8, 1.8, 1.8])

    for j, h in enumerate(headers):
        _set_docx_cell_font(table.rows[0].cells[j], h, bold=True, align="center")

    for i, r in enumerate(pair_rows, start=1):
        n1 = _pairwise_doc_label(r.get("path1", r.get("name1", "")))
        n2 = _pairwise_doc_label(r.get("path2", r.get("name2", "")))
        pair_txt = f"{n1} -- {n2}"
        row = table.add_row().cells
        _set_docx_cell_font(row[0], str(i), align="center")
        _set_docx_cell_font(row[1], pair_txt, align="left")
        _set_docx_cell_font(row[2], _rank_score(r.get("rank_cos_raw", r.get("rank")), r.get("cos_raw")), align="center")
        _set_docx_cell_font(row[3], _rank_score(r.get("rank_tfidf"), r.get("tfidf")), align="center")
        _set_docx_cell_font(row[4], _rank_score(r.get("rank_jaccard"), r.get("jaccard")), align="center")
        _set_docx_cell_font(row[5], _rank_score(r.get("rank_embed"), r.get("embed")), align="center")
        _set_docx_cell_font(row[6], _rank_score(r.get("rank_reranker_raw"), r.get("reranker_raw")), align="center")

    doc.add_heading("Default Node Colors", level=2)
    seen_labels: set = set()
    for raw_name in doc_names:
        lb = _pairwise_color_label(raw_name)
        if not lb or lb in seen_labels:
            continue
        seen_labels.add(lb)
        color = _node_color(str(raw_name))
        doc.add_paragraph(f"{lb} [color={color}]", style="Normal")

    doc.save(out_docx_path)


def _variant_unit_label(a_seg: str, b_seg: str) -> str:
    la = len(str(a_seg or ""))
    lb = len(str(b_seg or ""))
    return "字" if max(la, lb) <= 1 else "句"


def _levenshtein_diff_ops(a: str, b: str) -> Tuple[int, List[Tuple[str, str, str]]]:
    """Exact minimum edit distance + merged char-level operations."""
    sa = str(a or "")
    sb = str(b or "")
    n = len(sa)
    m = len(sb)
    dp = [[0] * (m + 1) for _ in range(n + 1)]
    for i in range(1, n + 1):
        dp[i][0] = i
    for j in range(1, m + 1):
        dp[0][j] = j

    for i in range(1, n + 1):
        ai = sa[i - 1]
        row = dp[i]
        row_prev = dp[i - 1]
        for j in range(1, m + 1):
            cost = 0 if ai == sb[j - 1] else 1
            row[j] = min(
                row_prev[j] + 1,          # delete
                row[j - 1] + 1,           # insert
                row_prev[j - 1] + cost,   # replace/equal
            )

    i, j = n, m
    rev_ops: List[Tuple[str, str, str]] = []
    while i > 0 or j > 0:
        if (
            i > 0
            and j > 0
            and sa[i - 1] == sb[j - 1]
            and dp[i][j] == dp[i - 1][j - 1]
        ):
            rev_ops.append(("equal", sa[i - 1], sb[j - 1]))
            i -= 1
            j -= 1
            continue
        if i > 0 and j > 0 and dp[i][j] == dp[i - 1][j - 1] + 1:
            rev_ops.append(("replace", sa[i - 1], sb[j - 1]))
            i -= 1
            j -= 1
            continue
        if i > 0 and dp[i][j] == dp[i - 1][j] + 1:
            rev_ops.append(("delete", sa[i - 1], ""))
            i -= 1
            continue
        if j > 0 and dp[i][j] == dp[i][j - 1] + 1:
            rev_ops.append(("insert", "", sb[j - 1]))
            j -= 1
            continue
        if i > 0 and j > 0:
            rev_ops.append(("replace", sa[i - 1], sb[j - 1]))
            i -= 1
            j -= 1
        elif i > 0:
            rev_ops.append(("delete", sa[i - 1], ""))
            i -= 1
        else:
            rev_ops.append(("insert", "", sb[j - 1]))
            j -= 1

    rev_ops.reverse()
    merged: List[Tuple[str, str, str]] = []
    for tag, aa, bb in rev_ops:
        if not merged or merged[-1][0] != tag:
            merged.append((tag, aa, bb))
        else:
            t, pa, pb = merged[-1]
            merged[-1] = (t, pa + aa, pb + bb)
    return int(dp[n][m]), merged


def _diff_counts(ops: List[Tuple[str, str, str]]) -> Tuple[int, int, int]:
    rep = 0
    ins = 0
    dele = 0
    for tag, aa, bb in ops:
        if tag == "replace":
            rep += len(aa)
        elif tag == "insert":
            ins += len(bb)
        elif tag == "delete":
            dele += len(aa)
    return int(rep), int(ins), int(dele)


def _diff_summary_han(a_text: str, b_text: str, max_items: int = 12) -> Dict[str, object]:
    dist, ops = _levenshtein_diff_ops(a_text, b_text)
    rep, ins, dele = _diff_counts(ops)
    notes: List[str] = []
    extra_notes = 0
    for tag, aa, bb in ops:
        if tag == "equal":
            continue
        if tag == "replace":
            msg = f"A作B：「{_clip_report_text(aa, 24)}」作「{_clip_report_text(bb, 24)}」"
        elif tag == "delete":
            msg = f"A无此{_variant_unit_label(aa, '')}：「{_clip_report_text(aa, 24)}」"
        else:  # insert
            msg = f"B增此{_variant_unit_label('', bb)}：「{_clip_report_text(bb, 24)}」"
        if len(notes) < int(max_items):
            notes.append(msg)
        else:
            extra_notes += 1

    norm_edit = float(dist / max(len(a_text), len(b_text), 1))
    return {
        "dist": int(dist),
        "rep": int(rep),
        "ins": int(ins),
        "dele": int(dele),
        "norm_edit": norm_edit,
        "notes": notes,
        "extra_notes": int(extra_notes),
    }


def _append_wordcount_section(doc: DocxWriter, doc_names: List[str], doc_texts: List[str]) -> None:
    doc.add_heading("文本字数统计", level=1)
    doc.add_paragraph("统计口径：仅保留中文汉字（去除标点、特殊符号、数字、英文）后计数。", style="Normal")
    rows: List[Dict[str, object]] = []
    for name, text in zip(doc_names, doc_texts):
        raw = str(text or "")
        han = preprocess_keep_han(raw)
        rows.append(
            {
                "name": str(name),
                "han_len": len(han),
                "raw_len": len(raw),
            }
        )
    if not rows:
        doc.add_paragraph("无可统计文本。", style="Normal")
        return

    rows.sort(key=lambda x: int(x["han_len"]), reverse=True)
    total_han = int(sum(int(r["han_len"]) for r in rows))
    total_raw = int(sum(int(r["raw_len"]) for r in rows))
    non_empty = int(sum(1 for r in rows if int(r["han_len"]) > 0))
    avg_han = float(total_han / max(1, len(rows)))
    p = doc.add_paragraph(style="Normal")
    p.add_run(f"文本数={len(rows)}；")
    p.add_run(f"非空文本={non_empty}；")
    p.add_run(f"Han总字数={total_han}；")
    p.add_run(f"原文总字符={total_raw}；")
    p.add_run(f"平均Han字数/文本={avg_han:.2f}")

    tb = doc.add_table(rows=1, cols=5)
    h = tb.rows[0].cells
    h[0].text = "rank"
    h[1].text = "name"
    h[2].text = "han_chars"
    h[3].text = "raw_chars"
    h[4].text = "han_ratio(%)"
    for i, r in enumerate(rows, start=1):
        rr = tb.add_row().cells
        han_len = int(r["han_len"])
        raw_len = int(r["raw_len"])
        ratio = (100.0 * han_len / raw_len) if raw_len > 0 else 0.0
        rr[0].text = str(i)
        rr[1].text = str(r["name"])
        rr[2].text = str(han_len)
        rr[3].text = str(raw_len)
        rr[4].text = f"{ratio:.2f}"
def _append_top_samples_section(
    doc: DocxWriter,
    pair_rows: List[Dict[str, Any]],
    doc_texts: List[str],
    topn: int = 50,
) -> None:
    n_show = int(min(max(0, topn), len(pair_rows)))
    doc.add_heading(f"Top{int(topn)}样本（按 cos raw）", level=1)
    if n_show <= 0:
        doc.add_paragraph("无候选样本。", style="Normal")
        return
    doc.add_paragraph(f"候选总数={len(pair_rows)}；展示前 {n_show} 对。", style="Normal")

    for i, r in enumerate(pair_rows[:n_show], start=1):
        idx1 = int(r.get("idx1", -1))
        idx2 = int(r.get("idx2", -1))
        text1 = str(doc_texts[idx1]) if 0 <= idx1 < len(doc_texts) else ""
        text2 = str(doc_texts[idx2]) if 0 <= idx2 < len(doc_texts) else ""
        han1 = preprocess_keep_han(text1)
        han2 = preprocess_keep_han(text2)
        doc.add_paragraph(
            (
                f"#{i} cos={_fmt_metric(r.get('cos_raw'))} "
                f"tfidf={_fmt_metric(r.get('tfidf'))} "
                f"jaccard={_fmt_metric(r.get('jaccard'))} "
                f"embed={_fmt_metric(r.get('embed'))} "
                f"reranker={_fmt_metric(r.get('reranker_raw'))} "
                f"ranks(c/t/j/e/r)=({r.get('rank_cos_raw')},{r.get('rank_tfidf')},{r.get('rank_jaccard')},{r.get('rank_embed')},{r.get('rank_reranker_raw')}) "
                f"ids={int(r.get('id1', 0))}-{int(r.get('id2', 0))}"
            ),
            style="Normal",
        )
        p = doc.add_paragraph(style="Normal")
        p.add_run("A: ")
        p.add_run(str(r.get("path1", r.get("name1", ""))))
        p.add_run("\nB: ")
        p.add_run(str(r.get("path2", r.get("name2", ""))))
        doc.add_paragraph(f"A正文（仅汉字）: {_clip_report_text(han1, 160)}", style="Normal")
        doc.add_paragraph(f"B正文（仅汉字）: {_clip_report_text(han2, 160)}", style="Normal")


def _append_top_collation_section(
    doc: DocxWriter,
    pair_rows: List[Dict[str, Any]],
    doc_texts: List[str],
    top_percent: float = 3.0,
    detail_norm_edit_max: float = 0.35,
) -> None:
    pct = float(top_percent) if float(top_percent) > 0 else 3.0
    n_show = int(max(1, round(len(pair_rows) * (pct / 100.0)))) if pair_rows else 0
    n_show = int(min(max(0, n_show), len(pair_rows)))
    doc.add_heading(f"Top{pct:.1f}%自动校勘（按 cos raw）", level=1)
    if n_show <= 0:
        doc.add_paragraph("无候选对可做自动校勘。", style="Normal")
        return
    doc.add_paragraph(
        (
            f"规则：仅展示（不参与评分）。在 Top{pct:.1f}% 候选中按纯汉字逐字比对；"
            f"跳过完全一致，并仅保留 norm_edit <= {float(detail_norm_edit_max):.2f} 的条目；"
            f"长度乘积超过 {MAX_COLLATION_DP_CELLS} 的长文本对自动跳过。"
        ),
        style="Normal",
    )

    skipped_empty = 0
    skipped_exact = 0
    skipped_long = 0
    skipped_threshold = 0
    shown = 0
    for r in pair_rows[:n_show]:
        idx1 = int(r.get("idx1", -1))
        idx2 = int(r.get("idx2", -1))
        text1 = str(doc_texts[idx1]) if 0 <= idx1 < len(doc_texts) else ""
        text2 = str(doc_texts[idx2]) if 0 <= idx2 < len(doc_texts) else ""
        han1 = preprocess_keep_han(text1)
        han2 = preprocess_keep_han(text2)
        if not han1 and not han2:
            skipped_empty += 1
            continue
        if han1 == han2:
            skipped_exact += 1
            continue
        if len(han1) * len(han2) > int(MAX_COLLATION_DP_CELLS):
            skipped_long += 1
            continue

        diff = _diff_summary_han(han1, han2, max_items=12)
        if float(diff["norm_edit"]) > float(detail_norm_edit_max):
            skipped_threshold += 1
            continue

        shown += 1
        doc.add_paragraph(
            (
                f"#{shown} cos={_fmt_metric(r.get('cos_raw'))} "
                f"norm_edit={float(diff['norm_edit']):.4f} "
                f"异文字数={int(diff['dist'])}（替换{int(diff['rep'])}/增{int(diff['ins'])}/删{int(diff['dele'])}）"
            ),
            style="Normal",
        )
        p = doc.add_paragraph(style="Normal")
        p.add_run("A: ")
        p.add_run(str(r.get("path1", r.get("name1", ""))))
        p.add_run("\nB: ")
        p.add_run(str(r.get("path2", r.get("name2", ""))))

        notes = list(diff.get("notes", []))
        extra_notes = int(diff.get("extra_notes", 0))
        if notes:
            note_text = "；".join(str(x) for x in notes)
            if extra_notes > 0:
                note_text += f"；... 其余 {extra_notes} 条略"
            doc.add_paragraph(f"校勘记：{note_text}", style="Normal")
        doc.add_paragraph(f"A正文（仅汉字）: {_clip_report_text(han1, 220)}", style="Normal")
        doc.add_paragraph(f"B正文（仅汉字）: {_clip_report_text(han2, 220)}", style="Normal")

    doc.add_paragraph(
        (
            f"校勘统计：候选={n_show}，过滤空文本={skipped_empty}，"
            f"过滤完全一致={skipped_exact}，超长跳过={skipped_long}，"
            f"阈值外={skipped_threshold}，展开={shown}。"
        ),
        style="Normal",
    )


def generate_report(out_docx_path, doc_names, feature_names, doc_vectors, sim_matrix,
                    top_ngrams_info, heatmap_path, doc_texts, doc_tokens_list=None,
                    candidate_pairs=None, include_collation: bool = True,
                    heading_title: str = "N-gram Similarity Report",
                    heatmap_tfidf_path: str = "",
                    heatmap_jaccard_path: str = "",
                    heatmap_embed_path: str = ""):
    doc = DocxWriter()
    _configure_word_builtin_styles(doc)
    pair_rows = _build_report_pair_rows(doc_names, sim_matrix, candidate_pairs=candidate_pairs)
    display_names = [_report_doc_label(n) for n in doc_names]

    # New front sections (all Heading 1): wordcount / top50 / top3%-collation
    _append_wordcount_section(doc, doc_names, doc_texts)
    _append_top_samples_section(doc, pair_rows, doc_texts, topn=50)
    if include_collation:
        _append_top_collation_section(doc, pair_rows, doc_texts, top_percent=3.0)

    doc.add_heading(str(heading_title or "N-gram Similarity Report"), level=1)

    # 新章节: Pairwise Similarity Weights (移到最前面)
    doc.add_heading('Pairwise Similarity Weights', level=2)
    for r in pair_rows:
        doc.add_paragraph(
            (
                f"{_report_doc_label(r.get('path1', r.get('name1')))} -- {_report_doc_label(r.get('path2', r.get('name2')))} "
                f"[cos={_fmt_metric(r.get('cos_raw'))}, "
                f"tfidf={_fmt_metric(r.get('tfidf'))}, "
                f"jaccard={_fmt_metric(r.get('jaccard'))}, "
                f"embed={_fmt_metric(r.get('embed'))}, "
                f"reranker={_fmt_metric(r.get('reranker_raw'))}; "
                f"ranks(c/t/j/e/r)=({r.get('rank_cos_raw')},{r.get('rank_tfidf')},{r.get('rank_jaccard')},{r.get('rank_embed')},{r.get('rank_reranker_raw')})]"
            ),
            style="Normal",
        )

    doc.add_paragraph('Documents analyzed: ' + ', '.join([
        _report_doc_label(n) for n in display_names
    ]), style="Normal")

    # 分词结果（用于相似度计算的预处理 token）
    doc.add_heading('Word Tokenization (Han-only for similarity)', level=2)
    # 展示完整 N-gram 列表
    if doc_tokens_list is None:
        doc_tokens_list = [list(str(t or "")) for t in doc_texts]

    for name, text, tokens in zip(display_names, doc_texts, doc_tokens_list):
        doc.add_paragraph(f"{_report_doc_label(name)}: " + ' | '.join(tokens), style="Normal")

    # 熱力圖
    doc.add_heading('Similarity Matrix', level=2)
    pic_items = [
        ("cos raw", heatmap_path),
        ("tfidf", heatmap_tfidf_path),
        ("jaccard", heatmap_jaccard_path),
        ("embed", heatmap_embed_path),
    ]
    for label, p in pic_items:
        if not p or not os.path.exists(p):
            continue
        doc.add_paragraph(label, style="Normal")
        doc.add_picture(p, width=Inches(6))

    # Top N-grams
    doc.add_heading('Top N-grams', level=2)
    table = doc.add_table(rows=1, cols=3)
    hdr = table.rows[0].cells
    hdr[0].text = 'Rank'
    hdr[1].text = 'N-gram'
    hdr[2].text = 'Frequency'
    for cell in hdr:
        _set_docx_cell_font(cell, cell.text, bold=True)
    for rank, info in enumerate(top_ngrams_info, start=1):
        row = table.add_row().cells
        row[0].text = str(rank)
        row[1].text = info['ngram']
        row[2].text = str(info['freq'])
        for cell in row:
            _set_docx_cell_font(cell, cell.text)

    # N-gram Vectors
    doc.add_heading('N-gram Vectors', level=2)
    table2 = doc.add_table(rows=1, cols=1+len(doc_names))
    hdr = table2.rows[0].cells
    hdr[0].text = 'N-gram \\ Document'
    for i, name in enumerate(display_names):
        hdr[i+1].text = _report_doc_label(name)
    for cell in hdr:
        _set_docx_cell_font(cell, cell.text, bold=True)
    for info in top_ngrams_info:
        row = table2.add_row().cells
        row[0].text = info['ngram']
        for i, name in enumerate(doc_names):
            row[i+1].text = str(info['doc_vector_values'].get(name, 0))
        for cell in row:
            _set_docx_cell_font(cell, cell.text)

    # 新部分: Top-10 N-grams Between Pairs of Documents
    doc.add_heading('Top-10 N-grams Between Pairs of Documents', level=2)
    # 计算每对文档的top-10 n-grams
    table3 = doc.add_table(rows=1, cols=3)
    hdr3 = table3.rows[0].cells
    hdr3[0].text = 'Document Pair'
    hdr3[1].text = 'Top-10 N-grams'
    hdr3[2].text = 'Counts (doc1, doc2)'
    for cell in hdr3:
        _set_docx_cell_font(cell, cell.text, bold=True)
    for idx1 in range(len(doc_names)):
        for idx2 in range(idx1 + 1, len(doc_names)):
            name1 = _report_doc_label(display_names[idx1])
            name2 = _report_doc_label(display_names[idx2])
            vec1 = doc_vectors[idx1]
            vec2 = doc_vectors[idx2]
            # 共同出现的 n-gram 及其总频率
            common_idxs = np.where((vec1 > 0) & (vec2 > 0))[0]
            if len(common_idxs) == 0:
                continue
            # 按两个文档的 ngram 频率之和排序
            common_ngrams = []
            for idx in common_idxs:
                ng = feature_names[idx]
                freq_sum = vec1[idx] + vec2[idx]
                common_ngrams.append((ng, freq_sum, vec1[idx], vec2[idx]))
            # 排序取前10
            top_common = sorted(common_ngrams, key=lambda x: x[1], reverse=True)[:10]
            ngram_strs = [f"{ng}" for ng,_,_,_ in top_common]
            counts_strs = [f"{a},{b}" for _,__,a,b in top_common]
            row = table3.add_row().cells
            row[0].text = f"{name1} & {name2}"
            row[1].text = ', '.join(ngram_strs)
            row[2].text = '; '.join(counts_strs)
            for cell in row:
                _set_docx_cell_font(cell, cell.text)

    doc.save(out_docx_path)


def generate_report_all(
    out_docx_path: str,
    results_by_g: Dict[str, Dict[str, Any]],
    plot_paths_by_g: Optional[Dict[str, Dict[str, str]]] = None,
    heading_title: str = "N-gram Similarity Report (ALL: h1/h2/normal)",
) -> None:
    doc = DocxWriter()
    _configure_word_builtin_styles(doc)
    doc.add_heading(str(heading_title or "N-gram Similarity Report (ALL)"), level=1)
    doc.add_paragraph(f"Generated at: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", style="Normal")
    doc.add_paragraph("ALL 模式会连续运行 h1 / h2 / normal 三种切分并写入同一份报告。", style="Normal")

    for g in ANALYSIS_GRANULARITIES:
        doc.add_heading(f"{g.upper()} 模式", level=1)
        res = results_by_g.get(g) or {}
        if not res:
            doc.add_paragraph("该模式无可用结果。", style="Normal")
            continue

        doc_names = list(res.get("doc_names") or [])
        doc_texts = list(res.get("doc_texts") or [])
        sim_raw = res.get("sim_matrix")
        if sim_raw is None:
            doc.add_paragraph("该模式缺少相似度矩阵。", style="Normal")
            continue
        pair_rows = list(res.get("candidate_pairs") or [])
        top_ngrams_info = list(res.get("top_ngrams_info") or [])
        semantic_enabled = bool(res.get("semantic_enabled", True))
        reranker_enabled = bool(res.get("reranker_enabled", True))
        embed_model = str(res.get("embed_model") or DEFAULT_EMBED_MODEL_SIZE)
        reranker_model = str(res.get("reranker_model") or DEFAULT_RERANK_MODEL_SIZE)
        rerank_top_pct = float(
            res.get("rerank_top_percent_effective")
            or res.get("rerank_top_percent_non_h1", RERANK_TOP_PERCENT_NON_H1)
        )
        rerank_top_pct_cfg = float(res.get("rerank_top_percent_configured", rerank_top_pct))
        rerank_auto_scaled = bool(res.get("rerank_auto_scaled", False))
        embed_method = str(res.get("embed_method") or ("maxsim" if g == "h1" else "avgpool"))
        embed_error = str(res.get("embed_error") or "").strip()
        rerank_error = str(res.get("reranker_error") or "").strip()
        doc.add_paragraph(
            f"units={len(doc_names)}；pairs={len(pair_rows)}；preprocessing=Han-only(body only)",
            style="Normal",
        )
        doc.add_paragraph(
            f"semantic={'on' if semantic_enabled else 'off'}({_embed_model_label(embed_model)}), "
            f"embed_method={embed_method}；"
            f"reranker={'on' if reranker_enabled else 'off'}({_reranker_model_label(reranker_model)}), "
            f"reranker_policy={'all(H1/H2)' if g in ('h1', 'h2') else f'top{rerank_top_pct:g}%(Normal)'}",
            style="Normal",
        )
        if rerank_auto_scaled and g == "normal":
            doc.add_paragraph(
                f"reranker auto-scale applied: configured={rerank_top_pct_cfg:g}% -> effective={rerank_top_pct:g}%, "
                f"cap={int(res.get('rerank_cap_candidates') or RERANK_AUTO_NON_H1_MAX_CANDIDATES)}",
                style="Normal",
            )
        if embed_error:
            doc.add_paragraph(f"embed unavailable: {embed_error}", style="Normal")
        if rerank_error:
            doc.add_paragraph(f"reranker unavailable: {rerank_error}", style="Normal")

        _append_wordcount_section(doc, doc_names, doc_texts)
        _append_top_samples_section(doc, pair_rows, doc_texts, topn=50)
        if g == "normal":
            _append_top_collation_section(doc, pair_rows, doc_texts, top_percent=3.0)

        plots = (plot_paths_by_g or {}).get(g) or {}
        pic_items = [
            ("cos raw", plots.get("heatmap_cos")),
            ("tfidf", plots.get("heatmap_tfidf")),
            ("jaccard", plots.get("heatmap_jaccard")),
            ("embed", plots.get("heatmap_embed")),
            ("top pairs", plots.get("top_pairs_img")),
        ]
        has_pic = any(p and os.path.exists(p) for _, p in pic_items)
        if has_pic:
            doc.add_heading("图片结果", level=2)
            for label, p in pic_items:
                if not p or not os.path.exists(p):
                    continue
                doc.add_paragraph(label, style="Normal")
                try:
                    doc.add_picture(p, width=Inches(6.8))
                except Exception:
                    doc.add_paragraph(f"无法插入图片: {p}", style="Normal")

        if top_ngrams_info:
            doc.add_heading("Top N-grams", level=2)
            tb = doc.add_table(rows=1, cols=3)
            hd = tb.rows[0].cells
            hd[0].text = "rank"
            hd[1].text = "ngram"
            hd[2].text = "freq"
            for cell in hd:
                _set_docx_cell_font(cell, cell.text, bold=True)
            for i, info in enumerate(top_ngrams_info[:50], start=1):
                rr = tb.add_row().cells
                rr[0].text = str(i)
                rr[1].text = str(info.get("ngram", ""))
                rr[2].text = str(info.get("freq", 0))
                for cell in rr:
                    _set_docx_cell_font(cell, cell.text)
    doc.save(out_docx_path)


def _get_builtin_style(doc: DocxWriter, names: List[str], fallback: str = "Normal"):
    for name in names:
        try:
            return doc.styles[name]
        except Exception:
            continue
    try:
        return doc.styles[fallback]
    except Exception:
        return doc.styles[0]


def _apply_builtin_style_format(
    style,
    cn_font: str,
    en_font: str,
    size_pt: float,
    bold: bool,
    centered: bool,
    first_line_indent_pt: Optional[float] = None,
) -> None:
    font = style.font
    font.name = str(en_font)
    font.size = Pt(float(size_pt))
    font.bold = bool(bold)
    font.color.rgb = RGBColor(0, 0, 0)

    rPr = style._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:eastAsia"), str(cn_font))
    rFonts.set(qn("w:ascii"), str(en_font))
    rFonts.set(qn("w:hAnsi"), str(en_font))

    pf = style.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.line_spacing = 1.0
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER if centered else WD_ALIGN_PARAGRAPH.LEFT
    if first_line_indent_pt is None:
        pf.first_line_indent = None
    else:
        pf.first_line_indent = Pt(float(first_line_indent_pt))

    # Disable extra spacing between consecutive paragraphs of the same style.
    try:
        pPr = style._element.get_or_add_pPr()
        # Keep at most one contextualSpacing tag.
        old = pPr.findall(qn("w:contextualSpacing"))
        for node in old:
            pPr.remove(node)
        pPr.append(OxmlElement("w:contextualSpacing"))
    except Exception:
        pass


def _configure_word_builtin_styles(doc: DocxWriter) -> Dict[str, str]:
    """统一设置 Word 内置样式（h1/h2/h3/normal）到论文格式。"""
    st_h1 = _get_builtin_style(doc, ["Heading 1", "标题 1"], fallback="Normal")
    st_h2 = _get_builtin_style(doc, ["Heading 2", "标题 2"], fallback="Normal")
    st_h3 = _get_builtin_style(doc, ["Heading 3", "标题 3"], fallback="Normal")
    st_n = _get_builtin_style(doc, ["Normal", "正文"], fallback="Normal")

    # h1: 宋体小三加粗居中
    _apply_builtin_style_format(st_h1, cn_font="宋体", en_font="宋体", size_pt=15, bold=True, centered=True)
    # h2: 宋体四号加粗居中
    _apply_builtin_style_format(st_h2, cn_font="宋体", en_font="宋体", size_pt=14, bold=True, centered=True)
    # h3: 宋体小四加粗居中
    _apply_builtin_style_format(st_h3, cn_font="宋体", en_font="宋体", size_pt=12, bold=True, centered=True)
    # 正文: 中文宋体五号、英文 Times New Roman、常规、首行缩进两字(10.5*2=21pt)
    _apply_builtin_style_format(
        st_n,
        cn_font="宋体",
        en_font="Times New Roman",
        size_pt=10.5,
        bold=False,
        centered=False,
        first_line_indent_pt=21.0,
    )

    return {
        "h1": st_h1.name,
        "h2": st_h2.name,
        "h3": st_h3.name,
        "normal": st_n.name,
    }


def _configure_preprocessed_monitor_styles(doc: DocxWriter) -> Dict[str, str]:
    return _configure_word_builtin_styles(doc)


def _build_preprocessed_units_payload(
    excel_path: str,
    results_by_g: Optional[Dict[str, Dict[str, Any]]] = None,
) -> Dict[str, Dict[str, Any]]:
    out: Dict[str, Dict[str, Any]] = {}
    rb = results_by_g or {}
    for g in ANALYSIS_GRANULARITIES:
        units: List[Unit] = []
        clean_texts: List[str] = []
        res = rb.get(g) or {}
        if res:
            units = list(res.get("units") or [])
            clean_texts = [str(x or "") for x in list(res.get("doc_texts_for_similarity") or [])]
        if not units:
            try:
                units, _ = build_units_from_excel(excel_path, g)
            except Exception:
                units = []
        if len(clean_texts) != len(units):
            clean_texts = [preprocess_keep_han(u.text or "") for u in units]
        out[g] = {
            "units": units,
            "clean_texts": clean_texts,
        }
    return out


def generate_preprocessed_monitor_doc(
    out_docx_path: str,
    payload_by_g: Dict[str, Dict[str, Any]],
    heading_title: str = "预处理文本监测（h1/h2/normal）",
) -> None:
    doc = DocxWriter()
    styles = _configure_preprocessed_monitor_styles(doc)

    doc.add_paragraph(str(heading_title), style=styles["h1"])
    doc.add_paragraph(f"Generated at: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", style=styles["normal"])

    for g in ANALYSIS_GRANULARITIES:
        payload = payload_by_g.get(g) or {}
        units: List[Unit] = list(payload.get("units") or [])
        clean_texts: List[str] = [str(x or "") for x in list(payload.get("clean_texts") or [])]
        if len(clean_texts) != len(units):
            clean_texts = [preprocess_keep_han(u.text or "") for u in units]

        doc.add_paragraph(f"切分方式：{g}", style=styles["h3"])
        if not units:
            doc.add_paragraph("（无可用单元）", style=styles["normal"])
            continue

        last_h1 = None
        last_h2 = None
        for i, u in enumerate(units):
            h1 = str(u.h1 or "").strip()
            h2 = str(u.h2 or "").strip()

            if h1 and h1 != last_h1:
                doc.add_paragraph(h1, style=styles["h1"])
                last_h1 = h1
                last_h2 = None
            if h2 and h2 != last_h2:
                doc.add_paragraph(h2, style=styles["h2"])
                last_h2 = h2

            unit_tag = str(unit_path(u) or f"{g}#{i+1}")
            doc.add_paragraph(unit_tag, style=styles["h3"])
            body = clean_texts[i] if i < len(clean_texts) else ""
            doc.add_paragraph(body if body else "（预处理后为空）", style=styles["normal"])

    doc.save(out_docx_path)

# ------------------- GUI -------------------
class NgramApp:
    def __init__(self, master):
        self.master = master
        master.title("Single-character N-gram Similarity Tool")
        master.geometry("1650x1200")

        self.excel_path_var = tk.StringVar(value="")
        try:
            desktop = os.path.join(os.path.expanduser("~"), "Desktop")
        except Exception:
            desktop = os.path.expanduser("~")
        self.out_dir_var = tk.StringVar(value=desktop)

        self.doc_texts = []
        self.doc_names = []
        self.sim_matrix_embed = None
        self.heatmap_path_embed = ""

        self.n_var = tk.IntVar(value=3)
        self.topk_edge_var = tk.IntVar(value=10)
        self.topk_ngram_var = tk.IntVar(value=10)
        self.granularity_var = tk.StringVar(value="h2")
        default_embed_ui = "8B" if "8B" in QWEN_EMBED_MODEL_DIRS else _normalize_embed_size_label(DEFAULT_EMBED_MODEL_SIZE)
        if default_embed_ui not in QWEN_EMBED_MODEL_DIRS:
            default_embed_ui = "0.6B"
        default_rerank_ui = "8B" if "8B" in QWEN_RERANK_MODEL_DIRS else _normalize_reranker_size_label(DEFAULT_RERANK_MODEL_SIZE)
        if default_rerank_ui not in QWEN_RERANK_MODEL_DIRS:
            default_rerank_ui = "0.6B"
        self.use_semantic_var = tk.BooleanVar(value=True)
        self.semantic_model_size_var = tk.StringVar(value=default_embed_ui)
        self.use_reranker_var = tk.BooleanVar(value=True)
        self.reranker_model_size_var = tk.StringVar(value=default_rerank_ui)
        self.normal_candidate_topk_var = tk.IntVar(value=int(DEFAULT_TOPK_CANDIDATES))
        self.rerank_top_percent_var = tk.DoubleVar(value=float(RERANK_TOP_PERCENT_NON_H1))
        self.model_dir_hint_var = tk.StringVar(value="")
        self.show_examples_var = tk.BooleanVar(value=True)  # 控制是否顯示並計算例句
        self.use_clustering_var = tk.BooleanVar(value=True)  # 聚類與中心性分析
        self.result_cache_reuse_var = tk.BooleanVar(value=True)
        self.result_cache_write_var = tk.BooleanVar(value=True)
        self.heatmap_title_var = tk.StringVar(value='圖1.1《論語》《史記》《家語》《衣鏡》（長度歸一化）餘弦相似度矩陣')
        self.strip_label_digits_var = tk.BooleanVar(value=False)
        self.heatmap_font_var = tk.StringVar(value=DEFAULT_HEATMAP_FONT_PATH)
        self.status_var = tk.StringVar(value="Ready.")
        self.progress_var = tk.StringVar(value="Steps: 0/0")
        self._progress_total = 0
        self._run_started_at: Optional[float] = None
        self._run_started_wall: Optional[datetime] = None
        self._run_log_lines: List[str] = []

        frame = ttk.Frame(master)
        frame.pack(fill='both', expand=True, padx=8, pady=8)

        # 输入：仅 Excel（与 pipeline13 一致）
        file_frame = ttk.LabelFrame(frame, text="Input Excel")
        file_frame.pack(fill='x', padx=4, pady=4)
        ttk.Entry(file_frame, textvariable=self.excel_path_var).pack(side='left', fill='x', expand=True, padx=6, pady=6)
        ttk.Button(file_frame, text="Browse Excel", command=self.browse_excel).pack(side='left', padx=6, pady=6)

        # 输出目录
        out_frame = ttk.LabelFrame(frame, text="Output Folder")
        out_frame.pack(fill='x', padx=4, pady=4)
        ttk.Entry(out_frame, textvariable=self.out_dir_var).pack(side='left', fill='x', expand=True, padx=6, pady=6)
        ttk.Button(out_frame, text="Change Folder", command=self.browse_out_dir).pack(side='left', padx=6, pady=6)

        # 參數
        param_frame = ttk.LabelFrame(frame, text="Parameters")
        param_frame.pack(fill='x', padx=4, pady=4)
        ttk.Label(param_frame, text="n (n-grams):").pack(side='left', padx=6)
        ttk.Entry(param_frame, width=6, textvariable=self.n_var).pack(side='left')
        ttk.Label(param_frame, text="Top-K edges:").pack(side='left', padx=6)
        ttk.Entry(param_frame, width=6, textvariable=self.topk_edge_var).pack(side='left')
        ttk.Label(param_frame, text="Top-K N-grams:").pack(side='left', padx=6)
        ttk.Entry(param_frame, width=6, textvariable=self.topk_ngram_var).pack(side='left')
        ttk.Label(param_frame, text="Excel split:").pack(side='left', padx=6)
        ttk.Combobox(
            param_frame,
            textvariable=self.granularity_var,
            values=GRANULARITY_CHOICES,
            width=8,
            state="readonly",
        ).pack(side='left')
        ttk.Checkbutton(param_frame, text="Show example sentences", variable=self.show_examples_var).pack(side='left', padx=6)
        ttk.Checkbutton(param_frame, text="Strip digits in labels", variable=self.strip_label_digits_var).pack(side='left', padx=6)
        ttk.Checkbutton(param_frame, text="聚類分析", variable=self.use_clustering_var).pack(side='left', padx=6)
        ttk.Checkbutton(param_frame, text="复用缓存", variable=self.result_cache_reuse_var).pack(side='left', padx=6)
        ttk.Checkbutton(param_frame, text="写入缓存", variable=self.result_cache_write_var).pack(side='left', padx=6)
        ttk.Label(param_frame, text="Heatmap Font (name/path):").pack(side='left', padx=6)
        ttk.Entry(param_frame, width=28, textvariable=self.heatmap_font_var).pack(side='left')
        ttk.Label(param_frame, text="Heatmap Title:").pack(side='left', padx=6)
        ttk.Entry(param_frame, width=50, textvariable=self.heatmap_title_var).pack(side='left')
        ttk.Button(param_frame, text="Show Top-K", command=self.show_topk_ngrams).pack(side='left', padx=6)

        sem_frame = ttk.LabelFrame(frame, text="Semantic Models (Qwen3)")
        sem_frame.pack(fill='x', padx=4, pady=4)
        ttk.Checkbutton(
            sem_frame,
            text="Enable embedding similarity",
            variable=self.use_semantic_var,
        ).grid(row=0, column=0, sticky="w", padx=8, pady=4)
        ttk.Label(sem_frame, text="Embedding model:").grid(row=0, column=1, sticky="e", padx=6, pady=4)
        sem_model_box = ttk.Combobox(
            sem_frame,
            textvariable=self.semantic_model_size_var,
            values=MODEL_SIZE_CHOICES,
            width=8,
            state="readonly",
        )
        sem_model_box.grid(row=0, column=2, sticky="w", padx=4, pady=4)
        ttk.Checkbutton(
            sem_frame,
            text="Enable reranker",
            variable=self.use_reranker_var,
        ).grid(row=1, column=0, sticky="w", padx=8, pady=4)
        ttk.Label(sem_frame, text="Reranker model:").grid(row=1, column=1, sticky="e", padx=6, pady=4)
        rerank_model_box = ttk.Combobox(
            sem_frame,
            textvariable=self.reranker_model_size_var,
            values=MODEL_SIZE_CHOICES,
            width=8,
            state="readonly",
        )
        rerank_model_box.grid(row=1, column=2, sticky="w", padx=4, pady=4)
        ttk.Label(sem_frame, text="Normal recall Top-K:").grid(row=1, column=3, sticky="e", padx=6, pady=4)
        ttk.Entry(sem_frame, textvariable=self.normal_candidate_topk_var, width=8).grid(row=1, column=4, sticky="w", padx=4, pady=4)
        ttk.Label(sem_frame, text="Normal rerank Top%:").grid(row=1, column=5, sticky="e", padx=6, pady=4)
        ttk.Entry(sem_frame, textvariable=self.rerank_top_percent_var, width=8).grid(row=1, column=6, sticky="w", padx=4, pady=4)
        ttk.Label(
            sem_frame,
            text=(
                "Qwen3 local cache: 优先环境变量 NGRAM_HF_HUB_DIR；"
                "未设置时使用当前用户 ~/.cache/huggingface/hub。"
                f" 召回策略: h1/h2=all, normal=topk。Reranker: h1/h2=all, normal=top%。"
                f" normal 大样本自动降比例（cap={int(RERANK_AUTO_NON_H1_MAX_CANDIDATES)}）。"
            ),
            foreground="gray",
        ).grid(row=2, column=0, columnspan=7, sticky="w", padx=8, pady=(2, 4))
        ttk.Label(
            sem_frame,
            textvariable=self.model_dir_hint_var,
            foreground="gray",
            justify="left",
        ).grid(row=3, column=0, columnspan=7, sticky="w", padx=8, pady=(0, 4))
        sem_model_box.bind("<<ComboboxSelected>>", self._refresh_model_dir_hint)
        rerank_model_box.bind("<<ComboboxSelected>>", self._refresh_model_dir_hint)
        self._refresh_model_dir_hint()

        # 操作按鈕
        action_frame = ttk.Frame(frame)
        action_frame.pack(fill='x', padx=4, pady=4)
        self.run_btn = ttk.Button(action_frame, text="Run", command=self.run_analysis_thread)
        self.run_btn.pack(side='left')
        self.prog = ttk.Progressbar(action_frame, mode="determinate", length=280)
        self.prog.pack(side='left', padx=12)
        ttk.Label(action_frame, textvariable=self.progress_var).pack(side='left', padx=8)
        ttk.Label(action_frame, textvariable=self.status_var).pack(side='left', padx=8)

        # 輸出區
        output_frame = ttk.LabelFrame(frame, text="Results")
        output_frame.pack(fill='both', expand=True, padx=4, pady=4)
        self.text_out = tk.Text(output_frame)
        self.text_out.pack(fill='both', expand=True)

        log_frame = ttk.LabelFrame(frame, text="Run Log")
        log_frame.pack(fill='both', expand=False, padx=4, pady=4)
        self.log_out = ScrolledText(log_frame, height=10)
        self.log_out.pack(fill='both', expand=True, padx=4, pady=4)
        self.log_out.configure(state="disabled")

    def show_topk_ngrams(self):
        if not hasattr(self, 'top_ngrams_info') or not self.top_ngrams_info:
            messagebox.showwarning("No data", "Please run analysis first.")
            return
        k = max(1, int(self.topk_ngram_var.get()))
        show_examples = bool(self.show_examples_var.get())
        self.text_out.delete('1.0', 'end')
        self.text_out.insert('end', f"Top-{k} N-grams:\n\n")
        for i, info in enumerate(self.top_ngrams_info[:k], start=1):
            self.text_out.insert('end', f"{i}. {info['ngram']} (freq={info['freq']})\n")
            if show_examples:
                ex = info.get("examples", {})
                for doc_name, sents in ex.items():
                    if not sents:
                        continue
                    self.text_out.insert('end', f"   - {doc_name}: {sents[0]}\n")
            self.text_out.insert('end', '\n')

    # ------------------- 文件操作 -------------------
    def browse_excel(self):
        path = filedialog.askopenfilename(
            title="Select Excel file",
            filetypes=[("Excel files", "*.xlsx *.xlsm *.xls"), ("All files", "*.*")],
        )
        if path:
            self.excel_path_var.set(path)

    def browse_out_dir(self):
        d = filedialog.askdirectory(title="Select output folder")
        if d:
            self.out_dir_var.set(d)

    def _refresh_model_dir_hint(self, _event=None):
        emb_size = _normalize_embed_size_label(self.semantic_model_size_var.get() or "0.6B")
        rr_size = _normalize_reranker_size_label(self.reranker_model_size_var.get() or "0.6B")
        emb_dir = QWEN_EMBED_MODEL_DIRS.get(emb_size, "")
        rr_dir = QWEN_RERANK_MODEL_DIRS.get(rr_size, "")
        self.model_dir_hint_var.set(
            f"Embed dir ({emb_size}): {emb_dir}\nReranker dir ({rr_size}): {rr_dir}"
        )

    def _validate_model_config(self) -> Optional[Dict[str, Any]]:
        try:
            use_semantic = bool(self.use_semantic_var.get())
            embed_model_size = _normalize_embed_size_label(self.semantic_model_size_var.get() or "0.6B")
            if embed_model_size not in QWEN_EMBED_MODEL_DIRS:
                raise ValueError(f"Unsupported embedding model size: {embed_model_size}")
            self.semantic_model_size_var.set(embed_model_size)

            use_reranker = bool(self.use_reranker_var.get())
            reranker_model_size = _normalize_reranker_size_label(self.reranker_model_size_var.get() or "0.6B")
            if reranker_model_size not in QWEN_RERANK_MODEL_DIRS:
                raise ValueError(f"Unsupported reranker model size: {reranker_model_size}")
            self.reranker_model_size_var.set(reranker_model_size)
            self._refresh_model_dir_hint()

            normal_candidate_topk = int(self.normal_candidate_topk_var.get())
            if normal_candidate_topk <= 0:
                raise ValueError("Normal recall Top-K must be >= 1.")

            rerank_top_percent = float(self.rerank_top_percent_var.get())
            if rerank_top_percent <= 0 or rerank_top_percent > 100:
                raise ValueError("Normal rerank Top% must be in (0, 100].")

            return {
                "use_semantic": use_semantic,
                "embed_model_size": embed_model_size,
                "use_reranker": use_reranker,
                "reranker_model_size": reranker_model_size,
                "normal_candidate_topk": int(normal_candidate_topk),
                "rerank_top_percent_non_h1": rerank_top_percent,
            }
        except Exception as e:
            messagebox.showerror("Error", f"Invalid semantic/reranker config: {e}")
            return None

    def _safe_filename(self, text: str) -> str:
        t = re.sub(r'[\\/:*?"<>|]+', "_", str(text or ""))
        t = re.sub(r"\s+", "_", t).strip("._")
        return t or "output"

    def _model_name_tag(
        self,
        use_semantic: bool,
        embed_model_size: str,
        use_reranker: bool,
        reranker_model_size: str,
    ) -> str:
        parts: List[str] = []
        if use_semantic:
            parts.append(_embed_model_label(embed_model_size))
        if use_reranker:
            parts.append(_reranker_model_label(reranker_model_size))
        if not parts:
            return "Lexical"
        return "+".join(parts)

    def _make_run_tag(
        self,
        excel_path: str,
        granularity: str,
        use_semantic: bool,
        embed_model_size: str,
        use_reranker: bool,
        reranker_model_size: str,
    ) -> str:
        ts_prefix = datetime.now().strftime("%Y%m%d%H%M")
        algo = "small pipe"
        g = normalize_granularity_choice(granularity)
        upload_base = self._safe_filename(os.path.splitext(os.path.basename(excel_path))[0])
        model_tag = self._safe_filename(
            self._model_name_tag(
                use_semantic=use_semantic,
                embed_model_size=embed_model_size,
                use_reranker=use_reranker,
                reranker_model_size=reranker_model_size,
            )
        )
        return f"{ts_prefix}＋{algo}＋{g}＋{upload_base}＋{model_tag}"

    def _resolve_output_run_dir(self, selected_out_dir: str, run_tag: str) -> str:
        """
        Keep named-run folder behavior, but avoid nesting the same folder twice:
        /base/run_tag/run_tag -> /base/run_tag
        """
        base = str(selected_out_dir or "").strip()
        if not base:
            return str(run_tag)
        try:
            tail = os.path.basename(os.path.normpath(base))
        except Exception:
            tail = os.path.basename(base)
        if str(tail) == str(run_tag):
            return base
        return os.path.join(base, run_tag)

    def _on_ui(self, fn, *args, **kwargs):
        try:
            if threading.current_thread() is threading.main_thread():
                fn(*args, **kwargs)
            else:
                self.master.after(0, lambda: fn(*args, **kwargs))
        except Exception:
            pass

    def _set_run_enabled(self, enabled: bool):
        def _do():
            try:
                self.run_btn.configure(state=("normal" if enabled else "disabled"))
            except Exception:
                pass
        self._on_ui(_do)

    def _status_set(self, msg: str):
        self._on_ui(self.status_var.set, str(msg))

    def _clear_run_log(self):
        self._run_log_lines = []

        def _do():
            try:
                self.log_out.configure(state="normal")
                self.log_out.delete("1.0", "end")
                self.log_out.configure(state="disabled")
            except Exception:
                pass

        self._on_ui(_do)

    def _reset_progress(self, total_steps: int):
        self._progress_total = max(0, int(total_steps))

        def _do():
            try:
                self.prog["maximum"] = max(1, int(total_steps))
                self.prog["value"] = 0
            except Exception:
                pass
            self.progress_var.set(f"Steps: 0/{max(0, int(total_steps))}")

        self._on_ui(_do)

    def _ui_progress(self, done: int, total: int, msg: str = ""):
        td = max(0, int(done))
        tt = max(1, int(total))

        def _do():
            try:
                self.prog["maximum"] = tt
                self.prog["value"] = min(td, tt)
            except Exception:
                pass
            self.progress_var.set(f"Steps: {min(td, tt)}/{tt}")
            if msg:
                self.status_var.set(str(msg))

        self._on_ui(_do)

    def _append_log(self, msg: str):
        now_tag = time.strftime("%H:%M:%S")
        if self._run_started_at is None:
            prefix = f"[{now_tag}] "
        else:
            elapsed = _fmt_elapsed(time.perf_counter() - float(self._run_started_at))
            prefix = f"[{now_tag} +{elapsed}] "
        line = prefix + str(msg)
        self._run_log_lines.append(line)

        def _do():
            try:
                self.log_out.configure(state="normal")
                self.log_out.insert("end", line + "\n")
                self.log_out.see("end")
                self.log_out.configure(state="disabled")
            except Exception:
                pass

        self._on_ui(_do)

    def _estimate_total_steps(self, granularity_choice: str) -> int:
        if normalize_granularity_choice(granularity_choice) == "all":
            # 3x analyze + 3x plot + excel + report + preprocessed_doc + log_report + finalize
            return 11
        # analyze + plot + excel + report + preprocessed_doc + log_report + finalize
        return 7

    def _write_run_log_report(
        self,
        out_dir: str,
        run_tag: str,
        config: Dict[str, Any],
        summary_lines: List[str],
        output_paths: List[str],
    ) -> str:
        os.makedirs(out_dir, exist_ok=True)
        log_path = os.path.join(out_dir, f"{run_tag}_runlog.txt")
        end_wall = datetime.now()
        start_wall = self._run_started_wall or end_wall
        elapsed = _fmt_elapsed(time.perf_counter() - float(self._run_started_at or time.perf_counter()))
        lines: List[str] = []
        lines.append("N-gram Similarity Run Log")
        lines.append("=" * 72)
        lines.append(f"Start:   {start_wall.strftime('%Y-%m-%d %H:%M:%S')}")
        lines.append(f"End:     {end_wall.strftime('%Y-%m-%d %H:%M:%S')}")
        lines.append(f"Elapsed: {elapsed}")
        lines.append("")
        lines.append("Config:")
        for k in (
            "excel_path",
            "output_base_dir",
            "out_dir",
            "output_name_tag",
            "granularity",
            "n",
            "topk_edge",
            "show_examples",
            "strip_label_digits",
            "heatmap_font_setting",
            "use_semantic",
            "embed_model_size",
            "use_reranker",
            "reranker_model_size",
            "normal_candidate_topk",
            "rerank_top_percent_non_h1",
            "result_cache_reuse",
            "result_cache_write",
            "result_cache_key",
        ):
            if k in config:
                lines.append(f"- {k}: {config[k]}")
        if summary_lines:
            lines.append("")
            lines.append("Summary:")
            lines.extend([f"- {x}" for x in summary_lines])
        if output_paths:
            lines.append("")
            lines.append("Outputs:")
            lines.extend([f"- {p}" for p in output_paths])
        lines.append("")
        lines.append("Timeline:")
        lines.extend(self._run_log_lines)
        lines.append("")
        with open(log_path, "w", encoding="utf-8") as f:
            f.write("\n".join(lines))
        return log_path

    def _result_cache_root(self) -> str:
        env_dir = str(os.environ.get("NGRAM_RESULT_CACHE_DIR", "") or "").strip()
        if env_dir:
            return os.path.expanduser(env_dir)
        return os.path.join(os.path.expanduser("~"), ".cache", "lex_small_pipe_result_cache")

    def _result_cache_dir(self, cache_key: str) -> str:
        return os.path.join(self._result_cache_root(), str(cache_key or ""))

    def _build_result_cache_signature(
        self,
        *,
        excel_path: str,
        granularity_choice: str,
        n: int,
        show_examples: bool,
        strip_label_digits: bool,
        heatmap_font_setting: str,
        heatmap_title: str,
        use_semantic: bool,
        embed_model_size: str,
        use_reranker: bool,
        reranker_model_size: str,
        normal_candidate_topk: int,
        rerank_top_percent_non_h1: float,
        use_clustering: bool,
    ) -> Tuple[str, Dict[str, Any]]:
        p = os.path.abspath(str(excel_path or ""))
        st_size = 0
        st_mtime_ns = 0
        try:
            st = os.stat(p)
            st_size = int(st.st_size)
            st_mtime_ns = int(st.st_mtime_ns)
        except Exception:
            pass
        excel_fp = {
            "path": p,
            "size": st_size,
            "mtime_ns": st_mtime_ns,
            "sha1": _file_sha1(p),
        }
        payload = {
            "cache_schema": RESULT_CACHE_SCHEMA_VERSION,
            "code_version": PIPELINE_CODE_VERSION,
            "script_file": os.path.abspath(__file__),
            "script_sha1": _script_sha1(),
            "excel": excel_fp,
            "params": {
                "granularity": normalize_granularity_choice(granularity_choice),
                "n": int(max(1, int(n))),
                "show_examples": bool(show_examples),
                "strip_label_digits": bool(strip_label_digits),
                "heatmap_font_setting": str(heatmap_font_setting or "").strip(),
                "heatmap_title": str(heatmap_title or "").strip(),
                "use_semantic": bool(use_semantic),
                "embed_model_size": _normalize_embed_size_label(embed_model_size or DEFAULT_EMBED_MODEL_SIZE),
                "use_reranker": bool(use_reranker),
                "reranker_model_size": _normalize_reranker_size_label(reranker_model_size or DEFAULT_RERANK_MODEL_SIZE),
                "normal_candidate_topk": int(max(1, int(normal_candidate_topk))),
                "rerank_top_percent_non_h1": float(rerank_top_percent_non_h1),
                "use_clustering": bool(use_clustering),
            },
        }
        # cache_key 仅用于「计算结果」复用：不随脚本样式/排版改动而变化。
        # 如分析算法有变更，请手工提升 RESULT_CACHE_ANALYSIS_VERSION。
        signature_payload = {
            "cache_schema": RESULT_CACHE_SCHEMA_VERSION,
            "analysis_signature_version": RESULT_CACHE_ANALYSIS_VERSION,
            "excel": {
                "path": str(excel_fp.get("path") or ""),
                "size": int(excel_fp.get("size") or 0),
                "mtime_ns": int(excel_fp.get("mtime_ns") or 0),
                "sha1": str(excel_fp.get("sha1") or ""),
            },
            "params": {
                "granularity": normalize_granularity_choice(payload["params"].get("granularity")),
                "n": int(payload["params"].get("n") or 1),
                "use_semantic": bool(payload["params"].get("use_semantic")),
                "embed_model_size": _normalize_embed_size_label(payload["params"].get("embed_model_size")),
                "use_reranker": bool(payload["params"].get("use_reranker")),
                "reranker_model_size": _normalize_reranker_size_label(payload["params"].get("reranker_model_size")),
                "normal_candidate_topk": int(payload["params"].get("normal_candidate_topk") or DEFAULT_TOPK_CANDIDATES),
                "rerank_top_percent_non_h1": float(payload["params"].get("rerank_top_percent_non_h1") or RERANK_TOP_PERCENT_NON_H1),
                "use_clustering": bool(payload["params"].get("use_clustering")),
            },
        }
        payload["analysis_signature_version"] = RESULT_CACHE_ANALYSIS_VERSION
        payload["signature_payload"] = signature_payload
        return _json_sha1(signature_payload), payload

    def _analysis_signature_from_cache_payload(self, cache_payload: Dict[str, Any]) -> Dict[str, Any]:
        payload = cache_payload if isinstance(cache_payload, dict) else {}
        excel = payload.get("excel") if isinstance(payload.get("excel"), dict) else {}
        params = payload.get("params") if isinstance(payload.get("params"), dict) else {}
        return {
            "cache_schema": RESULT_CACHE_SCHEMA_VERSION,
            "analysis_signature_version": str(
                payload.get("analysis_signature_version")
                or RESULT_CACHE_ANALYSIS_VERSION
            ),
            "excel": {
                "path": str(excel.get("path") or ""),
                "size": int(excel.get("size") or 0),
                "mtime_ns": int(excel.get("mtime_ns") or 0),
                "sha1": str(excel.get("sha1") or ""),
            },
            "params": {
                "granularity": normalize_granularity_choice(params.get("granularity")),
                "n": int(params.get("n") or 1),
                "use_semantic": bool(params.get("use_semantic")),
                "embed_model_size": _normalize_embed_size_label(params.get("embed_model_size")),
                "use_reranker": bool(params.get("use_reranker")),
                "reranker_model_size": _normalize_reranker_size_label(params.get("reranker_model_size")),
                "normal_candidate_topk": int(params.get("normal_candidate_topk") or DEFAULT_TOPK_CANDIDATES),
                "rerank_top_percent_non_h1": float(params.get("rerank_top_percent_non_h1") or RERANK_TOP_PERCENT_NON_H1),
                "use_clustering": bool(params.get("use_clustering")),
            },
        }

    def _find_compatible_result_cache_key(
        self,
        *,
        expected_mode: str,
        cache_payload: Dict[str, Any],
    ) -> str:
        """在现有缓存中查找与当前分析参数匹配的最新 key（忽略样式/脚本版本差异）。"""
        root = self._result_cache_root()
        if not os.path.isdir(root):
            return ""
        target_sig = self._analysis_signature_from_cache_payload(cache_payload)
        target_sig_hash = _json_sha1(target_sig)
        best_key = ""
        best_ts = float("-inf")
        try:
            entries = sorted(os.listdir(root))
        except Exception:
            return ""
        for key in entries:
            cache_dir = os.path.join(root, key)
            manifest_path = os.path.join(cache_dir, "manifest.json")
            analysis_path = os.path.join(cache_dir, "analysis.pkl")
            outputs_dir = os.path.join(cache_dir, "outputs")
            if not (os.path.isfile(manifest_path) and os.path.isfile(analysis_path) and os.path.isdir(outputs_dir)):
                continue
            try:
                with open(manifest_path, "r", encoding="utf-8") as f:
                    manifest = json.load(f)
            except Exception:
                continue
            if not isinstance(manifest, dict):
                continue
            if str(manifest.get("mode") or "") != str(expected_mode):
                continue
            cp = manifest.get("cache_payload") if isinstance(manifest.get("cache_payload"), dict) else {}
            sig_hash = _json_sha1(self._analysis_signature_from_cache_payload(cp))
            if sig_hash != target_sig_hash:
                continue
            ts = float(manifest.get("generated_ts") or 0.0)
            if ts >= best_ts:
                best_ts = ts
                best_key = str(key)
        return best_key

    def _save_result_cache_bundle(
        self,
        *,
        cache_key: str,
        cache_payload: Dict[str, Any],
        mode: str,
        analysis_payload: Dict[str, Any],
        output_paths: List[str],
        out_dir: str,
    ) -> Tuple[bool, str, int]:
        cache_dir = self._result_cache_dir(cache_key)
        outputs_dir = os.path.join(cache_dir, "outputs")
        analysis_path = os.path.join(cache_dir, "analysis.pkl")
        manifest_path = os.path.join(cache_dir, "manifest.json")

        try:
            if os.path.isdir(cache_dir):
                shutil.rmtree(cache_dir, ignore_errors=True)
            os.makedirs(outputs_dir, exist_ok=True)

            copied_entries: List[Dict[str, Any]] = []
            out_root = os.path.abspath(str(out_dir or ""))
            for raw_src in output_paths or []:
                src = str(raw_src or "").strip()
                if not src or (not os.path.isfile(src)):
                    continue
                src_abs = os.path.abspath(src)
                try:
                    rel = os.path.relpath(src_abs, out_root)
                    if rel.startswith(".."):
                        rel = os.path.basename(src_abs)
                except Exception:
                    rel = os.path.basename(src_abs)
                rel = str(rel).replace("\\", "/").lstrip("/")
                if not rel:
                    rel = os.path.basename(src_abs)
                if ".." in rel.split("/"):
                    rel = os.path.basename(rel)
                dst = os.path.join(outputs_dir, rel)
                os.makedirs(os.path.dirname(dst), exist_ok=True)
                shutil.copy2(src_abs, dst)
                copied_entries.append(
                    {
                        "rel": rel,
                        "name": os.path.basename(rel),
                        "size": int(os.path.getsize(dst)) if os.path.isfile(dst) else 0,
                    }
                )

            with open(analysis_path, "wb") as f:
                pickle.dump(
                    {
                        "mode": str(mode),
                        "analysis": analysis_payload or {},
                    },
                    f,
                    protocol=pickle.HIGHEST_PROTOCOL,
                )

            manifest = {
                "cache_schema": RESULT_CACHE_SCHEMA_VERSION,
                "generated_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                "generated_ts": float(time.time()),
                "code_version": PIPELINE_CODE_VERSION,
                "script_file": os.path.abspath(__file__),
                "script_sha1": _script_sha1(),
                "cache_key": str(cache_key),
                "mode": str(mode),
                "cache_payload": cache_payload or {},
                "analysis_file": "analysis.pkl",
                "outputs_dir": "outputs",
                "outputs": copied_entries,
            }
            with open(manifest_path, "w", encoding="utf-8") as f:
                json.dump(manifest, f, ensure_ascii=False, indent=2)
            return True, cache_dir, len(copied_entries)
        except Exception as e:
            return False, str(e), 0

    def _restore_result_cache_bundle(
        self,
        *,
        cache_key: str,
        expected_mode: str,
        out_dir: str,
    ) -> Optional[Dict[str, Any]]:
        cache_dir = self._result_cache_dir(cache_key)
        manifest_path = os.path.join(cache_dir, "manifest.json")
        analysis_path = os.path.join(cache_dir, "analysis.pkl")
        outputs_dir = os.path.join(cache_dir, "outputs")
        if not (os.path.isfile(manifest_path) and os.path.isfile(analysis_path) and os.path.isdir(outputs_dir)):
            return None

        try:
            with open(manifest_path, "r", encoding="utf-8") as f:
                manifest = json.load(f)
            if not isinstance(manifest, dict):
                return None
        except Exception:
            return None

        mode = str(manifest.get("mode") or "")
        if mode != str(expected_mode):
            return None

        try:
            with open(analysis_path, "rb") as f:
                payload = pickle.load(f)
            if not isinstance(payload, dict):
                return None
            analysis_obj = payload.get("analysis")
            if not isinstance(analysis_obj, dict):
                return None
        except Exception:
            return None

        restored_paths: List[str] = []
        for item in list(manifest.get("outputs") or []):
            if not isinstance(item, dict):
                continue
            rel = str(item.get("rel") or "").replace("\\", "/").lstrip("/")
            if not rel:
                continue
            if ".." in rel.split("/"):
                rel = os.path.basename(rel)
            src = os.path.join(outputs_dir, rel)
            if not os.path.isfile(src):
                continue
            dst = os.path.join(out_dir, rel)
            try:
                os.makedirs(os.path.dirname(dst), exist_ok=True)
                shutil.copy2(src, dst)
                if os.path.isfile(dst):
                    restored_paths.append(dst)
            except Exception:
                continue

        if not restored_paths:
            return None

        return {
            "cache_dir": cache_dir,
            "manifest": manifest,
            "analysis": analysis_obj,
            "output_paths": restored_paths,
        }

    def _post_warning(self, title: str, message: str):
        self.master.after(0, lambda t=title, m=message: messagebox.showwarning(t, m))

    def _post_error(self, title: str, message: str):
        self.master.after(0, lambda t=title, m=message: messagebox.showerror(t, m))

    def _post_info(self, title: str, message: str):
        self.master.after(0, lambda t=title, m=message: messagebox.showinfo(t, m))

    def _collect_output_lines(self) -> List[str]:
        out_list = list(getattr(self, "output_paths", []) or [])
        if out_list:
            return out_list
        keys = [
            "excel",
            "report",
            "preprocessed_monitor",
            "run_log",
            "pairwise_weights",
            "heatmap_cos",
            "heatmap_tfidf",
            "heatmap_jaccard",
            "heatmap_embed",
            "top_pairs_img",
            "cluster_dendrogram",
            "cluster_mds",
            "cluster_centrality",
            "cluster_pca",
            "cluster_network",
            "cluster_report",
        ]
        out: List[str] = []
        for k in keys:
            v = str(self.auto_output_paths.get(k, "")).strip() if hasattr(self, "auto_output_paths") else ""
            if v:
                out.append(v)
        return out

    def _render_analysis_output(self, k_edges: int):
        self.show_topk_ngrams()
        self.text_out.insert('end', f"\n\n{self.input_desc}\n")
        pairs = list(getattr(self, "candidate_pairs", []) or [])
        if not pairs:
            self.text_out.insert('end', "\nNo candidate pairs.\n")
            return
        k_edges = min(max(1, int(k_edges)), len(pairs))
        self.text_out.insert('end', f"\nTop-{k_edges} candidate pairs (cos/raw sorted):\n")
        for row in pairs[:k_edges]:
            self.text_out.insert(
                'end',
                (
                    f"{int(row['rank'])}. {row.get('path1', row.get('name1', ''))} <> {row.get('path2', row.get('name2', ''))} | "
                    f"cos={row['cos_raw']:.4f}, tfidf={row['tfidf']:.4f}, jaccard={row['jaccard']:.4f}, "
                    f"embed={_fmt_metric(row.get('embed'))}, reranker={_fmt_metric(row.get('reranker_raw'))}\n"
            ),
        )
        self.text_out.insert(
            'end',
            (
                "\nAll outputs:\n"
                + "\n".join([f"- {p}" for p in self._collect_output_lines()])
                + "\n"
            ),
        )

    def _render_analysis_output_all(self, results_by_g: Dict[str, Dict[str, Any]], skipped_by_g: Dict[str, str], k_edges: int):
        self.text_out.delete('1.0', 'end')
        self.text_out.insert('end', "ALL mode (h1/h2/normal)\n")
        self.text_out.insert('end', f"{self.input_desc}\n")
        for g in ANALYSIS_GRANULARITIES:
            if g in skipped_by_g:
                self.text_out.insert('end', f"\n[{g}] skipped: {skipped_by_g[g]}\n")
                continue
            res = results_by_g.get(g) or {}
            pairs = list(res.get("candidate_pairs") or [])
            names = list(res.get("doc_names") or [])
            self.text_out.insert('end', f"\n[{g}] units={len(names)} pairs={len(pairs)}\n")
            show_k = min(max(1, int(k_edges)), len(pairs)) if pairs else 0
            for row in pairs[:show_k]:
                self.text_out.insert(
                    'end',
                    (
                        f"  {int(row['rank'])}. {row.get('path1', row.get('name1', ''))} <> {row.get('path2', row.get('name2', ''))} | "
                        f"cos={row['cos_raw']:.4f}, tfidf={row['tfidf']:.4f}, jaccard={row['jaccard']:.4f}, "
                        f"embed={_fmt_metric(row.get('embed'))}, reranker={_fmt_metric(row.get('reranker_raw'))}\n"
                    ),
                )
        self.text_out.insert(
            'end',
            "\nAll outputs:\n" + "\n".join([f"- {p}" for p in self._collect_output_lines()]) + "\n",
        )

    def _write_matrix_sheet(self, ws, labels: List[str], matrix: Optional[np.ndarray]):
        if matrix is None or not isinstance(matrix, np.ndarray) or matrix.ndim != 2:
            ws.append(["Matrix unavailable"])
            return
        ws.append(["Document"] + labels)
        for i, name in enumerate(labels):
            ws.append([name] + [float(matrix[i, j]) for j in range(len(labels))])

    def _write_excel_workbook_single(self, save_path: str, result: Dict[str, Any]):
        labels = [str(n) for n in (result.get("doc_names") or [])]
        sim_raw = result["sim_matrix"]
        sim_tfidf = result["sim_matrix_tfidf"]
        sim_jaccard = result["sim_matrix_jaccard"]
        sim_embed = result.get("sim_matrix_embed")
        pairs = list(result.get("candidate_pairs") or [])

        wb = openpyxl.Workbook()
        ws_pairs = wb.active
        ws_pairs.title = "pairs_total"
        ws_pairs.append(
            [
                "rank_cos_raw",
                "rank_tfidf",
                "rank_jaccard",
                "rank_embed",
                "rank_reranker_raw",
                "id1",
                "id2",
                "Path1",
                "Path2",
                "cos_raw",
                "tfidf",
                "jaccard",
                "embed",
                "reranker_raw",
            ]
        )
        for r in pairs:
            ws_pairs.append([
                int(r.get("rank_cos_raw", r.get("rank", 0))),
                (None if r.get("rank_tfidf") is None else int(r.get("rank_tfidf"))),
                (None if r.get("rank_jaccard") is None else int(r.get("rank_jaccard"))),
                (None if r.get("rank_embed") is None else int(r.get("rank_embed"))),
                (None if r.get("rank_reranker_raw") is None else int(r.get("rank_reranker_raw"))),
                int(r.get("id1", 0)),
                int(r.get("id2", 0)),
                str(r.get("path1", r.get("name1", ""))),
                str(r.get("path2", r.get("name2", ""))),
                float(r.get("cos_raw", 0.0)),
                float(r.get("tfidf", 0.0)),
                float(r.get("jaccard", 0.0)),
                (None if r.get("embed") is None else float(r.get("embed"))),
                (None if r.get("reranker_raw") is None else float(r.get("reranker_raw"))),
            ])

        ws_raw = wb.create_sheet(title="matrix_cos_raw")
        self._write_matrix_sheet(ws_raw, labels, sim_raw)
        ws_tfidf = wb.create_sheet(title="matrix_tfidf")
        self._write_matrix_sheet(ws_tfidf, labels, sim_tfidf)
        ws_jaccard = wb.create_sheet(title="matrix_jaccard")
        self._write_matrix_sheet(ws_jaccard, labels, sim_jaccard)
        ws_embed = wb.create_sheet(title="matrix_embed")
        self._write_matrix_sheet(ws_embed, labels, sim_embed)

        # 聚類中心性分表（single 模式下若已計算）
        cdf = result.get("centrality_df")
        if cdf is not None and len(cdf) > 0:
            ws_cl = wb.create_sheet(title="centrality")
            ws_cl.append(["rank", "name", "book", "strength", "eigenvector", "betweenness",
                          "rank_eigenvector", "rank_betweenness"])
            for _, row in cdf.iterrows():
                ws_cl.append([
                    int(row["rank_strength"]),
                    str(row["name"]),
                    str(row["book"]),
                    float(row["strength"]),
                    float(row["eigenvector"]),
                    float(row["betweenness"]),
                    int(row["rank_eigenvector"]),
                    int(row["rank_betweenness"]),
                ])
        wb.save(save_path)

    def _write_excel_workbook_all(self, save_path: str, results_by_g: Dict[str, Dict[str, Any]]):
        wb = openpyxl.Workbook()
        ws0 = wb.active
        ws0.title = "h1_pairs"
        created_first = False

        for g in ANALYSIS_GRANULARITIES:
            res = results_by_g.get(g)
            if not res:
                continue
            labels = [str(n) for n in (res.get("doc_names") or [])]
            sim_raw = res["sim_matrix"]
            sim_tfidf = res["sim_matrix_tfidf"]
            sim_jaccard = res["sim_matrix_jaccard"]
            sim_embed = res.get("sim_matrix_embed")
            pairs = list(res.get("candidate_pairs") or [])
            top_ngrams = list(res.get("top_ngrams_info") or [])

            if not created_first:
                ws_pairs = ws0
                ws_pairs.title = f"{g}_pairs"
                created_first = True
            else:
                ws_pairs = wb.create_sheet(title=f"{g}_pairs")
            ws_pairs.append(
                [
                    "rank_cos_raw",
                    "rank_tfidf",
                    "rank_jaccard",
                    "rank_embed",
                    "rank_reranker_raw",
                    "id1",
                    "id2",
                    "Path1",
                    "Path2",
                    "cos_raw",
                    "tfidf",
                    "jaccard",
                    "embed",
                    "reranker_raw",
                ]
            )
            for r in pairs:
                ws_pairs.append([
                    int(r.get("rank_cos_raw", r.get("rank", 0))),
                    (None if r.get("rank_tfidf") is None else int(r.get("rank_tfidf"))),
                    (None if r.get("rank_jaccard") is None else int(r.get("rank_jaccard"))),
                    (None if r.get("rank_embed") is None else int(r.get("rank_embed"))),
                    (None if r.get("rank_reranker_raw") is None else int(r.get("rank_reranker_raw"))),
                    int(r.get("id1", 0)),
                    int(r.get("id2", 0)),
                    str(r.get("path1", r.get("name1", ""))),
                    str(r.get("path2", r.get("name2", ""))),
                    float(r.get("cos_raw", 0.0)),
                    float(r.get("tfidf", 0.0)),
                    float(r.get("jaccard", 0.0)),
                    (None if r.get("embed") is None else float(r.get("embed"))),
                    (None if r.get("reranker_raw") is None else float(r.get("reranker_raw"))),
                ])

            ws_raw = wb.create_sheet(title=f"{g}_matrix_cos")
            self._write_matrix_sheet(ws_raw, labels, sim_raw)
            ws_tfidf = wb.create_sheet(title=f"{g}_matrix_tfidf")
            self._write_matrix_sheet(ws_tfidf, labels, sim_tfidf)
            ws_jaccard = wb.create_sheet(title=f"{g}_matrix_jaccard")
            self._write_matrix_sheet(ws_jaccard, labels, sim_jaccard)
            ws_embed = wb.create_sheet(title=f"{g}_matrix_embed")
            self._write_matrix_sheet(ws_embed, labels, sim_embed)

            ws_ng = wb.create_sheet(title=f"{g}_top_ngrams")
            ws_ng.append(["rank", "ngram", "freq"])
            for i, info in enumerate(top_ngrams[:200], start=1):
                ws_ng.append([i, str(info.get("ngram", "")), int(info.get("freq", 0))])

            # 聚類中心性分表（如果已計算）
            cdf = res.get("centrality_df")
            if cdf is not None and len(cdf) > 0:
                ws_cl = wb.create_sheet(title=f"{g}_centrality")
                ws_cl.append(["rank", "name", "book", "strength", "eigenvector", "betweenness",
                              "rank_eigenvector", "rank_betweenness"])
                for _, row in cdf.iterrows():
                    ws_cl.append([
                        int(row["rank_strength"]),
                        str(row["name"]),
                        str(row["book"]),
                        float(row["strength"]),
                        float(row["eigenvector"]),
                        float(row["betweenness"]),
                        int(row["rank_eigenvector"]),
                        int(row["rank_betweenness"]),
                    ])

        if not created_first:
            ws0.title = "empty"
            ws0.append(["No available results"])
        wb.save(save_path)

    def _build_top_ngrams_info(
        self,
        doc_names: List[str],
        doc_texts: List[str],
        doc_vectors: np.ndarray,
        feature_names: np.ndarray,
        show_examples: bool,
    ) -> List[Dict[str, Any]]:
        if doc_vectors is None or len(doc_vectors) == 0 or len(feature_names) == 0:
            return []
        total_counts = np.sum(doc_vectors, axis=0)
        if total_counts is None or len(total_counts) == 0:
            return []
        top_idx = np.argsort(total_counts)[::-1][:50]
        split_sents_by_doc: Dict[str, List[Tuple[str, str]]] = {}
        if show_examples:
            split_sents_by_doc = {
                name: [
                    (s.strip(), preprocess_keep_han(s))
                    for s in str(text or "").split('。')
                    if s and s.strip()
                ]
                for name, text in zip(doc_names, doc_texts)
            }

        top_ngrams_info: List[Dict[str, Any]] = []
        for idx in top_idx:
            ngram = feature_names[idx]
            freq = int(total_counts[idx])
            examples = defaultdict(list)
            if show_examples:
                for name in doc_names:
                    sents = split_sents_by_doc.get(name, [])
                    found = [raw for raw, han in sents if str(ngram) in han]
                    if found:
                        examples[name] = found[:3]
            doc_vals = {name: int(vec[idx]) for name, vec in zip(doc_names, doc_vectors)}
            top_ngrams_info.append(
                {
                    'ngram': str(ngram),
                    'freq': freq,
                    'examples': examples,
                    'doc_vector_values': doc_vals,
                }
            )
        return top_ngrams_info

    def _analyze_one_granularity(
        self,
        excel_path: str,
        granularity: str,
        n: int,
        show_examples: bool,
        use_semantic: bool = True,
        embed_model_size: str = DEFAULT_EMBED_MODEL_SIZE,
        use_reranker: bool = True,
        reranker_model_size: str = DEFAULT_RERANK_MODEL_SIZE,
        normal_candidate_topk: int = DEFAULT_TOPK_CANDIDATES,
        rerank_top_percent_non_h1: float = RERANK_TOP_PERCENT_NON_H1,
    ) -> Dict[str, Any]:
        g = normalize_style(granularity)
        units, input_desc = build_units_from_excel(excel_path, g)
        if len(units) < 2:
            raise ValueError(f"units<{2} ({len(units)})")

        doc_texts = [u.text for u in units]
        doc_names = [unit_display_name(u) for u in units]
        # Scheme-1: n-gram generation is segment-internal (no cross-chapter boundary).
        doc_texts_clean, doc_tokens_list, all_ngram_tokens = _build_lexical_inputs_by_segments(units, n)

        sim_pack = compute_similarity_from_ngrams(all_ngram_tokens)
        sim_raw = sim_pack["cos_raw"]
        sim_tfidf = sim_pack["tfidf"]
        sim_jaccard = sim_pack["jaccard"]
        feature_names = sim_pack["feature_names"]
        doc_vectors = sim_pack["doc_vectors"]
        sim_embed: Optional[np.ndarray] = None
        semantic_enabled = bool(use_semantic)
        embed_model_used = _normalize_embed_size_label(embed_model_size or DEFAULT_EMBED_MODEL_SIZE)
        embed_method = ("maxsim" if g == "h1" else "avgpool") if semantic_enabled else "disabled"
        embed_error = ""
        if semantic_enabled:
            try:
                sim_embed, embed_method = _compute_embed_similarity_matrix(
                    units=units,
                    granularity=g,
                    model_size=embed_model_used,
                    unit_texts_clean=doc_texts_clean,
                )
            except Exception as ee:
                embed_error = str(ee)
                sim_embed = None

        candidate_mode, candidate_topk = _candidate_policy_for_granularity(
            g,
            len(units),
            normal_topk=int(max(1, int(normal_candidate_topk))),
        )
        candidate_pairs = build_candidate_pairs(
            doc_names,
            sim_raw,
            sim_tfidf,
            sim_jaccard,
            sim_embed=sim_embed,
            candidate_mode=candidate_mode,
            topk=(candidate_topk if candidate_topk is not None else DEFAULT_TOPK_CANDIDATES),
        )
        full_pairs_total = int(len(doc_names) * (len(doc_names) - 1) // 2)
        reranker_enabled = bool(use_reranker)
        reranker_model_used = _normalize_reranker_size_label(reranker_model_size or DEFAULT_RERANK_MODEL_SIZE)
        reranker_error = ""
        rerank_meta = {
            "configured_percent": float(rerank_top_percent_non_h1),
            "effective_percent": float(rerank_top_percent_non_h1),
            "auto_scaled": False,
            "max_candidates_cap": int(max(1, int(RERANK_AUTO_NON_H1_MAX_CANDIDATES))),
            "total_rows": int(len(candidate_pairs)),
        }
        reranker_plan_candidates = int(
            _estimate_reranker_candidate_count(
                len(candidate_pairs),
                granularity=g,
                top_percent_non_h1=float(rerank_top_percent_non_h1),
            )
        )
        if reranker_enabled:
            try:
                rerank_effective_pct, rerank_meta = _effective_reranker_percent_non_h1(
                    total_rows=len(candidate_pairs),
                    granularity=g,
                    top_percent_non_h1=float(rerank_top_percent_non_h1),
                )
                reranker_plan_candidates = int(
                    _estimate_reranker_candidate_count(
                        len(candidate_pairs),
                        granularity=g,
                        top_percent_non_h1=rerank_effective_pct,
                    )
                )
                reranker_model_used = _annotate_reranker_for_pairs(
                    candidate_pairs,
                    units=units,
                    granularity=g,
                    model_size=reranker_model_used,
                    top_percent_non_h1=rerank_effective_pct,
                    unit_texts_clean=doc_texts_clean,
                ) or reranker_model_used
            except Exception as re_err:
                reranker_error = str(re_err)

        top_ngrams_info = self._build_top_ngrams_info(
            doc_names=doc_names,
            doc_texts=doc_texts,
            doc_vectors=doc_vectors,
            feature_names=feature_names,
            show_examples=show_examples,
        )

        return {
            "granularity": g,
            "units": units,
            "input_desc": input_desc,
            "doc_names": doc_names,
            "doc_texts": doc_texts,
            "doc_texts_for_similarity": doc_texts_clean,
            "ngram_boundary_policy": "segment_internal_no_cross",
            "doc_tokens_list": doc_tokens_list,
            "all_ngram_tokens": all_ngram_tokens,
            "sim_matrix": sim_raw,
            "sim_matrix_tfidf": sim_tfidf,
            "sim_matrix_jaccard": sim_jaccard,
            "sim_matrix_embed": sim_embed,
            "feature_names": feature_names,
            "doc_vectors": doc_vectors,
            "candidate_pairs": candidate_pairs,
            "candidate_mode": str(candidate_mode),
            "candidate_topk": (None if candidate_topk is None else int(candidate_topk)),
            "candidate_pairs_count": int(len(candidate_pairs)),
            "candidate_pairs_universe": int(full_pairs_total),
            "top_ngrams_info": top_ngrams_info,
            "semantic_enabled": semantic_enabled,
            "embed_method": embed_method,
            "embed_model": embed_model_used,
            "reranker_enabled": reranker_enabled,
            "reranker_model": reranker_model_used,
            "reranker_plan_candidates": int(reranker_plan_candidates),
            "rerank_top_percent_configured": float(rerank_meta.get("configured_percent", rerank_top_percent_non_h1)),
            "rerank_top_percent_effective": float(rerank_meta.get("effective_percent", rerank_top_percent_non_h1)),
            "rerank_auto_scaled": bool(rerank_meta.get("auto_scaled", False)),
            "rerank_cap_candidates": int(rerank_meta.get("max_candidates_cap", RERANK_AUTO_NON_H1_MAX_CANDIDATES)),
            "rerank_top_percent_non_h1": float(rerank_top_percent_non_h1),
            "embed_error": embed_error,
            "reranker_error": reranker_error,
        }

    def _plot_result_images(
        self,
        result: Dict[str, Any],
        out_dir: str,
        run_tag_prefix: str,
        title_base: str,
        strip_label_digits: bool,
        heatmap_font_setting: str,
    ) -> Dict[str, str]:
        g = str(result.get("granularity", "normal"))
        labels = list(result.get("doc_names") or [])
        sim_raw = result["sim_matrix"]
        sim_tfidf = result["sim_matrix_tfidf"]
        sim_jaccard = result["sim_matrix_jaccard"]
        sim_embed = result.get("sim_matrix_embed")
        candidate_pairs = list(result.get("candidate_pairs") or [])
        embed_method = str(result.get("embed_method") or "")

        heatmap_path = ""
        path_tfidf = ""
        path_jaccard = ""
        path_embed = ""
        path_top_pairs = os.path.join(out_dir, f"{run_tag_prefix}_{g}_top_pairs_cos_top10pct.png")

        # Matrix heatmaps are only generated for h1/h2 (skip normal for performance).
        if g in ("h1", "h2"):
            heatmap_path = os.path.join(out_dir, f"{run_tag_prefix}_{g}_cos_raw.png")
            path_tfidf = os.path.join(out_dir, f"{run_tag_prefix}_{g}_tfidf.png")
            path_jaccard = os.path.join(out_dir, f"{run_tag_prefix}_{g}_jaccard.png")
            path_embed = os.path.join(out_dir, f"{run_tag_prefix}_{g}_embed.png")
            plot_heatmap(
                sim_raw,
                labels,
                heatmap_path,
                title=f"{title_base} ({g} | cos raw)",
                strip_label_digits=strip_label_digits,
                font_setting=heatmap_font_setting,
            )
            plot_heatmap(
                sim_tfidf,
                labels,
                path_tfidf,
                title=f"{title_base} ({g} | tfidf)",
                strip_label_digits=strip_label_digits,
                font_setting=heatmap_font_setting,
            )
            plot_heatmap(
                sim_jaccard,
                labels,
                path_jaccard,
                title=f"{title_base} ({g} | jaccard)",
                strip_label_digits=strip_label_digits,
                font_setting=heatmap_font_setting,
            )
            if isinstance(sim_embed, np.ndarray) and sim_embed.ndim == 2 and sim_embed.shape[0] == len(labels):
                em_method = embed_method or ("maxsim" if g == "h1" else "avgpool")
                plot_heatmap(
                    sim_embed,
                    labels,
                    path_embed,
                    title=f"{title_base} ({g} | embed:{em_method})",
                    strip_label_digits=strip_label_digits,
                    font_setting=heatmap_font_setting,
                )
            else:
                path_embed = ""
        if candidate_pairs:
            plot_top_candidates_image(
                candidate_pairs,
                path_top_pairs,
                recall_ratio=_TOP_CANDIDATES_RECALL_RATIO,
                title=_TOP_CANDIDATES_TITLE,
                font_setting=heatmap_font_setting,
                max_rows_single_page=_TOP_CANDIDATES_MAX_ROWS_SINGLE_PAGE,
            )
        else:
            path_top_pairs = ""

        return {
            "heatmap_cos": heatmap_path,
            "heatmap_tfidf": path_tfidf,
            "heatmap_jaccard": path_jaccard,
            "heatmap_embed": path_embed,
            "top_pairs_img": path_top_pairs,
        }

    def _apply_primary_result(self, result: Dict[str, Any], plot_paths: Dict[str, str]):
        self.doc_texts = list(result.get("doc_texts") or [])
        self.doc_names = list(result.get("doc_names") or [])
        self.doc_texts_for_similarity = list(result.get("doc_texts_for_similarity") or [])
        self.doc_tokens_list = list(result.get("doc_tokens_list") or [])
        self.sim_matrix = result.get("sim_matrix")
        self.sim_matrix_tfidf = result.get("sim_matrix_tfidf")
        self.sim_matrix_jaccard = result.get("sim_matrix_jaccard")
        self.sim_matrix_embed = result.get("sim_matrix_embed")
        self.feature_names = result.get("feature_names")
        self.doc_vectors = result.get("doc_vectors")
        self.candidate_pairs = list(result.get("candidate_pairs") or [])
        self.top_ngrams_info = list(result.get("top_ngrams_info") or [])
        self.heatmap_path = str(plot_paths.get("heatmap_cos", "") or "")
        self.heatmap_path_tfidf = str(plot_paths.get("heatmap_tfidf", "") or "")
        self.heatmap_path_jaccard = str(plot_paths.get("heatmap_jaccard", "") or "")
        self.heatmap_path_embed = str(plot_paths.get("heatmap_embed", "") or "")
        self.top_pairs_img_path = str(plot_paths.get("top_pairs_img", "") or "")

    # ------------------- 分析 -------------------
    def run_analysis_thread(self):
        excel_path = self.excel_path_var.get().strip()
        if not excel_path:
            messagebox.showwarning("No file", "Please select one Excel file.")
            return
        if not os.path.isfile(excel_path):
            messagebox.showwarning("File not found", f"File does not exist:\n{excel_path}")
            return
        if os.path.splitext(excel_path)[1].lower() not in EXCEL_EXTS:
            messagebox.showerror("Error", "Only Excel input is supported (*.xlsx, *.xlsm, *.xls).")
            return

        try:
            n = max(1, int(self.n_var.get()))
            topk_edge = max(1, int(self.topk_edge_var.get()))
        except Exception:
            messagebox.showerror("Error", "n and Top-K edges must be integers.")
            return

        try:
            _ensure_ml_stack()
        except Exception as e_dep:
            messagebox.showerror("Dependency error", f"Cannot load sklearn/scipy dependencies:\n{e_dep}")
            return

        model_cfg = self._validate_model_config()
        if model_cfg is None:
            return

        out_dir = self.out_dir_var.get().strip()
        if not out_dir:
            try:
                out_dir = os.path.join(os.path.expanduser("~"), "Desktop")
            except Exception:
                out_dir = os.path.expanduser("~")

        config = {
            "excel_path": excel_path,
            "out_dir": out_dir,
            "n": n,
            "topk_edge": topk_edge,
            "granularity": normalize_granularity_choice(self.granularity_var.get()),
            "show_examples": bool(self.show_examples_var.get()),
            "strip_label_digits": bool(self.strip_label_digits_var.get()),
            "heatmap_font_setting": self.heatmap_font_var.get().strip(),
            "heatmap_title": self.heatmap_title_var.get().strip(),
            "use_semantic": bool(model_cfg["use_semantic"]),
            "embed_model_size": str(model_cfg["embed_model_size"]),
            "use_reranker": bool(model_cfg["use_reranker"]),
            "reranker_model_size": str(model_cfg["reranker_model_size"]),
            "normal_candidate_topk": int(model_cfg["normal_candidate_topk"]),
            "rerank_top_percent_non_h1": float(model_cfg["rerank_top_percent_non_h1"]),
            "use_clustering": bool(self.use_clustering_var.get()),
            "result_cache_reuse": bool(self.result_cache_reuse_var.get()),
            "result_cache_write": bool(self.result_cache_write_var.get()),
        }
        total_steps = self._estimate_total_steps(str(config.get("granularity", "normal")))
        self._run_started_at = time.perf_counter()
        self._run_started_wall = datetime.now()
        self._clear_run_log()
        self._reset_progress(total_steps)
        self._status_set("Run started...")
        self._append_log("Run started.")
        self._append_log(f"Input: {excel_path}")
        self._append_log(f"Output base dir: {out_dir}")
        self._set_run_enabled(False)
        t = threading.Thread(target=self.run_analysis, args=(config,), daemon=True)
        t.start()

    def run_analysis(self, config: Dict[str, object]):
        run_t0 = time.perf_counter()
        summary_lines: List[str] = []
        try:
            excel_path = str(config["excel_path"])
            out_dir = str(config["out_dir"])
            n = int(config["n"])
            topk_edge = int(config["topk_edge"])
            granularity_choice = normalize_granularity_choice(str(config.get("granularity", "normal")))
            show_examples = bool(config["show_examples"])
            strip_label_digits = bool(config["strip_label_digits"])
            heatmap_font_setting = str(config["heatmap_font_setting"])
            heatmap_title = str(config["heatmap_title"])
            use_semantic = bool(config.get("use_semantic", True))
            embed_model_size = _normalize_embed_size_label(str(config.get("embed_model_size", DEFAULT_EMBED_MODEL_SIZE)))
            use_reranker = bool(config.get("use_reranker", True))
            reranker_model_size = _normalize_reranker_size_label(str(config.get("reranker_model_size", DEFAULT_RERANK_MODEL_SIZE)))
            normal_candidate_topk = int(max(1, int(config.get("normal_candidate_topk", DEFAULT_TOPK_CANDIDATES))))
            rerank_top_percent_non_h1 = float(config.get("rerank_top_percent_non_h1", RERANK_TOP_PERCENT_NON_H1))
            use_clustering = bool(config.get("use_clustering", True))
            result_cache_reuse = bool(config.get("result_cache_reuse", True))
            result_cache_write = bool(config.get("result_cache_write", True))
            run_tag = self._make_run_tag(
                excel_path=excel_path,
                granularity=granularity_choice,
                use_semantic=use_semantic,
                embed_model_size=embed_model_size,
                use_reranker=use_reranker,
                reranker_model_size=reranker_model_size,
            )
            out_base_dir = str(out_dir)
            out_dir = self._resolve_output_run_dir(out_base_dir, run_tag)
            config["output_base_dir"] = out_base_dir
            config["out_dir"] = out_dir
            config["output_name_tag"] = run_tag
            os.makedirs(out_dir, exist_ok=True)
            total_steps = max(1, int(self._progress_total or self._estimate_total_steps(granularity_choice)))
            step_done = 0

            def _step(status_msg: str, log_msg: str = ""):
                nonlocal step_done
                step_done += 1
                self._ui_progress(step_done, total_steps, status_msg)
                if log_msg:
                    self._append_log(log_msg)

            if os.path.splitext(excel_path)[1].lower() not in EXCEL_EXTS:
                self._post_error("Error", "Only Excel input is supported.")
                return

            title_base = heatmap_title or '圖1.1《論語》《史記》《家語》《衣鏡》（長度歸一化）餘弦相似度矩陣'
            self.output_paths = []
            self.auto_output_paths = {}
            self._status_set("Preparing...")
            self._append_log(
                f"Config: granularity={granularity_choice}, n={n}, semantic={'on' if use_semantic else 'off'}"
                f"({_embed_model_label(embed_model_size)}), reranker={'on' if use_reranker else 'off'}"
                f"({_reranker_model_label(reranker_model_size)}), recall(h1/h2=all,normal=topk{normal_candidate_topk}),"
                f" rerank(h1/h2=all,normal=top{rerank_top_percent_non_h1:g}%)."
            )
            self._append_log(f"Output run dir: {out_dir}")
            self._append_log(
                f"Result cache policy: reuse={'on' if result_cache_reuse else 'off'}, "
                f"write={'on' if result_cache_write else 'off'}."
            )

            cache_key = ""
            cache_payload: Dict[str, Any] = {}
            if result_cache_reuse or result_cache_write:
                cache_key, cache_payload = self._build_result_cache_signature(
                    excel_path=excel_path,
                    granularity_choice=granularity_choice,
                    n=n,
                    show_examples=show_examples,
                    strip_label_digits=strip_label_digits,
                    heatmap_font_setting=heatmap_font_setting,
                    heatmap_title=heatmap_title,
                    use_semantic=use_semantic,
                    embed_model_size=embed_model_size,
                    use_reranker=use_reranker,
                    reranker_model_size=reranker_model_size,
                    normal_candidate_topk=normal_candidate_topk,
                    rerank_top_percent_non_h1=rerank_top_percent_non_h1,
                    use_clustering=use_clustering,
                )
                config["result_cache_key"] = cache_key
                self._append_log(f"Result cache key: {cache_key[:16]}...")

            expected_cache_mode = ("all" if granularity_choice == "all" else "single")
            if result_cache_reuse and cache_key:
                self._status_set("Checking result cache...")
                restored = self._restore_result_cache_bundle(
                    cache_key=cache_key,
                    expected_mode=expected_cache_mode,
                    out_dir=out_dir,
                )
                if restored is None:
                    # 兼容旧缓存键：按分析签名回溯查找，不受样式改动影响。
                    fallback_key = self._find_compatible_result_cache_key(
                        expected_mode=expected_cache_mode,
                        cache_payload=cache_payload,
                    )
                    if fallback_key and fallback_key != cache_key:
                        self._append_log(
                            f"Result cache exact-key miss; fallback compatible key: {fallback_key[:16]}..."
                        )
                        restored = self._restore_result_cache_bundle(
                            cache_key=fallback_key,
                            expected_mode=expected_cache_mode,
                            out_dir=out_dir,
                        )
                        if restored is not None:
                            cache_key = str(fallback_key)
                            config["result_cache_key"] = cache_key
                if restored is not None:
                    cache_dir = str(restored.get("cache_dir") or "")
                    self._append_log(f"Result cache hit: {cache_dir}")
                    summary_lines.append(f"cache_hit: {cache_key}")
                    # 不再 early return；恢復分析數據後 fall-through 到繪圖/聚類/報告流程
                    # 這樣可以在不重新跑 embedding/reranker 的情況下迭代圖片格式
                    self.auto_output_paths = {"result_cache_dir": cache_dir}
                    analysis_obj = restored.get("analysis") or {}
                    if not isinstance(analysis_obj, dict):
                        analysis_obj = {}
                    self.input_desc = str(
                        analysis_obj.get("input_desc")
                        or f"Cache hit: {cache_key[:12]}..., mode={expected_cache_mode}"
                    )
                    self._append_log(
                        "Cache hit: restoring analysis data, will regenerate ALL images..."
                    )

                    if expected_cache_mode == "all":
                        results_by_g = analysis_obj.get("results_by_g") or {}
                        skipped_by_g = analysis_obj.get("skipped_by_g") or {}
                        # plot_paths_by_g 將由下方繪圖代碼重新填充
                        _step("Cache restored (all)", "Cache hit (all mode): data restored, proceeding to re-plot.")
                    else:
                        result = analysis_obj.get("result") or {}
                        # plot_paths 將由下方繪圖代碼重新填充
                        _step("Cache restored (single)", "Cache hit (single mode): data restored, proceeding to re-plot.")

            # _cache_hit_data: True 時跳過 _analyze_one_granularity（數據已從 cache 恢復），
            # 但仍執行全部繪圖、聚類、Excel、Word 流程。
            _cache_hit_data = (result_cache_reuse and cache_key and restored is not None)

            if granularity_choice == "all":
                if not _cache_hit_data:
                    # 非 cache hit：正常跑 _analyze_one_granularity
                    results_by_g: Dict[str, Dict[str, Any]] = {}
                    skipped_by_g: Dict[str, str] = {}
                    for g in ANALYSIS_GRANULARITIES:
                        self._status_set(f"[{g}] analyzing...")
                        t_g = time.perf_counter()
                        try:
                            results_by_g[g] = self._analyze_one_granularity(
                                excel_path=excel_path,
                                granularity=g,
                                n=n,
                                show_examples=show_examples,
                                use_semantic=use_semantic,
                                embed_model_size=embed_model_size,
                                use_reranker=use_reranker,
                                reranker_model_size=reranker_model_size,
                                normal_candidate_topk=normal_candidate_topk,
                                rerank_top_percent_non_h1=rerank_top_percent_non_h1,
                            )
                            rows_n = len(results_by_g[g].get("candidate_pairs") or [])
                            units_n = len(results_by_g[g].get("doc_names") or [])
                            rr_plan_n = int(results_by_g[g].get("reranker_plan_candidates") or 0)
                            rr_eff_pct = float(results_by_g[g].get("rerank_top_percent_effective") or rerank_top_percent_non_h1)
                            rr_cfg_pct = float(results_by_g[g].get("rerank_top_percent_configured") or rerank_top_percent_non_h1)
                            rr_auto = bool(results_by_g[g].get("rerank_auto_scaled") or False)
                            if normalize_style(g) in ("h1", "h2"):
                                rr_suffix = f", reranker_plan={rr_plan_n}, rerank=all"
                            else:
                                rr_suffix = (
                                    f", reranker_plan={rr_plan_n}, rerank_pct={rr_eff_pct:g}%"
                                    + (f"(auto from {rr_cfg_pct:g}%)" if rr_auto else "")
                                )
                            summary_lines.append(f"{g}: units={units_n}, pairs={rows_n}{rr_suffix}")
                            _step(
                                f"[{g}] analyzed",
                                f"[{g}] analyze done: units={units_n}, pairs={rows_n}{rr_suffix}, time={_fmt_elapsed(time.perf_counter() - t_g)}.",
                            )
                        except Exception as ee:
                            skipped_by_g[g] = str(ee)
                            summary_lines.append(f"{g}: skipped ({ee})")
                            _step(
                                f"[{g}] skipped",
                                f"[{g}] skipped: {ee} (time={_fmt_elapsed(time.perf_counter() - t_g)}).",
                            )
                # else: _cache_hit_data == True → results_by_g / skipped_by_g 已從 cache 恢復

                if not results_by_g:
                    self._status_set("No valid units.")
                    self._append_log("No valid granularity results: h1/h2/normal all skipped.")
                    self._post_warning("Not enough units", "h1/h2/normal 均少于2个可计算文本单位。")
                    return

                plot_paths_by_g: Dict[str, Dict[str, str]] = {}
                output_lines: List[str] = []
                for g in ANALYSIS_GRANULARITIES:
                    res = results_by_g.get(g)
                    if not res:
                        _step(f"[{g}] plots skipped", f"[{g}] no result; plot stage skipped.")
                        continue
                    self._status_set(f"[{g}] plotting...")
                    t_p = time.perf_counter()
                    plot_paths = self._plot_result_images(
                        result=res,
                        out_dir=out_dir,
                        run_tag_prefix=run_tag,
                        title_base=title_base,
                        strip_label_digits=strip_label_digits,
                        heatmap_font_setting=heatmap_font_setting,
                    )
                    plot_paths_by_g[g] = plot_paths
                    for k in ("heatmap_cos", "heatmap_tfidf", "heatmap_jaccard", "heatmap_embed", "top_pairs_img"):
                        p = str(plot_paths.get(k, "") or "").strip()
                        if p:
                            output_lines.append(p)
                    _step(
                        f"[{g}] plots ready",
                        f"[{g}] plot bundle done in {_fmt_elapsed(time.perf_counter() - t_p)}.",
                    )

                # ---- 聚類與中心性分析（在 Excel 前計算，以便寫入分表） ----
                clustering_output_lines = []
                if use_clustering:
                    self._status_set("Running clustering & centrality analysis...")
                    t_cl = time.perf_counter()
                    for g in ANALYSIS_GRANULARITIES:
                        res = results_by_g.get(g)
                        if not res:
                            continue
                        sim_raw_g = res.get("sim_matrix")
                        doc_names_g = res.get("doc_names")
                        if sim_raw_g is None or doc_names_g is None:
                            continue
                        if len(doc_names_g) < 3:
                            continue
                        cl_prefix = f"{run_tag}_{g}"
                        path_dendro = os.path.join(out_dir, f"{cl_prefix}_Ward樹狀圖.png")
                        path_mds = os.path.join(out_dir, f"{cl_prefix}_MDS散點圖.png")
                        path_centr = os.path.join(out_dir, f"{cl_prefix}_Strength中心性.png")

                        cdf = compute_centrality_metrics(sim_raw_g, doc_names_g)
                        plot_cluster_dendrogram(sim_raw_g, doc_names_g, path_dendro, heatmap_font_setting)
                        plot_mds_scatter(sim_raw_g, doc_names_g, path_mds, heatmap_font_setting)
                        plot_centrality_bars(cdf, path_centr, heatmap_font_setting)
                        clustering_output_lines.extend([path_dendro, path_mds, path_centr])
                        res["centrality_df"] = cdf

                        # PCA biplot（需要 doc_vectors 和 feature_names）
                        doc_vectors_g = res.get("doc_vectors")
                        feature_names_g = res.get("feature_names")
                        if doc_vectors_g is not None and feature_names_g is not None and len(doc_names_g) >= 3:
                            path_pca = os.path.join(out_dir, f"{cl_prefix}_PCA_TF-IDF.png")
                            try:
                                plot_pca_biplot(doc_vectors_g, feature_names_g, doc_names_g,
                                                path_pca, heatmap_font_setting)
                                if os.path.isfile(path_pca):
                                    clustering_output_lines.append(path_pca)
                                else:
                                    self._append_log(f"[PCA:{g}] skipped: figure not generated.")
                            except Exception as e_pca:
                                self._append_log(f"[PCA:{g}] skipped: {e_pca}")

                        # 網絡關係圖（ctext 風格）
                        path_net = os.path.join(out_dir, f"{cl_prefix}_Network圖.png")
                        try:
                            plot_network_graph(sim_raw_g, doc_names_g, path_net,
                                               heatmap_font_setting, score_key="cos_raw")
                            if os.path.isfile(path_net):
                                clustering_output_lines.append(path_net)
                            else:
                                self._append_log(f"[Network:{g}] skipped: figure not generated.")
                        except Exception as e_net:
                            self._append_log(f"[Network:{g}] skipped: {e_net}")

                    # 生成聚类专属报告
                    if clustering_output_lines:
                        cl_res = results_by_g.get("normal") or results_by_g.get("h2") or results_by_g.get("h1")
                        if cl_res:
                            cl_sim = cl_res["sim_matrix"]
                            cl_names = cl_res["doc_names"]
                            cl_g = cl_res["granularity"]
                            cl_prefix_r = f"{run_tag}_{cl_g}"
                            cl_plot_paths = {
                                "dendrogram": os.path.join(out_dir, f"{cl_prefix_r}_Ward樹狀圖.png"),
                                "mds": os.path.join(out_dir, f"{cl_prefix_r}_MDS散點圖.png"),
                                "centrality": os.path.join(out_dir, f"{cl_prefix_r}_Strength中心性.png"),
                                "pca": os.path.join(out_dir, f"{cl_prefix_r}_PCA_TF-IDF.png"),
                                "network": os.path.join(out_dir, f"{cl_prefix_r}_Network圖.png"),
                            }
                            from docx import Document as DocxDoc
                            cl_doc = DocxDoc()
                            append_clustering_section(cl_doc, cl_sim, cl_names, cl_plot_paths,
                                                      cl_res.get("centrality_df"), heatmap_font_setting)
                            cl_report = os.path.join(out_dir, f"{run_tag}_聚類與中心性分析報告.docx")
                            cl_doc.save(cl_report)
                            clustering_output_lines.append(cl_report)
                    _step("Clustering done", f"Clustering & centrality done in {_fmt_elapsed(time.perf_counter() - t_cl)}.")

                out_excel = os.path.join(out_dir, f"{run_tag}.xlsx")
                self._status_set("Writing Excel...")
                t_x = time.perf_counter()
                self._write_excel_workbook_all(out_excel, results_by_g)
                _step("Excel written", f"Excel written: {out_excel} ({_fmt_elapsed(time.perf_counter() - t_x)}).")
                out_report = os.path.join(out_dir, f"{run_tag}.docx")
                self._status_set("Writing Word report...")
                t_w = time.perf_counter()
                generate_report_all(
                    out_report,
                    results_by_g=results_by_g,
                    plot_paths_by_g=plot_paths_by_g,
                    heading_title="N-gram Similarity Report (ALL: h1/h2/normal)",
                )
                _step("Word report written", f"Word report written: {out_report} ({_fmt_elapsed(time.perf_counter() - t_w)}).")

                out_preprocessed = os.path.join(out_dir, f"{run_tag}_preprocessed_monitor.docx")
                self._status_set("Writing preprocessed monitor Word...")
                t_pp = time.perf_counter()
                payload_by_g = _build_preprocessed_units_payload(excel_path=excel_path, results_by_g=results_by_g)
                generate_preprocessed_monitor_doc(
                    out_preprocessed,
                    payload_by_g=payload_by_g,
                    heading_title="预处理文本监测（h1/h2/normal）",
                )
                _step(
                    "Preprocessed monitor written",
                    f"Preprocessed monitor written: {out_preprocessed} ({_fmt_elapsed(time.perf_counter() - t_pp)}).",
                )

                pairwise_output_lines: List[str] = []
                self._status_set("Writing pairwise weights Word...")
                t_pw = time.perf_counter()
                # ALL 模式仅输出 1 份：优先 normal，其次 h2/h1。
                pw_res = results_by_g.get("normal") or results_by_g.get("h2") or results_by_g.get("h1") or {}
                doc_names_pw = list(pw_res.get("doc_names") or [])
                sim_pw = pw_res.get("sim_matrix")
                pairs_pw = list(pw_res.get("candidate_pairs") or [])
                pw_g = str(pw_res.get("granularity") or "normal")
                pw_doc = os.path.join(out_dir, f"{run_tag}_pairwise_similarity_weights.docx")
                if doc_names_pw and sim_pw is not None:
                    try:
                        generate_pairwise_weights_doc(
                            out_docx_path=pw_doc,
                            doc_names=doc_names_pw,
                            sim_matrix=sim_pw,
                            candidate_pairs=pairs_pw,
                            heading_title=f"Pairwise Similarity Weights ({pw_g}, cos raw tf)",
                        )
                        if os.path.isfile(pw_doc):
                            pairwise_output_lines.append(pw_doc)
                    except Exception as e_pw:
                        self._append_log(f"[Pairwise:{pw_g}] skipped: {e_pw}")
                self._append_log(
                    f"Pairwise weights docs written: {len(pairwise_output_lines)} ({_fmt_elapsed(time.perf_counter() - t_pw)})."
                )

                output_lines = [out_excel, out_report, out_preprocessed] + output_lines + clustering_output_lines + pairwise_output_lines
                self.output_paths = output_lines
                self.auto_output_paths = {
                    "excel": out_excel,
                    "report": out_report,
                    "preprocessed_monitor": out_preprocessed,
                    "pairwise_weights": (pairwise_output_lines[0] if pairwise_output_lines else ""),
                }

                primary = results_by_g.get("normal") or next(iter(results_by_g.values()))
                primary_plot = plot_paths_by_g.get(primary.get("granularity", "normal"), {})
                self._apply_primary_result(primary, primary_plot)
                err_bits: List[str] = []
                for g in ANALYSIS_GRANULARITIES:
                    rr = results_by_g.get(g) or {}
                    ee = str(rr.get("embed_error") or "").strip()
                    re_err = str(rr.get("reranker_error") or "").strip()
                    if ee:
                        err_bits.append(f"{g}.embed_fail")
                    if re_err:
                        err_bits.append(f"{g}.rerank_fail")
                self.input_desc = (
                    f"Excel mode: ALL(h1/h2/normal), preprocessing=Han-only(body only), "
                    f"semantic({'on' if use_semantic else 'off'}:{_embed_model_label(embed_model_size)};h1=maxsim,h2/normal=avgpool), "
                    f"reranker({'on' if use_reranker else 'off'}:{_reranker_model_label(reranker_model_size)};"
                    f"h1/h2=all,normal=top{rerank_top_percent_non_h1:g}% with auto-scale cap={int(RERANK_AUTO_NON_H1_MAX_CANDIDATES)}), "
                    f"available={','.join([g for g in ANALYSIS_GRANULARITIES if g in results_by_g])}"
                )
                mode_bits: List[str] = []
                rr_bits: List[str] = []
                for gg in ANALYSIS_GRANULARITIES:
                    rr = results_by_g.get(gg) or {}
                    if not rr:
                        continue
                    cm = str(rr.get("candidate_mode") or "all")
                    ck = rr.get("candidate_topk")
                    c_cnt = int(rr.get("candidate_pairs_count") or len(rr.get("candidate_pairs") or []))
                    c_uni = int(rr.get("candidate_pairs_universe") or 0)
                    if ck is None:
                        mode_bits.append(f"{gg}:{cm}({c_cnt}/{c_uni})")
                    else:
                        mode_bits.append(f"{gg}:{cm}{int(ck)}({c_cnt}/{c_uni})")
                    if normalize_style(gg) in ("h1", "h2"):
                        rr_plan = int(rr.get("reranker_plan_candidates") or 0)
                        rr_bits.append(f"{gg}:all[{rr_plan}]")
                    else:
                        rr_eff = float(rr.get("rerank_top_percent_effective") or rerank_top_percent_non_h1)
                        rr_cfg = float(rr.get("rerank_top_percent_configured") or rerank_top_percent_non_h1)
                        rr_auto = bool(rr.get("rerank_auto_scaled") or False)
                        rr_plan = int(rr.get("reranker_plan_candidates") or 0)
                        rr_bits.append(
                            f"{gg}:top{rr_eff:g}%"
                            + (f"(auto from {rr_cfg:g}%)" if rr_auto else "")
                            + f"[{rr_plan}]"
                        )
                if mode_bits:
                    self.input_desc += f", candidate_policy={' | '.join(mode_bits)}"
                if rr_bits:
                    self.input_desc += f", reranker_policy={' | '.join(rr_bits)}"
                if err_bits:
                    self.input_desc += f", warnings={';'.join(err_bits)}"
                    summary_lines.append(f"warnings={';'.join(err_bits)}")

                self._status_set("Writing run log report...")
                t_lr = time.perf_counter()
                log_report = self._write_run_log_report(
                    out_dir=out_dir,
                    run_tag=run_tag,
                    config={k: config[k] for k in config},
                    summary_lines=summary_lines,
                    output_paths=self.output_paths,
                )
                self.auto_output_paths["run_log"] = log_report
                self.output_paths.append(log_report)
                self.output_paths = [p for p in self.output_paths if p]
                _step(
                    "Run log report written",
                    f"Run log report written: {log_report} ({_fmt_elapsed(time.perf_counter() - t_lr)}).",
                )

                if result_cache_write and cache_key and not _cache_hit_data:
                    ok_cache, cache_msg, cache_files = self._save_result_cache_bundle(
                        cache_key=cache_key,
                        cache_payload=cache_payload,
                        mode="all",
                        analysis_payload={
                            "results_by_g": results_by_g,
                            "skipped_by_g": skipped_by_g,
                            "plot_paths_by_g": plot_paths_by_g,
                            "input_desc": self.input_desc,
                        },
                        output_paths=self.output_paths,
                        out_dir=out_dir,
                    )
                    if ok_cache:
                        self._append_log(f"Result cache saved: {cache_msg} (files={cache_files}).")
                    else:
                        self._append_log(f"Result cache save skipped: {cache_msg}")

                self.master.after(
                    0,
                    lambda rb=results_by_g, sb=skipped_by_g, k=topk_edge: self._render_analysis_output_all(rb, sb, k),
                )
                total_elapsed = _fmt_elapsed(time.perf_counter() - run_t0)
                _step("Done ✅", f"Done ✅ total_elapsed={total_elapsed}.")
                self._post_info("Done", "Outputs generated:\n" + "\n".join(self._collect_output_lines()))
                return

            # single granularity mode
            if not _cache_hit_data:
                # 非 cache hit：正常跑分析
                self._status_set(f"[{granularity_choice}] analyzing...")
                t_g = time.perf_counter()
                result = self._analyze_one_granularity(
                    excel_path=excel_path,
                    granularity=granularity_choice,
                    n=n,
                    show_examples=show_examples,
                    use_semantic=use_semantic,
                    embed_model_size=embed_model_size,
                    use_reranker=use_reranker,
                    reranker_model_size=reranker_model_size,
                    normal_candidate_topk=normal_candidate_topk,
                    rerank_top_percent_non_h1=rerank_top_percent_non_h1,
                )
                units_n = len(result.get("doc_names") or [])
                rows_n = len(result.get("candidate_pairs") or [])
                rr_plan_n = int(result.get("reranker_plan_candidates") or 0)
                rr_eff_pct = float(result.get("rerank_top_percent_effective") or rerank_top_percent_non_h1)
                rr_cfg_pct = float(result.get("rerank_top_percent_configured") or rerank_top_percent_non_h1)
                rr_auto = bool(result.get("rerank_auto_scaled") or False)
                if granularity_choice in ("h1", "h2"):
                    rr_suffix = f", reranker_plan={rr_plan_n}, rerank=all"
                else:
                    rr_suffix = (
                        f", reranker_plan={rr_plan_n}, rerank_pct={rr_eff_pct:g}%"
                        + (f"(auto from {rr_cfg_pct:g}%)" if rr_auto else "")
                    )
                summary_lines.append(f"{granularity_choice}: units={units_n}, pairs={rows_n}{rr_suffix}")
                _step(
                    f"[{granularity_choice}] analyzed",
                    f"[{granularity_choice}] analyze done: units={units_n}, pairs={rows_n}{rr_suffix}, time={_fmt_elapsed(time.perf_counter() - t_g)}.",
                )
            # else: _cache_hit_data == True → result 已從 cache 恢復
            self._status_set(f"[{granularity_choice}] plotting...")
            t_p = time.perf_counter()
            plot_paths = self._plot_result_images(
                result=result,
                out_dir=out_dir,
                run_tag_prefix=run_tag,
                title_base=title_base,
                strip_label_digits=strip_label_digits,
                heatmap_font_setting=heatmap_font_setting,
            )
            _step(
                f"[{granularity_choice}] plots ready",
                f"[{granularity_choice}] plot bundle done in {_fmt_elapsed(time.perf_counter() - t_p)}.",
            )
            clustering_output_lines: List[str] = []
            clustering_paths: Dict[str, str] = {}
            if use_clustering:
                self._status_set("Running clustering & centrality analysis...")
                t_cl = time.perf_counter()
                sim_raw_g = result.get("sim_matrix")
                doc_names_g = result.get("doc_names")
                if sim_raw_g is None or doc_names_g is None or len(doc_names_g) < 3:
                    self._append_log(f"[{granularity_choice}] clustering skipped: need >=3 docs.")
                else:
                    try:
                        cl_prefix = f"{run_tag}_{granularity_choice}"
                        path_dendro = os.path.join(out_dir, f"{cl_prefix}_Ward樹狀圖.png")
                        path_mds = os.path.join(out_dir, f"{cl_prefix}_MDS散點圖.png")
                        path_centr = os.path.join(out_dir, f"{cl_prefix}_Strength中心性.png")

                        cdf = compute_centrality_metrics(sim_raw_g, doc_names_g)
                        result["centrality_df"] = cdf
                        plot_cluster_dendrogram(sim_raw_g, doc_names_g, path_dendro, heatmap_font_setting)
                        plot_mds_scatter(sim_raw_g, doc_names_g, path_mds, heatmap_font_setting)
                        plot_centrality_bars(cdf, path_centr, heatmap_font_setting)
                        clustering_output_lines.extend([path_dendro, path_mds, path_centr])
                        clustering_paths.update({
                            "dendrogram": path_dendro,
                            "mds": path_mds,
                            "centrality": path_centr,
                        })

                        # PCA biplot（需要 doc_vectors 和 feature_names）
                        doc_vectors_g = result.get("doc_vectors")
                        feature_names_g = result.get("feature_names")
                        if doc_vectors_g is not None and feature_names_g is not None:
                            path_pca = os.path.join(out_dir, f"{cl_prefix}_PCA_TF-IDF.png")
                            try:
                                plot_pca_biplot(doc_vectors_g, feature_names_g, doc_names_g,
                                                path_pca, heatmap_font_setting)
                                if os.path.isfile(path_pca):
                                    clustering_output_lines.append(path_pca)
                                    clustering_paths["pca"] = path_pca
                                else:
                                    self._append_log(f"[PCA:{granularity_choice}] skipped: figure not generated.")
                            except Exception as e_pca:
                                self._append_log(f"[PCA:{granularity_choice}] skipped: {e_pca}")

                        # 網絡關係圖（ctext 風格）
                        path_net = os.path.join(out_dir, f"{cl_prefix}_Network圖.png")
                        try:
                            plot_network_graph(sim_raw_g, doc_names_g, path_net,
                                               heatmap_font_setting, score_key="cos_raw")
                            if os.path.isfile(path_net):
                                clustering_output_lines.append(path_net)
                                clustering_paths["network"] = path_net
                            else:
                                self._append_log(f"[Network:{granularity_choice}] skipped: figure not generated.")
                        except Exception as e_net:
                            self._append_log(f"[Network:{granularity_choice}] skipped: {e_net}")

                        # 聚類專屬報告
                        try:
                            from docx import Document as DocxDoc
                            cl_doc = DocxDoc()
                            append_clustering_section(
                                cl_doc,
                                sim_raw_g,
                                doc_names_g,
                                clustering_paths,
                                cdf,
                                heatmap_font_setting,
                            )
                            cl_report = os.path.join(out_dir, f"{run_tag}_聚類與中心性分析報告.docx")
                            cl_doc.save(cl_report)
                            clustering_output_lines.append(cl_report)
                            clustering_paths["cluster_report"] = cl_report
                        except Exception as e_clr:
                            self._append_log(f"[ClusterReport:{granularity_choice}] skipped: {e_clr}")
                    except Exception as e_cl:
                        self._append_log(f"[Clustering:{granularity_choice}] failed: {e_cl}")
                _step(
                    "Clustering done",
                    f"[{granularity_choice}] clustering stage done in {_fmt_elapsed(time.perf_counter() - t_cl)}.",
                )
            self._apply_primary_result(result, plot_paths)
            if not _cache_hit_data:
                rr_policy = "all" if granularity_choice in ("h1", "h2") else (
                    f"top{rr_eff_pct:g}%"
                    + (f"(auto from {rr_cfg_pct:g}%)" if rr_auto else "")
                )
                self.input_desc = (
                    f"Excel mode: {result.get('input_desc')}, granularity={granularity_choice}, "
                    f"units={len(result.get('doc_names') or [])}, preprocessing=Han-only(body only), "
                    f"candidates={int(result.get('candidate_pairs_count') or len(result.get('candidate_pairs') or []))}/"
                    f"{int(result.get('candidate_pairs_universe') or 0)}"
                    f"({result.get('candidate_mode')}{'' if result.get('candidate_topk') is None else f':{int(result.get('candidate_topk'))}'}), "
                    f"semantic({'on' if use_semantic else 'off'}:{_embed_model_label(result.get('embed_model') or embed_model_size)};"
                    f"{'maxsim' if granularity_choice == 'h1' else 'avgpool'}), "
                    f"reranker({'on' if use_reranker else 'off'}:{_reranker_model_label(result.get('reranker_model') or reranker_model_size)};"
                    f"{rr_policy})"
                )
                if str(result.get("embed_error") or "").strip():
                    self.input_desc += ", embed_fail"
                    summary_lines.append("embed_fail")
                if str(result.get("reranker_error") or "").strip():
                    self.input_desc += ", rerank_fail"
                    summary_lines.append("rerank_fail")
            # else: self.input_desc 已在 cache hit 恢復時設置

            out_excel = os.path.join(out_dir, f"{run_tag}.xlsx")
            self._status_set("Writing Excel...")
            t_x = time.perf_counter()
            self._write_excel_workbook_single(out_excel, result)
            _step("Excel written", f"Excel written: {out_excel} ({_fmt_elapsed(time.perf_counter() - t_x)}).")
            out_report = os.path.join(out_dir, f"{run_tag}.docx")
            self._status_set("Writing Word report...")
            t_w = time.perf_counter()
            generate_report(
                out_report,
                result["doc_names"],
                result["feature_names"],
                result["doc_vectors"],
                result["sim_matrix"],
                result["top_ngrams_info"],
                plot_paths.get("heatmap_cos", ""),
                result["doc_texts"],
                result["doc_tokens_list"],
                candidate_pairs=result["candidate_pairs"],
                include_collation=(granularity_choice == "normal"),
                heading_title=f"N-gram Similarity Report ({granularity_choice})",
                heatmap_tfidf_path=str(plot_paths.get("heatmap_tfidf", "") or ""),
                heatmap_jaccard_path=str(plot_paths.get("heatmap_jaccard", "") or ""),
                heatmap_embed_path=str(plot_paths.get("heatmap_embed", "") or ""),
            )
            _step("Word report written", f"Word report written: {out_report} ({_fmt_elapsed(time.perf_counter() - t_w)}).")
            out_preprocessed = os.path.join(out_dir, f"{run_tag}_preprocessed_monitor.docx")
            self._status_set("Writing preprocessed monitor Word...")
            t_pp = time.perf_counter()
            payload_by_g = _build_preprocessed_units_payload(
                excel_path=excel_path,
                results_by_g={granularity_choice: result},
            )
            generate_preprocessed_monitor_doc(
                out_preprocessed,
                payload_by_g=payload_by_g,
                heading_title="预处理文本监测（h1/h2/normal）",
            )
            _step(
                "Preprocessed monitor written",
                f"Preprocessed monitor written: {out_preprocessed} ({_fmt_elapsed(time.perf_counter() - t_pp)}).",
            )

            pairwise_doc = os.path.join(out_dir, f"{run_tag}_pairwise_similarity_weights.docx")
            self._status_set("Writing pairwise weights Word...")
            t_pw = time.perf_counter()
            try:
                generate_pairwise_weights_doc(
                    out_docx_path=pairwise_doc,
                    doc_names=list(result.get("doc_names") or []),
                    sim_matrix=result.get("sim_matrix"),
                    candidate_pairs=list(result.get("candidate_pairs") or []),
                    heading_title=f"Pairwise Similarity Weights ({granularity_choice}, cos raw tf)",
                )
                if not os.path.isfile(pairwise_doc):
                    pairwise_doc = ""
            except Exception as e_pw:
                pairwise_doc = ""
                self._append_log(f"[Pairwise:{granularity_choice}] skipped: {e_pw}")
            self._append_log(
                f"Pairwise weights doc {'written' if pairwise_doc else 'skipped'} ({_fmt_elapsed(time.perf_counter() - t_pw)})."
            )

            self.auto_output_paths = {
                "excel": out_excel,
                "report": out_report,
                "preprocessed_monitor": out_preprocessed,
                "pairwise_weights": pairwise_doc,
                "heatmap_cos": str(plot_paths.get("heatmap_cos", "") or ""),
                "heatmap_tfidf": str(plot_paths.get("heatmap_tfidf", "") or ""),
                "heatmap_jaccard": str(plot_paths.get("heatmap_jaccard", "") or ""),
                "heatmap_embed": str(plot_paths.get("heatmap_embed", "") or ""),
                "top_pairs_img": str(plot_paths.get("top_pairs_img", "") or ""),
                "cluster_dendrogram": str(clustering_paths.get("dendrogram", "") or ""),
                "cluster_mds": str(clustering_paths.get("mds", "") or ""),
                "cluster_centrality": str(clustering_paths.get("centrality", "") or ""),
                "cluster_pca": str(clustering_paths.get("pca", "") or ""),
                "cluster_network": str(clustering_paths.get("network", "") or ""),
                "cluster_report": str(clustering_paths.get("cluster_report", "") or ""),
            }
            plot_output_lines = [
                str(plot_paths.get(k, "") or "")
                for k in ("heatmap_cos", "heatmap_tfidf", "heatmap_jaccard", "heatmap_embed", "top_pairs_img")
            ]
            self.output_paths = [
                p for p in ([out_excel, out_report, out_preprocessed, pairwise_doc] + plot_output_lines + clustering_output_lines)
                if str(p).strip()
            ]
            self._status_set("Writing run log report...")
            t_lr = time.perf_counter()
            log_report = self._write_run_log_report(
                out_dir=out_dir,
                run_tag=run_tag,
                config={k: config[k] for k in config},
                summary_lines=summary_lines,
                output_paths=self.output_paths,
            )
            self.auto_output_paths["run_log"] = log_report
            self.output_paths.append(log_report)
            self.output_paths = [p for p in self.output_paths if p]
            _step(
                "Run log report written",
                f"Run log report written: {log_report} ({_fmt_elapsed(time.perf_counter() - t_lr)}).",
            )

            if result_cache_write and cache_key and not _cache_hit_data:
                ok_cache, cache_msg, cache_files = self._save_result_cache_bundle(
                    cache_key=cache_key,
                    cache_payload=cache_payload,
                    mode="single",
                    analysis_payload={
                        "result": result,
                        "plot_paths": plot_paths,
                        "input_desc": self.input_desc,
                    },
                    output_paths=self.output_paths,
                    out_dir=out_dir,
                )
                if ok_cache:
                    self._append_log(f"Result cache saved: {cache_msg} (files={cache_files}).")
                else:
                    self._append_log(f"Result cache save skipped: {cache_msg}")
            self.master.after(0, lambda k=topk_edge: self._render_analysis_output(k))
            total_elapsed = _fmt_elapsed(time.perf_counter() - run_t0)
            _step("Done ✅", f"Done ✅ total_elapsed={total_elapsed}.")
            self._post_info(
                "Done",
                "Outputs generated:\n" + "\n".join(self._collect_output_lines()),
            )
        except Exception as e:
            self._status_set("Error.")
            self._append_log(f"ERROR: {e}")
            self._post_error("Error", f"Analysis failed: {e}")
        finally:
            self._set_run_enabled(True)
            self._run_started_at = None
            self._run_started_wall = None

# ------------------- 運行 -------------------
if len(sys.argv) > 1 and sys.argv[1] == "--cluster":
    _csv = sys.argv[2] if len(sys.argv) > 2 else ""
    _out = sys.argv[3] if len(sys.argv) > 3 else ""
    if not _csv or not os.path.isfile(_csv):
        print("用法: python 'samll pipe 4.py' --cluster <csv_path> [out_dir]")
        sys.exit(1)
    if not _out:
        _out = os.path.join(os.path.expanduser("~"), "Desktop")
    run_clustering_standalone(_csv, _out)
    sys.exit(0)

if __name__ == '__main__':
    root = tk.Tk()
    app = NgramApp(root)
    root.mainloop()
