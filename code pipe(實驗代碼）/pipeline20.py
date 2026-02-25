import os
import re
import json
import time
import csv
import hashlib
import argparse
import sys
import pathlib
import threading
import signal
import warnings
import math
import gc
from dataclasses import dataclass, field
from typing import List, Optional, Dict, Tuple, Any

# Keep matplotlib cache in writable location and force non-GUI backend.
if "MPLCONFIGDIR" not in os.environ:
    _mpl_cache_dir = os.path.join("/tmp", "matplotlib-codex-cache")
    try:
        os.makedirs(_mpl_cache_dir, exist_ok=True)
        os.environ["MPLCONFIGDIR"] = _mpl_cache_dir
    except Exception:
        pass

# Runtime stability / noise control for local desktop runs.
os.environ.setdefault("TOKENIZERS_PARALLELISM", "false")
os.environ.setdefault("TRANSFORMERS_VERBOSITY", "error")
os.environ.setdefault("JOBLIB_MULTIPROCESSING", "0")
warnings.filterwarnings(
    "ignore",
    message=r"resource_tracker: There appear to be \d+ leaked semaphore objects to clean up at shutdown:.*",
    category=UserWarning,
)

import tkinter as tk
from tkinter import ttk, filedialog, messagebox
from tkinter.scrolledtext import ScrolledText

import numpy as np
from sklearn.feature_extraction.text import CountVectorizer, TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from sklearn.preprocessing import normalize as sk_normalize

import openpyxl
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
def _import_matplotlib_safe():
    try:
        import matplotlib
        try:
            matplotlib.use("Agg", force=True)
        except Exception:
            pass

        if sys.platform == "darwin":
            # On macOS, font discovery shells out to `system_profiler`, which can hang.
            import subprocess
            import plistlib

            _orig_check_output = subprocess.check_output
            _empty_fonts_plist = plistlib.dumps([])

            def _check_output_with_timeout(*args, **kwargs):
                cmd = args[0] if args else kwargs.get("args")
                if isinstance(cmd, (list, tuple)) and cmd and cmd[0] == "system_profiler":
                    kwargs.setdefault("timeout", 8)
                    try:
                        return _orig_check_output(*args, **kwargs)
                    except Exception:
                        return _empty_fonts_plist
                return _orig_check_output(*args, **kwargs)

            subprocess.check_output = _check_output_with_timeout
            try:
                import matplotlib.pyplot as _plt
                from matplotlib.backends.backend_pdf import PdfPages as _pdf_pages
            finally:
                subprocess.check_output = _orig_check_output
        else:
            import matplotlib.pyplot as _plt
            from matplotlib.backends.backend_pdf import PdfPages as _pdf_pages

        return _plt, _pdf_pages
    except Exception:
        return None, None


plt, PdfPages = _import_matplotlib_safe()

_PLOT_ZH_READY = False

try:
    from datasketch import MinHash
except Exception:
    MinHash = None


# =========================
# Data structures
# =========================
@dataclass
class Para:
    pid: int
    text: str
    style: str  # h1/h2/h3/normal


@dataclass
class Unit:
    uid: int
    level: str              # h1/h2/h3/normal (chosen granularity)
    h1: str = ""
    h2: str = ""
    h3: str = ""
    text: str = ""
    para_ids: List[int] = field(default_factory=list)


# =========================
# Utilities
# =========================
def desktop_dir() -> str:
    home = pathlib.Path.home()
    cand = home / "Desktop"
    return str(cand if cand.exists() else home)


def normalize_style(s: str) -> str:
    s = (s or "").strip().lower()
    if s in ("h1", "title", "heading1", "heading 1"):
        return "h1"
    if s in ("h2", "heading2", "heading 2"):
        return "h2"
    if s in ("h3", "heading3", "heading 3"):
        return "h3"
    return "normal"


def _semantic_maxsim_enabled(granularity: str) -> bool:
    g = normalize_style(granularity)
    if g in SEMANTIC_MAXSIM_ENABLE_BY_G:
        return bool(SEMANTIC_MAXSIM_ENABLE_BY_G[g])
    return bool(SEMANTIC_MAXSIM_ENABLE_BY_G.get("default", False))


def _merge_tail_slope(granularity: str) -> float:
    g = normalize_style(granularity)
    v = MERGE_ROBUST_TAIL_SLOPE.get(g, MERGE_ROBUST_TAIL_SLOPE.get("default", 1.0))
    try:
        s = float(v)
    except Exception:
        s = 1.0
    return max(0.0, s)


def _merge_softclip_k(granularity: str) -> float:
    g = normalize_style(granularity)
    v = MERGE_ROBUST_SOFTCLIP_K.get(g, MERGE_ROBUST_SOFTCLIP_K.get("default", 0.10))
    try:
        s = float(v)
    except Exception:
        s = 0.10
    return max(0.001, s)

# ---- Text preprocessing for similarity: keep ONLY Chinese Han characters ----
# Includes CJK Unified Ideographs + Extensions commonly used in classical texts.
_HAN_RE = re.compile(
    r"[\u3400-\u4DBF\u4E00-\u9FFF\U00020000-\U0002A6DF\U0002A700-\U0002B73F\U0002B740-\U0002B81F\U0002B820-\U0002CEAF\U0002CEB0-\U0002EBEF]"
)

def preprocess_keep_han(text: str) -> str:
    """For similarity only: remove punctuation/symbols and keep ONLY Han characters."""
    if not text:
        return ""
    return "".join(_HAN_RE.findall(str(text)))


def char_ngrams(text: str, n: int) -> List[str]:
    if text is None:
        return []
    t = text
    if len(t) < n:
        return []
    return [t[i:i+n] for i in range(len(t) - n + 1)]


def jaccard(a: set, b: set) -> float:
    if not a and not b:
        return 1.0
    if not a or not b:
        return 0.0
    inter = len(a & b)
    union = len(a | b)
    return inter / union if union else 0.0


def containment_short_in_long(a: set, b: set) -> Tuple[float, int]:
    """
    containment: |S(short) ∩ S(long)| / |S(short)|
      flag=0: A 更短或等长 => short=A, long=B
      flag=1: B 更短        => short=B, long=A
    """
    if not a and not b:
        return 1.0, 0
    if not a or not b:
        return 0.0, 0
    if len(a) <= len(b):
        short, long = a, b
        flag = 0
    else:
        short, long = b, a
        flag = 1
    denom = len(short)
    if denom == 0:
        return 0.0, flag
    inter = len(short & long)
    return inter / denom, flag


def minhash_jaccard(mh_a, mh_b) -> Optional[float]:
    if mh_a is None or mh_b is None:
        return None
    return float(mh_a.jaccard(mh_b))


# =========================
# Semantic embedding (local models)
# =========================
_HF_HUB_DIR = str(pathlib.Path.home() / ".cache" / "huggingface" / "hub")

QWEN_EMBED_MODEL_DIRS: Dict[str, str] = {
    "0.6B": os.path.join(_HF_HUB_DIR, "models--Qwen--Qwen3-Embedding-0.6B"),
    "4B": os.path.join(_HF_HUB_DIR, "models--Qwen--Qwen3-Embedding-4B"),
    "8B": os.path.join(_HF_HUB_DIR, "models--Qwen--Qwen3-Embedding-8B"),
    "Qwen2-7B": os.path.join(_HF_HUB_DIR, "gte-qwen2-7b"),
}

_SEMANTIC_BACKEND_CACHE: Dict[str, dict] = {}
_RERANK_BACKEND_CACHE: Dict[str, dict] = {}
_SCRIPT_SHA1_CACHE: Optional[str] = None

QWEN_RERANK_MODEL_DIRS: Dict[str, str] = {
    "0.6B": os.path.join(_HF_HUB_DIR, "models--Qwen--Qwen3-Reranker-0.6B"),
    "4B": os.path.join(_HF_HUB_DIR, "models--Qwen--Qwen3-Reranker-4B"),
    "8B": os.path.join(_HF_HUB_DIR, "models--Qwen--Qwen3-Reranker-8B"),
}

# ── BGE 对照实验模型 ──────────────────────────────────────────────
BGE_EMBED_MODEL_DIR = os.path.join(_HF_HUB_DIR, "bge-m3")           # git clone 目录
BGE_RERANK_MODEL_DIR = os.path.join(_HF_HUB_DIR, "models--BAAI--bge-reranker-base")


def _is_bge_model(model_size: str) -> bool:
    """判断 model_size 是否属于 BGE 族 (embedding 或 reranker)."""
    return str(model_size or "").strip().upper().startswith("BGE")


MERGE_BASELINE = {
    "topk_total": 20,          # lex=14, sem=6 at 0.70:0.30
    "recall_lex_ratio": 0.70,
    "recall_sem_ratio": 0.30,
    "weight_lex": 0.70,
    "weight_sem": 0.30,
}

# Reranker defaults (sidecar only): focus on high-value pairs and cap runtime.
RERANK_DEFAULT_TOPN_CAP = 1200
RERANK_DEFAULT_TOP_PERCENT_NON_H1 = 3.0

# Coverage defaults are centralized for easier tuning/ablation.
COVERAGE_DEFAULTS = {
    "seg_topk": 10,
    "min_cos": 0.20,
    "min_jaccard": 0.02,
    "hits_max": 6,
}

# Semantic MaxSim defaults (H1/H2 only; Normal keeps single-vector cosine).
SEMANTIC_MAXSIM_CHUNK_SIZE = 500
SEMANTIC_MAXSIM_CHUNK_STRIDE = 400
SEMANTIC_MAXSIM_H1_LAMBDA = 0.00
SEMANTIC_MAXSIM_H2_LAMBDA = 0.15
SEMANTIC_MAXSIM_H2_TOP_PERCENT = 3.0
SEMANTIC_MAXSIM_H2_TOPN_MIN = 500
SEMANTIC_MAXSIM_H2_TOPN_MAX = 1500
# Layer policy switch:
# final current rollout = H1 use MaxSim, H2/Normal keep avgpool cosine.
SEMANTIC_MAXSIM_ENABLE_BY_G = {
    "h1": True,
    "h2": False,
    "normal": False,
    "default": False,
}

# For normal-level direct TF score, avoid over-penalizing extreme length ratios.
NORMAL_LEN_BALANCE_FLOOR = 0.30

# Shared open-interval normalization epsilon for value-based score mapping.
SCORE_NORM_EPS = 1e-6

# Robust normalization quantiles used at merge stage.
# H1 is small-sample; use wider guard bands to reduce endpoint jitter.
MERGE_ROBUST_QUANTILES = {
    "h1": (0.10, 0.90),
    "h2": (0.05, 0.95),
    "normal": (0.05, 0.95),
    "default": (0.05, 0.95),
}

# Robust-map safety fallback:
# when sample is too small or quantile span collapses, return a neutral midpoint.
MERGE_ROBUST_MIN_SAMPLES = 10

# Merge-stage normalization policy:
# - robust_hard: quantile min-max + hard clipping
# - robust_soft: quantile min-max + piecewise tanh soft-clip (bounded, monotonic)
# - robust_linear: quantile min-max + linear tail extrapolation (no hard clipping)
# - minmax: global min-max
MERGE_NORM_METHOD = "robust_soft"
MERGE_ROBUST_TAIL_SLOPE = {
    "h1": 1.0,
    "h2": 1.0,
    "normal": 1.0,
    "default": 1.0,
}
MERGE_ROBUST_SOFTCLIP_K = {
    "h1": 0.20,
    "h2": 0.10,
    "normal": 0.10,
    "default": 0.10,
}

# Display label: this is an in-layer percentile reference, not a probability.
RERANK_REF_01_HEADER = "reranker_ref_01（层内百分位参考分，非概率）"

LEXICAL_WEIGHTS_DEFAULT_H = {
    "cosine_tf_raw": 0.35,
    "coverage": 0.35,
    "cosine_tfidf": 0.20,
    "jaccard": 0.10,
}
LEXICAL_WEIGHTS_DEFAULT_N = {
    "cosine_tf_raw": 0.60,
    "coverage": 0.00,
    "cosine_tfidf": 0.20,
    "jaccard": 0.20,
}

# Pair-length-aware Jaccard adjustment:
# short texts keep near-raw Jaccard; long texts get a mild lift.
JACCARD_LEN_ADJUST = {
    "short_len_lo": 200,      # <= this: keep raw
    "short_len_hi": 1200,     # >= this: full long-text adjustment
    "gamma_long": 0.80,       # long-text map: j_eff = jaccard ** gamma_long
    "weight_boost_long": 0.20 # long-text: up to +20% on jaccard weight
}

# Shadow audit defaults: build a unified candidate universe per granularity.
# 0 means using all rows from that channel.
SHADOW_AUDIT_TOPK_DEFAULT = {
    "h1": 0,
    "h2": 50,
    "normal": 50,
}

# Normal collation report (display-only, does NOT affect scoring/ranking).
NORMAL_COLLATION_DEFAULTS = {
    "topn_merge_normal": 1000,
    "detail_norm_edit_max": 0.35,
    "skip_han_exact": True,
    "max_variant_items_per_pair": 0,  # 0 = no limit (show all variant lines)
}

FRAMEWORK_ORDER: List[Tuple[str, str]] = [
    ("lexical", "字面"),
    ("semantic", "语义"),
    ("merge", "融合"),
]
FRAMEWORK_PLOT_ORDER: List[Tuple[str, str, str]] = [
    ("lexical", "字面", "#D1495B"),
    ("semantic", "语义", "#EDA63B"),
    ("merge", "融合", "#8A5A44"),
]
FRAMEWORK_PREFIX_MAP: Dict[str, str] = dict(FRAMEWORK_ORDER)
FRAMEWORK_LABEL_MAP: Dict[str, str] = {
    "lexical": "lex",
    "semantic": "sem",
    "merge": "merge",
}
VALID_FRAMEWORKS = set(list(FRAMEWORK_PREFIX_MAP.keys()) + ["all"])

RERANK_PROMPT_DEFAULT = (
    "你是上古汉语专家，面对早期孔门文本，判断整体文本相似度（综合字面与语义 两个方面）。"
    "字面相似度包括是否含有互见的相同相似的文本，和整体的文体文类风格两个方面。"
    "整体分数为（0-1）。如果两个文本完全一样，一字不差，默认相似度为1。"
)

# Local code tag for cache metadata (manual bump when major pipeline logic changes).
PIPELINE_CODE_VERSION = "pipeline14_robust_soft_tanh_v1"

_INVALID_FILENAME_CHARS_RE = re.compile(r'[\\/:*?"<>|]+')
_INVALID_SHEET_CHARS_RE = re.compile(r'[:\\/?*\[\]]+')


def _safe_name_part(s: str, fallback: str = "NA") -> str:
    t = _INVALID_FILENAME_CHARS_RE.sub(" ", str(s or "").strip())
    t = re.sub(r"\s+", " ", t).strip()
    return t if t else fallback


def _excel_safe_sheet_name(name: str, fallback: str = "Sheet") -> str:
    t = _INVALID_SHEET_CHARS_RE.sub(" ", str(name or "").strip())
    t = re.sub(r"\s+", " ", t).strip().strip("'")
    if not t:
        t = fallback
    return t[:31]


def _framework_cn(framework: str) -> str:
    fw = str(framework or "").strip().lower()
    if fw == "lexical":
        return "字面"
    if fw == "semantic":
        return "语义"
    if fw == "merge":
        return "融合"
    if fw == "all":
        return "字面-语义-融合"
    return "融合"


def _granularity_name(granularity: str) -> str:
    g = str(granularity or "").strip().lower()
    if g in ("h1", "h2", "h3", "normal"):
        return g
    if g in ("all", "h1-h2-normal"):
        return "h1-h2-normal"
    return normalize_style(g)


def _normalize_embed_size_label(model_size: str) -> str:
    s = str(model_size or "").strip().upper().replace("Ｂ", "B")
    if s.startswith("BGE"):
        return "BGE-M3"
    alias = {
        "0.6": "0.6B",
        "0.6B": "0.6B",
        "4": "4B",
        "4B": "4B",
        "8": "8B",
        "8B": "8B",
        "0.8B": "8B",  # backward compatibility
        "QWEN2-7B": "Qwen2-7B",
        "QWEN2 7B": "Qwen2-7B",
        "GTE-QWEN2-7B": "Qwen2-7B",
    }
    return alias.get(s, s if s else "0.6B")


def _embed_model_label(model_size: str) -> str:
    if _is_bge_model(model_size):
        return "BGE-M3"
    m = _normalize_embed_size_label(model_size or "0.6B")
    if m == "Qwen2-7B":
        return "gte-qwen2-7b"
    return f"Qwen3-Embedding-{m}"


def _reranker_model_label(model_size: str) -> str:
    if _is_bge_model(model_size):
        return "BGE-reranker-base"
    m = _normalize_reranker_size_label(model_size or "0.6B")
    return f"Qwen3-Reranker-{m}"


def _model_label_for_framework(
    framework: str,
    semantic_model: str,
    reranker_model: str,
    reranker_enabled: bool,
) -> str:
    fw = str(framework or "").strip().lower()
    if fw == "lexical":
        return "Lexical"
    if fw == "semantic":
        return _embed_model_label(semantic_model)
    if fw == "merge":
        if reranker_enabled:
            return f"{_embed_model_label(semantic_model)} + {_reranker_model_label(reranker_model)}"
        return _embed_model_label(semantic_model)
    if reranker_enabled:
        return f"{_embed_model_label(semantic_model)} + {_reranker_model_label(reranker_model)}"
    return _embed_model_label(semantic_model)


def _make_output_name_tag(
    framework: str,
    granularity: str,
    upload_base: str,
    semantic_model: str,
    reranker_model: str,
    reranker_enabled: bool,
    time_prefix: str = "",
) -> str:
    fw_name = _safe_name_part(_framework_cn(framework), fallback="融合")
    g_name = _safe_name_part(_granularity_name(granularity), fallback="h1-h2-normal")
    b_name = _safe_name_part(upload_base, fallback="input")
    m_name = _safe_name_part(
        _model_label_for_framework(framework, semantic_model, reranker_model, reranker_enabled),
        fallback="Model",
    )
    p = _safe_name_part(time_prefix, fallback="")
    stem = f"{fw_name}_{g_name}_{b_name}_{m_name}"
    return f"{p}_{stem}" if p else stem


def _make_dual_all_output_name_tag(
    upload_base: str,
    semantic_model: str,
    reranker_model: str,
    reranker_enabled: bool,
    time_prefix: str = "",
) -> str:
    """Naming for framework=all + granularity=all: all_all_<upload>_<model>."""
    b_name = _safe_name_part(upload_base, fallback="input")
    m_name = _safe_name_part(
        _model_label_for_framework("all", semantic_model, reranker_model, reranker_enabled),
        fallback="Model",
    )
    p = _safe_name_part(time_prefix, fallback="")
    stem = f"all_all_{b_name}_{m_name}"
    return f"{p}_{stem}" if p else stem


def _score_header_with_model(metric_name: str, model_label: Optional[str]) -> str:
    """Return plain metric_name as Excel header (no bracket suffix).

    Previously appended '[Qwen3-Embedding-8B]' etc., causing external scripts
    to fail when looking up simple keys like 'semantic_cosine'.
    Model info is already stored in dedicated 'semantic_model' / 'reranker_model'
    columns, so the bracket suffix is redundant and harmful for interoperability.
    """
    return str(metric_name or "").strip()


def _fmt_score_with_model(v: Any, model_label: Optional[str]) -> Any:
    if v is None:
        return None
    try:
        return float(v)
    except Exception:
        return v


def strip_header_model_suffix(header: str) -> str:
    """Strip model bracket suffix from Excel header for uniform lookup.

    e.g. 'semantic_cosine[Qwen3-Embedding-8B]' → 'semantic_cosine'
         'score_reranker_raw[Qwen3-Reranker-8B]' → 'score_reranker_raw'
         'score_final' → 'score_final' (unchanged)
    """
    h = str(header or "").strip()
    idx = h.find("[")
    if idx > 0 and h.endswith("]"):
        return h[:idx]
    return h


def normalize_sheet_headers(headers: List[str]) -> Dict[str, int]:
    """Build a lookup {simple_name: col_index} from Excel headers.

    Both the original header AND its bracket-stripped form map to the same index.
    This allows code to use r.get('semantic_cosine') regardless of whether
    the Excel header is 'semantic_cosine' or 'semantic_cosine[Qwen3-Embedding-8B]'.
    """
    mapping: Dict[str, int] = {}
    for i, h in enumerate(headers):
        h_str = str(h or "").strip()
        if h_str:
            mapping[h_str] = i
            simple = strip_header_model_suffix(h_str)
            if simple != h_str and simple not in mapping:
                mapping[simple] = i
    return mapping


def _normalize_reranker_size_label(model_size: str) -> str:
    s = str(model_size or "").strip().upper().replace("Ｂ", "B")
    if s.startswith("BGE"):
        return "BGE-reranker"
    # gte-qwen2-7b 没有配套 reranker，显式拦截
    if s in ("QWEN2-7B", "QWEN2 7B", "GTE-QWEN2-7B"):
        raise ValueError(
            "gte-qwen2-7b 没有配套 reranker 模型。"
            "请选择 Qwen3 系列 (0.6B/4B/8B) 或 BGE-reranker 作为 reranker。"
        )
    alias = {
        "0.6": "0.6B",
        "0.6B": "0.6B",
        "4": "4B",
        "4B": "4B",
        "8": "8B",
        "8B": "8B",
        "0.8B": "8B",  # backward compatibility
    }
    return alias.get(s, s)


def _fmt_elapsed(seconds: float) -> str:
    s = max(0.0, float(seconds))
    if s < 60.0:
        return f"{s:.1f}s"
    total = int(round(s))
    m, sec = divmod(total, 60)
    h, m = divmod(m, 60)
    if h > 0:
        return f"{h}h{m:02d}m{sec:02d}s"
    return f"{m}m{sec:02d}s"


def _estimate_reranker_candidate_count(
    total_rows: int,
    granularity: str,
    topn: Optional[int],
    top_percent_non_h1: float,
) -> int:
    n_rows = int(max(0, int(total_rows)))
    if n_rows <= 0:
        return 0
    g = normalize_style(granularity)
    if g == "h1":
        n_sel = n_rows
    else:
        pct = float(top_percent_non_h1)
        if pct <= 0:
            pct = float(RERANK_DEFAULT_TOP_PERCENT_NON_H1)
        n_sel = int(max(1, round(n_rows * (pct / 100.0))))
    if topn is not None:
        n_sel = int(max(1, min(int(topn), n_sel, n_rows)))
    else:
        n_sel = int(max(1, min(n_sel, n_rows)))
    return int(n_sel)


def _quiet_transformers_logging() -> None:
    try:
        from transformers.utils import logging as hf_logging
        hf_logging.set_verbosity_error()
    except Exception:
        pass


def _ensure_padding_token(tokenizer, model=None) -> bool:
    """
    Ensure tokenizer/model have a usable padding token id for batched inference.
    Returns True when pad token id is available after best-effort setup.
    """
    pad_id = getattr(tokenizer, "pad_token_id", None)
    if pad_id is None:
        eos_tok = getattr(tokenizer, "eos_token", None)
        if eos_tok is not None:
            try:
                tokenizer.pad_token = eos_tok
            except Exception:
                pass
        pad_id = getattr(tokenizer, "pad_token_id", None)

    if pad_id is None:
        eos_id = getattr(tokenizer, "eos_token_id", None)
        if eos_id is not None:
            try:
                tokenizer.pad_token_id = int(eos_id)
            except Exception:
                pass
        pad_id = getattr(tokenizer, "pad_token_id", None)

    if pad_id is not None and model is not None:
        try:
            if getattr(model.config, "pad_token_id", None) is None:
                model.config.pad_token_id = int(pad_id)
        except Exception:
            pass
    return pad_id is not None


def _is_torch_oom_error(err: Exception) -> bool:
    msg = str(err or "").lower()
    pats = (
        "out of memory",
        "cuda out of memory",
        "mps backend out of memory",
        "cudnn_status_alloc_failed",
    )
    return any(p in msg for p in pats)


def _release_torch_memory(torch_mod, device: str) -> None:
    try:
        gc.collect()
    except Exception:
        pass
    dev = str(device or "").strip().lower()
    try:
        if dev == "cuda":
            if getattr(torch_mod, "cuda", None) is not None and torch_mod.cuda.is_available():
                torch_mod.cuda.empty_cache()
        elif dev == "mps":
            mps_mod = getattr(torch_mod, "mps", None)
            if mps_mod is not None and hasattr(mps_mod, "empty_cache"):
                mps_mod.empty_cache()
    except Exception:
        pass


def _resolve_hf_snapshot_dir(model_cache_dir: str) -> str:
    """Resolve HF cache dir to a concrete snapshot path (offline/local only)."""
    p = pathlib.Path(model_cache_dir)
    if not p.exists() or not p.is_dir():
        raise RuntimeError(f"Model dir not found: {model_cache_dir}")

    snaps = p / "snapshots"
    if not snaps.exists() or not snaps.is_dir():
        return str(p)

    # Prefer refs/main when present.
    ref_main = p / "refs" / "main"
    if ref_main.exists():
        try:
            commit = ref_main.read_text(encoding="utf-8").strip()
            cand = snaps / commit
            if cand.exists() and cand.is_dir():
                return str(cand)
        except Exception:
            pass

    # Fallback: latest snapshot by mtime.
    children = [x for x in snaps.iterdir() if x.is_dir()]
    if not children:
        raise RuntimeError(f"No snapshots found under: {snaps}")
    children.sort(key=lambda x: x.stat().st_mtime, reverse=True)
    return str(children[0])


def _read_rope_theta_from_config_json(model_dir: str, default: float = 1000000.0) -> float:
    """Best-effort load rope_theta from model_dir/config.json."""
    cfg_path = os.path.join(str(model_dir or "").strip(), "config.json")
    try:
        with open(cfg_path, "r", encoding="utf-8") as f:
            obj = json.load(f)
        v = obj.get("rope_theta") if isinstance(obj, dict) else None
        if v is None:
            return float(default)
        v = float(v)
        if not np.isfinite(v) or v <= 0:
            return float(default)
        return v
    except Exception:
        return float(default)


def _load_bge_m3_backend() -> dict:
    """Lazy-load BGE-M3 embedding backend (BAAI/bge-m3, XLM-RoBERTa, hidden=1024)."""
    model_dir = BGE_EMBED_MODEL_DIR
    cache_key = f"BGE-M3:{model_dir}"
    if cache_key in _SEMANTIC_BACKEND_CACHE:
        return _SEMANTIC_BACKEND_CACHE[cache_key]

    _quiet_transformers_logging()
    try:
        import torch
        from transformers import AutoTokenizer, AutoModel
    except Exception as e:
        raise RuntimeError(f"BGE-M3 backend requires torch+transformers: {e}")

    if torch.cuda.is_available():
        device = "cuda"
    elif getattr(torch.backends, "mps", None) is not None and torch.backends.mps.is_available():
        device = "mps"
    else:
        device = "cpu"

    tokenizer = AutoTokenizer.from_pretrained(model_dir, local_files_only=True)
    _ensure_padding_token(tokenizer)

    load_kwargs = {"local_files_only": True}
    if device == "cuda":
        load_kwargs["torch_dtype"] = torch.float16

    try:
        model = AutoModel.from_pretrained(model_dir, **load_kwargs)
    except Exception:
        model = AutoModel.from_pretrained(model_dir, local_files_only=True)
    _ensure_padding_token(tokenizer, model=model)

    model.to(device)
    model.eval()

    backend = {
        "torch": torch,
        "tokenizer": tokenizer,
        "model": model,
        "device": device,
        "model_dir": model_dir,
        "model_size": "BGE-M3",
        "pooling": "cls",  # BGE-M3 uses CLS token for dense retrieval
    }
    _SEMANTIC_BACKEND_CACHE[cache_key] = backend
    print(f"[BGE-M3] Loaded embedding model from {model_dir} → {device}")
    return backend


def _load_semantic_backend(model_size: str) -> dict:
    """Lazy-load local embedding backend; cached by resolved model path."""
    model_size = _normalize_embed_size_label(model_size or "0.6B")
    if _is_bge_model(model_size):
        return _load_bge_m3_backend()
    if model_size not in QWEN_EMBED_MODEL_DIRS:
        raise RuntimeError(f"Unsupported embedding model size: {model_size}")

    model_dir = _resolve_hf_snapshot_dir(QWEN_EMBED_MODEL_DIRS[model_size])
    cache_key = f"{model_size}:{model_dir}"
    if cache_key in _SEMANTIC_BACKEND_CACHE:
        return _SEMANTIC_BACKEND_CACHE[cache_key]

    _quiet_transformers_logging()
    try:
        import torch
        from transformers import AutoTokenizer, AutoModel, AutoConfig
    except Exception as e:
        raise RuntimeError(f"Semantic backend requires torch+transformers: {e}")

    if torch.cuda.is_available():
        device = "cuda"
    elif getattr(torch.backends, "mps", None) is not None and torch.backends.mps.is_available():
        device = "mps"
    else:
        device = "cpu"

    tokenizer = AutoTokenizer.from_pretrained(model_dir, trust_remote_code=True, local_files_only=True)
    _ensure_padding_token(tokenizer)

    load_kwargs = {
        "trust_remote_code": True,
        "local_files_only": True,
    }
    model_cfg = None
    try:
        model_cfg = AutoConfig.from_pretrained(
            model_dir,
            trust_remote_code=True,
            local_files_only=True,
        )
    except Exception:
        model_cfg = None
    # Compatibility patch for some transformers builds whose Qwen2Config lacks rope_theta.
    if model_size == "Qwen2-7B" and model_cfg is not None and not hasattr(model_cfg, "rope_theta"):
        try:
            model_cfg.rope_theta = _read_rope_theta_from_config_json(model_dir, default=1000000.0)
            print(f"[gte-qwen2-7b] patched missing config.rope_theta={float(model_cfg.rope_theta):.1f}")
        except Exception:
            pass
    if model_cfg is not None:
        load_kwargs["config"] = model_cfg
    # Use half-precision on GPU devices to save memory.
    # Qwen3 is trained in bfloat16; float32 only adds useless zero-padding.
    if device == "cuda":
        load_kwargs["torch_dtype"] = torch.float16
    elif device == "mps":
        load_kwargs["torch_dtype"] = torch.bfloat16

    try:
        model = AutoModel.from_pretrained(model_dir, **load_kwargs)
    except Exception:
        # Fallback without dtype hint.
        fb_kwargs = {
            "trust_remote_code": True,
            "local_files_only": True,
        }
        if model_cfg is not None:
            fb_kwargs["config"] = model_cfg
        model = AutoModel.from_pretrained(model_dir, **fb_kwargs)
    _ensure_padding_token(tokenizer, model=model)

    # gte-qwen2-7b ships custom modeling_qwen.py that expects old DynamicCache APIs
    # (from_legacy_cache/get_usable_length/to_legacy_cache). On transformers>=5.x,
    # disable cache path to keep embedding forward compatible.
    if model_size == "Qwen2-7B":
        try:
            model.config.use_cache = False
        except Exception:
            pass

    model.to(device)
    model.eval()

    # ── pooling 策略：gte-qwen2 用 last_token，Qwen3 用 avgpool ──
    _pooling = "last_token" if model_size == "Qwen2-7B" else "avgpool"

    backend = {
        "torch": torch,
        "tokenizer": tokenizer,
        "model": model,
        "device": device,
        "model_dir": model_dir,
        "model_size": model_size,
        "pooling": _pooling,
    }
    _SEMANTIC_BACKEND_CACHE[cache_key] = backend
    print(f"[{_embed_model_label(model_size)}] pooling={_pooling}, device={device}")
    return backend


def _embed_text_for_semantic(text: str) -> str:
    """Use original text (minus whitespace) for semantic model."""
    t = str(text or "")
    t = re.sub(r"\s+", "", t)
    if t:
        return t
    return preprocess_keep_han(text or "")


def _semantic_emb_cache_path(
    out_dir: str,
    granularity: str,
    model_size: str,
    texts: List[str],
    max_length: int,
) -> str:
    """Build cache path for semantic embeddings with corpus fingerprint."""
    g = normalize_style(granularity)
    m = _normalize_embed_size_label(model_size or "8B")
    ml = int(max(64, int(max_length)))
    h = hashlib.sha1()
    h.update(f"g={g}|m={m}|ml={ml}|n={len(texts)}".encode("utf-8"))
    for t in texts:
        b = str(t or "").encode("utf-8", errors="ignore")
        h.update(b"|")
        h.update(str(len(b)).encode("utf-8"))
        h.update(b":")
        h.update(b)
    digest = h.hexdigest()[:20]
    cache_dir = os.path.join(str(out_dir or "").strip(), "_cache_semantic_embeddings")
    return os.path.join(cache_dir, f"sem_emb_{g}_{m}_ml{ml}_{digest}.npy")


def _semantic_emb_meta_path(cache_path: str) -> str:
    return f"{cache_path}.meta.json"


def _script_sha1() -> str:
    global _SCRIPT_SHA1_CACHE
    if _SCRIPT_SHA1_CACHE:
        return str(_SCRIPT_SHA1_CACHE)
    try:
        with open(__file__, "rb") as f:
            data = f.read()
        _SCRIPT_SHA1_CACHE = hashlib.sha1(data).hexdigest()
    except Exception:
        _SCRIPT_SHA1_CACHE = ""
    return str(_SCRIPT_SHA1_CACHE)


def _build_semantic_emb_cache_meta(
    cache_path: str,
    input_file: str,
    granularity: str,
    model_size: str,
    max_length: int,
    text_count: int,
    embs: np.ndarray,
) -> Dict[str, Any]:
    arr = np.asarray(embs)
    rows = int(arr.shape[0]) if arr.ndim >= 1 else 0
    dim = int(arr.shape[1]) if arr.ndim >= 2 else 0
    src = os.path.abspath(str(input_file or "")) if input_file else ""
    return {
        "cache_schema": "semantic_embedding_cache_v1",
        "generated_at": time.strftime("%Y-%m-%d %H:%M:%S"),
        "generated_ts": float(time.time()),
        "code_version": str(PIPELINE_CODE_VERSION),
        "script_file": os.path.abspath(__file__),
        "script_name": os.path.basename(__file__),
        "script_sha1": _script_sha1(),
        "input_file": src,
        "input_name": (os.path.basename(src) if src else ""),
        "granularity": normalize_style(granularity),
        "model_size": _normalize_embed_size_label(model_size or "8B"),
        "max_length": int(max(64, int(max_length))),
        "text_count": int(max(0, int(text_count))),
        "embedding_rows": rows,
        "embedding_dim": dim,
        "cache_file": os.path.abspath(cache_path),
    }


def _load_semantic_emb_meta(cache_path: str) -> Dict[str, Any]:
    path = _semantic_emb_meta_path(cache_path)
    if not os.path.isfile(path):
        return {}
    try:
        with open(path, "r", encoding="utf-8") as f:
            obj = json.load(f)
        return obj if isinstance(obj, dict) else {}
    except Exception:
        return {}


def _load_semantic_emb_cache(path: str, expected_n: int) -> Optional[np.ndarray]:
    if not path or (not os.path.isfile(path)):
        return None
    try:
        arr = np.load(path, allow_pickle=False)
        arr = np.asarray(arr, dtype=np.float32)
    except Exception:
        return None
    if arr.ndim != 2:
        return None
    if int(arr.shape[0]) != int(max(0, expected_n)):
        return None
    if int(arr.shape[1]) <= 0:
        return None
    try:
        if not np.isfinite(arr).all():
            return None
    except Exception:
        return None
    return arr


def _save_semantic_emb_cache(path: str, embs: np.ndarray) -> bool:
    try:
        os.makedirs(os.path.dirname(path), exist_ok=True)
        np.save(path, np.asarray(embs, dtype=np.float32), allow_pickle=False)
        return True
    except Exception:
        return False


def _save_semantic_emb_meta(cache_path: str, meta: Dict[str, Any]) -> bool:
    try:
        os.makedirs(os.path.dirname(cache_path), exist_ok=True)
        with open(_semantic_emb_meta_path(cache_path), "w", encoding="utf-8") as f:
            json.dump(meta or {}, f, ensure_ascii=False, indent=2)
        return True
    except Exception:
        return False


def _encode_semantic_embeddings(
    texts: List[str],
    model_size: str,
    batch_size: int = 8,
    max_length: int = 8192,
    stride_pct: int = 75,
) -> np.ndarray:
    """Return l2-normalized embeddings [N, D].

    For texts that fit within *max_length* tokens: direct avgpool encoding.
    For texts that exceed *max_length* tokens: sliding-window chunking with
    token-count-weighted mean aggregation, then L2 normalization.  This
    guarantees **zero truncation** — every token in the original text
    contributes to the final embedding.

    Sliding-window parameters (only for overlong texts):
        window  = max_length
        stride  = max_length * stride_pct // 100

    Runtime safety:
        If torch OOM occurs (common on MPS with long context), the encoder
        auto-retries with smaller batch size and then smaller max_length.
    """
    backend = _load_semantic_backend(model_size)
    torch = backend["torch"]
    tokenizer = backend["tokenizer"]
    model = backend["model"]
    device = backend["device"]

    if not texts:
        return np.zeros((0, 1), dtype=np.float32)

    ml_req = int(max(64, max_length))
    _spct = int(max(10, min(99, stride_pct)))  # clamp to [10, 99]

    # ── helper: encode a single batch of token dicts → pooled, UN-normalized vectors ──
    _pooling_mode = backend.get("pooling", "avgpool")  # cls | last_token | avgpool

    def _forward_pool(tok_dict):
        """Run model forward + pool.  Returns [B, D] float32 tensor (NOT L2-normed).
        Pooling strategy: cls (BGE-M3), last_token (gte-qwen2), avgpool (Qwen3)."""
        tok_dict = {k: v.to(device) for k, v in tok_dict.items()}
        out = model(**tok_dict)
        try:
            last_hidden = None
            if isinstance(out, dict):
                last_hidden = out.get("last_hidden_state")
            else:
                last_hidden = getattr(out, "last_hidden_state", None)
                if last_hidden is None and isinstance(out, (tuple, list)) and len(out) > 0:
                    last_hidden = out[0]
            if last_hidden is None:
                raise RuntimeError("Embedding model output missing hidden states.")
            # ── BGE-M3: CLS token ──
            if _pooling_mode == "cls":
                return last_hidden[:, 0].detach().float()
            # ── gte-qwen2: last non-padding token ──
            if _pooling_mode == "last_token":
                attn = tok_dict.get("attention_mask")
                if attn is not None:
                    # sequence_lengths[i] = last index where mask==1
                    sequence_lengths = attn.sum(dim=1) - 1  # [B]
                    sequence_lengths = sequence_lengths.clamp(min=0)
                    batch_idx = torch.arange(last_hidden.size(0), device=last_hidden.device)
                    return last_hidden[batch_idx, sequence_lengths].detach().float()
                # fallback: rightmost token (no padding)
                return last_hidden[:, -1].detach().float()
            # ── Qwen3: attention-mask-aware mean pooling ──
            pooler = None
            if isinstance(out, dict):
                pooler = out.get("pooler_output")
            else:
                pooler = getattr(out, "pooler_output", None)
            if pooler is not None:
                return pooler.detach().float()
            attn = tok_dict.get("attention_mask")
            if attn is None:
                return last_hidden[:, 0].detach().float()
            mask = attn.unsqueeze(-1).expand_as(last_hidden).float()
            emb = (last_hidden * mask).sum(dim=1) / torch.clamp(mask.sum(dim=1), min=1e-9)
            return emb.detach().float()
        finally:
            # ── 及时释放大张量，防止内存累积 ──
            del tok_dict, out
            try:
                del last_hidden, mask, pooler, attn
            except (NameError, UnboundLocalError):
                pass

    # ── helper: sliding-window encode for a single overlong text ──
    def _encode_overlong(text: str, ml_cur: int, stride_cur: int) -> np.ndarray:
        """Sliding-window + token-weighted mean for one text.  Returns [1, D] L2-normed."""
        # tokenize without truncation to get full token ids
        full_ids = tokenizer.encode(text, add_special_tokens=False)
        n_full = len(full_ids)
        # build windows
        windows = []
        for start in range(0, n_full, stride_cur):
            end = min(start + ml_cur, n_full)
            win_ids = full_ids[start:end]
            if len(win_ids) == 0:
                break
            windows.append(win_ids)
            if end >= n_full:
                break
        # encode each window (batch_size=1 for overlong — rare path)
        weighted_sum = None
        total_tokens = 0
        for win_ids in windows:
            n_tok = len(win_ids)
            # re-encode via tokenizer to get proper special tokens & attention mask
            win_text = tokenizer.decode(win_ids, skip_special_tokens=True)
            toks = tokenizer(
                [win_text],
                padding=False,
                truncation=True,
                max_length=ml_cur,
                return_tensors="pt",
            )
            with torch.no_grad():
                vec = _forward_pool(toks)  # [1, D], un-normalized
            vec_np = vec.cpu().numpy().astype(np.float32)[0]  # [D]
            if weighted_sum is None:
                weighted_sum = vec_np * n_tok
            else:
                weighted_sum += vec_np * n_tok
            total_tokens += n_tok
        # weighted mean → L2 normalize
        mean_vec = weighted_sum / max(total_tokens, 1)
        norm = np.linalg.norm(mean_vec)
        if norm > 1e-12:
            mean_vec /= norm
        return mean_vec.reshape(1, -1)

    # Pre-calc token counts once to avoid repeated tokenizer scans on retries.
    token_counts: List[int] = []
    for t in texts:
        try:
            token_counts.append(int(len(tokenizer.encode(t, add_special_tokens=False))))
        except Exception:
            token_counts.append(0)

    bs_req = int(max(1, batch_size))
    if getattr(tokenizer, "pad_token_id", None) is None or getattr(getattr(model, "config", None), "pad_token_id", None) is None:
        bs_req = 1

    def _encode_once(ml_cur: int, bs_cur: int) -> np.ndarray:
        stride_cur = max(1, int(ml_cur * _spct // 100))
        vecs: List[np.ndarray] = []
        overlong_indices = {idx for idx, n_tok in enumerate(token_counts) if int(n_tok) > int(ml_cur)}

        i = 0
        while i < len(texts):
            if i in overlong_indices:
                vec = _encode_overlong(texts[i], ml_cur=ml_cur, stride_cur=stride_cur)
                vecs.append(vec)
                i += 1
            else:
                j = i
                while j < len(texts) and j not in overlong_indices and (j - i) < bs_cur:
                    j += 1
                batch = texts[i:j]
                toks = tokenizer(
                    batch,
                    padding=(len(batch) > 1),
                    truncation=True,
                    max_length=ml_cur,
                    return_tensors="pt",
                )
                with torch.no_grad():
                    emb = _forward_pool(toks)
                    emb = torch.nn.functional.normalize(emb, p=2, dim=1)
                    # bfloat16 tensors cannot always convert to NumPy directly.
                    vecs.append(emb.cpu().numpy().astype(np.float32, copy=False))
                i = j

            # MPS unified memory can stay fragmented across steps; trim cache eagerly.
            if str(device).lower() == "mps":
                _release_torch_memory(torch, device)

        return np.vstack(vecs)

    def _uniq_ints(vals: List[int]) -> List[int]:
        out: List[int] = []
        seen = set()
        for v in vals:
            try:
                iv = int(v)
            except Exception:
                continue
            if iv < 1:
                continue
            if iv in seen:
                continue
            seen.add(iv)
            out.append(iv)
        return out

    # OOM retry: only reduce batch_size, never reduce max_length.
    # Reducing max_length truncates H2 texts (P95=2476 tokens) and damages
    # embedding quality.  Reducing batch has zero quality impact — identical
    # vectors, just slower.  With bfloat16 model (~8 GB) + batch=1 peak
    # activation ~5.3 GB, total ~13.3 GB — fits even with other apps open.
    bs_candidates = _uniq_ints([bs_req, 4, 2, 1])

    last_oom: Optional[Exception] = None
    for bs_try in bs_candidates:
        if bs_try > bs_req:
            continue
        try:
            arr = _encode_once(ml_cur=ml_req, bs_cur=bs_try)
            if bs_try != bs_req:
                print(
                    f"[semantic] WARN: OOM at batch_size={bs_req}, "
                    f"succeeded with batch_size={bs_try} on {device}. "
                    f"max_length={ml_req} preserved (no truncation)."
                )
            return arr
        except Exception as e:
            if _is_torch_oom_error(e):
                last_oom = e
                _release_torch_memory(torch, device)
                continue
            raise

    if last_oom is not None:
        tried_bs = ",".join(str(x) for x in bs_candidates if x <= bs_req)
        raise RuntimeError(
            f"Semantic embedding OOM after all batch retries on device={device}. "
            f"max_length={ml_req}, tried batch_size=[{tried_bs}]. "
            f"Last error: {last_oom}"
        )
    raise RuntimeError("Semantic embedding failed: no valid retry configuration.")


def build_semantic_channel(
    units: List[Unit],
    model_size: str,
    topk_neighbors: int,
    embed_max_length: int = 8192,
    embed_stride_pct: int = 75,
) -> Tuple[Optional[np.ndarray], Dict[int, List[int]]]:
    """Stage-2 semantic channel: embeddings + topk neighbor map for candidate expansion."""
    if not units or len(units) < 2:
        return None, {}

    texts = [_embed_text_for_semantic(u.text or "") for u in units]
    embs = _encode_semantic_embeddings(texts, model_size=model_size, max_length=embed_max_length, stride_pct=embed_stride_pct)

    neighbors = build_semantic_neighbors(embs, topk_neighbors=topk_neighbors)
    return embs, neighbors


def build_semantic_neighbors(embs: np.ndarray, topk_neighbors: int) -> Dict[int, List[int]]:
    """Build top-k neighbor index from normalized semantic embeddings."""
    if embs is None or not isinstance(embs, np.ndarray):
        return {}
    n = int(embs.shape[0]) if embs.ndim >= 2 else 0
    if n < 2:
        return {}

    k = int(max(1, min(int(topk_neighbors), n - 1)))
    neighbors: Dict[int, List[int]] = {}
    for i in range(n):
        sims = np.dot(embs, embs[i])
        sims[i] = -1.0
        if k >= n - 1:
            idxs = np.argsort(-sims)
        else:
            idxs = np.argpartition(-sims, k)[:k]
            idxs = idxs[np.argsort(-sims[idxs])]
        neighbors[i] = [int(j) for j in idxs if int(j) != i]
    return neighbors


def _split_semantic_chunks(text: str, chunk_size: int = SEMANTIC_MAXSIM_CHUNK_SIZE, stride: int = SEMANTIC_MAXSIM_CHUNK_STRIDE) -> List[str]:
    """Split semantic text into overlapped chunks."""
    t = _embed_text_for_semantic(text or "")
    if not t:
        return []
    c = int(max(80, chunk_size))
    s = int(max(1, stride))
    if s > c:
        s = c
    if len(t) <= c:
        return [t]
    out: List[str] = []
    i = 0
    n = len(t)
    while i < n:
        seg = t[i: i + c]
        if seg:
            out.append(seg)
        if i + c >= n:
            break
        i += s
    return out


def _l2_normalize_vec(v: np.ndarray) -> np.ndarray:
    arr = np.asarray(v, dtype=np.float32)
    den = float(np.linalg.norm(arr))
    if den <= 1e-12:
        return arr.astype(np.float32, copy=False)
    return (arr / den).astype(np.float32, copy=False)


def build_semantic_maxsim_context(
    units: List[Unit],
    model_size: str,
    chunk_size: int = SEMANTIC_MAXSIM_CHUNK_SIZE,
    chunk_stride: int = SEMANTIC_MAXSIM_CHUNK_STRIDE,
    embed_max_length: int = 8192,
    embed_stride_pct: int = 75,
) -> dict:
    """Build mean-vec recall embeddings + per-unit chunk embeddings for MaxSim refine."""
    n = len(units or [])
    if n <= 0:
        return {
            "mean_embeddings": np.zeros((0, 1), dtype=np.float32),
            "chunk_embeddings": {},
            "chunk_counts": [],
            "chunk_total": 0,
            "chunk_size": int(chunk_size),
            "chunk_stride": int(chunk_stride),
        }

    all_chunks: List[str] = []
    unit_spans: List[Tuple[int, int]] = []
    for u in units:
        chunks = _split_semantic_chunks(u.text or "", chunk_size=chunk_size, stride=chunk_stride)
        if not chunks:
            fallback = _embed_text_for_semantic(u.text or "")
            chunks = [fallback] if fallback else ["空"]
        st = len(all_chunks)
        all_chunks.extend(chunks)
        unit_spans.append((st, len(chunks)))

    chunk_emb = _encode_semantic_embeddings(all_chunks, model_size=model_size, max_length=embed_max_length, stride_pct=embed_stride_pct)
    if chunk_emb.ndim != 2 or chunk_emb.shape[0] <= 0:
        raise RuntimeError("Failed to build semantic chunk embeddings.")

    d = int(chunk_emb.shape[1])
    mean_emb = np.zeros((n, d), dtype=np.float32)
    chunk_map: Dict[int, np.ndarray] = {}
    chunk_counts: List[int] = [0] * n
    for uid, (st, cnt) in enumerate(unit_spans):
        arr = chunk_emb[st: st + cnt]
        if arr.ndim != 2 or arr.shape[0] <= 0:
            arr = np.zeros((1, d), dtype=np.float32)
        arr = arr.astype(np.float32, copy=False)
        chunk_map[int(uid)] = arr
        chunk_counts[uid] = int(arr.shape[0])
        mean_emb[uid] = _l2_normalize_vec(arr.mean(axis=0))

    return {
        "mean_embeddings": mean_emb.astype(np.float32, copy=False),
        "chunk_embeddings": chunk_map,
        "chunk_counts": chunk_counts,
        "chunk_total": int(len(all_chunks)),
        "chunk_size": int(chunk_size),
        "chunk_stride": int(chunk_stride),
    }


def _maxsim_bidirectional(a: np.ndarray, b: np.ndarray) -> float:
    """Bidirectional MaxSim score in [0,1] from chunk embedding matrices."""
    if a is None or b is None:
        return 0.0
    A = np.asarray(a, dtype=np.float32)
    B = np.asarray(b, dtype=np.float32)
    if A.ndim != 2 or B.ndim != 2 or A.shape[0] <= 0 or B.shape[0] <= 0:
        return 0.0
    sim = np.matmul(A, B.T)  # cosine matrix (embeddings are L2-normalized)
    if sim.size <= 0:
        return 0.0
    row_max = np.max(sim, axis=1)
    col_max = np.max(sim, axis=0)
    score = 0.5 * (float(np.mean(row_max)) + float(np.mean(col_max)))
    return float(max(0.0, min(1.0, score)))


def _estimate_semantic_maxsim_refine_count(total_rows: int, granularity: str) -> int:
    """Select candidate count for MaxSim refine by granularity policy."""
    n = int(max(0, int(total_rows)))
    if n <= 0:
        return 0
    g = normalize_style(granularity)
    if not _semantic_maxsim_enabled(g):
        return 0
    if g == "h1":
        return n
    if g == "h2":
        raw = int(max(1, round(n * (float(SEMANTIC_MAXSIM_H2_TOP_PERCENT) / 100.0))))
        raw = max(int(SEMANTIC_MAXSIM_H2_TOPN_MIN), raw)
        raw = min(int(SEMANTIC_MAXSIM_H2_TOPN_MAX), raw)
        return int(max(1, min(raw, n)))
    return 0


def apply_semantic_maxsim_refine(rows: List[dict], ctx: Optional[dict], granularity: str) -> Dict[str, Any]:
    """Apply H1/H2 semantic MaxSim refine in-place on row['semantic_cosine']."""
    out = {
        "applied": 0,
        "selected": 0,
        "total": int(len(rows or [])),
        "lambda": None,
        "policy": "skip",
    }
    if not rows:
        return out
    g = normalize_style(granularity)
    if g not in ("h1", "h2"):
        return out
    if not _semantic_maxsim_enabled(g):
        out["policy"] = "disabled"
        return out
    if not isinstance(ctx, dict):
        return out

    mean_emb = ctx.get("mean_embeddings")
    chunk_map = ctx.get("chunk_embeddings") or {}
    if not isinstance(mean_emb, np.ndarray) or mean_emb.ndim != 2:
        return out

    lam = float(SEMANTIC_MAXSIM_H1_LAMBDA if g == "h1" else SEMANTIC_MAXSIM_H2_LAMBDA)
    lam = max(0.0, min(1.0, lam))

    n_sel = _estimate_semantic_maxsim_refine_count(len(rows), g)
    order = sorted(
        range(len(rows)),
        key=lambda idx: float(rows[idx].get("semantic_cosine") or 0.0),
        reverse=True,
    )
    sel = set(order[:n_sel]) if g == "h2" else set(range(len(rows)))
    out["selected"] = int(len(sel))
    out["lambda"] = float(lam)
    out["policy"] = ("all" if g == "h1" else f"top{SEMANTIC_MAXSIM_H2_TOP_PERCENT:.1f}%")

    applied = 0
    for i, r in enumerate(rows):
        try:
            id1 = int(r.get("id1"))
            id2 = int(r.get("id2"))
        except Exception:
            continue
        if id1 < 0 or id2 < 0 or id1 >= mean_emb.shape[0] or id2 >= mean_emb.shape[0]:
            continue

        doc_cos = r.get("semantic_cosine")
        try:
            doc = float(doc_cos) if doc_cos is not None else float(np.dot(mean_emb[id1], mean_emb[id2]))
        except Exception:
            doc = float(np.dot(mean_emb[id1], mean_emb[id2]))
        doc = float(max(0.0, min(1.0, doc)))

        r["semantic_doc_cos"] = doc
        r["semantic_mix_lambda"] = float(1.0 if i not in sel else lam)
        r["semantic_maxsim_bi"] = None
        r["semantic_maxsim_applied"] = 0

        if i not in sel:
            r["semantic_cosine"] = float(doc)
            continue

        A = chunk_map.get(int(id1))
        B = chunk_map.get(int(id2))
        ms = _maxsim_bidirectional(A, B)
        mix = float(lam * doc + (1.0 - lam) * ms)
        mix = float(max(0.0, min(1.0, mix)))
        r["semantic_maxsim_bi"] = float(ms)
        r["semantic_maxsim_applied"] = 1
        r["semantic_cosine"] = float(mix)
        applied += 1

    out["applied"] = int(applied)
    return out


def _load_bge_reranker_backend() -> dict:
    """Lazy-load BGE-reranker-base (BAAI/bge-reranker-base, RoBERTa, hidden=768)."""
    model_dir = _resolve_hf_snapshot_dir(BGE_RERANK_MODEL_DIR)
    cache_key = f"BGE-reranker:{model_dir}"
    if cache_key in _RERANK_BACKEND_CACHE:
        return _RERANK_BACKEND_CACHE[cache_key]

    _quiet_transformers_logging()
    try:
        import torch
        from transformers import AutoTokenizer, AutoModelForSequenceClassification
    except Exception as e:
        raise RuntimeError(f"BGE-reranker backend requires torch+transformers: {e}")

    if torch.cuda.is_available():
        device = "cuda"
    elif getattr(torch.backends, "mps", None) is not None and torch.backends.mps.is_available():
        device = "mps"
    else:
        device = "cpu"

    tokenizer = AutoTokenizer.from_pretrained(model_dir, local_files_only=True)

    load_kwargs = {"local_files_only": True}
    if device == "cuda":
        load_kwargs["torch_dtype"] = torch.float16

    try:
        model = AutoModelForSequenceClassification.from_pretrained(model_dir, **load_kwargs)
    except Exception:
        model = AutoModelForSequenceClassification.from_pretrained(
            model_dir, local_files_only=True,
        )
    _ensure_padding_token(tokenizer, model=model)

    model.to(device)
    model.eval()

    backend = {
        "torch": torch,
        "tokenizer": tokenizer,
        "model": model,
        "device": device,
        "model_dir": model_dir,
        "model_size": "BGE-reranker",
        "prompt_format": "raw",  # BGE-reranker: no prompt prefix needed
    }
    _RERANK_BACKEND_CACHE[cache_key] = backend
    print(f"[BGE-reranker] Loaded reranker from {model_dir} → {device}")
    return backend


def _load_reranker_backend(model_size: str) -> dict:
    """Lazy-load local reranker backend; cached by resolved model path."""
    model_size = _normalize_reranker_size_label(model_size or "0.6B")
    if _is_bge_model(model_size):
        return _load_bge_reranker_backend()
    if model_size not in QWEN_RERANK_MODEL_DIRS:
        raise RuntimeError(f"Unsupported reranker model size: {model_size}")

    model_key = model_size
    try:
        model_dir = _resolve_hf_snapshot_dir(QWEN_RERANK_MODEL_DIRS[model_key])
    except Exception:
        # Prefer robust baseline: fallback to local 0.6B when selected model is absent.
        fallback_key = "0.6B"
        if model_key != fallback_key:
            model_key = fallback_key
            model_dir = _resolve_hf_snapshot_dir(QWEN_RERANK_MODEL_DIRS[model_key])
        else:
            raise

    cache_key = f"{model_size}:{model_dir}"
    if cache_key in _RERANK_BACKEND_CACHE:
        return _RERANK_BACKEND_CACHE[cache_key]

    _quiet_transformers_logging()
    try:
        import torch
        from transformers import AutoTokenizer, AutoModelForSequenceClassification
    except Exception as e:
        raise RuntimeError(f"Reranker backend requires torch+transformers: {e}")

    if torch.cuda.is_available():
        device = "cuda"
    elif getattr(torch.backends, "mps", None) is not None and torch.backends.mps.is_available():
        device = "mps"
    else:
        device = "cpu"

    tokenizer = AutoTokenizer.from_pretrained(model_dir, trust_remote_code=True, local_files_only=True)

    load_kwargs = {
        "trust_remote_code": True,
        "local_files_only": True,
    }
    if device == "cuda":
        load_kwargs["torch_dtype"] = torch.float16

    try:
        model = AutoModelForSequenceClassification.from_pretrained(model_dir, **load_kwargs)
    except Exception:
        model = AutoModelForSequenceClassification.from_pretrained(
            model_dir,
            trust_remote_code=True,
            local_files_only=True,
        )
    _ensure_padding_token(tokenizer, model=model)

    model.to(device)
    model.eval()

    backend = {
        "torch": torch,
        "tokenizer": tokenizer,
        "model": model,
        "device": device,
        "model_dir": model_dir,
        "model_size": model_key,
    }
    _RERANK_BACKEND_CACHE[cache_key] = backend
    return backend


def _truncate_for_rerank(text: str, max_chars: int = 1200) -> str:
    t = _embed_text_for_semantic(text or "")
    m = int(max(64, max_chars))
    if len(t) <= m:
        return t
    return t[:m]


def _score_reranker_pairs(
    pairs: List[Tuple[str, str]],
    model_size: str,
    batch_size: int = 8,
    max_length: int = 512,
    prompt: str = "",
) -> Tuple[List[float], str]:
    backend = _load_reranker_backend(model_size)
    torch = backend["torch"]
    tokenizer = backend["tokenizer"]
    model = backend["model"]
    device = backend["device"]
    model_used = str(backend.get("model_size") or model_size)

    if not pairs:
        return [], model_used

    out_scores: List[float] = []
    bs = int(max(1, batch_size))
    if getattr(tokenizer, "pad_token_id", None) is None or getattr(model.config, "pad_token_id", None) is None:
        # Some sequence-classification backbones reject batch>1 without config pad_token_id.
        bs = 1
    ml = int(max(64, max_length))

    for i in range(0, len(pairs), bs):
        chunk = pairs[i: i + bs]
        a_texts = [x[0] for x in chunk]
        b_texts = [x[1] for x in chunk]
        if prompt and backend.get("prompt_format") != "raw":
            p = str(prompt).strip()
            a_texts = [f"{p}\n文本A：{a}" for a in a_texts]
            b_texts = [f"文本B：{b}" for b in b_texts]
        with torch.no_grad():
            toks = tokenizer(
                a_texts,
                b_texts,
                padding=(len(chunk) > 1),
                truncation=True,
                max_length=ml,
                return_tensors="pt",
            )
            toks = {k: v.to(device) for k, v in toks.items()}
            out = model(**toks)

            logits = None
            if isinstance(out, dict):
                logits = out.get("logits")
            else:
                logits = getattr(out, "logits", None)
                if logits is None and isinstance(out, (tuple, list)) and len(out) > 0:
                    logits = out[0]
            if logits is None:
                raise RuntimeError("Reranker output missing logits.")

            if logits.dim() == 1:
                scores = logits
            elif logits.size(-1) == 1:
                scores = logits.squeeze(-1)
            else:
                # binary/multi-class fallback: use first logit as relevance proxy
                scores = logits[:, 0]

            # bfloat16 logits may fail at direct `.numpy()` conversion.
            out_scores.extend([float(x) for x in scores.detach().float().cpu().tolist()])

    return out_scores, model_used


def annotate_reranker_sidecar(
    rows: List[dict],
    units: List[Unit],
    enabled: bool,
    model_size: str,
    topn: Optional[int],
    granularity: str = "normal",
    top_percent_non_h1: float = float(RERANK_DEFAULT_TOP_PERCENT_NON_H1),
    prompt: str = "",
) -> None:
    """Compute reranker diagnostics without changing final ranking.

    `reranker_ref_01` is an in-layer percentile reference score on reranked rows,
    not a probability.
    """
    if not rows:
        return

    for r in rows:
        try:
            qb = float(r.get("score_final") or 0.0)
        except Exception:
            qb = 0.0
        qb = max(0.0, min(1.0, qb))
        r["score_base_q"] = float(qb)
        r["score_reranker_raw"] = None
        r["reranker_ref_01"] = None
        r["rank_rerank_raw"] = None
        r["rerank_applied"] = 0
        r["reranker_model"] = None

    if not enabled:
        return

    n_sel = _estimate_reranker_candidate_count(
        total_rows=len(rows),
        granularity=granularity,
        topn=topn,
        top_percent_non_h1=top_percent_non_h1,
    )

    order0 = sorted(range(len(rows)), key=lambda idx: float(rows[idx].get("score_final") or 0.0), reverse=True)
    sel = order0[:n_sel]

    pairs: List[Tuple[str, str]] = []
    for idx in sel:
        r = rows[idx]
        u1 = units[int(r["id1"])]
        u2 = units[int(r["id2"])]
        pairs.append((
            _truncate_for_rerank(u1.text or "", max_chars=1200),
            _truncate_for_rerank(u2.text or "", max_chars=1200),
        ))

    raw_scores, model_used = _score_reranker_pairs(pairs, model_size=model_size, prompt=prompt)
    if len(raw_scores) != len(sel):
        raise RuntimeError(f"reranker score size mismatch: got={len(raw_scores)} expect={len(sel)}")

    rerank_vals = [(int(sel[i]), float(raw_scores[i])) for i in range(len(sel))]
    rank_rer = _rank_order(rerank_vals, descending=True)
    ref_rer = _rank_map(rerank_vals)

    for i, idx in enumerate(sel):
        r = rows[idx]
        rr = float(raw_scores[i])
        r["score_reranker_raw"] = rr
        r["reranker_ref_01"] = float(ref_rer.get(idx, 0.0)) if idx in ref_rer else None
        r["rank_rerank_raw"] = int(rank_rer.get(idx, 0)) if idx in rank_rer else None
        r["rerank_applied"] = 1
        r["reranker_model"] = str(model_used)


# =========================
# IO: Excel reading
# =========================
def read_excel_as_units_table(path: str) -> Optional[List[Unit]]:
    """
    Units table format (recommended):
      Header includes at least: text
      Optional: level / h1 / h2 / h3 / chapter(章节号)
    Each row is already a unit; granularity later filters by level.
    """
    wb = openpyxl.load_workbook(path)
    ws = wb.active
    headers = [str(c.value).strip().lower() if c.value is not None else "" for c in ws[1]]
    # Units-table requires an explicit `level` column; otherwise it is treated as a source table.
    if "level" not in headers:
        return None

    # accept either English `text` or Chinese `正文` as the text column
    if "text" in headers:
        idx_text = headers.index("text")
    elif "正文" in headers:
        idx_text = headers.index("正文")
    else:
        return None

    idx_level = headers.index("level") if "level" in headers else None
    idx_h1 = headers.index("h1") if "h1" in headers else None
    idx_h2 = headers.index("h2") if "h2" in headers else None
    idx_h3 = headers.index("h3") if "h3" in headers else None
    idx_ch = None
    for name in ("章节号", "章節號", "chapter", "章节"):
        if name in headers:
            idx_ch = headers.index(name)
            break

    units: List[Unit] = []
    uid = 0
    for r in ws.iter_rows(min_row=2, values_only=True):
        txt = r[idx_text] if idx_text is not None else ""
        if txt is None:
            txt = ""
        lvl = r[idx_level] if idx_level is not None else "h2"
        lvl = normalize_style(str(lvl))
        h1 = str(r[idx_h1]) if idx_h1 is not None and r[idx_h1] is not None else ""
        h2 = str(r[idx_h2]) if idx_h2 is not None and r[idx_h2] is not None else ""
        h3 = str(r[idx_h3]) if idx_h3 is not None and r[idx_h3] is not None else ""
        ch = str(r[idx_ch]) if idx_ch is not None and idx_ch < len(r) and r[idx_ch] is not None else ""
        h3_final = h3.strip() if h3.strip() else ch.strip()

        units.append(Unit(uid=uid, level=lvl, h1=h1.strip(), h2=h2.strip(), h3=h3_final, text=str(txt), para_ids=[]))
        uid += 1
    return units


def read_excel_as_paragraphs_table(path: str) -> List[Para]:
    """
    Paragraphs table format:
      Column A = text
      Column B = style (h1/h2/h3/normal)
    """
    wb = openpyxl.load_workbook(path)
    ws = wb.active
    paras: List[Para] = []
    pid = 0
    for r in ws.iter_rows(min_row=1, values_only=True):
        if r is None:
            continue
        txt = r[0] if len(r) >= 1 else ""
        sty = r[1] if len(r) >= 2 else "normal"
        if txt is None:
            txt = ""
        paras.append(Para(pid=pid, text=str(txt), style=normalize_style(str(sty))))
        pid += 1
    return paras


# =========================
# Source table (user format) support
# =========================

def read_excel_as_source_table(path: str) -> Optional[List[dict]]:
    """Source table format (user): headers like 序号/书名/篇名/章节号/正文.
    Returns list of rows with keys: h1,h2,h3,chapter,text.
    `h3` is path-ready and prefers chapter number when provided.
    """
    wb = openpyxl.load_workbook(path)
    ws = wb.active

    # Keep original headers (do not lower-case Chinese)
    raw_headers = [str(c.value).strip() if c.value is not None else "" for c in ws[1]]
    # Also prepare a lower-case copy for English keys
    low_headers = [h.lower() for h in raw_headers]

    def find_idx(names: List[str]) -> Optional[int]:
        for n in names:
            if n in raw_headers:
                return raw_headers.index(n)
            if n.lower() in low_headers:
                return low_headers.index(n.lower())
        return None

    idx_text = find_idx(["正文", "text"])
    idx_h1 = find_idx(["书名", "h1"])
    idx_h2 = find_idx(["篇名", "h2"])
    idx_h3 = find_idx(["h3", "小题", "小標", "小标题"])  # optional
    idx_ch = find_idx(["章节号", "章節號", "chapter", "章节"])

    # Must have text and at least h1/h2 to be considered this source-table
    if idx_text is None or idx_h1 is None or idx_h2 is None:
        return None

    out: List[dict] = []
    for r in ws.iter_rows(min_row=2, values_only=True):
        if r is None:
            continue
        txt = r[idx_text] if idx_text < len(r) else ""
        if txt is None:
            txt = ""
        h1 = r[idx_h1] if idx_h1 is not None and idx_h1 < len(r) else ""
        h2 = r[idx_h2] if idx_h2 is not None and idx_h2 < len(r) else ""
        h3 = r[idx_h3] if idx_h3 is not None and idx_h3 < len(r) else ""
        ch = r[idx_ch] if idx_ch is not None and idx_ch < len(r) else ""

        h3_s = "" if h3 is None else str(h3).strip()
        ch_s = "" if ch is None else str(ch).strip()
        # For normal-level path display, prefer chapter number when provided.
        h3_for_path = ch_s if ch_s else h3_s

        out.append({
            "h1": "" if h1 is None else str(h1).strip(),
            "h2": "" if h2 is None else str(h2).strip(),
            "h3": h3_for_path,
            "chapter": ch_s,
            "text": str(txt),
        })
    return out


def build_units_from_source_rows(rows: List[dict], granularity: str) -> List[Unit]:
    """Build comparison units from source rows (书名/篇名/正文...).

    - normal: each row -> one unit
    - h2: group by (h1,h2)
    - h1: group by (h1)
    - h3: group by (h1,h2,h3/chapter) if exists, else fallback to h2

    The unit text is concatenation of row texts with '\n'.
    """
    g = normalize_style(granularity)

    # Build base normal units (one row = one unit) for later wordcount stats
    normal_units: List[Unit] = []
    for i, r in enumerate(rows):
        h3_row = str(r.get("h3") or "").strip()
        ch_row = str(r.get("chapter") or "").strip()
        h3_for_path = ch_row if ch_row else h3_row
        normal_units.append(Unit(
            uid=i,
            level="normal",
            h1=r.get("h1", ""),
            h2=r.get("h2", ""),
            h3=h3_for_path,
            text=(r.get("text") or ""),
            para_ids=[i],
        ))

    if g == "normal":
        return normal_units

    # group
    groups: Dict[Tuple, List[Unit]] = {}
    for u in normal_units:
        if g == "h1":
            key = (u.h1 or "(empty)",)
        elif g == "h2":
            key = (u.h1 or "(empty)", u.h2 or "(empty)")
        elif g == "h3":
            # if no h3, fallback to h2 grouping
            if (u.h3 or "").strip():
                key = (u.h1 or "(empty)", u.h2 or "(empty)", u.h3 or "(empty)")
            else:
                key = (u.h1 or "(empty)", u.h2 or "(empty)")
        else:
            key = (u.h1 or "(empty)", u.h2 or "(empty)")
        groups.setdefault(key, []).append(u)

    out: List[Unit] = []
    uid = 0
    for key, items in groups.items():
        h1 = items[0].h1
        h2 = items[0].h2
        h3 = items[0].h3
        txt = "\n".join([(it.text or "").strip() for it in items if (it.text or "").strip()]) + "\n"
        out.append(Unit(uid=uid, level=g, h1=h1, h2=h2, h3=h3, text=txt, para_ids=[it.para_ids[0] for it in items if it.para_ids]))
        uid += 1

    return out


# =========================
# Unit segmentation
# =========================
def segment_units_from_paragraphs(paras: List[Para], granularity: str) -> List[Unit]:
    """
    Aggregate NORMAL paragraphs under selected granularity:
      - granularity=h1: each h1 starts a new unit
      - granularity=h2: each h2 starts a new unit (under current h1)
      - granularity=h3: each h3 starts a new unit (under current h1/h2)
    """
    granularity = normalize_style(granularity)
    # Special case: paragraph-level units (NORMAL)
    # 每个 normal 段落单独成一个 unit；h1/h2/h3 仅用于给该段落挂载路径信息
    if granularity == "normal":
        units: List[Unit] = []
        cur_h1 = ""
        cur_h2 = ""
        cur_h3 = ""
        uid = 0
        for p in paras:
            if p.style == "h1":
                cur_h1 = p.text.strip()
                cur_h2 = ""
                cur_h3 = ""
                continue
            if p.style == "h2":
                cur_h2 = p.text.strip()
                cur_h3 = ""
                continue
            if p.style == "h3":
                cur_h3 = p.text.strip()
                continue
            # normal
            txt = (p.text or "")
            if txt.strip():
                units.append(Unit(uid=uid, level="normal", h1=cur_h1, h2=cur_h2, h3=cur_h3, text=txt + "\n", para_ids=[p.pid]))
                uid += 1
        return units

    units: List[Unit] = []

    cur_h1 = ""
    cur_h2 = ""
    cur_h3 = ""

    uid = 0

    def new_unit() -> Unit:
        nonlocal uid
        u = Unit(uid=uid, level=granularity, h1=cur_h1, h2=cur_h2, h3=cur_h3, text="", para_ids=[])
        uid += 1
        return u

    def finalize(u: Optional[Unit]):
        if u is None:
            return
        if u.text.strip():
            units.append(u)

    current_unit: Optional[Unit] = None

    for p in paras:
        if p.style == "h1":
            finalize(current_unit)
            cur_h1 = p.text.strip()
            cur_h2 = ""
            cur_h3 = ""
            current_unit = new_unit() if granularity == "h1" else None
            continue

        if p.style == "h2":
            finalize(current_unit)
            cur_h2 = p.text.strip()
            cur_h3 = ""
            current_unit = new_unit() if granularity in ("h2",) else None
            continue

        if p.style == "h3":
            finalize(current_unit)
            cur_h3 = p.text.strip()
            current_unit = new_unit() if granularity in ("h3",) else None
            continue

        # normal text
        if current_unit is None:
            # if no current unit due to higher granularity, create under current headings
            current_unit = new_unit()
        current_unit.text += p.text + "\n"
        current_unit.para_ids.append(p.pid)

    finalize(current_unit)
    return units


def select_units_from_units_table(units: List[Unit], granularity: str) -> List[Unit]:
    """
    如果 Excel 是“单位表”，每行已经是一个 unit。
    这里按用户选择的 granularity 过滤：
      - granularity=h1: 只取 level=h1 的行
      - granularity=h2: 只取 level=h2 的行
      - granularity=h3: 只取 level=h3 的行
    """
    g = normalize_style(granularity)
    return [u for u in units if normalize_style(u.level) == g]


# =========================
# Representations
# =========================
def build_representations(units: List[Unit], ngram_n: int, use_tfidf: bool):
    # Similarity is computed on preprocessed text (Han-only). Raw text is preserved in Unit.text for output.
    texts = [preprocess_keep_han(u.text or "") for u in units]

    tf_vec = CountVectorizer(
        analyzer="char",
        ngram_range=(ngram_n, ngram_n),
        lowercase=False
    )
    X_tf = tf_vec.fit_transform(texts)

    X_tfidf = None
    tfidf_vec = None
    if use_tfidf:
        tfidf_vec = TfidfVectorizer(
            analyzer="char",
            ngram_range=(ngram_n, ngram_n),
            lowercase=False
        )
        X_tfidf = tfidf_vec.fit_transform(texts)

    S_set = [set(char_ngrams(t, ngram_n)) for t in texts]

    mh_sigs = None
    if MinHash is not None:
        mh_sigs = []
        for s in S_set:
            mh = MinHash(num_perm=128)
            for tok in s:
                mh.update(tok.encode("utf-8"))
            mh_sigs.append(mh)

    return {"X_tf": X_tf, "X_tfidf": X_tfidf, "S_set": S_set, "mh_sigs": mh_sigs, "tf_vec": tf_vec, "tfidf_vec": tfidf_vec}


def _split_cn_title_and_tail(s: str) -> Tuple[Optional[str], str]:
    t = str(s or "").strip()
    if not t:
        return None, ""
    m = re.match(r"^\s*《([^》]+)》\s*(.*?)\s*$", t)
    if not m:
        return None, t
    return (m.group(1) or "").strip(), (m.group(2) or "").strip()


def _merge_book_article_path(h1: str, h2: str) -> str:
    b = str(h1 or "").strip()
    p = str(h2 or "").strip()
    if not b and not p:
        return ""
    if not b:
        return p
    if not p:
        return b

    bt, bs = _split_cn_title_and_tail(b)
    pt, ps = _split_cn_title_and_tail(p)

    # Preferred style: 《书名·篇名》<尾注>
    if bt and pt:
        merged = f"《{bt}·{pt}》"
        tail = f"{bs}{ps}".strip()
        return f"{merged}{tail}" if tail else merged

    # Generic fallback keeps middle-dot merge without slash.
    return f"{b}·{p}"


def _normalize_chapter_path_part(s: str) -> str:
    t = str(s or "").strip()
    if not t:
        return ""
    # Normalize decimal-like dot to full-width: 1. 3 -> 1．3
    t = re.sub(r"\s*[\.．]\s*", "．", t)
    t = re.sub(r"\s+", " ", t).strip()
    return t


def _unit_path_parts(u: Unit) -> List[str]:
    lvl = normalize_style(u.level)
    h1 = str(u.h1 or "").strip()
    h2 = str(u.h2 or "").strip()
    h3 = _normalize_chapter_path_part(u.h3 or "")
    merged_h1_h2 = _merge_book_article_path(h1, h2)

    if lvl == "h1":
        raw = [h1 or h2]
    elif lvl == "h2":
        raw = [merged_h1_h2]
    else:
        # h3 / normal: keep merged book-article and chapter.
        raw = [merged_h1_h2, h3]
    return [p.strip() for p in raw if (p or "").strip()]


def unit_path(u: Unit) -> str:
    parts = _unit_path_parts(u)
    return " / ".join(parts) if parts else f"unit_{u.uid}"


# =========================
# Length utility
# =========================
def _eff_len_chars(s: str) -> int:
    """Effective length: count Han characters only."""
    return len(preprocess_keep_han(s or ""))


def _topk_desc_indices(arr: np.ndarray, k: int) -> np.ndarray:
    if arr.size <= 0 or k <= 0:
        return np.array([], dtype=int)
    if k >= arr.size:
        return np.argsort(-arr)
    idx = np.argpartition(-arr, k - 1)[:k]
    return idx[np.argsort(-arr[idx])]


def _unique_keep_order(vals: List[int]) -> List[int]:
    seen = set()
    out: List[int] = []
    for v in vals:
        if v in seen:
            continue
        seen.add(v)
        out.append(v)
    return out


def _split_han_chunks(text: str, chunk_size: int) -> List[str]:
    t = preprocess_keep_han(text or "")
    if not t:
        return []
    c = int(max(80, chunk_size))
    if len(t) <= c:
        return [t]
    out: List[str] = []
    i = 0
    while i < len(t):
        out.append(t[i:i + c])
        i += c
    return out


def _safe_int(x: Any) -> Optional[int]:
    try:
        return int(x)
    except Exception:
        return None


def _build_coverage_context(
    units: List[Unit],
    normal_units: List[Unit],
    reps: dict,
    ngram_n: int,
) -> dict:
    """Build segment-level context for coverage scoring.

    Priority:
      1) Use normal_units mapped by para_ids/headings.
      2) If mapping unavailable, fallback to pseudo chunks from unit.text.
    """
    tf_vec = reps.get("tf_vec")
    tfidf_vec = reps.get("tfidf_vec")
    if tf_vec is None:
        raise RuntimeError("Coverage requires tf_vec in representations.")

    seg_texts: List[str] = []
    seg_lens: List[int] = []
    seg_sets: List[set] = []
    seg_paths: List[str] = []

    unit_seg_ids: Dict[int, List[int]] = {u.uid: [] for u in units}

    normals = [u for u in (normal_units or []) if normalize_style(u.level) == "normal"]
    para_to_sids: Dict[int, List[int]] = {}
    h1_to_sids: Dict[str, List[int]] = {}
    h2_to_sids: Dict[Tuple[str, str], List[int]] = {}

    for nu in normals:
        t = preprocess_keep_han(nu.text or "")
        if not t:
            continue
        sid = len(seg_texts)
        seg_texts.append(t)
        seg_lens.append(len(t))
        seg_sets.append(set(char_ngrams(t, ngram_n)))
        seg_paths.append(unit_path(nu))

        if nu.para_ids:
            for pid in nu.para_ids:
                iv = _safe_int(pid)
                if iv is not None:
                    para_to_sids.setdefault(iv, []).append(sid)
        else:
            iv = _safe_int(nu.uid)
            if iv is not None:
                para_to_sids.setdefault(iv, []).append(sid)

        k1 = (nu.h1 or "").strip()
        k2 = ((nu.h1 or "").strip(), (nu.h2 or "").strip())
        h1_to_sids.setdefault(k1, []).append(sid)
        h2_to_sids.setdefault(k2, []).append(sid)

    for u in units:
        got: List[int] = []

        for pid in (u.para_ids or []):
            iv = _safe_int(pid)
            if iv is None:
                continue
            got.extend(para_to_sids.get(iv, []))

        if not got:
            lvl = normalize_style(u.level)
            if lvl == "h1":
                got.extend(h1_to_sids.get((u.h1 or "").strip(), []))
            elif lvl == "h2":
                got.extend(h2_to_sids.get(((u.h1 or "").strip(), (u.h2 or "").strip()), []))

        got = _unique_keep_order(got)

        # fallback: pseudo chunks from this unit itself
        if not got:
            chunk_size = 500 if normalize_style(u.level) == "h1" else 300
            chunks = _split_han_chunks(u.text or "", chunk_size)
            for ci, ch in enumerate(chunks):
                sid = len(seg_texts)
                seg_texts.append(ch)
                seg_lens.append(len(ch))
                seg_sets.append(set(char_ngrams(ch, ngram_n)))
                seg_paths.append(f"{unit_path(u)} / pseudo_{ci}")
                got.append(sid)

        unit_seg_ids[u.uid] = got

    X_seg_tf = tf_vec.transform(seg_texts) if seg_texts else None
    X_seg_tfidf = tfidf_vec.transform(seg_texts) if (seg_texts and tfidf_vec is not None) else None

    return {
        "unit_seg_ids": unit_seg_ids,
        "seg_lens": seg_lens,
        "seg_sets": seg_sets,
        "seg_paths": seg_paths,
        "X_seg_tf": X_seg_tf,
        "X_seg_tfidf": X_seg_tfidf,
        "pair_cache": {},
    }


def _coverage_pair_score(
    id1: int,
    id2: int,
    ctx: dict,
    seg_topk: int = 10,
    min_cos: float = 0.20,
    min_jaccard: float = 0.02,
    hits_max: int = 6,
) -> dict:
    key = (id1, id2) if id1 < id2 else (id2, id1)
    cache = ctx.get("pair_cache")
    if isinstance(cache, dict) and key in cache:
        return cache[key]

    sids_a = ctx["unit_seg_ids"].get(id1, [])
    sids_b = ctx["unit_seg_ids"].get(id2, [])
    seg_lens = ctx["seg_lens"]
    seg_sets = ctx["seg_sets"]
    seg_paths = ctx["seg_paths"]
    X_seg_tf = ctx.get("X_seg_tf")
    X_seg_tfidf = ctx.get("X_seg_tfidf")

    total_a = int(sum(seg_lens[s] for s in sids_a)) if sids_a else 0
    total_b = int(sum(seg_lens[s] for s in sids_b)) if sids_b else 0

    out = {
        "coverage_score": 0.0,
        "coverage_f1": 0.0,
        "coverage_cov1": 0.0,
        "coverage_cov2": 0.0,
        "coverage_q_tf": 0.0,
        "coverage_q_tfidf": None,
        "coverage_overlap_chars": 0,
        "coverage_match_edges": 0,
        "coverage_candidate_edges": 0,
        "coverage_hits": "",
    }
    if not sids_a or not sids_b or total_a <= 0 or total_b <= 0 or X_seg_tf is None:
        if isinstance(cache, dict):
            cache[key] = out
        return out

    XA = X_seg_tf[sids_a]
    XB = X_seg_tf[sids_b]
    if XA.shape[0] == 0 or XB.shape[0] == 0:
        if isinstance(cache, dict):
            cache[key] = out
        return out

    sim_tf = cosine_similarity(XA, XB)
    sim_tfidf = None
    if X_seg_tfidf is not None:
        try:
            sim_tfidf = cosine_similarity(X_seg_tfidf[sids_a], X_seg_tfidf[sids_b])
        except Exception:
            sim_tfidf = None

    edges: List[dict] = []
    k_each = int(max(1, seg_topk))
    for ai in range(sim_tf.shape[0]):
        row = sim_tf[ai]
        top_idx = _topk_desc_indices(row, min(k_each, row.size))
        for bj in top_idx:
            ctf = float(row[int(bj)])
            if ctf < float(min_cos):
                continue
            sid_a = sids_a[ai]
            sid_b = sids_b[int(bj)]
            jac = float(jaccard(seg_sets[sid_a], seg_sets[sid_b]))
            if jac < float(min_jaccard):
                continue
            cap = int(min(seg_lens[sid_a], seg_lens[sid_b]))
            if cap <= 0:
                continue
            ctfidf = None
            if sim_tfidf is not None:
                ctfidf = float(sim_tfidf[ai, int(bj)])
            edges.append({
                "ai": ai,
                "bj": int(bj),
                "sid_a": sid_a,
                "sid_b": sid_b,
                "cap": cap,
                "tf": ctf,
                "tfidf": ctfidf,
                "jac": jac,
                "w": ctf * cap,
            })

    out["coverage_candidate_edges"] = int(len(edges))
    if not edges:
        if isinstance(cache, dict):
            cache[key] = out
        return out

    edges.sort(key=lambda e: (e["w"], e["tf"], e["jac"]), reverse=True)

    rem_a = [int(seg_lens[s]) for s in sids_a]
    rem_b = [int(seg_lens[s]) for s in sids_b]

    overlap = 0.0
    q_tf_num = 0.0
    q_tfidf_num = 0.0
    q_tfidf_den = 0.0
    matched: List[dict] = []

    for e in edges:
        ai = e["ai"]
        bj = e["bj"]
        if ai >= len(rem_a) or bj >= len(rem_b):
            continue
        flow = float(min(rem_a[ai], rem_b[bj], e["cap"]))
        if flow <= 0:
            continue
        rem_a[ai] -= int(flow)
        rem_b[bj] -= int(flow)
        overlap += flow
        q_tf_num += e["tf"] * flow
        if e["tfidf"] is not None:
            q_tfidf_num += float(e["tfidf"]) * flow
            q_tfidf_den += flow
        matched.append({
            "sid_a": e["sid_a"],
            "sid_b": e["sid_b"],
            "flow": flow,
            "tf": e["tf"],
            "tfidf": e["tfidf"],
            "jac": e["jac"],
            "contrib": e["tf"] * flow,
        })

    if overlap <= 0:
        if isinstance(cache, dict):
            cache[key] = out
        return out

    cov1 = float(overlap / max(1, total_a))
    cov2 = float(overlap / max(1, total_b))
    f1 = float(0.0 if (cov1 + cov2) <= 0 else (2.0 * cov1 * cov2 / (cov1 + cov2)))
    q_tf = float(q_tf_num / overlap)
    q_tfidf = None if q_tfidf_den <= 0 else float(q_tfidf_num / q_tfidf_den)
    score = float(f1 * q_tf)

    matched.sort(key=lambda x: x["contrib"], reverse=True)
    hit_lines: List[str] = []
    for h in matched[: max(1, int(hits_max))]:
        pa = str(seg_paths[h["sid_a"]]).replace("\n", " ").strip()
        pb = str(seg_paths[h["sid_b"]]).replace("\n", " ").strip()
        if len(pa) > 40:
            pa = pa[:37] + "..."
        if len(pb) > 40:
            pb = pb[:37] + "..."
        hit_lines.append(
            f"A[{h['sid_a']}:{pa}]<->B[{h['sid_b']}:{pb}] "
            f"len={int(h['flow'])} tf={h['tf']:.3f} jac={h['jac']:.3f}"
        )

    out = {
        "coverage_score": score,
        "coverage_f1": f1,
        "coverage_cov1": cov1,
        "coverage_cov2": cov2,
        "coverage_q_tf": q_tf,
        "coverage_q_tfidf": q_tfidf,
        "coverage_overlap_chars": int(round(overlap)),
        "coverage_match_edges": int(len(matched)),
        "coverage_candidate_edges": int(len(edges)),
        "coverage_hits": " || ".join(hit_lines),
    }

    if isinstance(cache, dict):
        cache[key] = out
    return out


# =========================
# Similarity computation
# =========================
def compute_similarities(
    units: List[Unit],
    reps: dict,
    mode: str = "topk",
    topk: int = 20,
    include_tfidf: bool = True,
    granularity: str = "h2",
    normal_units: Optional[List[Unit]] = None,
    ngram_n: int = 3,
    coverage_seg_topk: int = int(COVERAGE_DEFAULTS["seg_topk"]),
    coverage_min_cos: float = float(COVERAGE_DEFAULTS["min_cos"]),
    coverage_min_jaccard: float = float(COVERAGE_DEFAULTS["min_jaccard"]),
    coverage_hits_max: int = int(COVERAGE_DEFAULTS["hits_max"]),
    semantic_embeddings: Optional[np.ndarray] = None,
    semantic_neighbors: Optional[Dict[int, List[int]]] = None,
    candidate_source: str = "lexical",
    framework: str = "lexical",
    progress_cb=None,
) -> List[dict]:
    """
    关键保证：
      - 只输出 unordered pair：id1 < id2
      - 不输出 self：id1 == id2 永远被过滤
      - topk 模式只用于“候选生成”，最后按 (id1,id2) 去重，绝不会出现 A-B 与 B-A
    """
    X_tf = reps["X_tf"]
    X_tfidf = reps["X_tfidf"]
    S_set = reps["S_set"]
    mh_sigs = reps["mh_sigs"]

    n = len(units)
    if n < 2:
        return []

    mode = (mode or "topk").lower()
    topk = max(1, int(topk))
    g = normalize_style(granularity)
    fw = str(framework or "lexical").strip().lower()
    if fw not in ("lexical", "semantic", "merge"):
        fw = "lexical"
    semantic_only = (fw == "semantic")

    sem_neighbors = semantic_neighbors or {}
    sem_emb = semantic_embeddings
    use_sem = sem_emb is not None and isinstance(sem_emb, np.ndarray) and sem_emb.shape[0] == n

    candidate_source = str(candidate_source or "lexical").strip().lower()
    if candidate_source not in ("lexical", "semantic", "hybrid"):
        candidate_source = "lexical"
    if semantic_only:
        candidate_source = "semantic"
    elif fw == "lexical":
        candidate_source = "lexical"

    use_coverage = (g in ("h1", "h2")) and (not semantic_only)

    if candidate_source in ("semantic", "hybrid") and not use_sem:
        raise RuntimeError(f"candidate_source={candidate_source} requires semantic embeddings.")

    X_tf_norm = None if semantic_only else sk_normalize(X_tf, norm="l2", copy=True)

    # Build self-like fingerprints once:
    #   1) exact same index (hard self)
    #   2) same uid (input duplicated logical id)
    #   3) same heading path + same Han-only text (exact duplicated unit content)
    unit_uid: List[Optional[int]] = []
    unit_path_sig: List[str] = []
    unit_text_sig: List[str] = []
    for u in units:
        try:
            uidv = int(u.uid)
        except Exception:
            uidv = None
        unit_uid.append(uidv)
        unit_path_sig.append(unit_path(u))
        unit_text_sig.append(preprocess_keep_han(u.text or ""))

    sem_k_hint = 0
    if sem_neighbors:
        try:
            sem_k_hint = max(len(v) for v in sem_neighbors.values())
        except Exception:
            sem_k_hint = 0

    if mode == "all":
        total_est = n * (n - 1) // 2
    else:
        if candidate_source == "semantic":
            total_est = n * min(topk, max(0, n - 1))
        elif candidate_source == "hybrid":
            total_est = n * min(topk + sem_k_hint, max(0, n - 1))
        else:
            total_est = n * min(topk, max(0, n - 1))

    def report(done: int, total: int, msg: str):
        if progress_cb is None:
            return
        try:
            progress_cb(int(done), int(total), msg)
        except Exception:
            pass

    def tfidf_cos(i: int, j: int) -> Optional[float]:
        if semantic_only:
            return None
        if not include_tfidf or X_tfidf is None:
            return None
        try:
            # TfidfVectorizer default outputs l2-normalized rows, so dot product is cosine.
            return float(X_tfidf[i].multiply(X_tfidf[j]).sum())
        except Exception:
            return float(cosine_similarity(X_tfidf[i], X_tfidf[j])[0, 0])

    def lex_cos_row(i: int) -> np.ndarray:
        if X_tf_norm is None:
            raise RuntimeError("Lexical matrix is unavailable for lex cosine row.")
        # Sparse cosine row from pre-normalized matrix.
        col = X_tf_norm @ X_tf_norm[i].T
        if hasattr(col, "toarray"):
            arr = col.toarray().ravel()
        else:
            arr = np.asarray(col).ravel()
        return arr.astype(np.float32, copy=False)

    coverage_ctx = None
    if use_coverage:
        report(0, max(1, total_est), "building coverage context...")
        coverage_ctx = _build_coverage_context(
            units=units,
            normal_units=(normal_units or []),
            reps=reps,
            ngram_n=max(1, int(ngram_n)),
        )

    # key=(id1,id2) -> row
    best: Dict[Tuple[int, int], dict] = {}
    skipped_self_like = 0

    def is_self_like(i: int, j: int) -> bool:
        if i == j:
            return True
        ui = unit_uid[i]
        uj = unit_uid[j]
        if ui is not None and uj is not None and ui == uj:
            return True
        if unit_path_sig[i] == unit_path_sig[j]:
            ti = unit_text_sig[i]
            tj = unit_text_sig[j]
            if ti and tj and ti == tj:
                return True
        return False

    def add_pair(
        i: int,
        j: int,
        cos_tf_raw: Optional[float],
        semantic_raw: Optional[float],
    ):
        nonlocal skipped_self_like
        if is_self_like(i, j):
            skipped_self_like += 1
            return

        id1, id2 = (i, j) if i < j else (j, i)
        key = (id1, id2)
        if key in best:
            return

        sem_pair = None
        if semantic_raw is not None:
            try:
                sem_pair = float(semantic_raw)
            except Exception:
                sem_pair = None
        elif use_sem:
            sem_pair = float(np.dot(sem_emb[id1], sem_emb[id2]))

        len_a = _eff_len_chars(units[id1].text or "")
        len_b = _eff_len_chars(units[id2].text or "")
        short_len = min(len_a, len_b) if (len_a > 0 and len_b > 0) else 0
        long_len = max(len_a, len_b)
        len_ratio = float(long_len / max(1, short_len)) if short_len > 0 else 0.0
        # Symmetric length-balance factor in (0,1], equals 1 when lengths are equal.
        len_balance = float(0.0 if (short_len + long_len) <= 0 else (2.0 * short_len / (short_len + long_len)))

        jac = None
        cont = None
        short_id = None
        long_id = None
        mh = None
        cos_tfidf = None
        cov = None
        score_mode = "semantic_cosine" if semantic_only else "direct_tf"
        score_main = float(sem_pair or 0.0) if semantic_only else float(cos_tf_raw or 0.0)

        if not semantic_only:
            jac = jaccard(S_set[id1], S_set[id2])
            cont, flag = containment_short_in_long(S_set[id1], S_set[id2])
            if flag == 0:
                short_id, long_id = id1, id2
            else:
                short_id, long_id = id2, id1

            if mh_sigs is not None:
                mh = minhash_jaccard(mh_sigs[id1], mh_sigs[id2])

            cos_tfidf = tfidf_cos(id1, id2)

            if use_coverage and coverage_ctx is not None:
                cov = _coverage_pair_score(
                    id1=id1,
                    id2=id2,
                    ctx=coverage_ctx,
                    seg_topk=int(max(1, coverage_seg_topk)),
                    min_cos=float(coverage_min_cos),
                    min_jaccard=float(coverage_min_jaccard),
                    hits_max=int(max(1, coverage_hits_max)),
                )
                score_mode = "coverage"
                score_main = float(cov.get("coverage_score") or 0.0)
            else:
                # normal-level primary score: direct TF with explicit length normalization.
                score_mode = "direct_tf_len_norm"
                lb_eff = (
                    max(float(NORMAL_LEN_BALANCE_FLOOR), float(len_balance))
                    if g == "normal"
                    else float(len_balance)
                )
                score_main = float(float(cos_tf_raw or 0.0) * lb_eff)

        row = {
            "id1": id1,
            "id2": id2,
            "path1": unit_path(units[id1]),
            "path2": unit_path(units[id2]),
            "han_chars1": int(len_a),
            "han_chars2": int(len_b),
            "len_ratio": float(len_ratio),
            "len_balance": float(len_balance),
            "score_final": None,
            "rank_final": None,
            "score_mode": score_mode,
            "fusion_mode": None,
            "cosine_tf": (None if semantic_only else float(score_main)),
            "cosine_tf_raw": (None if semantic_only or cos_tf_raw is None else float(cos_tf_raw)),
            "score_lex": None,
            "rank_lex": None,
            "semantic_cosine": (None if sem_pair is None else float(sem_pair)),
            "semantic_q": None,
            "rank_semantic": None,
            "semantic_doc_cos": None,
            "semantic_maxsim_bi": None,
            "semantic_mix_lambda": None,
            "semantic_maxsim_applied": 0,
            "score_base_q": None,
            "score_reranker_raw": None,
            "reranker_ref_01": None,
            "rank_rerank_raw": None,
            "rerank_applied": 0,
            "reranker_model": None,
            "coverage_score": (None if cov is None else float(cov.get("coverage_score") or 0.0)),
            "coverage_gate": None,
            "coverage_f1": (None if cov is None else float(cov.get("coverage_f1") or 0.0)),
            "coverage_cov1": (None if cov is None else float(cov.get("coverage_cov1") or 0.0)),
            "coverage_cov2": (None if cov is None else float(cov.get("coverage_cov2") or 0.0)),
            "coverage_q_tf": (None if cov is None else float(cov.get("coverage_q_tf") or 0.0)),
            "coverage_q_tfidf": (None if cov is None or cov.get("coverage_q_tfidf") is None else float(cov.get("coverage_q_tfidf"))),
            "coverage_overlap_chars": (None if cov is None else int(cov.get("coverage_overlap_chars") or 0)),
            "coverage_match_edges": (None if cov is None else int(cov.get("coverage_match_edges") or 0)),
            "coverage_candidate_edges": (None if cov is None else int(cov.get("coverage_candidate_edges") or 0)),
            "coverage_hits": ("" if cov is None else str(cov.get("coverage_hits") or "")),
            "jaccard": (None if jac is None else float(jac)),
            "minhash_jaccard": (None if mh is None else float(mh)),
            "cosine_tfidf": (None if cos_tfidf is None else float(cos_tfidf)),
            "short_id": (None if short_id is None else int(short_id)),
            "long_id": (None if long_id is None else int(long_id)),
            "contain_short_in_long": (None if cont is None else float(cont)),
        }

        best[key] = row

    if mode == "all":
        for i in range(n):
            sem_sims = None
            if use_sem:
                sem_sims = np.dot(sem_emb, sem_emb[i]).astype(np.float32)
                sem_sims[i] = -1.0

            if semantic_only:
                for j in range(i + 1, n):
                    add_pair(i, j, None, float(sem_sims[j]))
            else:
                lex_sims = lex_cos_row(i)
                for j in range(i + 1, n):
                    sem_raw = None if sem_sims is None else float(sem_sims[j])
                    add_pair(i, j, float(lex_sims[j]), sem_raw)
            done_pairs = (i + 1) * (n - 1) - (i * (i + 1)) // 2
            report(done_pairs, total_est, f"computing pairs (all): {done_pairs}/{total_est}")
    else:
        if candidate_source == "hybrid":
            k_eff = min(topk + sem_k_hint, max(0, n - 1))
        else:
            k_eff = min(topk, max(0, n - 1))

        for i in range(n):
            lex_sims = None
            if not semantic_only:
                lex_sims = lex_cos_row(i)
                lex_sims[i] = -1.0
            sem_sims = None
            if use_sem:
                sem_sims = np.dot(sem_emb, sem_emb[i]).astype(np.float32)
                sem_sims[i] = -1.0

            # Candidate generation source
            if candidate_source == "semantic":
                work_sims = sem_sims
            else:
                work_sims = lex_sims

            if work_sims is None:
                raise RuntimeError("Candidate generation similarities are unavailable.")

            if topk >= n - 1:
                idxs = np.argsort(-work_sims)
            else:
                idxs = np.argpartition(-work_sims, topk)[:topk]
                idxs = idxs[np.argsort(-work_sims[idxs])]
            cand_set = {int(j) for j in idxs if int(j) != i}
            if candidate_source == "hybrid":
                for j in sem_neighbors.get(i, []):
                    j = int(j)
                    if j != i:
                        cand_set.add(j)

            sort_sims = sem_sims if (candidate_source == "semantic" and sem_sims is not None) else lex_sims
            idxs_all = sorted(list(cand_set), key=lambda jj: float(sort_sims[jj]), reverse=True)
            for j in idxs_all:
                j = int(j)
                if j == i:
                    continue
                lex_raw = None if lex_sims is None else float(lex_sims[j])
                sem_raw = None if sem_sims is None else float(sem_sims[j])
                add_pair(i, j, lex_raw, sem_raw)
            done_pairs = (i + 1) * k_eff
            report(done_pairs, total_est, f"computing candidates (topk): {done_pairs}/{total_est}")

    rows = list(best.values())
    rows = [r for r in rows if r["id1"] < r["id2"] and r["id1"] != r["id2"]]

    report(
        total_est,
        total_est,
        f"done. unique unordered pairs={len(rows)}; skipped self-like={int(skipped_self_like)}",
    )
    return rows


def _pair_key_from_row(r: dict) -> Optional[Tuple[int, int]]:
    try:
        a = int(r.get("id1"))
        b = int(r.get("id2"))
    except Exception:
        return None
    if a == b:
        return None
    return (a, b) if a < b else (b, a)


def join_channel_rows_for_merge(
    lexical_rows: List[dict],
    semantic_rows: List[dict],
) -> List[dict]:
    """Join lexical/semantic channel rows by pair key; no new channel computation."""
    by_pair: Dict[Tuple[int, int], dict] = {}

    for r in lexical_rows or []:
        k = _pair_key_from_row(r)
        if k is None:
            continue
        rr = dict(r)
        rr["has_lexical_channel"] = 1
        rr["has_semantic_channel"] = 0
        by_pair[k] = rr

    for r in semantic_rows or []:
        k = _pair_key_from_row(r)
        if k is None:
            continue
        if k in by_pair:
            cur = by_pair[k]
            cur["has_semantic_channel"] = 1
            cur["semantic_cosine"] = r.get("semantic_cosine")
            if not cur.get("path1"):
                cur["path1"] = r.get("path1")
            if not cur.get("path2"):
                cur["path2"] = r.get("path2")
            if cur.get("han_chars1") is None:
                cur["han_chars1"] = r.get("han_chars1")
            if cur.get("han_chars2") is None:
                cur["han_chars2"] = r.get("han_chars2")
            if cur.get("len_ratio") is None:
                cur["len_ratio"] = r.get("len_ratio")
            if cur.get("len_balance") is None:
                cur["len_balance"] = r.get("len_balance")
        else:
            rr = dict(r)
            rr["has_lexical_channel"] = 0
            rr["has_semantic_channel"] = 1
            # Non-destructive join: keep semantic row's original field values
            # (currently all None for lexical fields).
            # Lexical fields will be filled by complete_merge_dual_scores().
            by_pair[k] = rr

    out = list(by_pair.values())
    out.sort(key=lambda r: (_pair_key_from_row(r) or (10**12, 10**12)))
    return out


#
# =========================
# Word styling + Similarity analysis report
# =========================

def _apply_bnu_word_styles(doc: Document) -> None:
    """Apply required Word built-in styles: Heading 1/2/3 + Normal."""
    # Heading 1: 宋体 小三 加粗 居中
    s1 = doc.styles["Heading 1"]
    s1.font.name = "宋体"
    s1._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    s1.font.size = Pt(16)  # 小三 approx
    s1.font.bold = True
    s1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s1.paragraph_format.space_before = Pt(0)
    s1.paragraph_format.space_after = Pt(0)
    s1.paragraph_format.line_spacing = 1.0

    # Heading 2: 宋体 四号 加粗 居中
    s2 = doc.styles["Heading 2"]
    s2.font.name = "宋体"
    s2._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    s2.font.size = Pt(14)  # 四号 approx
    s2.font.bold = True
    s2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s2.paragraph_format.space_before = Pt(0)
    s2.paragraph_format.space_after = Pt(0)
    s2.paragraph_format.line_spacing = 1.0

    # Heading 3: 宋体 小四 加粗 居中
    s3 = doc.styles["Heading 3"]
    s3.font.name = "宋体"
    s3._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    s3.font.size = Pt(12)  # 小四 approx
    s3.font.bold = True
    s3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s3.paragraph_format.space_before = Pt(0)
    s3.paragraph_format.space_after = Pt(0)
    s3.paragraph_format.line_spacing = 1.0

    # Normal: 中文宋体五号 + 英文 Times New Roman；段前段后0；单倍
    sn = doc.styles["Normal"]
    sn.font.name = "Times New Roman"
    sn._element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')
    sn._element.rPr.rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    sn._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    sn.font.size = Pt(10.5)  # 五号
    sn.font.bold = False
    sn.paragraph_format.space_before = Pt(0)
    sn.paragraph_format.space_after = Pt(0)
    sn.paragraph_format.line_spacing = 1.0


def _percentile(values: List[float], q: float) -> Optional[float]:
    """q in [0,1]."""
    if not values:
        return None
    v = sorted(values)
    if q <= 0:
        return float(v[0])
    if q >= 1:
        return float(v[-1])
    pos = q * (len(v) - 1)
    lo = int(pos)
    hi = min(len(v) - 1, lo + 1)
    frac = pos - lo
    return float(v[lo] * (1 - frac) + v[hi] * frac)


def _rank_map(values: List[Tuple[int, float]]) -> Dict[int, float]:
    """Return index->rank score in (0,1], higher means better.

    q = (N - rank + 0.5) / N, where rank is 1-based (competition rank with tie averaging).
    """
    if not values:
        return {}
    vals = sorted(values, key=lambda x: x[1], reverse=True)
    n = len(vals)
    if n == 1:
        return {vals[0][0]: 1.0}
    out: Dict[int, float] = {}
    i = 0
    while i < n:
        j = i
        while j + 1 < n and vals[j + 1][1] == vals[i][1]:
            j += 1
        avg_rank = (i + j) / 2.0 + 1.0  # 1-based
        norm = (float(n) - float(avg_rank) + 0.5) / float(n)
        for k in range(i, j + 1):
            out[int(vals[k][0])] = float(norm)
        i = j + 1
    return out


def _value_open_map(values: List[Tuple[int, float]], eps: float = SCORE_NORM_EPS) -> Dict[int, float]:
    """Return index->value score in (0,1) from min-max scaling (order-preserving)."""
    if not values:
        return {}
    vmin = min(float(v) for _, v in values)
    vmax = max(float(v) for _, v in values)
    lo = max(0.0, float(eps))
    hi = max(lo, 1.0 - lo)
    if vmax <= vmin:
        mid = 0.5 * (lo + hi)
        return {int(i): float(mid) for i, _ in values}
    span = float(vmax - vmin)
    out: Dict[int, float] = {}
    for i, v in values:
        z = (float(v) - vmin) / span  # [0,1]
        out[int(i)] = float(lo + z * (hi - lo))  # (eps, 1-eps)
    return out


def _flat_open_map(
    values: List[Tuple[int, float]],
    value: float = 0.5,
    eps: float = SCORE_NORM_EPS,
) -> Dict[int, float]:
    """Return a constant open-interval score map for all indices."""
    if not values:
        return {}
    lo = max(0.0, float(eps))
    hi = max(lo, 1.0 - lo)
    x = float(value)
    if x < lo:
        x = lo
    elif x > hi:
        x = hi
    return {int(i): float(x) for i, _ in values}


def _merge_quantile_bounds(granularity: str) -> Tuple[float, float]:
    """Return robust quantile endpoints for merge normalization by granularity."""
    g = normalize_style(granularity)
    lo_q, hi_q = MERGE_ROBUST_QUANTILES.get(g, MERGE_ROBUST_QUANTILES["default"])
    try:
        lo = float(lo_q)
    except Exception:
        lo = 0.05
    try:
        hi = float(hi_q)
    except Exception:
        hi = 0.95
    lo = min(0.99, max(0.0, lo))
    hi = min(1.0, max(0.01, hi))
    if hi <= lo:
        lo, hi = 0.05, 0.95
    return float(lo), float(hi)


def _value_robust_map(
    values: List[Tuple[int, float]],
    lo_q: float,
    hi_q: float,
    eps: float = SCORE_NORM_EPS,
    hard_clip: bool = True,
    tail_slope: float = 1.0,
    softclip_k: float = 0.0,
) -> Dict[int, float]:
    """Return robust quantile min-max map with hard-clip / soft-clip / linear tails."""
    if not values:
        return {}
    if len(values) < int(max(1, MERGE_ROBUST_MIN_SAMPLES)):
        # Too few samples: robust endpoints are unstable; use neutral score.
        return _flat_open_map(values, value=0.5, eps=eps)
    nums = [float(v) for _, v in values]
    p_lo = _percentile(nums, lo_q)
    p_hi = _percentile(nums, hi_q)
    if p_lo is None or p_hi is None or float(p_hi) <= float(p_lo):
        # Endpoint collapse / invalid span: avoid divide-by-zero and extreme mapping.
        return _flat_open_map(values, value=0.5, eps=eps)

    lo = max(0.0, float(eps))
    hi = max(lo, 1.0 - lo)
    span = float(p_hi - p_lo)
    ts = max(0.0, float(tail_slope))
    try:
        k_soft = float(softclip_k)
    except Exception:
        k_soft = 0.0
    use_softclip = (not hard_clip) and (k_soft > 0.0)
    out: Dict[int, float] = {}
    for i, v in values:
        z = (float(v) - float(p_lo)) / span
        if hard_clip:
            if z < 0.0:
                z = 0.0
            elif z > 1.0:
                z = 1.0
        elif use_softclip:
            # Piecewise tanh soft-clip: linear in [0,1], monotonic bounded tails.
            if z < 0.0:
                z = -math.tanh(-z * k_soft)
            elif z > 1.0:
                z = 1.0 + math.tanh((z - 1.0) * k_soft)
            # z_out in (-1,2): map to (eps, 1-eps) via (z+1)/3.
        else:
            if z < 0.0:
                z = z * ts
            elif z > 1.0:
                z = 1.0 + (z - 1.0) * ts
        if use_softclip:
            mapped = float(lo + ((z + 1.0) / 3.0) * (hi - lo))
        else:
            mapped = float(lo + z * (hi - lo))
        if not np.isfinite(mapped):
            mapped = float(0.5 * (lo + hi))
        out[int(i)] = mapped
    return out


def _linear01(v: float, lo: float, hi: float) -> float:
    """Map v to [0,1] with linear ramp between lo and hi."""
    try:
        x = float(v)
    except Exception:
        x = 0.0
    if hi <= lo:
        return 1.0 if x >= hi else 0.0
    if x <= lo:
        return 0.0
    if x >= hi:
        return 1.0
    return float((x - lo) / (hi - lo))


def _jaccard_adjust_for_pair(jacc_raw: float, row: dict) -> Tuple[float, float, float]:
    """Return (jacc_eff, len_factor, weight_gain) for a pair."""
    try:
        len1 = int(row.get("han_chars1") or 0)
    except Exception:
        len1 = 0
    try:
        len2 = int(row.get("han_chars2") or 0)
    except Exception:
        len2 = 0

    if len1 > 0 and len2 > 0:
        short_len = min(len1, len2)
    else:
        short_len = max(len1, len2)

    lo = float(JACCARD_LEN_ADJUST["short_len_lo"])
    hi = float(JACCARD_LEN_ADJUST["short_len_hi"])
    len_factor = _linear01(float(short_len), lo, hi)

    j = max(0.0, min(1.0, float(jacc_raw)))
    gamma_long = float(JACCARD_LEN_ADJUST["gamma_long"])
    gamma_long = min(1.0, max(0.05, gamma_long))
    gamma = float(1.0 - (1.0 - gamma_long) * len_factor)
    jacc_eff = float(pow(j, gamma))

    weight_boost_long = max(0.0, float(JACCARD_LEN_ADJUST["weight_boost_long"]))
    weight_gain = float(1.0 + weight_boost_long * len_factor)
    return jacc_eff, len_factor, weight_gain


def _semantic_q_map(values: List[Tuple[int, float]], granularity: str) -> Dict[int, float]:
    """Semantic normalization map in (0,1): value-based only, same depth as lexical."""
    _ = granularity
    return _value_open_map(values, eps=SCORE_NORM_EPS)


def _normalize_weights(raw: Dict[str, float], fallback: Dict[str, float]) -> Dict[str, float]:
    out: Dict[str, float] = {}
    for k in ["cosine_tf_raw", "coverage", "cosine_tfidf", "jaccard"]:
        try:
            v = float(raw.get(k, 0.0))
        except Exception:
            v = 0.0
        out[k] = max(0.0, v)
    s = sum(out.values())
    if s <= 0:
        return dict(fallback)
    return {k: float(v / s) for k, v in out.items()}


def _rank_order(values: List[Tuple[int, float]], descending: bool = True) -> Dict[int, int]:
    """Return index -> 1-based competition rank."""
    if not values:
        return {}
    vals = sorted(values, key=lambda x: x[1], reverse=descending)
    out: Dict[int, int] = {}
    rank = 0
    prev = None
    for pos, (idx, v) in enumerate(vals, start=1):
        if prev is None or v != prev:
            rank = pos
        out[int(idx)] = int(rank)
        prev = v
    return out


def _rrf_from_rank(rank: Optional[int], k: int) -> float:
    if rank is None or rank <= 0:
        return 0.0
    return float(1.0 / float(int(max(1, k)) + int(rank)))


def annotate_semantic_only_scores(rows: List[dict], granularity: str = "normal") -> None:
    """Semantic-only scoring: raw cosine -> shared value normalization -> final rank."""
    if not rows:
        return

    sem_vals: List[Tuple[int, float]] = []
    for i, r in enumerate(rows):
        sv = r.get("semantic_cosine")
        try:
            svf = float(sv) if sv is not None else 0.0
        except Exception:
            svf = 0.0
        svf = max(0.0, min(1.0, svf))
        sem_vals.append((i, svf))

    q_sem = _semantic_q_map(sem_vals, granularity)
    rank_sem = _rank_order(sem_vals, descending=True)

    for i, r in enumerate(rows):
        qs = float(q_sem.get(i, 0.0))
        r["score_lex"] = None
        r["rank_lex"] = None
        r["semantic_q"] = qs
        r["rank_semantic"] = int(rank_sem.get(i, 0)) if i in rank_sem else None
        r["score_fused"] = qs
        r["score_final"] = qs
        r["fusion_mode"] = "semantic_only_norm"

    order = sorted(
        range(len(rows)),
        key=lambda idx: float(rows[idx].get("score_final") or 0.0),
        reverse=True,
    )
    rank = 0
    prev = None
    for pos, idx in enumerate(order, start=1):
        cur = float(rows[idx].get("score_final") or 0.0)
        if prev is None or cur != prev:
            rank = pos
        rows[idx]["rank_final"] = int(rank)
        prev = cur


def annotate_lexical_only_scores(
    rows: List[dict],
    granularity: str,
    weights_h: Optional[Dict[str, float]] = None,
    weights_n: Optional[Dict[str, float]] = None,
) -> None:
    """Set final ranking by lexical composite score only."""
    if not rows:
        return

    annotate_composite_scores(rows, granularity, weights_h=weights_h, weights_n=weights_n)
    lex_vals: List[Tuple[int, float]] = []
    for i, r in enumerate(rows):
        try:
            lv = float(r.get("score_final") or 0.0)
        except Exception:
            lv = 0.0
        lex_vals.append((i, lv))
    rank_lex = _rank_order(lex_vals, descending=True)

    for i, r in enumerate(rows):
        r["score_lex"] = float(r.get("score_final") or 0.0)
        r["rank_lex"] = int(rank_lex.get(i, 0)) if i in rank_lex else None
        r["semantic_q"] = None
        r["rank_semantic"] = None
        r["score_fused"] = float(r.get("score_final") or 0.0)
        r["fusion_mode"] = "lexical_only"


def annotate_fused_scores(
    rows: List[dict],
    granularity: str,
    semantic_alpha_h: float = 0.15,
    semantic_alpha_n: float = 0.10,
    rrf_k_h1: int = 60,
    use_semantic: bool = True,
) -> None:
    """Three-stage fusion:

    1) lexical score from existing score_final (or cosine_tf fallback)
    2) semantic score from embedding cosine
    3) fusion: linear blend for h1/h2/normal in normalized value space.
    """
    if not rows:
        return

    g = normalize_style(granularity)
    _ = rrf_k_h1

    # stage-1 lexical base
    lex_vals: List[Tuple[int, float]] = []
    sem_vals: List[Tuple[int, float]] = []
    for i, r in enumerate(rows):
        lv = r.get("score_final")
        if lv is None:
            lv = r.get("cosine_tf")
        try:
            lvf = float(lv)
        except Exception:
            lvf = 0.0
        r["score_lex"] = float(lvf)
        lex_vals.append((i, lvf))

        sv = r.get("semantic_cosine")
        if sv is None:
            continue
        try:
            svf = float(sv)
        except Exception:
            continue
        svf = max(0.0, min(1.0, svf))
        sem_vals.append((i, svf))

    q_lex = _value_open_map(lex_vals, eps=SCORE_NORM_EPS)
    rank_lex = _rank_order(lex_vals, descending=True)
    q_sem = _semantic_q_map(sem_vals, g) if (use_semantic and sem_vals) else {}
    rank_sem = _rank_order(sem_vals, descending=True) if (use_semantic and sem_vals) else {}

    # Semantic alpha per layer
    alpha = float(semantic_alpha_h if g in ("h1", "h2") else semantic_alpha_n)
    alpha = min(1.0, max(0.0, alpha))

    for i, r in enumerate(rows):
        r["rank_lex"] = int(rank_lex.get(i, 0)) if i in rank_lex else None
        r["semantic_q"] = float(q_sem.get(i, 0.0)) if q_sem else 0.0
        r["rank_semantic"] = int(rank_sem.get(i, 0)) if i in rank_sem else None

        if not use_semantic or not q_sem:
            fused = float(q_lex.get(i, 0.0))
            r["fusion_mode"] = "lex_only"
        else:
            fused = float((1.0 - alpha) * float(q_lex.get(i, 0.0)) + alpha * float(q_sem.get(i, 0.0)))
            r["fusion_mode"] = "linear_norm"

        r["score_fused"] = float(fused)

    # Assign final score/rank from fused score (descending)
    order = sorted(
        range(len(rows)),
        key=lambda idx: float(rows[idx].get("score_fused") or 0.0),
        reverse=True,
    )
    rank = 0
    prev = None
    for pos, idx in enumerate(order, start=1):
        cur = float(rows[idx].get("score_fused") or 0.0)
        if prev is None or cur != prev:
            rank = pos
        rows[idx]["score_final"] = cur
        rows[idx]["rank_final"] = int(rank)
        prev = cur


def _has_lex_signal(r: dict, g: str) -> bool:
    """Check whether a row has any usable lexical signal for scoring.

    Used by both ``complete_merge_dual_scores`` (missing-lex counting) and
    ``annotate_composite_scores`` (scoring gate) to guarantee identical logic.
    """
    if g in ("h1", "h2"):
        cos_v = r.get("cosine_tf_raw")
    else:
        cos_v = r.get("cosine_tf")
        if cos_v is None:
            cos_v = r.get("cosine_tf_raw")
    return (
        cos_v is not None
        or r.get("cosine_tfidf") is not None
        or r.get("jaccard") is not None
        or r.get("coverage_score") is not None
    )


def annotate_composite_scores(
    rows: List[dict],
    granularity: str,
    weights_h: Optional[Dict[str, float]] = None,
    weights_n: Optional[Dict[str, float]] = None,
) -> None:
    """Lexical scoring in raw space (no internal normalization)."""
    if not rows:
        return

    g = normalize_style(granularity)
    if g in ("h1", "h2"):
        w = _normalize_weights(weights_h or LEXICAL_WEIGHTS_DEFAULT_H, LEXICAL_WEIGHTS_DEFAULT_H)
    else:
        w = _normalize_weights(weights_n or LEXICAL_WEIGHTS_DEFAULT_N, LEXICAL_WEIGHTS_DEFAULT_N)
    if g not in ("h1", "h2"):
        # Coverage is undefined at normal granularity in this pipeline.
        w["coverage"] = 0.0
        s = float(sum(w.values()))
        if s > 0.0:
            w = {k: float(v / s) for k, v in w.items()}

    def _clip01(v: Any) -> float:
        try:
            x = float(v)
        except Exception:
            x = 0.0
        return max(0.0, min(1.0, x))

    lex_vals: List[Tuple[int, float]] = []
    has_signal_flags: List[bool] = [False] * len(rows)

    for i, r in enumerate(rows):
        # H1/H2 keep using raw lexical cosine; normal uses length-normalized cosine
        # prepared in compute_similarities (cosine_tf = cos_tf_raw * lb_eff).
        if g in ("h1", "h2"):
            cos_feature_val = r.get("cosine_tf_raw")
        else:
            cos_feature_val = r.get("cosine_tf")
            if cos_feature_val is None:
                cos_feature_val = r.get("cosine_tf_raw")
        # Gate by value existence (not by has_lexical_channel flag).
        # has_lexical_channel is kept as an audit marker but no longer gates scoring.
        has_lex_signal = _has_lex_signal(r, g)
        has_signal_flags[i] = bool(has_lex_signal)
        if not has_lex_signal:
            r["q_cov_star"] = 0.0
            r["q_cosine_tf_raw"] = 0.0
            r["q_cosine_tfidf"] = 0.0
            r["q_jaccard"] = 0.0
            r["jaccard_effective"] = 0.0
            r["jaccard_len_factor"] = 0.0
            r["jaccard_weight_effective"] = 0.0
            r["coverage_gate"] = 0.0
            r["score_lex_raw"] = None
            r["score_final"] = 0.0
            continue

        cov_raw = _clip01(r.get("coverage_score"))
        cos_raw = _clip01(cos_feature_val)
        tfidf_raw = _clip01(r.get("cosine_tfidf"))
        jacc_raw = _clip01(r.get("jaccard"))
        jacc_eff, jacc_len_factor, jacc_weight_gain = _jaccard_adjust_for_pair(jacc_raw, r)

        # Keep legacy q_* fields for report compatibility; values are raw-space here.
        r["q_cov_star"] = cov_raw
        r["q_cosine_tf_raw"] = cos_raw
        r["q_cosine_tfidf"] = tfidf_raw
        r["q_jaccard"] = jacc_eff
        r["jaccard_effective"] = jacc_eff
        r["jaccard_len_factor"] = jacc_len_factor

        s_raw = 0.0
        wsum = 0.0
        mw_cos = float(w.get("cosine_tf_raw", 0.0))
        mw_cov = float(w.get("coverage", 0.0))
        mw_tfidf = float(w.get("cosine_tfidf", 0.0))
        mw_j = float(w.get("jaccard", 0.0))
        mw_j_eff = float(mw_j * jacc_weight_gain)
        r["jaccard_weight_effective"] = mw_j_eff

        s_raw += mw_cos * cos_raw
        wsum += mw_cos
        s_raw += mw_cov * cov_raw
        wsum += mw_cov
        s_raw += mw_tfidf * tfidf_raw
        wsum += mw_tfidf
        s_raw += mw_j_eff * jacc_eff
        wsum += mw_j_eff

        # Keep lexical score as pure explicit weighted sum (no multiplicative gate).
        r["coverage_gate"] = 1.0
        score_lex_raw = float(s_raw / wsum) if wsum > 0 else 0.0
        r["score_lex_raw"] = score_lex_raw
        lex_vals.append((i, score_lex_raw))

    # No lexical internal normalization; score_final is raw lexical composite.
    for i, r in enumerate(rows):
        if not has_signal_flags[i]:
            r["score_final"] = 0.0
        else:
            r["score_final"] = float(r.get("score_lex_raw") or 0.0)

    # Assign rank_final (1-based). Ties share rank (competition ranking).
    order = sorted(
        range(len(rows)),
        key=lambda idx: float(rows[idx].get("score_final") or 0.0),
        reverse=True,
    )
    rank = 0
    prev = None
    for pos, idx in enumerate(order, start=1):
        cur = float(rows[idx].get("score_final") or 0.0)
        if prev is None or cur != prev:
            rank = pos
        rows[idx]["rank_final"] = int(rank)
        prev = cur


# ---------------------------------------------------------------------------
# Semantic cosine backfill for merge channel
# ---------------------------------------------------------------------------
# Problem: pairs recalled only by the lexical topk have semantic_cosine=None,
# which annotate_merge_scores treats as 0.0 → ~66.7% of pairs get a floor
# semantic_q ≈ 0.333.  This helper computes the true cosine from sem_emb
# (L2-normalised [N, D] np.ndarray) and fills it in BEFORE scoring.
# ---------------------------------------------------------------------------

def backfill_semantic_cosine(
    rows: List[dict],
    sem_emb: Optional[np.ndarray],
) -> Dict[str, float]:
    """Backfill missing semantic_cosine from sem_emb and return diagnostics.

    Returns a dict with keys:
        total, missing_before, filled, missing_after,
        native_count, backfilled_count,
        mean_native, std_native,
        mean_backfilled, std_backfilled,
        mean_all_after, std_all_after.
    """
    out: Dict[str, float] = {
        "total": int(len(rows or [])),
        "missing_before": 0,
        "filled": 0,
        "missing_after": 0,
        "native_count": 0,
        "backfilled_count": 0,
        "mean_native": 0.0,
        "std_native": 0.0,
        "mean_backfilled": 0.0,
        "std_backfilled": 0.0,
        "mean_all_after": 0.0,
        "std_all_after": 0.0,
    }
    if not rows:
        return out

    # sem_emb expected: np.ndarray of shape [N, D], L2-normalised
    if sem_emb is None or not isinstance(sem_emb, np.ndarray) or sem_emb.ndim < 2:
        # still stamp semantic_source for audit consistency
        for r in rows:
            if r.get("semantic_cosine") is not None and not r.get("semantic_source"):
                r["semantic_source"] = "native"
        out["missing_before"] = int(sum(1 for r in rows if r.get("semantic_cosine") is None))
        out["missing_after"] = out["missing_before"]
        return out

    n = int(sem_emb.shape[0])
    native_vals: List[float] = []
    fill_vals: List[float] = []

    for r in rows:
        sv = r.get("semantic_cosine")
        if sv is None:
            out["missing_before"] += 1
            # --- attempt backfill ---
            try:
                id1 = int(r.get("id1"))
                id2 = int(r.get("id2"))
            except Exception:
                continue
            if id1 < 0 or id2 < 0 or id1 >= n or id2 >= n or id1 == id2:
                continue
            cos_val = float(np.dot(sem_emb[id1], sem_emb[id2]))
            if not np.isfinite(cos_val):
                continue
            cos_val = max(0.0, min(1.0, cos_val))
            r["semantic_cosine"] = cos_val
            if r.get("cross_sem_score") is None:
                r["cross_sem_score"] = cos_val
            r["semantic_source"] = "backfill"
            fill_vals.append(cos_val)
            out["filled"] += 1
        else:
            # native value — stamp audit label
            try:
                svf = float(sv)
            except Exception:
                svf = None
            if svf is not None and np.isfinite(svf):
                svf = max(0.0, min(1.0, svf))
                native_vals.append(svf)
            r.setdefault("semantic_source", "native")

    # --- second pass: compute all_after stats and missing_after ---
    all_after_vals: List[float] = []
    for r in rows:
        sv = r.get("semantic_cosine")
        if sv is None:
            out["missing_after"] += 1
            continue
        try:
            svf = float(sv)
        except Exception:
            out["missing_after"] += 1
            continue
        if not np.isfinite(svf):
            out["missing_after"] += 1
            continue
        all_after_vals.append(max(0.0, min(1.0, svf)))

    out["native_count"] = int(len(native_vals))
    out["backfilled_count"] = int(len(fill_vals))
    if native_vals:
        arr = np.asarray(native_vals, dtype=np.float64)
        out["mean_native"] = float(arr.mean())
        out["std_native"] = float(arr.std(ddof=0))
    if fill_vals:
        arr = np.asarray(fill_vals, dtype=np.float64)
        out["mean_backfilled"] = float(arr.mean())
        out["std_backfilled"] = float(arr.std(ddof=0))
    if all_after_vals:
        arr = np.asarray(all_after_vals, dtype=np.float64)
        out["mean_all_after"] = float(arr.mean())
        out["std_all_after"] = float(arr.std(ddof=0))
    return out


def complete_merge_dual_scores(
    rows: List[dict],
    reps: dict,
    sem_emb: Optional[np.ndarray],
    units: List,
    granularity: str,
    normal_units: Optional[List] = None,
    ngram_n: int = 3,
    coverage_cfg: Optional[dict] = None,
) -> Dict[str, int]:
    """Fill in missing lexical/semantic fields for merge-channel rows.

    After ``join_channel_rows_for_merge`` produces a union of lexical-only and
    semantic-only pairs, this function computes any missing scores so that every
    row has complete dual-channel data.

    Principle: only fill ``None`` fields — never overwrite existing values.
    Exception: ``score_mode`` is allowed to be overwritten (audit/display field)
    when lexical fields are filled, to reflect the final scoring semantics.

    Returns a stats dict with keys: total, filled_lex, filled_sem,
    missing_lex, missing_sem, missing_lex_fields.
    """
    import time as _time
    t0 = _time.perf_counter()

    g = normalize_style(granularity) if callable(normalize_style) else str(granularity).strip().lower()
    cfg = coverage_cfg or COVERAGE_DEFAULTS

    out: Dict[str, int] = {
        "total": len(rows),
        "filled_lex": 0,
        "filled_sem": 0,
        "missing_lex": 0,
        "missing_sem": 0,
        "missing_lex_fields": 0,
    }
    if not rows:
        return out

    # ── build lexical caches (same matrices used in compute_similarities) ──
    X_tf = reps.get("X_tf")
    X_tfidf = reps.get("X_tfidf")
    S_set = reps.get("S_set")
    mh_sigs = reps.get("mh_sigs")

    X_tf_norm = None
    if X_tf is not None:
        X_tf_norm = sk_normalize(X_tf, norm="l2", copy=True)

    # ── build coverage context for H1/H2 ──
    coverage_ctx = None
    if g in ("h1", "h2"):
        if normal_units is None:
            raise RuntimeError(
                f"complete_merge_dual_scores: H1/H2 ({g}) requires normal_units "
                f"for coverage (weight=35%), but normal_units is None."
            )
        # coverage_ctx construction must succeed for H1/H2;
        # failure would introduce systematic bias (coverage weight = 35%).
        coverage_ctx = _build_coverage_context(units, normal_units, reps, ngram_n)

    # ── semantic embedding validation ──
    sem_ok = (
        sem_emb is not None
        and isinstance(sem_emb, np.ndarray)
        and sem_emb.ndim == 2
    )

    # ── iterate rows ──
    for r in rows:
        # shared id boundary check (backfill L3541-style)
        try:
            id1 = int(r.get("id1"))
            id2 = int(r.get("id2"))
        except Exception:
            continue
        if id1 < 0 or id2 < 0 or id1 == id2:
            continue

        filled_any_lex = False
        filled_any_sem = False

        # --- (a) semantic_cosine ---
        if r.get("semantic_cosine") is None and sem_ok:
            if sem_emb.shape[0] > max(id1, id2):
                cos_val = float(np.dot(sem_emb[id1], sem_emb[id2]))
                if np.isfinite(cos_val):
                    cos_val = max(0.0, min(1.0, cos_val))
                    r["semantic_cosine"] = cos_val
                    r["semantic_source"] = "completion"
                    if r.get("cross_sem_score") is None:
                        r["cross_sem_score"] = cos_val
                    filled_any_sem = True

        # --- (b) cosine_tf_raw ---
        if r.get("cosine_tf_raw") is None and X_tf_norm is not None:
            if X_tf_norm.shape[0] > max(id1, id2):
                val = float(X_tf_norm[id1].multiply(X_tf_norm[id2]).sum())
                r["cosine_tf_raw"] = val
                filled_any_lex = True

        # --- (c) cosine_tfidf ---
        if r.get("cosine_tfidf") is None and X_tfidf is not None:
            if X_tfidf.shape[0] > max(id1, id2):
                val = float(X_tfidf[id1].multiply(X_tfidf[id2]).sum())
                r["cosine_tfidf"] = val
                filled_any_lex = True

        # --- (d) jaccard ---
        if r.get("jaccard") is None and S_set is not None:
            if len(S_set) > max(id1, id2):
                sa, sb = S_set[id1], S_set[id2]
                union_len = len(sa | sb)
                r["jaccard"] = float(len(sa & sb) / max(1, union_len))
                filled_any_lex = True

        # --- (e) minhash_jaccard ---
        if r.get("minhash_jaccard") is None and mh_sigs is not None:
            if len(mh_sigs) > max(id1, id2):
                mh_val = minhash_jaccard(mh_sigs[id1], mh_sigs[id2])
                r["minhash_jaccard"] = mh_val
                filled_any_lex = True

        # --- (f) coverage (H1/H2, must precede cosine_tf) ---
        if g in ("h1", "h2") and coverage_ctx is not None:
            need_cov = (
                r.get("coverage_score") is None
                or r.get("coverage_f1") is None
                or r.get("coverage_cov1") is None
            )
            if need_cov:
                cov = _coverage_pair_score(
                    id1=id1,
                    id2=id2,
                    ctx=coverage_ctx,
                    seg_topk=int(cfg.get("seg_topk", 10)),
                    min_cos=float(cfg.get("min_cos", 0.20)),
                    min_jaccard=float(cfg.get("min_jaccard", 0.02)),
                    hits_max=int(cfg.get("hits_max", 6)),
                )
                r["coverage_score"] = float(cov.get("coverage_score") or 0.0)
                r["coverage_f1"] = float(cov.get("coverage_f1") or 0.0)
                r["coverage_cov1"] = float(cov.get("coverage_cov1") or 0.0)
                r["coverage_cov2"] = float(cov.get("coverage_cov2") or 0.0)
                r["coverage_q_tf"] = float(cov.get("coverage_q_tf") or 0.0)
                r["coverage_q_tfidf"] = (
                    None if cov.get("coverage_q_tfidf") is None
                    else float(cov.get("coverage_q_tfidf"))
                )
                r["coverage_overlap_chars"] = int(cov.get("coverage_overlap_chars") or 0)
                r["coverage_match_edges"] = int(cov.get("coverage_match_edges") or 0)
                r["coverage_candidate_edges"] = int(cov.get("coverage_candidate_edges") or 0)
                r["coverage_hits"] = str(cov.get("coverage_hits") or "")
                filled_any_lex = True

        # --- (g) cosine_tf (last — depends on coverage & cosine_tf_raw) ---
        if r.get("cosine_tf") is None:
            ctf_raw = r.get("cosine_tf_raw")
            if ctf_raw is not None:
                ctf_raw = float(ctf_raw)
                if g in ("h1", "h2"):
                    cov_s = r.get("coverage_score")
                    if cov_s is not None:
                        r["cosine_tf"] = float(cov_s)
                    else:
                        lb = float(r.get("len_balance") or 0.0)
                        r["cosine_tf"] = ctf_raw * lb
                else:
                    # Normal: length-normalized
                    lb = float(r.get("len_balance") or 0.0)
                    lb_eff = max(float(NORMAL_LEN_BALANCE_FLOOR), lb)
                    r["cosine_tf"] = ctf_raw * lb_eff
                filled_any_lex = True

        # --- audit marks ---
        if filled_any_lex and filled_any_sem:
            r["completion_source"] = "both"
        elif filled_any_lex:
            r["completion_source"] = "lex"
        elif filled_any_sem:
            r["completion_source"] = "sem"

        # score_mode: overwrite allowed for audit/display (exception to "don't overwrite")
        if filled_any_lex:
            if g in ("h1", "h2") and r.get("coverage_score") is not None:
                r["score_mode"] = "coverage"
            else:
                r["score_mode"] = "direct_tf_len_norm"

        if filled_any_lex:
            out["filled_lex"] += 1
        if filled_any_sem:
            out["filled_sem"] += 1

    # ── missing counts ──
    for r in rows:
        if r.get("semantic_cosine") is None:
            out["missing_sem"] += 1
        if not _has_lex_signal(r, g):
            out["missing_lex"] += 1
        # field-level completeness: core required fields only.
        # cosine_tfidf and minhash_jaccard are excluded because they depend
        # on optional components (use_tfidf flag / MinHash library) and may
        # legitimately be None.
        field_missing = (
            r.get("cosine_tf_raw") is None
            or r.get("cosine_tf") is None
            or r.get("jaccard") is None
        )
        if g in ("h1", "h2"):
            field_missing = field_missing or r.get("coverage_score") is None
        if field_missing:
            out["missing_lex_fields"] += 1

    out["elapsed_s"] = round(_time.perf_counter() - t0, 3)
    return out


def annotate_merge_scores(
    rows: List[dict],
    granularity: str,
    weights_h: Optional[Dict[str, float]] = None,
    weights_n: Optional[Dict[str, float]] = None,
    w_lex: float = 0.70,
    w_sem: float = 0.30,
    rrf_k_h1: int = 30,
) -> None:
    """Merge with a single normalization stage at fusion only."""
    if not rows:
        return

    g = normalize_style(granularity)
    _ = rrf_k_h1  # kept for backward-compatible signature; not used in linear mode.

    # Build lexical normalized score from raw lexical metrics.
    annotate_composite_scores(rows, g, weights_h=weights_h, weights_n=weights_n)

    lex_vals: List[Tuple[int, float]] = []
    sem_vals: List[Tuple[int, float]] = []
    for i, r in enumerate(rows):
        try:
            lv = float(r.get("score_lex_raw") if r.get("score_lex_raw") is not None else 0.0)
        except Exception:
            lv = 0.0
        lex_vals.append((i, lv))
        sv = r.get("semantic_cosine")
        try:
            svf = float(sv) if sv is not None else 0.0
        except Exception:
            svf = 0.0
        svf = max(0.0, min(1.0, svf))
        sem_vals.append((i, svf))

    # Single normalization stage at merge only.
    q_lo, q_hi = _merge_quantile_bounds(g)
    norm_method = str(MERGE_NORM_METHOD or "").strip().lower()
    if norm_method == "minmax":
        q_lex = _value_open_map(lex_vals, eps=SCORE_NORM_EPS)
        q_sem = _value_open_map(sem_vals, eps=SCORE_NORM_EPS)
    elif norm_method == "robust_hard":
        q_lex = _value_robust_map(lex_vals, lo_q=q_lo, hi_q=q_hi, eps=SCORE_NORM_EPS, hard_clip=True)
        q_sem = _value_robust_map(sem_vals, lo_q=q_lo, hi_q=q_hi, eps=SCORE_NORM_EPS, hard_clip=True)
    elif norm_method == "robust_soft":
        sc_k = _merge_softclip_k(g)
        q_lex = _value_robust_map(
            lex_vals,
            lo_q=q_lo,
            hi_q=q_hi,
            eps=SCORE_NORM_EPS,
            hard_clip=False,
            softclip_k=sc_k,
        )
        q_sem = _value_robust_map(
            sem_vals,
            lo_q=q_lo,
            hi_q=q_hi,
            eps=SCORE_NORM_EPS,
            hard_clip=False,
            softclip_k=sc_k,
        )
    else:
        ts = _merge_tail_slope(g)
        q_lex = _value_robust_map(
            lex_vals,
            lo_q=q_lo,
            hi_q=q_hi,
            eps=SCORE_NORM_EPS,
            hard_clip=False,
            tail_slope=ts,
        )
        q_sem = _value_robust_map(
            sem_vals,
            lo_q=q_lo,
            hi_q=q_hi,
            eps=SCORE_NORM_EPS,
            hard_clip=False,
            tail_slope=ts,
        )
    rank_lex = _rank_order(lex_vals, descending=True)
    rank_sem = _rank_order(sem_vals, descending=True)

    wl = float(max(0.0, w_lex))
    ws = float(max(0.0, w_sem))
    if wl + ws <= 0:
        wl, ws = 0.7, 0.3
    den = wl + ws
    wl_n = wl / den
    ws_n = ws / den

    for i, r in enumerate(rows):
        ql = float(q_lex.get(i, 0.0))
        qs = float(q_sem.get(i, 0.0))
        merged = float(wl_n * ql + ws_n * qs)

        r["rank_lex"] = int(rank_lex.get(i, 0)) if i in rank_lex else None
        r["rank_semantic"] = int(rank_sem.get(i, 0)) if i in rank_sem else None
        r["score_lex"] = float(ql)
        r["semantic_q"] = float(qs)
        r["coverage_gate_merge"] = 1.0
        if norm_method == "minmax":
            norm_tag = "minmax"
        elif norm_method == "robust_hard":
            norm_tag = f"robust_hard_q{int(round(q_lo * 100.0))}_q{int(round(q_hi * 100.0))}"
        elif norm_method == "robust_soft":
            sc_k = _merge_softclip_k(g)
            norm_tag = (
                f"robust_soft_q{int(round(q_lo * 100.0))}_q{int(round(q_hi * 100.0))}"
                f"_k{sc_k:.2f}"
            )
        else:
            ts = _merge_tail_slope(g)
            norm_tag = (
                f"robust_linear_q{int(round(q_lo * 100.0))}_q{int(round(q_hi * 100.0))}"
                f"_ts{ts:.2f}"
            )
        r["fusion_mode"] = f"merge_linear_{norm_tag}_{g}_{wl_n:.2f}_{ws_n:.2f}"
        r["score_fused"] = merged
        r["score_final"] = merged

    order = sorted(
        range(len(rows)),
        key=lambda idx: float(rows[idx].get("score_final") or 0.0),
        reverse=True,
    )
    rank = 0
    prev = None
    for pos, idx in enumerate(order, start=1):
        cur = float(rows[idx].get("score_final") or 0.0)
        if prev is None or cur != prev:
            rank = pos
        rows[idx]["rank_final"] = int(rank)
        prev = cur


def build_similarity_analysis(rows: List[dict], top_percent: float) -> dict:
    """Select top k% by cosine_tf and compute summary statistics for an auto-written report."""
    total_pairs = len(rows)
    k = float(top_percent)
    if total_pairs <= 0:
        return {"total_pairs": 0, "k": k, "top_n": 0, "top_rows": [], "stats": {}}

    # sanitize
    if k <= 0:
        k = 0.1
    if k > 100:
        k = 100.0

    top_n = max(1, int(round(total_pairs * (k / 100.0))))
    top_rows = rows[:top_n]
    has_final = any(r.get("score_final") is not None for r in rows)
    score_col = "score_final" if has_final else "cosine_tf"
    score_mode = str(rows[0].get("score_mode") or "direct_tf")
    score_label = "score_final(norm-fused)" if has_final else ("coverage_score(covF1*q_tf)" if score_mode == "coverage" else "cosine_tf")

    def col(name: str, data: List[dict]) -> List[float]:
        out = []
        for r in data:
            v = r.get(name)
            if v is None:
                continue
            try:
                out.append(float(v))
            except Exception:
                pass
        return out

    cos_all = col(score_col, rows)
    cos_top = col(score_col, top_rows)
    jac_top = col("jaccard", top_rows)

    stats = {
        "cos_all_mean": float(np.mean(cos_all)) if cos_all else None,
        "cos_all_p50": _percentile(cos_all, 0.50),
        "cos_all_p95": _percentile(cos_all, 0.95),
        "cos_top_mean": float(np.mean(cos_top)) if cos_top else None,
        "cos_top_min": float(min(cos_top)) if cos_top else None,
        "cos_top_max": float(max(cos_top)) if cos_top else None,
        "cos_top_p50": _percentile(cos_top, 0.50),
        "jac_top_mean": float(np.mean(jac_top)) if jac_top else None,
    }

    return {
        "total_pairs": total_pairs,
        "k": k,
        "top_n": top_n,
        "top_rows": top_rows,
        "score_mode": score_mode,
        "score_label": score_label,
        "score_col": score_col,
        "stats": stats,
    }


def write_similarity_analysis_section(doc: Document, analysis: dict, units: List[Unit], show_examples: int = 20) -> None:
    """Append an auto-written similarity analysis report section to the Word doc."""
    total_pairs = analysis.get("total_pairs", 0)
    k = analysis.get("k", 1.0)
    top_n = analysis.get("top_n", 0)
    score_label = str(analysis.get("score_label") or "cosine_tf")
    score_col = str(analysis.get("score_col") or "cosine_tf")
    stats = analysis.get("stats", {}) or {}
    top_rows = analysis.get("top_rows", []) or []

    doc.add_paragraph("相似度分析报告", style="Heading 1")

    p = doc.add_paragraph(style="Normal")
    p.add_run(f"本报告基于本次相似度计算结果（unordered pairs, no self），共得到 {total_pairs} 对文本对比。")
    p.add_run(f"我们选取相似度最高的前 {k:g}%（共 {top_n} 对）作为‘高相似样本’，用于统计与快速质检。\n")

    # Summary stats paragraph
    def fmt(x, nd=4):
        return "" if x is None else f"{x:.{nd}f}"

    p2 = doc.add_paragraph(style="Normal")
    p2.add_run(f"整体分布（{score_label}）参考：")
    p2.add_run(f"mean={fmt(stats.get('cos_all_mean'))}，")
    p2.add_run(f"p50={fmt(stats.get('cos_all_p50'))}，")
    p2.add_run(f"p95={fmt(stats.get('cos_all_p95'))}。\n")
    p2.add_run("高相似样本（Top k%）统计：")
    p2.add_run(f"mean={fmt(stats.get('cos_top_mean'))}，")
    p2.add_run(f"median={fmt(stats.get('cos_top_p50'))}，")
    p2.add_run(f"min={fmt(stats.get('cos_top_min'))}，")
    p2.add_run(f"max={fmt(stats.get('cos_top_max'))}。")

    jac_mean = stats.get("jac_top_mean")
    p3 = doc.add_paragraph(style="Normal")
    if jac_mean is not None:
        p3.add_run(f"在 Top k% 样本内，Jaccard(set) 的平均值为 {jac_mean:.4f}，可作为字面重合的参照指标。")

    # Heuristic interpretation
    p4 = doc.add_paragraph(style="Normal")
    p4.add_run("解读建议：\n")
    p4.add_run("(1) 主评分高且 jaccard 也高：往往是明显互见/抄录/公式化重用；\n")
    p4.add_run("(2) 主评分高但 jaccard 中等：可能是结构相似或同义改写。")

    # Examples table
    doc.add_paragraph(f"Top {min(show_examples, len(top_rows))} 高相似样本（按 {score_label} 降序）", style="Heading 2")
    table = doc.add_table(rows=1, cols=6)
    hdr = table.rows[0].cells
    hdr[0].text = "rank"
    hdr[1].text = "id1-id2"
    hdr[2].text = "score"
    hdr[3].text = "jaccard"
    hdr[4].text = "path1"
    hdr[5].text = "path2"

    m = min(show_examples, len(top_rows))
    for i in range(m):
        r = top_rows[i]
        rr = table.add_row().cells
        rr[0].text = str(i + 1)
        rr[1].text = f"{r['id1']}-{r['id2']}"
        rr[2].text = f"{float(r.get(score_col) or 0.0):.6f}"
        jac = r.get("jaccard")
        rr[3].text = f"{float(jac):.6f}" if jac is not None else "n/a"
        rr[4].text = str(r.get("path1") or "")
        rr[5].text = str(r.get("path2") or "")


# =========================
# Reports
# =========================
def write_excel_report_fulltext(
    path: str,
    units: List[Unit],
    rows: List[dict],
    framework: str = "merge",
    granularity: str = "h2",
    semantic_model_label: Optional[str] = None,
    reranker_model_label: Optional[str] = None,
):
    wb = openpyxl.Workbook()
    ws = wb.active
    sem_label = semantic_model_label
    rer_label = reranker_model_label
    if sem_label is None:
        for r in rows:
            if r.get("semantic_model_label"):
                sem_label = str(r.get("semantic_model_label"))
                break
    if rer_label is None:
        for r in rows:
            if r.get("reranker_model_label"):
                rer_label = str(r.get("reranker_model_label"))
                break

    h_sem_q = _score_header_with_model("semantic_q", sem_label)
    h_sem_cos = _score_header_with_model("semantic_cosine", sem_label)
    h_rr_raw = _score_header_with_model("score_reranker_raw", rer_label)
    h_rr_ref = RERANK_REF_01_HEADER

    fw_norm = str(framework or "").strip().lower()
    if fw_norm == "lexical":
        sheet_model = "Lexical"
    elif fw_norm == "semantic":
        sheet_model = sem_label or "SemanticModel"
    elif fw_norm == "merge":
        sheet_model = (f"{sem_label}+{rer_label}" if (sem_label and rer_label) else (sem_label or rer_label or "MergeModel"))
    else:
        sheet_model = sem_label or rer_label or "Model"
    ws.title = _excel_safe_sheet_name(
        f"{_framework_cn(framework)}_{_granularity_name(granularity)}_{sheet_model}",
        fallback="similarity",
    )
    g = normalize_style(granularity)
    include_full_text = (g in ("normal", "h2"))
    include_cov_columns = (g in ("h1", "h2"))
    coverage_headers = [
        "coverage_score",
        "coverage_gate",
        "coverage_gate_merge",
        "coverage_f1",
        "coverage_cov1",
        "coverage_cov2",
        "coverage_q_tf",
        "coverage_q_tfidf",
        "coverage_overlap_chars",
        "coverage_match_edges",
        "coverage_candidate_edges",
        "coverage_hits",
    ]

    headers = [
        # headline ranking/scores first
        "rank_final",
        "rank_rerank_raw",
        "rank_lex",
        "rank_semantic",
        "score_final",
        h_rr_raw,
        h_rr_ref,
        "score_lex",
        h_sem_q,
        "semantic_model",

        # base info
        "id1", "id2",
        "path1", "path2",

        # primary details
        "score_mode",
        "fusion_mode",
        "score_fused",
        h_sem_cos,
        "score_base_q",
        "rerank_applied",
        "reranker_model_label",
        "reranker_model",
        "cosine_tf",

        # other raw/reference metrics
        "cosine_tf_raw",
        "cosine_tfidf",
        "jaccard",
        "minhash_jaccard",
        "short_id", "long_id", "contain_short_in_long",
    ]
    if include_cov_columns:
        cov_insert_idx = headers.index("cosine_tf") + 1
        headers = headers[:cov_insert_idx] + coverage_headers + headers[cov_insert_idx:]
    full_text_insert_idx = headers.index("path2") + 1
    if include_full_text:
        headers = headers[:full_text_insert_idx] + ["text1_full", "text2_full"] + headers[full_text_insert_idx:]
    ws.append(headers)

    def _fmt_path(u: Unit) -> str:
        parts = _unit_path_parts(u)
        parts.append(f"uid={u.uid}")
        return " / ".join(parts) if parts else f"unit_{u.uid}"

    for r in rows:
        u1 = units[r["id1"]]
        u2 = units[r["id2"]]
        # full body text (normal only)
        text1_full = (u1.text or "").rstrip()
        text2_full = (u2.text or "").rstrip()

        # readable path (include book/section headings + uid)
        path1 = _fmt_path(u1)
        path2 = _fmt_path(u2)

        row_head = [
            r.get("rank_final"),
            r.get("rank_rerank_raw"),
            r.get("rank_lex"),
            r.get("rank_semantic"),
            r.get("score_final"),
            _fmt_score_with_model(r.get("score_reranker_raw"), rer_label),
            r.get("reranker_ref_01"),
            r.get("score_lex"),
            _fmt_score_with_model(r.get("semantic_q"), sem_label),
            (r.get("semantic_model_label") or sem_label),

            r["id1"], r["id2"],
            path1, path2,

            r.get("score_mode"),
            r.get("fusion_mode"),
            r.get("score_fused"),
            _fmt_score_with_model(r.get("semantic_cosine"), sem_label),
            r.get("score_base_q"),
            r.get("rerank_applied"),
            (r.get("reranker_model_label") or rer_label),
            r.get("reranker_model"),
            r.get("cosine_tf"),
        ]
        row_cov = [
            r.get("coverage_score"),
            r.get("coverage_gate"),
            r.get("coverage_gate_merge"),
            r.get("coverage_f1"),
            r.get("coverage_cov1"),
            r.get("coverage_cov2"),
            r.get("coverage_q_tf"),
            r.get("coverage_q_tfidf"),
            r.get("coverage_overlap_chars"),
            r.get("coverage_match_edges"),
            r.get("coverage_candidate_edges"),
            r.get("coverage_hits", ""),
        ]
        row_tail = [
            r.get("cosine_tf_raw"),
            r.get("cosine_tfidf"),
            r.get("jaccard"),
            r.get("minhash_jaccard"),
            r.get("short_id"), r.get("long_id"), r.get("contain_short_in_long"),
        ]
        row_vals = row_head + (row_cov if include_cov_columns else []) + row_tail
        if include_full_text:
            row_vals = row_vals[:full_text_insert_idx] + [text1_full, text2_full] + row_vals[full_text_insert_idx:]
        ws.append(row_vals)

    ws.freeze_panes = "A2"

    # 美化：列宽 + 表头
    for col in range(1, ws.max_column + 1):
        ws.cell(row=1, column=col).alignment = openpyxl.styles.Alignment(horizontal="center", vertical="center")

    # 给关键列设置宽度（按列名映射，避免顺序调整后错位）
    col_idx = {h: i + 1 for i, h in enumerate(headers)}
    col_widths_by_name = {
        "id1": 6, "id2": 6,
        "path1": 32, "path2": 32,
        "text1_full": 90, "text2_full": 90,
        "score_mode": 14, "fusion_mode": 12, "rank_final": 10, "score_final": 10,
        "score_lex": 10, "score_fused": 10,
        h_sem_cos: 22, h_sem_q: 22, "semantic_model": 18,
        "rank_rerank_raw": 12, "rank_lex": 10, "rank_semantic": 12,
        "score_base_q": 10, h_rr_raw: 24, h_rr_ref: 14,
        "rerank_applied": 10, "reranker_model_label": 22, "reranker_model": 18,
        "cosine_tf": 10,
        "coverage_score": 12, "coverage_gate": 12, "coverage_gate_merge": 14, "coverage_f1": 10,
        "coverage_cov1": 10, "coverage_cov2": 10,
        "coverage_q_tf": 10, "coverage_q_tfidf": 12,
        "coverage_overlap_chars": 14,
        "coverage_match_edges": 12, "coverage_candidate_edges": 14,
        "coverage_hits": 90,
        "cosine_tf_raw": 10, "cosine_tfidf": 10,
        "jaccard": 10, "minhash_jaccard": 12,
        "short_id": 8, "long_id": 8, "contain_short_in_long": 18,
    }
    for name, w in col_widths_by_name.items():
        c = col_idx.get(name)
        if c is not None:
            ws.column_dimensions[openpyxl.utils.get_column_letter(c)].width = w

    wb.save(path)


def _levenshtein_diff_ops(a: str, b: str) -> Tuple[int, List[Tuple[str, str, str]]]:
    """
    Compute minimum edit distance and merged char-level op blocks.
    Returns: (edit_distance, [(tag, a_seg, b_seg), ...])
    tag in {"equal","replace","delete","insert"}.
    """
    sa = str(a or "")
    sb = str(b or "")
    n = len(sa)
    m = len(sb)
    dp = [[0] * (m + 1) for _ in range(n + 1)]
    for i in range(1, n + 1):
        dp[i][0] = i
    for j in range(1, m + 1):
        dp[0][j] = j

    for i in range(1, n + 1):
        ai = sa[i - 1]
        row = dp[i]
        row_prev = dp[i - 1]
        for j in range(1, m + 1):
            cost = 0 if ai == sb[j - 1] else 1
            row[j] = min(
                row_prev[j] + 1,      # delete
                row[j - 1] + 1,       # insert
                row_prev[j - 1] + cost,  # replace/equal
            )

    # Backtrace with deterministic tie-break: equal/replace > delete > insert.
    i, j = n, m
    rev_ops: List[Tuple[str, str, str]] = []
    while i > 0 or j > 0:
        if (
            i > 0
            and j > 0
            and sa[i - 1] == sb[j - 1]
            and dp[i][j] == dp[i - 1][j - 1]
        ):
            rev_ops.append(("equal", sa[i - 1], sb[j - 1]))
            i -= 1
            j -= 1
            continue
        if i > 0 and j > 0 and dp[i][j] == dp[i - 1][j - 1] + 1:
            rev_ops.append(("replace", sa[i - 1], sb[j - 1]))
            i -= 1
            j -= 1
            continue
        if i > 0 and dp[i][j] == dp[i - 1][j] + 1:
            rev_ops.append(("delete", sa[i - 1], ""))
            i -= 1
            continue
        if j > 0 and dp[i][j] == dp[i][j - 1] + 1:
            rev_ops.append(("insert", "", sb[j - 1]))
            j -= 1
            continue
        # Defensive fallback for rare tie/state inconsistencies.
        if i > 0 and j > 0:
            rev_ops.append(("replace", sa[i - 1], sb[j - 1]))
            i -= 1
            j -= 1
        elif i > 0:
            rev_ops.append(("delete", sa[i - 1], ""))
            i -= 1
        else:
            rev_ops.append(("insert", "", sb[j - 1]))
            j -= 1

    rev_ops.reverse()

    # Merge consecutive same tags.
    merged: List[Tuple[str, str, str]] = []
    for tag, aa, bb in rev_ops:
        if not merged or merged[-1][0] != tag:
            merged.append((tag, aa, bb))
        else:
            t, pa, pb = merged[-1]
            merged[-1] = (t, pa + aa, pb + bb)
    return int(dp[n][m]), merged


def _diff_counts(ops: List[Tuple[str, str, str]]) -> Tuple[int, int, int]:
    rep = 0
    ins = 0
    dele = 0
    for tag, aa, bb in ops:
        if tag == "replace":
            rep += len(aa)
        elif tag == "insert":
            ins += len(bb)
        elif tag == "delete":
            dele += len(aa)
    return int(rep), int(ins), int(dele)


def _clip_diff_token(s: str, max_chars: int = 30) -> str:
    t = str(s or "")
    m = int(max(1, max_chars))
    if len(t) <= m:
        return t
    return t[: m - 3] + "..."


def _variant_unit_label(a_seg: str, b_seg: str) -> str:
    """Basic label for collation wording: single-char -> 字, else -> 句."""
    la = len(str(a_seg or ""))
    lb = len(str(b_seg or ""))
    return "字" if max(la, lb) <= 1 else "句"


def _add_diff_colored_line(
    doc: Document,
    prefix: str,
    ops: List[Tuple[str, str, str]],
    side: str,
) -> None:
    p = doc.add_paragraph(style="Normal")
    p.add_run(prefix)
    for tag, aa, bb in ops:
        txt = aa if side == "a" else bb
        if not txt:
            continue
        run = p.add_run(txt)
        is_diff = (
            (side == "a" and tag in ("replace", "delete"))
            or (side == "b" and tag in ("replace", "insert"))
        )
        if is_diff:
            run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
        if side == "a" and tag == "delete":
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
        if side == "b" and tag == "insert":
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def write_normal_collation_section(
    doc: Document,
    units: List[Unit],
    rows: List[dict],
    topn: int = 1000,
    detail_norm_edit_max: float = 0.35,
    skip_han_exact: bool = True,
    max_variant_items_per_pair: int = 0,
) -> None:
    """Append display-only normal collation section (merge channel)."""
    if not rows or not units:
        return

    n_show = int(max(1, min(int(topn), len(rows))))
    cand = list(rows[:n_show])

    doc.add_paragraph(f"融合 Normal 自动校勘（详细展开 Top {n_show}）", style="Heading 1")
    doc.add_paragraph(
        (
            "规则：仅展示（不参与评分）。过滤汉字完全一致文本；"
            f"仅收录 norm_edit <= {float(detail_norm_edit_max):.2f} 的 Normal 对，并执行红字与增删高亮。"
        ),
        style="Normal",
    )

    skipped_exact = 0
    skipped_threshold = 0
    detailed_items: List[Dict[str, Any]] = []
    for r in cand:
        try:
            id1 = int(r.get("id1"))
            id2 = int(r.get("id2"))
        except Exception:
            continue
        if id1 < 0 or id2 < 0 or id1 >= len(units) or id2 >= len(units):
            continue
        u1 = units[id1]
        u2 = units[id2]
        han1 = preprocess_keep_han(u1.text or "")
        han2 = preprocess_keep_han(u2.text or "")
        if skip_han_exact and han1 == han2:
            skipped_exact += 1
            continue

        dist, ops = _levenshtein_diff_ops(han1, han2)
        denom = max(len(han1), len(han2), 1)
        norm_edit = float(dist / denom)
        if norm_edit > float(detail_norm_edit_max):
            skipped_threshold += 1
            continue
        rep, ins, dele = _diff_counts(ops)
        detailed_items.append(
            {
                "row": r,
                "id1": id1,
                "id2": id2,
                "u1": u1,
                "u2": u2,
                "han1": han1,
                "han2": han2,
                "ops": ops,
                "dist": int(dist),
                "norm_edit": norm_edit,
                "rep": rep,
                "ins": ins,
                "dele": dele,
            }
        )

    if not detailed_items:
        doc.add_paragraph("无满足详细校勘阈值的 Normal 对。", style="Normal")
    for shown, item in enumerate(detailed_items, start=1):
        r = item["row"]
        u1 = item["u1"]
        u2 = item["u2"]
        ops = item["ops"]
        score = float(r.get("score_final") or 0.0)
        doc.add_paragraph(
            (
                f"#{shown} score={score:.4f} ids={item['id1']}-{item['id2']} "
                f"HanChars(A-B)={len(item['han1'])}-{len(item['han2'])} "
                f"norm_edit={item['norm_edit']:.4f} "
                f"异文字数={int(item['dist'])}（替换{item['rep']}/增{item['ins']}/删{item['dele']}）"
            ),
            style="Heading 3",
        )

        pm = doc.add_paragraph(style="Normal")
        pm.add_run("A: ")
        pm.add_run(str(r.get("path1") or unit_path(u1)))
        pm.add_run("\nB: ")
        pm.add_run(str(r.get("path2") or unit_path(u2)))

        var_lines: List[str] = []
        for tag, aa, bb in ops:
            if tag == "equal":
                continue
            if tag == "replace":
                var_lines.append(f"A作B：\"{str(aa or '')}\"作\"{str(bb or '')}\"")
            elif tag == "delete":
                kind = _variant_unit_label(aa, "")
                var_lines.append(f"A无此{kind}：\"{str(aa or '')}\"")
            elif tag == "insert":
                kind = _variant_unit_label("", bb)
                var_lines.append(f"B增此{kind}：\"{str(bb or '')}\"")

        if var_lines:
            doc.add_paragraph("校勘记：", style="Normal")
            for ln in var_lines:
                doc.add_paragraph(f"  - {ln}", style="Normal")

        _add_diff_colored_line(doc, "A 正文（差异标红；删字高亮）: ", ops, side="a")
        _add_diff_colored_line(doc, "B 正文（差异标红；增字高亮）: ", ops, side="b")
        doc.add_paragraph(style="Normal")

    doc.add_paragraph(
        (
            f"校勘段落统计：候选={n_show}，过滤完全一致={skipped_exact}，"
            f"阈值外={skipped_threshold}，详细展开={len(detailed_items)}。"
        ),
        style="Normal",
    )
    doc.add_paragraph(style="Normal")


def write_word_report(
    path: str,
    units: List[Unit],
    normal_units: List[Unit],
    rows: List[dict],
    title: str,
    max_rows: int = 100,
    top_percent: float = 1.0,
    framework: str = "",
    granularity: str = "",
):
    # Word 放太多行会巨慢/巨大：默认只放前 max_rows（Excel 里有全量）
    doc = Document()
    _apply_bnu_word_styles(doc)
    g_norm = normalize_style(granularity)
    show_cov_meta = (g_norm in ("h1", "h2"))
    doc.add_paragraph(title, style="Heading 1")
    doc.add_paragraph(f"Generated at: {time.strftime('%Y-%m-%d %H:%M:%S')}", style="Normal")
    doc.add_paragraph(f"Pairs (unordered, no self): {len(rows)}", style="Normal")
    doc.add_paragraph("Containment is directional: short->long = |S(short)∩S(long)|/|S(short)|", style="Normal")

    # --- Wordcount stats (h1/h2/normal) integrated into this report ---
    write_wordcount_section(doc, normal_units, show_top=20)

    # --- Similarity analysis report (Top k%) ---
    analysis = build_similarity_analysis(rows, top_percent)
    write_similarity_analysis_section(doc, analysis, units, show_examples=min(20, analysis.get("top_n", 0) or 0))

    # --- Top high-sim samples (readable paragraph style; preprocessed Han-only text) ---
    doc.add_paragraph(f"Top {min(max_rows, len(rows))} 高相似样本（按综合分降序；正文为预处理汉字版）", style="Heading 2")

    show = rows[:max_rows]
    for i, r in enumerate(show, start=1):
        u1 = units[int(r["id1"])]
        u2 = units[int(r["id2"])]

        score_mode = str(r.get("score_mode") or "direct_tf")
        score_final = r.get("score_final")
        cosv = float((score_final if score_final is not None else (r.get("cosine_tf") or 0.0)))
        jacv = float(r.get("jaccard") or 0.0)
        if score_final is not None:
            score_label = "score_final"
        elif show_cov_meta and score_mode == "coverage":
            score_label = "score(covF1*q_tf)"
        else:
            score_label = "cos_tf"

        # Each sample as a small subsection
        doc.add_paragraph(
            f"#{i}  {score_label}={cosv:.4f}  jaccard={jacv:.4f}  {r['id1']}-{r['id2']}",
            style="Heading 3",
        )

        # Paths
        pmeta = doc.add_paragraph(style="Normal")
        pmeta.add_run("A: ")
        pmeta.add_run(str(r.get("path1") or unit_path(u1)))
        pmeta.add_run("\nB: ")
        pmeta.add_run(str(r.get("path2") or unit_path(u2)))
        if show_cov_meta and score_mode == "coverage":
            pmeta.add_run("\ncoverage: ")
            pmeta.add_run(
                f"cov1={float(r.get('coverage_cov1') or 0.0):.4f}, "
                f"cov2={float(r.get('coverage_cov2') or 0.0):.4f}, "
                f"f1={float(r.get('coverage_f1') or 0.0):.4f}, "
                f"q_tf={float(r.get('coverage_q_tf') or 0.0):.4f}, "
                f"overlap_chars={int(r.get('coverage_overlap_chars') or 0)}"
            )
            hits = str(r.get("coverage_hits") or "").strip()
            if hits:
                pmeta.add_run("\ncoverage_hits: ")
                pmeta.add_run(hits)

        # Full正文 in preprocessed Han-only form (no punctuation)
        doc.add_paragraph("A 正文（预处理：仅汉字）", style="Normal")
        doc.add_paragraph(preprocess_keep_han(u1.text or ""), style="Normal")
        doc.add_paragraph("B 正文（预处理：仅汉字）", style="Normal")
        doc.add_paragraph(preprocess_keep_han(u2.text or ""), style="Normal")

        # Spacer
        doc.add_paragraph(style="Normal")

    fw = str(framework or "").strip().lower()
    gg = normalize_style(granularity)
    if fw == "merge" and gg == "normal":
        write_normal_collation_section(
            doc,
            units=units,
            rows=rows,
            topn=int(NORMAL_COLLATION_DEFAULTS["topn_merge_normal"]),
            detail_norm_edit_max=float(NORMAL_COLLATION_DEFAULTS["detail_norm_edit_max"]),
            skip_han_exact=bool(NORMAL_COLLATION_DEFAULTS["skip_han_exact"]),
            max_variant_items_per_pair=int(NORMAL_COLLATION_DEFAULTS["max_variant_items_per_pair"]),
        )

    doc.save(path)


def _han_excerpt(s: str, max_chars: int) -> str:
    t = preprocess_keep_han(s or "")
    m = int(max(1, max_chars))
    if len(t) <= m:
        return t
    return t[:m] + "..."


def write_wordcount_section_all(
    doc: Document,
    results_by_g: Dict[str, Dict[str, Any]],
) -> None:
    """Append all-report wordcount stats for h1/h2/normal.

    Counting rule: preprocess to Han-only (remove punctuation/special symbols/whitespace),
    then count characters.
    """
    doc.add_paragraph("字数统计（all）", style="Heading 1")
    p0 = doc.add_paragraph(style="Normal")
    p0.add_run("统计口径说明：字数统计与字面相似度计算统一使用预处理文本。")
    p0.add_run("预处理规则为仅保留中文汉字，去除所有标点符号、空白和其他特殊符号。")

    for g in ["h1", "h2", "normal"]:
        item = results_by_g.get(g) or {}
        units = item.get("units") or []

        unit_rows: List[Tuple[int, str, int]] = []
        for u in units:
            try:
                uid = int(u.uid)
            except Exception:
                uid = 0
            path = unit_path(u)
            L = _eff_len_chars(u.text or "")
            unit_rows.append((uid, path, int(L)))

        unit_rows.sort(key=lambda x: x[0])
        total_chars = int(sum(x[2] for x in unit_rows))
        avg_chars = float(total_chars / max(1, len(unit_rows)))
        nonzero = int(sum(1 for x in unit_rows if x[2] > 0))

        doc.add_paragraph(f"{g.upper()} 层级统计", style="Heading 2")
        p = doc.add_paragraph(style="Normal")
        p.add_run(f"units={len(unit_rows)}；")
        p.add_run(f"non_empty={nonzero}；")
        p.add_run(f"Total Han chars={total_chars}；")
        p.add_run(f"Avg Han chars / unit={avg_chars:.2f}")

        if not unit_rows:
            doc.add_paragraph("该层级无可统计 unit。", style="Normal")
            continue

        t = doc.add_table(rows=1, cols=4)
        h = t.rows[0].cells
        h[0].text = "rank"
        h[1].text = "unit_id"
        h[2].text = "path"
        h[3].text = "han_chars"
        for i, (uid, path, L) in enumerate(unit_rows, start=1):
            r = t.add_row().cells
            r[0].text = str(i)
            r[1].text = str(uid)
            r[2].text = str(path)
            r[3].text = str(L)
        doc.add_paragraph(style="Normal")


def _setup_plot_zh_font() -> None:
    """Best-effort Chinese font setup for matplotlib figures."""
    global _PLOT_ZH_READY
    if _PLOT_ZH_READY or plt is None:
        return
    try:
        from matplotlib import font_manager
        # Try registering common CJK font files on macOS/Linux first.
        font_paths = [
            "/System/Library/Fonts/PingFang.ttc",
            "/System/Library/Fonts/Hiragino Sans GB.ttc",
            "/System/Library/Fonts/STHeiti Light.ttc",
            "/System/Library/Fonts/STHeiti Medium.ttc",
            "/System/Library/Fonts/Supplemental/Songti.ttc",
            "/System/Library/Fonts/Supplemental/STSong.ttf",
            "/Library/Fonts/Arial Unicode.ttf",
            "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
        ]
        for p in font_paths:
            try:
                if os.path.exists(p):
                    font_manager.fontManager.addfont(p)
            except Exception:
                pass
        candidates = [
            "Noto Sans CJK SC",
            "Noto Serif CJK SC",
            "PingFang SC",
            "Hiragino Sans GB",
            "Songti SC",
            "Heiti SC",
            "STHeiti",
            "STSong",
            "SimHei",
            "Microsoft YaHei",
            "WenQuanYi Zen Hei",
            "Arial Unicode MS",
        ]
        installed = {f.name for f in font_manager.fontManager.ttflist}
        chosen = None
        for name in candidates:
            if name in installed:
                chosen = name
                break
        if chosen:
            old = list(plt.rcParams.get("font.sans-serif", []))
            plt.rcParams["font.family"] = ["sans-serif"]
            plt.rcParams["font.sans-serif"] = [chosen] + [x for x in old if x != chosen]
        else:
            plt.rcParams["font.family"] = ["sans-serif"]
            plt.rcParams["font.sans-serif"] = candidates
        # Keep monospaced text CJK-capable as well (for compact table-like text blocks).
        plt.rcParams["font.monospace"] = [chosen] if chosen else candidates
        plt.rcParams["axes.unicode_minus"] = False
    except Exception:
        pass
    _PLOT_ZH_READY = True


def _plot_top10_scores_for_g(
    rows: List[dict],
    granularity: str,
    out_path: str,
    framework_key: str,
    weights_h: Optional[Dict[str, float]] = None,
    weights_n: Optional[Dict[str, float]] = None,
    merge_w_lex: float = 0.70,
    merge_w_sem: float = 0.30,
    rrf_k_h1: int = 30,
) -> Optional[str]:
    if plt is None or not rows:
        return None
    _setup_plot_zh_font()

    top = rows[:10]
    if not top:
        return None

    def _v(r: dict, key: str, default: float = 0.0) -> float:
        try:
            x = r.get(key)
            return float(default if x is None else x)
        except Exception:
            return float(default)

    def _weights_for_g(gg: str) -> Dict[str, float]:
        if normalize_style(gg) in ("h1", "h2"):
            return _normalize_weights(weights_h or LEXICAL_WEIGHTS_DEFAULT_H, LEXICAL_WEIGHTS_DEFAULT_H)
        return _normalize_weights(weights_n or LEXICAL_WEIGHTS_DEFAULT_N, LEXICAL_WEIGHTS_DEFAULT_N)

    def _formula_note() -> str:
        fw = str(framework_key or "").strip().lower()
        g = normalize_style(granularity)
        if fw == "lexical":
            w = _weights_for_g(g)
            if g in ("h1", "h2"):
                return (
                    f"字面通道（{g.upper()}）\n"
                    "score_lex_raw = w_cos·cos + w_cov·cov + w_tfidf·tfidf + w_jacc_eff·jacc_eff\n"
                    "字面通道内部不做跨指标归一化；仅线性加权\n"
                    "jacc_eff = jaccard^gamma(len_short)，长文本 gamma 更小（温和抬升）\n\n"
                    f"w_cos={w['cosine_tf_raw']:.2f}, "
                    f"w_cov={w['coverage']:.2f}, "
                    f"w_tfidf={w['cosine_tfidf']:.2f}, "
                    f"w_jacc={w['jaccard']:.2f}\n\n"
                    "指标解释：\n"
                    "cov: 证据覆盖(evidence)\n"
                    "cos: 文体/体裁近似(style/genre)\n"
                    "tfidf: 去噪抑制高频虚词\n"
                    "jaccard: 长度校正后辅助判别"
                )
            return (
                f"字面通道（{g.upper()}）\n"
                "score_lex_raw = w_cos·cos + w_tfidf·tfidf + w_jacc_eff·jacc_eff\n"
                "字面通道内部不做跨指标归一化；仅线性加权\n"
                "jacc_eff = jaccard^gamma(len_short)，长文本 gamma 更小（温和抬升）\n\n"
                f"w_cos={w['cosine_tf_raw']:.2f}, "
                f"w_tfidf={w['cosine_tfidf']:.2f}, "
                f"w_jacc={w['jaccard']:.2f}\n\n"
                "指标解释：\n"
                "cos: 文体/体裁近似(style/genre)\n"
                "tfidf: 去噪抑制高频虚词\n"
                "jaccard: 长度校正后辅助判别"
            )
        if fw == "semantic":
            if g in ("h1", "h2") and _semantic_maxsim_enabled(g):
                lam = SEMANTIC_MAXSIM_H1_LAMBDA if g == "h1" else SEMANTIC_MAXSIM_H2_LAMBDA
                return (
                    f"语义通道（{g.upper()}）\n"
                    "doc_cos = cosine(mean_vec(A), mean_vec(B))\n"
                    "maxsim_bi = 0.5·(mean(max_row(S)) + mean(max_col(S)))\n"
                    f"score_sem_raw = {lam:.2f}·doc_cos + {1.0 - float(lam):.2f}·maxsim_bi\n"
                    f"H1: 全量 MaxSim（lambda={SEMANTIC_MAXSIM_H1_LAMBDA:.2f})；"
                    f"H2: Top{SEMANTIC_MAXSIM_H2_TOP_PERCENT:.1f}% MaxSim 精排（{SEMANTIC_MAXSIM_H2_TOPN_MIN}~{SEMANTIC_MAXSIM_H2_TOPN_MAX}）\n\n"
                    "本通道不使用字面分量\n"
                    "(cos/tfidf/jaccard/cov 均不参与)"
                )
            msg = (
                f"语义通道（{g.upper()}）\n"
                "score_sem_raw = clip(cosine(mean_vec(A), mean_vec(B)), 0, 1)\n"
                "score_sem = Norm(score_sem_raw) （语义通道内归一化）\n\n"
                "本通道不使用字面分量\n"
                "(cos/tfidf/jaccard/cov 均不参与)"
            )
            if g == "h2":
                msg += "\nH2 当前配置: avgpool（不启用 MaxSim）"
            return msg

        wl = max(0.0, float(merge_w_lex))
        ws = max(0.0, float(merge_w_sem))
        den = wl + ws if (wl + ws) > 0 else 1.0
        wl_n = wl / den
        ws_n = ws / den
        q_lo, q_hi = _merge_quantile_bounds(g)
        norm_method = str(MERGE_NORM_METHOD or "").strip().lower()
        if norm_method == "minmax":
            norm_note = "q_lex = MinMax(score_lex_raw), q_sem = MinMax(score_sem_raw)"
        elif norm_method == "robust_hard":
            norm_note = (
                f"q_lex = RobustNormHard(score_lex_raw; q{int(round(q_lo*100))}/q{int(round(q_hi*100))}), "
                f"q_sem = RobustNormHard(score_sem_raw; q{int(round(q_lo*100))}/q{int(round(q_hi*100))})"
            )
        elif norm_method == "robust_soft":
            sc_k = _merge_softclip_k(g)
            norm_note = (
                f"q_lex = RobustSoft(score_lex_raw; q{int(round(q_lo*100))}/q{int(round(q_hi*100))}, k={sc_k:.2f}), "
                f"q_sem = RobustSoft(score_sem_raw; q{int(round(q_lo*100))}/q{int(round(q_hi*100))}, k={sc_k:.2f})"
            )
        else:
            ts = _merge_tail_slope(g)
            norm_note = (
                f"q_lex = RobustLinear(score_lex_raw; q{int(round(q_lo*100))}/q{int(round(q_hi*100))}, tail={ts:.2f}), "
                f"q_sem = RobustLinear(score_sem_raw; q{int(round(q_lo*100))}/q{int(round(q_hi*100))}, tail={ts:.2f})"
            )
        return (
            f"融合通道（{g.upper()}）\n"
            "仅在融合阶段做一次归一化：\n"
            f"{norm_note}\n"
            f"score = {wl_n:.2f}·q_lex + {ws_n:.2f}·q_sem"
        )

    x = list(range(1, len(top) + 1))
    fw = str(framework_key or "").strip().lower()
    fig = plt.figure(figsize=(16, 9))
    gs = fig.add_gridspec(1, 2, width_ratios=[3.6, 1.4], wspace=0.06)
    ax = fig.add_subplot(gs[0, 0])
    ax_note = fig.add_subplot(gs[0, 1])
    ax_note.axis("off")

    if fw == "lexical":
        g_norm = normalize_style(granularity)
        y_total = [_v(r, "score_final") for r in top]
        y_cos = [_v(r, "q_cosine_tf_raw") for r in top]
        y_tfidf = [_v(r, "q_cosine_tfidf") for r in top]
        y_jacc = [_v(r, "q_jaccard") for r in top]
        ax.plot(x, y_total, marker="o", linewidth=2.4, color="#D1495B", label="总分 score_lex(raw)")
        if g_norm in ("h1", "h2"):
            y_cov = [_v(r, "q_cov_star") for r in top]
            ax.plot(x, y_cov, marker="s", linewidth=1.8, color="#4F7CAC", label="cov_raw")
        ax.plot(x, y_cos, marker="^", linewidth=1.8, color="#3B8A5B", label="cos_raw")
        ax.plot(x, y_tfidf, marker="D", linewidth=1.8, color="#A46D1F", label="tfidf_raw")
        ax.plot(x, y_jacc, marker="v", linewidth=1.8, color="#7D4E57", label="jacc_eff")
        title_prefix = "LEX"
    elif fw == "semantic":
        y_sem = [_v(r, "score_final") for r in top]
        y_raw = [_v(r, "semantic_cosine") for r in top]
        ax.plot(x, y_sem, marker="o", linewidth=2.4, color="#2F6EA8", label="score_sem(final)")
        ax.plot(x, y_raw, marker="s", linewidth=1.8, color="#7FB3D5", label="cosine(embed)")
        title_prefix = "SEM"
    else:
        y_total = [_v(r, "score_final") for r in top]
        y_lex = [_v(r, "score_lex") for r in top]
        y_sem = [_v(r, "semantic_q") for r in top]
        ax.plot(x, y_total, marker="o", linewidth=2.4, color="#8A5A44", label="融合总分")
        ax.plot(x, y_lex, marker="s", linewidth=1.9, color="#D1495B", label="字面分 q_lex")
        ax.plot(x, y_sem, marker="^", linewidth=1.9, color="#EDA63B", label="语义分 q_sem")
        title_prefix = "MERGE"

    ax.set_title(f"{title_prefix} | {granularity.upper()} Top10", fontsize=16, fontweight="bold")
    ax.set_xlabel("相似度排名（1=最相似）", fontsize=12)
    ax.set_ylabel("分数", fontsize=12)
    ax.set_xticks(x)
    ax.set_ylim(0.0, 1.02)
    ax.grid(True, alpha=0.25, linestyle="--")
    ax.legend(loc="lower left", fontsize=10)
    fig.patch.set_facecolor("#FFF6E9")
    ax.set_facecolor("#FFFDF8")
    ax_note.set_facecolor("#FFF6E9")
    ax_note.text(
        0.02,
        0.98,
        _formula_note(),
        va="top",
        ha="left",
        fontsize=10,
        linespacing=1.35,
        bbox=dict(boxstyle="round,pad=0.5", facecolor="#FFFDF8", edgecolor="#C39A73", linewidth=1.0),
    )

    os.makedirs(os.path.dirname(out_path), exist_ok=True)
    fig.tight_layout()
    fig.savefig(out_path, dpi=600)
    plt.close(fig)
    return out_path


def build_top10_plot_bundle(
    results_by_g: Dict[str, Dict[str, Any]],
    out_dir: str,
    framework_key: str,
    file_stem: str,
    weights_h: Optional[Dict[str, float]] = None,
    weights_n: Optional[Dict[str, float]] = None,
    merge_w_lex: float = 0.70,
    merge_w_sem: float = 0.30,
    rrf_k_h1: int = 30,
) -> Dict[str, str]:
    out: Dict[str, str] = {}
    for g in ["h1", "h2", "normal"]:
        item = results_by_g.get(g) or {}
        rows = item.get("rows") or []
        p = os.path.join(out_dir, f"{file_stem}_{g}_top10.png")
        got = _plot_top10_scores_for_g(
            rows,
            g,
            p,
            framework_key=framework_key,
            weights_h=weights_h,
            weights_n=weights_n,
            merge_w_lex=merge_w_lex,
            merge_w_sem=merge_w_sem,
            rrf_k_h1=rrf_k_h1,
        )
        if got:
            out[g] = got
    return out


def build_framework_comparison_plot(
    results_by_fw: Dict[str, Dict[str, Dict[str, Any]]],
    out_dir: str,
    file_stem: str,
) -> Optional[str]:
    """Draw one comparison figure: 3 granularities x 3 frameworks (top10 score curves)."""
    if plt is None:
        return None
    _setup_plot_zh_font()

    gs = ["h1", "h2", "normal"]

    fig, axes = plt.subplots(1, 3, figsize=(16, 9), sharey=True)
    if not isinstance(axes, np.ndarray):
        axes = np.array([axes])

    for ax, g in zip(axes, gs):
        x = list(range(1, 11))
        for fw, lbl, color in FRAMEWORK_PLOT_ORDER:
            rows = ((results_by_fw.get(fw) or {}).get(g) or {}).get("rows") or []
            top = rows[:10]
            y = [float(r.get("score_final") or 0.0) for r in top]
            if len(y) < 10:
                y.extend([0.0] * (10 - len(y)))
            ax.plot(x, y, marker="o", linewidth=2.0, color=color, label=lbl)
        ax.set_title(f"{g.upper()} Top10", fontsize=13)
        ax.set_xticks(x)
        ax.grid(True, alpha=0.25, linestyle="--")
        ax.set_facecolor("#FFFDF8")

    axes[0].set_ylabel("排名得分", fontsize=12)
    fig.suptitle("三种模式在三层级 Top10 排名得分比较", fontsize=16)
    handles, labels = axes[0].get_legend_handles_labels()
    fig.legend(handles, labels, loc="upper center", ncol=3)
    fig.patch.set_facecolor("#FFF6E9")
    fig.tight_layout(rect=[0, 0, 1, 0.93])

    os.makedirs(out_dir, exist_ok=True)
    out_path = os.path.join(out_dir, f"{file_stem}_framework_compare_top10.png")
    fig.savefig(out_path, dpi=600)
    plt.close(fig)
    return out_path


def _pair_key(id1: int, id2: int) -> Tuple[int, int]:
    a = int(id1)
    b = int(id2)
    return (a, b) if a < b else (b, a)


def enrich_cross_channel_scores(
    results_by_fw: Dict[str, Dict[str, Dict[str, Any]]],
    reps_cache: Dict[str, dict],
    sem_emb_cache: Dict[str, np.ndarray],
    granularities: List[str],
    sem_model: str = "",
) -> None:
    """Post-hoc cross audit enrichment.

    Keep channel isolation unchanged. After all channels finish, compute:
      - cross_sem_score: semantic cosine for any pair
      - cross_lex_score: lexical cosine for any pair
    and attach them to existing row dicts.
    """
    for g in granularities:
        reps = reps_cache.get(g)
        _ck = f"{g}:{sem_model}" if sem_model else g
        sem_emb = sem_emb_cache.get(_ck)
        if reps is None or sem_emb is None or not isinstance(sem_emb, np.ndarray):
            continue
        if sem_emb.ndim < 2 or sem_emb.shape[0] < 2:
            continue
        X_tf = reps.get("X_tf") if isinstance(reps, dict) else None
        if X_tf is None:
            continue

        try:
            X_tf_norm = sk_normalize(X_tf, norm="l2", copy=True)
        except Exception:
            continue

        pair_set: set = set()
        for fw in ("lexical", "semantic", "merge"):
            rows = (((results_by_fw.get(fw) or {}).get(g) or {}).get("rows") or [])
            for r in rows:
                try:
                    id1 = int(r.get("id1"))
                    id2 = int(r.get("id2"))
                except Exception:
                    continue
                if id1 == id2:
                    continue
                pair_set.add(_pair_key(id1, id2))

        if not pair_set:
            continue

        n = int(sem_emb.shape[0])
        cross_sem: Dict[Tuple[int, int], float] = {}
        cross_lex: Dict[Tuple[int, int], float] = {}
        for id1, id2 in pair_set:
            if id1 < 0 or id2 < 0 or id1 >= n or id2 >= n:
                continue
            cross_sem[(id1, id2)] = float(np.dot(sem_emb[id1], sem_emb[id2]))
            try:
                cross_lex[(id1, id2)] = float(X_tf_norm[id1].multiply(X_tf_norm[id2]).sum())
            except Exception:
                try:
                    cross_lex[(id1, id2)] = float(cosine_similarity(X_tf_norm[id1], X_tf_norm[id2])[0, 0])
                except Exception:
                    pass

        for fw in ("lexical", "semantic", "merge"):
            rows = (((results_by_fw.get(fw) or {}).get(g) or {}).get("rows") or [])
            for r in rows:
                try:
                    id1 = int(r.get("id1"))
                    id2 = int(r.get("id2"))
                except Exception:
                    continue
                if id1 == id2:
                    continue
                key = _pair_key(id1, id2)
                r["cross_sem_score"] = cross_sem.get(key)
                r["cross_lex_score"] = cross_lex.get(key)


def _write_excel_all_like_sheet(
    ws,
    framework: str,
    granularity: str,
    units: List[Unit],
    rows: List[dict],
    weights_h: Optional[Dict[str, float]] = None,
    weights_n: Optional[Dict[str, float]] = None,
    semantic_model_label: Optional[str] = None,
    reranker_model_label: Optional[str] = None,
) -> None:
    """Write one sheet using the same structure as write_excel_report_all."""
    def _sheet_weights(g: str) -> Dict[str, float]:
        gg = normalize_style(g)
        if gg in ("h1", "h2"):
            return _normalize_weights(weights_h or LEXICAL_WEIGHTS_DEFAULT_H, LEXICAL_WEIGHTS_DEFAULT_H)
        return _normalize_weights(weights_n or LEXICAL_WEIGHTS_DEFAULT_N, LEXICAL_WEIGHTS_DEFAULT_N)

    def _wlabel(metric_name: str, w: float) -> str:
        return f"{metric_name}[w={100.0 * float(w):.1f}%]"

    g = normalize_style(granularity)
    sem_label = semantic_model_label
    rer_label = reranker_model_label
    if sem_label is None:
        for r in rows:
            if r.get("semantic_model_label"):
                sem_label = str(r.get("semantic_model_label"))
                break
    if rer_label is None:
        for r in rows:
            if r.get("reranker_model_label"):
                rer_label = str(r.get("reranker_model_label"))
                break

    h_sem_q = _score_header_with_model("semantic_q", sem_label)
    h_sem_cos = _score_header_with_model("semantic_cosine", sem_label)
    h_cross_sem = _score_header_with_model("cross_sem_score", sem_label)
    h_rr_raw = _score_header_with_model("score_reranker_raw", rer_label)
    h_rr_ref = RERANK_REF_01_HEADER

    include_full_text = (g in ("normal", "h2"))
    include_cov_columns = (g in ("h1", "h2"))
    w = _sheet_weights(g)
    metric_order = [
        ("cosine_tf_raw", float(w.get("cosine_tf_raw", 0.0))),
        ("cosine_tfidf", float(w.get("cosine_tfidf", 0.0))),
        ("jaccard", float(w.get("jaccard", 0.0))),
    ]
    if include_cov_columns:
        metric_order.append(("coverage_score", float(w.get("coverage", 0.0))))
    metric_order.sort(key=lambda x: x[1], reverse=True)
    metric_headers = [(k, _wlabel(k, wk)) for k, wk in metric_order]

    headers = ["id1", "id2", "path1", "path2"]
    if include_full_text:
        headers.extend(["text1_full", "text2_full"])
    headers += [
        "rank_final", "rank_rerank_raw", "rank_lex", "rank_semantic",
        "score_final", h_rr_raw, h_rr_ref, "score_lex", h_sem_q, "semantic_model",
        "cross_lex_score", h_cross_sem,
    ]
    headers += [
        "score_mode", "fusion_mode", "score_fused", h_sem_cos,
        "score_base_q", "rerank_applied", "reranker_model_label", "reranker_model",
    ] + [h for _, h in metric_headers] + [
        "minhash_jaccard",
        "han_chars_A", "han_chars_B",
    ]
    if include_cov_columns:
        headers += [
            "q_cov_star", "coverage_gate", "coverage_gate_merge",
            "coverage_f1", "coverage_cov1", "coverage_cov2",
            "coverage_q_tf", "coverage_q_tfidf",
            "coverage_overlap_chars", "coverage_match_edges", "coverage_candidate_edges", "coverage_hits",
        ]
    ws.append(headers)
    for r in rows:
        try:
            id1 = int(r["id1"])
            id2 = int(r["id2"])
        except Exception:
            continue
        if id1 < 0 or id2 < 0 or id1 >= len(units) or id2 >= len(units):
            continue
        u1 = units[id1]
        u2 = units[id2]
        han1 = preprocess_keep_han(u1.text or "")
        han2 = preprocess_keep_han(u2.text or "")
        text1_full = (u1.text or "").rstrip()
        text2_full = (u2.text or "").rstrip()
        row_vals = [
            id1, id2,
            str(r.get("path1") or unit_path(u1)),
            str(r.get("path2") or unit_path(u2)),
        ]
        if include_full_text:
            row_vals.extend([text1_full, text2_full])
        row_vals.extend([
            r.get("rank_final"),
            r.get("rank_rerank_raw"),
            r.get("rank_lex"),
            r.get("rank_semantic"),
            r.get("score_final"),
            _fmt_score_with_model(r.get("score_reranker_raw"), rer_label),
            r.get("reranker_ref_01"),
            r.get("score_lex"),
            _fmt_score_with_model(r.get("semantic_q"), sem_label),
            (r.get("semantic_model_label") or sem_label),
            r.get("cross_lex_score"),
            _fmt_score_with_model(r.get("cross_sem_score"), sem_label),
        ])
        row_vals.extend([
            r.get("score_mode"),
            r.get("fusion_mode"),
            r.get("score_fused"),
            _fmt_score_with_model(r.get("semantic_cosine"), sem_label),
            r.get("score_base_q"),
            r.get("rerank_applied"),
            (r.get("reranker_model_label") or rer_label),
            r.get("reranker_model"),
        ])
        for key, _ in metric_headers:
            row_vals.append(r.get(key))

        row_vals.extend([
            r.get("minhash_jaccard"),
            len(han1),
            len(han2),
        ])
        if include_cov_columns:
            row_vals.extend([
                r.get("q_cov_star"),
                r.get("coverage_gate"),
                r.get("coverage_gate_merge"),
                r.get("coverage_f1"),
                r.get("coverage_cov1"),
                r.get("coverage_cov2"),
                r.get("coverage_q_tf"),
                r.get("coverage_q_tfidf"),
                r.get("coverage_overlap_chars"),
                r.get("coverage_match_edges"),
                r.get("coverage_candidate_edges"),
                r.get("coverage_hits", ""),
            ])
        ws.append(row_vals)

    for c, h in enumerate(headers, start=1):
        wcol = 10
        if h in ("id1", "id2"):
            wcol = 6
        elif h in ("path1", "path2"):
            wcol = 34
        elif h in ("text1_full", "text2_full"):
            wcol = 90
        elif h in ("score_mode",):
            wcol = 16
        elif h in ("fusion_mode",):
            wcol = 12
        elif h in ("rank_final", "rank_rerank_raw", "score_final"):
            wcol = 12
        elif h in ("score_lex", "score_fused"):
            wcol = 10
        elif h.startswith("semantic_cosine") or h.startswith("semantic_q") or h.startswith("cross_sem_score"):
            wcol = 24
        elif h in ("cross_lex_score",):
            wcol = 14
        elif h in ("semantic_model",):
            wcol = 18
        elif h in ("rank_lex", "rank_semantic"):
            wcol = 12
        elif h == "score_base_q":
            wcol = 12
        elif h.startswith("score_reranker_raw"):
            wcol = 26
        elif h == h_rr_ref:
            wcol = 14
        elif h in ("rerank_applied",):
            wcol = 10
        elif h in ("reranker_model_label", "reranker_model"):
            wcol = 20
        elif h.startswith("cosine_tf_raw"):
            wcol = 14
        elif h.startswith("coverage_score"):
            wcol = 14
        elif h in ("coverage_gate", "coverage_gate_merge"):
            wcol = 14
        elif h.startswith("cosine_tfidf"):
            wcol = 14
        elif h.startswith("jaccard"):
            wcol = 12
        elif h in ("coverage_hits",):
            wcol = 95
        elif h in ("han_chars_A", "han_chars_B", "coverage_overlap_chars"):
            wcol = 14
        elif h in ("coverage_match_edges", "coverage_candidate_edges"):
            wcol = 14
        ws.column_dimensions[openpyxl.utils.get_column_letter(c)].width = wcol

    for c in range(1, len(headers) + 1):
        ws.cell(row=1, column=c).alignment = openpyxl.styles.Alignment(horizontal="center", vertical="center")
    ws.freeze_panes = "A2"


def write_excel_report_framework_all(
    path: str,
    results_by_fw: Dict[str, Dict[str, Dict[str, Any]]],
    top_rows_per_g: Optional[int] = None,
    weights_h: Optional[Dict[str, float]] = None,
    weights_n: Optional[Dict[str, float]] = None,
    semantic_model_label: Optional[str] = None,
    reranker_model_label: Optional[str] = None,
) -> None:
    """Write framework=all total workbook with 9 sheets: (字面/语义/融合) x (h1/h2/normal)."""
    wb = openpyxl.Workbook()
    if wb.active is not None:
        wb.remove(wb.active)

    for fw, fw_cn in FRAMEWORK_ORDER:
        g_map = results_by_fw.get(fw) or {}
        if fw == "lexical":
            model_tag = "Lexical"
            sem_label_fw = None
            rer_label_fw = None
        elif fw == "semantic":
            model_tag = semantic_model_label or "SemanticModel"
            sem_label_fw = semantic_model_label
            rer_label_fw = None
        else:
            if semantic_model_label and reranker_model_label:
                model_tag = f"{semantic_model_label}+{reranker_model_label}"
            else:
                model_tag = semantic_model_label or reranker_model_label or "MergeModel"
            sem_label_fw = semantic_model_label
            rer_label_fw = reranker_model_label
        for g in ["h1", "h2", "normal"]:
            item = g_map.get(g) or {}
            units = item.get("units") or []
            rows = list(item.get("rows") or [])
            if top_rows_per_g is not None:
                rows = rows[: max(0, int(top_rows_per_g))]
            ws = wb.create_sheet(title=_excel_safe_sheet_name(f"{fw_cn}_{g}_{model_tag}", fallback=f"{fw_cn}_{g}"))
            _write_excel_all_like_sheet(
                ws,
                framework=fw,
                granularity=g,
                units=units,
                rows=rows,
                weights_h=weights_h,
                weights_n=weights_n,
                semantic_model_label=sem_label_fw,
                reranker_model_label=rer_label_fw,
            )

    wb.save(path)


def _slice_topk_rows(rows: List[dict], topk: int) -> List[dict]:
    if not rows:
        return []
    k = int(topk)
    if k <= 0 or k >= len(rows):
        return list(rows)
    return list(rows[:k])


def _pair_row_map(rows: List[dict]) -> Dict[Tuple[int, int], dict]:
    out: Dict[Tuple[int, int], dict] = {}
    for r in rows or []:
        try:
            k = _pair_key(int(r.get("id1")), int(r.get("id2")))
        except Exception:
            continue
        if k not in out:
            out[k] = r
    return out


def _pick_units_for_g(results_by_fw: Dict[str, Dict[str, Dict[str, Any]]], g: str) -> List[Unit]:
    for fw in ("merge", "lexical", "semantic"):
        item = ((results_by_fw.get(fw) or {}).get(g) or {})
        units = item.get("units") or []
        if units:
            return units
    return []


def _first_non_none(vals: List[Any]) -> Any:
    for v in vals:
        if v is not None:
            return v
    return None


def build_shadow_audit_rows(
    results_by_fw: Dict[str, Dict[str, Dict[str, Any]]],
    topk_by_g: Optional[Dict[str, int]] = None,
    reranker_enabled: bool = False,
    reranker_model: str = "0.6B",
    reranker_prompt: str = "",
    log_cb=None,
) -> Dict[str, List[dict]]:
    """Build unified-universe (U) cross-audit rows per granularity.

    U = TopK(lexical) ∪ TopK(semantic) ∪ TopK(merge), keyed by (id1,id2).
    Reranker is sidecar-only: raw score + rank for reference, no rank overwrite.
    """
    out: Dict[str, List[dict]] = {}
    def _log(msg: str) -> None:
        if log_cb is None:
            return
        try:
            log_cb(str(msg))
        except Exception:
            pass
    kcfg = dict(SHADOW_AUDIT_TOPK_DEFAULT)
    if isinstance(topk_by_g, dict):
        for kk in ("h1", "h2", "normal"):
            if kk in topk_by_g:
                try:
                    kcfg[kk] = int(topk_by_g[kk])
                except Exception:
                    pass

    for g in ["h1", "h2", "normal"]:
        lex_item = ((results_by_fw.get("lexical") or {}).get(g) or {})
        sem_item = ((results_by_fw.get("semantic") or {}).get(g) or {})
        merge_item = ((results_by_fw.get("merge") or {}).get(g) or {})
        lex_rows = list(lex_item.get("rows") or [])
        sem_rows = list(sem_item.get("rows") or [])
        merge_rows = list(merge_item.get("rows") or [])
        units = _pick_units_for_g(results_by_fw, g)

        k = int(kcfg.get(g, 0))
        lex_top = _slice_topk_rows(lex_rows, k)
        sem_top = _slice_topk_rows(sem_rows, k)
        merge_top = _slice_topk_rows(merge_rows, k)

        full_map_lex = _pair_row_map(lex_rows)
        full_map_sem = _pair_row_map(sem_rows)
        full_map_merge = _pair_row_map(merge_rows)
        top_set_lex = set(_pair_row_map(lex_top).keys())
        top_set_sem = set(_pair_row_map(sem_top).keys())
        top_set_merge = set(_pair_row_map(merge_top).keys())

        universe = sorted(list(top_set_lex | top_set_sem | top_set_merge))
        n_u = len(universe)
        if n_u <= 0:
            out[g] = []
            continue

        rows_out: List[dict] = []
        for key in universe:
            id1, id2 = int(key[0]), int(key[1])
            rl = full_map_lex.get(key)
            rs = full_map_sem.get(key)
            rm = full_map_merge.get(key)

            rank_lex_native = None
            rank_sem_native = None
            rank_merge_native = None
            try:
                rank_lex_native = (None if rl is None else (int(rl.get("rank_final")) if rl.get("rank_final") is not None else None))
            except Exception:
                rank_lex_native = None
            try:
                rank_sem_native = (None if rs is None else (int(rs.get("rank_final")) if rs.get("rank_final") is not None else None))
            except Exception:
                rank_sem_native = None
            try:
                rank_merge_native = (None if rm is None else (int(rm.get("rank_final")) if rm.get("rank_final") is not None else None))
            except Exception:
                rank_merge_native = None

            score_lex_raw_native = None
            semantic_raw_native = None
            merge_score_final = None
            try:
                score_lex_raw_native = (None if rl is None else rl.get("score_lex_raw"))
            except Exception:
                score_lex_raw_native = None
            try:
                semantic_raw_native = (None if rs is None else rs.get("semantic_cosine"))
            except Exception:
                semantic_raw_native = None
            try:
                merge_score_final = (None if rm is None else rm.get("score_final"))
            except Exception:
                merge_score_final = None

            cross_lex = _first_non_none([
                (rm.get("cross_lex_score") if rm is not None else None),
                (rl.get("cross_lex_score") if rl is not None else None),
                (rs.get("cross_lex_score") if rs is not None else None),
            ])
            cross_sem = _first_non_none([
                (rm.get("cross_sem_score") if rm is not None else None),
                (rs.get("cross_sem_score") if rs is not None else None),
                (rl.get("cross_sem_score") if rl is not None else None),
            ])

            score_lex_ref = score_lex_raw_native if score_lex_raw_native is not None else cross_lex
            sem_ref = semantic_raw_native if semantic_raw_native is not None else cross_sem
            lex_source = "native" if score_lex_raw_native is not None else ("cross" if cross_lex is not None else "missing")
            sem_source = "native" if semantic_raw_native is not None else ("cross" if cross_sem is not None else "missing")

            rank_lex_u = int(rank_lex_native) if rank_lex_native is not None else int(n_u + 1)
            rank_sem_u = int(rank_sem_native) if rank_sem_native is not None else int(n_u + 1)
            rank_merge_u = int(rank_merge_native) if rank_merge_native is not None else int(n_u + 1)
            gap_ls_rank = int(abs(rank_lex_u - rank_sem_u))
            merge_mid = 0.5 * (float(rank_lex_u) + float(rank_sem_u))
            merge_mid_dev_rank = float(abs(float(rank_merge_u) - merge_mid))

            path1 = ""
            path2 = ""
            for rr in (rm, rl, rs):
                if rr is None:
                    continue
                p1 = str(rr.get("path1") or "").strip()
                p2 = str(rr.get("path2") or "").strip()
                if not path1 and p1:
                    path1 = p1
                if not path2 and p2:
                    path2 = p2
            if (not path1 or not path2) and units and id1 < len(units) and id2 < len(units):
                if not path1:
                    path1 = unit_path(units[id1])
                if not path2:
                    path2 = unit_path(units[id2])

            row = {
                "id1": id1,
                "id2": id2,
                "pair_key": f"{id1}-{id2}",
                "path1": path1,
                "path2": path2,
                "in_lex_topk": int(key in top_set_lex),
                "in_sem_topk": int(key in top_set_sem),
                "in_merge_topk": int(key in top_set_merge),
                "rank_lex": rank_lex_native,
                "rank_semantic": rank_sem_native,
                "rank_merge": rank_merge_native,
                "rank_lex_u": rank_lex_u,
                "rank_sem_u": rank_sem_u,
                "rank_merge_u": rank_merge_u,
                "score_lex_raw": score_lex_raw_native,
                "cross_lex_score": cross_lex,
                "lex_source": lex_source,
                "semantic_cosine": semantic_raw_native,
                "cross_sem_score": cross_sem,
                "sem_source": sem_source,
                "score_final": merge_score_final,
                "score_lex_ref": score_lex_ref,
                "semantic_ref": sem_ref,
                "gap_ls_rank": gap_ls_rank,
                "merge_mid_dev_rank": merge_mid_dev_rank,
                "reranker_raw": None,
                "reranker_ref_01": None,
                "rank_rerank_raw": None,
                "dist_r_lex": None,
                "dist_r_sem": None,
                "dist_r_merge": None,
                "audit_ref": None,
                "reranker_model": None,
            }
            rows_out.append(row)

        if reranker_enabled and rows_out and units:
            pairs: List[Tuple[str, str]] = []
            idxs: List[int] = []
            for i, rr in enumerate(rows_out):
                try:
                    id1 = int(rr["id1"])
                    id2 = int(rr["id2"])
                except Exception:
                    continue
                if id1 < 0 or id2 < 0 or id1 >= len(units) or id2 >= len(units):
                    continue
                idxs.append(i)
                pairs.append((
                    _truncate_for_rerank(units[id1].text or "", max_chars=1200),
                    _truncate_for_rerank(units[id2].text or "", max_chars=1200),
                ))

            if pairs:
                _log(f"[shadow/{g}] reranker candidates={len(pairs)}/{len(rows_out)} (universe U).")
                t_rer = time.perf_counter()
                raw_scores, model_used = _score_reranker_pairs(
                    pairs,
                    model_size=reranker_model,
                    prompt=reranker_prompt,
                )
                _log(f"[shadow/{g}] reranker scored {len(pairs)} pairs in {_fmt_elapsed(time.perf_counter() - t_rer)}.")
                if len(raw_scores) == len(idxs):
                    vals: List[Tuple[int, float]] = []
                    for pos, ridx in enumerate(idxs):
                        rv = float(raw_scores[pos])
                        rows_out[ridx]["reranker_raw"] = rv
                        rows_out[ridx]["reranker_model"] = str(model_used)
                        vals.append((int(ridx), rv))
                    rank_r = _rank_order(vals, descending=True)
                    ref_r = _rank_map(vals)
                    for ridx, _ in vals:
                        rr = rows_out[int(ridx)]
                        rr["reranker_ref_01"] = float(ref_r.get(int(ridx), 0.0)) if int(ridx) in ref_r else None
                        rrk = int(rank_r.get(int(ridx), 0)) if int(ridx) in rank_r else None
                        rr["rank_rerank_raw"] = rrk
                        if rrk is not None:
                            rr["dist_r_lex"] = abs(int(rrk) - int(rr["rank_lex_u"]))
                            rr["dist_r_sem"] = abs(int(rrk) - int(rr["rank_sem_u"]))
                            rr["dist_r_merge"] = abs(int(rrk) - int(rr["rank_merge_u"]))
                            dmap = {
                                "lexical": rr["dist_r_lex"],
                                "semantic": rr["dist_r_sem"],
                                "merge": rr["dist_r_merge"],
                            }
                            rr["audit_ref"] = min(dmap, key=lambda kx: float(dmap[kx]))

        rows_out.sort(
            key=lambda x: (
                int(x.get("gap_ls_rank") or 0),
                float(x.get("merge_mid_dev_rank") or 0.0),
                float(x.get("score_final") or 0.0),
            ),
            reverse=True,
        )
        out[g] = rows_out
    return out


def write_excel_shadow_audit(
    path: str,
    audit_by_g: Dict[str, List[dict]],
    semantic_model_label: Optional[str] = None,
    reranker_model_label: Optional[str] = None,
) -> None:
    """Write unified-universe shadow audit workbook (3 sheets)."""
    wb = openpyxl.Workbook()
    if wb.active is not None:
        wb.remove(wb.active)

    h_sem = _score_header_with_model("semantic_cosine", semantic_model_label)
    h_cross_sem = _score_header_with_model("cross_sem_score", semantic_model_label)
    h_rr = _score_header_with_model("reranker_raw", reranker_model_label)
    h_rr_ref = RERANK_REF_01_HEADER
    headers = [
        "pair_key", "id1", "id2", "path1", "path2",
        "in_lex_topk", "in_sem_topk", "in_merge_topk",
        "rank_lex", "rank_semantic", "rank_merge",
        "rank_lex_u", "rank_sem_u", "rank_merge_u",
        "score_lex_raw", "cross_lex_score", "lex_source",
        h_sem, h_cross_sem, "sem_source",
        "score_final",
        "gap_ls_rank", "merge_mid_dev_rank",
        h_rr, h_rr_ref, "rank_rerank_raw",
        "dist_r_lex", "dist_r_sem", "dist_r_merge", "audit_ref", "reranker_model",
    ]

    for g in ["h1", "h2", "normal"]:
        ws = wb.create_sheet(title=_excel_safe_sheet_name(f"audit_{g}", fallback=f"audit_{g}"))
        ws.append(headers)
        rows = list(audit_by_g.get(g) or [])
        for r in rows:
            ws.append([
                r.get("pair_key"), r.get("id1"), r.get("id2"), r.get("path1"), r.get("path2"),
                r.get("in_lex_topk"), r.get("in_sem_topk"), r.get("in_merge_topk"),
                r.get("rank_lex"), r.get("rank_semantic"), r.get("rank_merge"),
                r.get("rank_lex_u"), r.get("rank_sem_u"), r.get("rank_merge_u"),
                r.get("score_lex_raw"), r.get("cross_lex_score"), r.get("lex_source"),
                _fmt_score_with_model(r.get("semantic_cosine"), semantic_model_label),
                _fmt_score_with_model(r.get("cross_sem_score"), semantic_model_label),
                r.get("sem_source"),
                r.get("score_final"),
                r.get("gap_ls_rank"), r.get("merge_mid_dev_rank"),
                _fmt_score_with_model(r.get("reranker_raw"), reranker_model_label),
                r.get("reranker_ref_01"),
                r.get("rank_rerank_raw"),
                r.get("dist_r_lex"), r.get("dist_r_sem"), r.get("dist_r_merge"),
                r.get("audit_ref"), r.get("reranker_model"),
            ])

        for c, h in enumerate(headers, start=1):
            wcol = 10
            if h in ("id1", "id2", "in_lex_topk", "in_sem_topk", "in_merge_topk"):
                wcol = 8
            elif h in ("pair_key",):
                wcol = 12
            elif h in ("path1", "path2"):
                wcol = 42
            elif h in ("lex_source", "sem_source", "audit_ref"):
                wcol = 12
            elif h in ("reranker_model",):
                wcol = 24
            elif h in ("score_lex_raw", "cross_lex_score", "score_final", "merge_mid_dev_rank"):
                wcol = 14
            elif h.startswith("semantic_cosine") or h.startswith("cross_sem_score") or h.startswith("reranker_raw"):
                wcol = 24
            elif h == h_rr_ref:
                wcol = 14
            elif h in ("gap_ls_rank", "rank_rerank_raw", "dist_r_lex", "dist_r_sem", "dist_r_merge"):
                wcol = 12
            ws.column_dimensions[openpyxl.utils.get_column_letter(c)].width = wcol

        for c in range(1, len(headers) + 1):
            ws.cell(row=1, column=c).alignment = openpyxl.styles.Alignment(horizontal="center", vertical="center")
        ws.freeze_panes = "A2"

    wb.save(path)


def _table_page(
    pdf: Any,
    title: str,
    columns: List[str],
    data_rows: List[List[Any]],
    footnote: Optional[str] = None,
    col_widths: Optional[List[float]] = None,
) -> None:
    fig, ax = plt.subplots(figsize=(16, 9))
    fig.patch.set_facecolor("#FFF6E9")
    ax.set_facecolor("#FFFDF8")
    ax.axis("off")
    ax.set_title(title, fontsize=17, fontweight="bold", pad=14)

    if not data_rows:
        data_rows = [["-"] * max(1, len(columns))]

    text_rows = [[str(x) for x in rr] for rr in data_rows]
    note_text = str(footnote or "").strip()
    note_bottom = 0.02
    note_top = note_bottom
    if note_text:
        # Estimate wrapped line count so table and note do not overlap.
        est_lines = 0
        for line in note_text.splitlines() or [note_text]:
            est_lines += max(1, (len(line) + 57) // 58)
        note_h = min(0.26, 0.04 + 0.012 * est_lines)
        note_top = note_bottom + note_h

    table_left = 0.05
    table_width = 0.90
    table_top = 0.89
    table_floor = (note_top + 0.02) if note_text else 0.05
    table_available_h = max(0.45, table_top - table_floor)
    table_target_h = 0.80 if note_text else 0.84
    table_h = min(table_target_h, table_available_h)
    table_y = table_floor + max(0.0, (table_available_h - table_h) / 2.0)
    table_bbox = [table_left, table_y, table_width, table_h]

    widths = None
    if col_widths and len(col_widths) == len(columns):
        try:
            ws = [max(1e-6, float(w)) for w in col_widths]
            s = float(sum(ws))
            widths = ([w / s for w in ws] if s > 0 else None)
        except Exception:
            widths = None
    tab = ax.table(
        cellText=text_rows,
        colLabels=columns,
        bbox=table_bbox,
        cellLoc="center",
        colWidths=widths,
    )
    tab.auto_set_font_size(False)
    base_font = 9.4
    if len(columns) <= 7:
        base_font = 10.0
    elif len(columns) >= 9:
        base_font = 9.0
    tab.set_fontsize(base_font)
    tab.scale(1.0, 1.08)

    for (ri, ci), cell in tab.get_celld().items():
        cell.set_edgecolor("#C39A73")
        cell.set_linewidth(0.6)
        if ri == 0:
            cell.set_facecolor("#D1495B")
            cell.get_text().set_color("white")
            cell.get_text().set_fontweight("bold")
            cell.get_text().set_fontsize(max(8.4, base_font - 0.2))
        else:
            cell.set_facecolor("#FFFDF8" if (ri % 2 == 1) else "#FFF1DC")
        cell.get_text().set_ha("center")
        cell.get_text().set_va("center")

    if note_text:
        fig.text(
            0.02,
            note_bottom,
            note_text,
            ha="left",
            va="bottom",
            fontsize=9.0,
            linespacing=1.3,
            wrap=True,
            bbox=dict(boxstyle="round,pad=0.4", facecolor="#FFFDF8", edgecolor="#C39A73", linewidth=0.8),
        )
    pdf.savefig(fig, dpi=600)
    plt.close(fig)


def _audit_sorted_for_framework(audit_rows: List[dict], framework: str) -> List[dict]:
    fw = str(framework or "").strip().lower()
    if fw == "lexical":
        rk = "rank_lex_u"
        score_key = "score_lex_ref"
    elif fw == "semantic":
        rk = "rank_sem_u"
        score_key = "semantic_ref"
    else:
        rk = "rank_merge_u"
        score_key = "score_final"
    rows = list(audit_rows or [])
    rows.sort(
        key=lambda r: (
            int(r.get(rk) if r.get(rk) is not None else 10**9),
            -float(r.get(score_key) or 0.0),
        )
    )
    return rows


def _audit_anomaly_sorted_for_framework(audit_rows: List[dict], framework: str) -> List[dict]:
    fw = str(framework or "").strip().lower()
    if fw == "lexical":
        dist_key = "dist_r_lex"
        fallback_key = "gap_ls_rank"
    elif fw == "semantic":
        dist_key = "dist_r_sem"
        fallback_key = "gap_ls_rank"
    else:
        dist_key = "dist_r_merge"
        fallback_key = "merge_mid_dev_rank"

    def _metric(r: dict) -> float:
        v = r.get(dist_key)
        if v is not None:
            return float(v)
        return float(r.get(fallback_key) or 0.0)

    rows = list(audit_rows or [])
    rows.sort(
        key=lambda r: (
            _metric(r),
            float(r.get("gap_ls_rank") or 0.0),
            float(r.get("merge_mid_dev_rank") or 0.0),
        ),
        reverse=True,
    )
    return rows


def _audit_score_for_framework(r: dict, framework: str) -> Any:
    fw = str(framework or "").strip().lower()
    if fw == "lexical":
        return r.get("score_lex_ref")
    if fw == "semantic":
        return r.get("semantic_ref")
    return r.get("score_final")


def build_framework_all_top10_pdf(
    out_pdf: str,
    audit_rows_by_g: Dict[str, List[dict]],
    weights_h: Optional[Dict[str, float]] = None,
    weights_n: Optional[Dict[str, float]] = None,
    merge_w_lex: float = 0.70,
    merge_w_sem: float = 0.30,
) -> Optional[str]:
    """Build one PDF containing 18 pages: 9 unified-U Top10 + 9 cross-audit Top10."""
    if plt is None or PdfPages is None:
        return None
    _setup_plot_zh_font()
    os.makedirs(os.path.dirname(out_pdf), exist_ok=True)

    def _weights_for_g(gg: str) -> Dict[str, float]:
        if normalize_style(gg) in ("h1", "h2"):
            return _normalize_weights(weights_h or LEXICAL_WEIGHTS_DEFAULT_H, LEXICAL_WEIGHTS_DEFAULT_H)
        return _normalize_weights(weights_n or LEXICAL_WEIGHTS_DEFAULT_N, LEXICAL_WEIGHTS_DEFAULT_N)

    def _granularity_cn(gg: str) -> str:
        g = normalize_style(gg)
        if g == "h1":
            return "书名（h1）"
        if g == "h2":
            return "篇名（h2）"
        return "章节（normal）"

    def _note_total(fw: str, gg: str) -> str:
        g = normalize_style(gg)
        g_cn = _granularity_cn(g)
        if fw == "lexical":
            w = _weights_for_g(g)
            if g in ("h1", "h2"):
                return (
                    f"字面通道（Lexical）公式 | {g_cn}: "
                    "score_lex_raw=w_cos·cos+w_cov·cov+w_tfidf·tfidf+w_jacc·jaccard\n"
                    "字面通道内部保持原始分（raw），不做归一化（Normalization）\n"
                    f"权重: cos={w['cosine_tf_raw']:.2f}, cov={w['coverage']:.2f}, tfidf={w['cosine_tfidf']:.2f}, jacc={w['jaccard']:.2f} | "
                    "cov=证据覆盖（coverage）, cos=文体/体裁, tfidf=去噪并抑制高频词, jaccard=短文本辅助"
                )
            return (
                f"字面通道（Lexical）公式 | {g_cn}: "
                "score_lex_raw=w_cos·cos+w_tfidf·tfidf+w_jacc·jaccard\n"
                "字面通道内部保持原始分（raw），不做归一化（Normalization）\n"
                f"权重: cos={w['cosine_tf_raw']:.2f}, tfidf={w['cosine_tfidf']:.2f}, jacc={w['jaccard']:.2f} | "
                "cos=文体/体裁, tfidf=去噪并抑制高频词, jaccard=短文本辅助"
            )
        if fw == "semantic":
            if g in ("h1", "h2") and _semantic_maxsim_enabled(g):
                lam = SEMANTIC_MAXSIM_H1_LAMBDA if g == "h1" else SEMANTIC_MAXSIM_H2_LAMBDA
                return (
                    f"语义通道（Semantic）公式 | {g_cn}: "
                    "doc_cos=cosine(mean_vecA,mean_vecB), "
                    "maxsim_bi=0.5·(mean(max_row(S))+mean(max_col(S))), "
                    f"score_sem_raw={lam:.2f}·doc_cos+{1.0 - float(lam):.2f}·maxsim_bi；"
                    f"H2 仅对 Top{SEMANTIC_MAXSIM_H2_TOP_PERCENT:.1f}% 候选做 MaxSim 精排"
                    f"（{SEMANTIC_MAXSIM_H2_TOPN_MIN}~{SEMANTIC_MAXSIM_H2_TOPN_MAX}）并保留层内归一化排序。"
                )
            msg = (
                f"语义通道（Semantic）公式 | {g_cn}: "
                "score_sem_raw=clip(cosine(mean_vecA,mean_vecB),0,1)，"
                "score_sem=Norm(score_sem_raw)，仅按语义通道排序。"
            )
            if g == "h2":
                msg += " H2 当前配置为 avgpool（不启用 MaxSim）。"
            return msg
        wl = max(0.0, float(merge_w_lex))
        ws = max(0.0, float(merge_w_sem))
        den = wl + ws if (wl + ws) > 0 else 1.0
        wl_n = wl / den
        ws_n = ws / den
        q_lo, q_hi = _merge_quantile_bounds(g)
        norm_method = str(MERGE_NORM_METHOD or "").strip().lower()
        if norm_method == "minmax":
            norm_note = "q_lex=MinMax(score_lex_raw), q_sem=MinMax(score_sem_raw)"
        elif norm_method == "robust_hard":
            norm_note = (
                f"q_lex=RobustNormHard(score_lex_raw;q{int(round(q_lo*100))}/q{int(round(q_hi*100))}), "
                f"q_sem=RobustNormHard(score_sem_raw;q{int(round(q_lo*100))}/q{int(round(q_hi*100))})"
            )
        elif norm_method == "robust_soft":
            sc_k = _merge_softclip_k(g)
            norm_note = (
                f"q_lex=RobustSoft(score_lex_raw;q{int(round(q_lo*100))}/q{int(round(q_hi*100))},k={sc_k:.2f}), "
                f"q_sem=RobustSoft(score_sem_raw;q{int(round(q_lo*100))}/q{int(round(q_hi*100))},k={sc_k:.2f})"
            )
        else:
            ts = _merge_tail_slope(g)
            norm_note = (
                f"q_lex=RobustLinear(score_lex_raw;q{int(round(q_lo*100))}/q{int(round(q_hi*100))},tail={ts:.2f}), "
                f"q_sem=RobustLinear(score_sem_raw;q{int(round(q_lo*100))}/q{int(round(q_hi*100))},tail={ts:.2f})"
            )
        return (
            f"融合通道（Merge）公式 | {g_cn}: 仅此一步做归一化（Normalization）"
            f"{norm_note}；"
            f"score={wl_n:.2f}·q_lex+{ws_n:.2f}·q_sem。"
        )

    def _fmt_path_short(s: str, max_len: int = 26) -> str:
        t = str(s or "").replace("\n", " ").strip()
        if len(t) <= max_len:
            return t
        return t[: max(1, max_len - 3)] + "..."

    with PdfPages(out_pdf) as pdf:
        # 9 pages: unified-U channel top10
        for fw, fw_cn in FRAMEWORK_ORDER:
            for g in ["h1", "h2", "normal"]:
                g_cn = _granularity_cn(g)
                audit_rows = list(audit_rows_by_g.get(g) or [])
                rows = _audit_sorted_for_framework(audit_rows, fw)[:10]
                table_rows: List[List[Any]] = []
                for idx, r in enumerate(rows, start=1):
                    pa = _fmt_path_short(str(r.get("path1") or ""), max_len=24)
                    pb = _fmt_path_short(str(r.get("path2") or ""), max_len=24)
                    score_ref = _audit_score_for_framework(r, fw)
                    table_rows.append([
                        idx,
                        f"{r.get('id1')}-{r.get('id2')}",
                        pa,
                        pb,
                        ("n/a" if score_ref is None else f"{float(score_ref):.4f}"),
                        r.get("rank_lex_u"),
                        r.get("rank_sem_u"),
                        r.get("rank_merge_u"),
                    ])
                _table_page(
                    pdf,
                    title=f"{fw_cn} | {g_cn} | 候选集合U通道 Top10",
                    columns=["序号(rank)", "ID对(pair)", "A路径", "B路径", "通道分(score)", "字面名次(r_lex)", "语义名次(r_sem)", "融合名次(r_merge)"],
                    data_rows=table_rows,
                    footnote=(
                        f"{_note_total(fw, g)}\n"
                        "候选集合（Candidate Set）: U = TopK(字面) ∪ TopK(语义) ∪ TopK(融合)。"
                    ),
                    # Give path columns more room; shrink rank/id and tail numeric columns.
                    col_widths=[0.06, 0.09, 0.24, 0.24, 0.11, 0.09, 0.09, 0.08],
                )

        # 9 pages: cross-audit top10
        for fw, fw_cn in FRAMEWORK_ORDER:
            for g in ["h1", "h2", "normal"]:
                g_cn = _granularity_cn(g)
                audit_rows = list(audit_rows_by_g.get(g) or [])
                rows = _audit_anomaly_sorted_for_framework(audit_rows, fw)[:10]
                table_rows = []
                for idx, r in enumerate(rows, start=1):
                    if fw == "lexical":
                        dist_key = "dist_r_lex"
                    elif fw == "semantic":
                        dist_key = "dist_r_sem"
                    else:
                        dist_key = "dist_r_merge"
                    dist_val = r.get(dist_key)
                    if dist_val is None:
                        dist_val = r.get("gap_ls_rank")
                    table_rows.append([
                        idx,
                        f"{r.get('id1')}-{r.get('id2')}",
                        _fmt_path_short(str(r.get("path1") or ""), max_len=30),
                        _fmt_path_short(str(r.get("path2") or ""), max_len=30),
                        r.get("gap_ls_rank"),
                        f"{float(r.get('merge_mid_dev_rank') or 0.0):.3f}",
                        ("无" if dist_val is None else f"{float(dist_val):.3f}"),
                        ("无" if r.get("reranker_raw") is None else f"{float(r.get('reranker_raw')):.4f}"),
                        (r.get("audit_ref") or "无"),
                    ])
                _table_page(
                    pdf,
                    title=f"{fw_cn} | {g_cn} | 交叉审计 Top10",
                    columns=[
                        "审计序号(audit_rank)",
                        "ID对(pair)",
                        "A路径",
                        "B路径",
                        "字面语义差",
                        "融合中位偏差",
                        "与重排差",
                        "重排分",
                        "审计参考",
                    ],
                    data_rows=table_rows,
                    footnote=(
                        "交叉审计算法（Cross Audit）\n"
                        "公式：\n"
                        "1) 字面语义差 = |rank_lex_u - rank_sem_u|\n"
                        "2) 融合中位偏差 = |rank_merge_u - mid(rank_lex_u, rank_sem_u)|\n"
                        "3) 与重排差 = |rank_rerank_raw - rank_{当前通道}|\n\n"
                        "解释：\n"
                        "字面语义差越大，表示字面与语义判断分歧越强；"
                        "融合中位偏差越大，表示融合名次偏离两通道共识更明显；"
                        "与重排差越大，表示重排模型与当前通道排序差异更大。"
                    ),
                    # Keep A/B path dominant while preserving readable numeric audit columns.
                    col_widths=[0.07, 0.10, 0.21, 0.21, 0.08, 0.10, 0.08, 0.08, 0.07],
                )

    return out_pdf


def build_framework_all_summary_image(
    out_img: str,
    audit_rows_by_g: Dict[str, List[dict]],
    dpi: int = 450,
) -> Optional[str]:
    """Build one A4 portrait summary image (3x3 cells).

    Columns: lexical / semantic / merge
    Rows: 书名（h1） / 篇名（h2） / 章节（normal）
    Each cell includes:
      - 统一U 通道Top10
      - 交叉审计 Top10
    with compact columns: rank / path / score.
    """
    if plt is None:
        return None
    _setup_plot_zh_font()
    os.makedirs(os.path.dirname(out_img), exist_ok=True)

    import matplotlib.patches as patches

    def _fmt_path_short(s: str, max_len: int = 18) -> str:
        t = str(s or "").replace("\n", " ").strip()
        if len(t) <= max_len:
            return t
        return t[: max(1, max_len - 3)] + "..."

    def _fmt_num(v: Any, d: int = 3) -> str:
        try:
            return f"{float(v):.{d}f}"
        except Exception:
            return "n/a"

    def _pair_path_audit(r: dict, max_side: int = 10) -> str:
        pa = str(r.get("path1") or "")
        pb = str(r.get("path2") or "")
        return f"{_fmt_path_short(pa, max_side)}<->{_fmt_path_short(pb, max_side)}"

    fw_color = {
        "lexical": "#C44536",
        "semantic": "#2F6EA8",
        "merge": "#7A4A2C",
    }

    fig = plt.figure(figsize=(8.27, 11.69))
    fig.patch.set_facecolor("#FFF7ED")
    gs = fig.add_gridspec(
        3, 3,
        left=0.025, right=0.975, top=0.945, bottom=0.03,
        wspace=0.03, hspace=0.045,
    )

    fig.suptitle("字面 / 语义 / 融合 × H1/H2/Normal 交叉审计总览（Top10）", fontsize=17, fontweight="bold", y=0.975)
    fig.text(
        0.5,
        0.955,
        "每格包含：统一U通道Top10 + 交叉审计Top10（rank | path | score）",
        ha="center",
        va="center",
        fontsize=10.2,
        color="#5A4A3A",
    )

    for ri, g in enumerate(["h1", "h2", "normal"]):
        for ci, (fw, fw_cn) in enumerate(FRAMEWORK_ORDER):
            ax = fig.add_subplot(gs[ri, ci])
            ax.axis("off")

            box = patches.FancyBboxPatch(
                (0.005, 0.005), 0.99, 0.99,
                boxstyle="round,pad=0.008,rounding_size=0.015",
                transform=ax.transAxes,
                facecolor="#FFFEFC",
                edgecolor="#CFAE8E",
                linewidth=0.9,
            )
            ax.add_patch(box)

            ax.text(
                0.02, 0.975,
                f"{fw_cn} | {g.upper()}",
                transform=ax.transAxes,
                ha="left", va="top",
                fontsize=11.2, fontweight="bold",
                color=fw_color.get(fw, "#333333"),
            )

            audit_rows = list(audit_rows_by_g.get(g) or [])
            rows_total = _audit_sorted_for_framework(audit_rows, fw)[:10]
            rows_audit = _audit_anomaly_sorted_for_framework(audit_rows, fw)[:10]

            lines: List[str] = []
            lines.append("统一U Top10")
            lines.append("rk | path | score")
            for j, r in enumerate(rows_total, start=1):
                ptxt = _pair_path_audit(r, max_side=8)
                vtxt = _fmt_num(_audit_score_for_framework(r, fw), 3)
                lines.append(f"{j:>2} | {ptxt} | {vtxt}")

            lines.append("")
            lines.append("交叉审计 Top10")
            lines.append("rk | path | score")
            for j, r in enumerate(rows_audit, start=1):
                ptxt = _pair_path_audit(r, max_side=8)
                if fw == "lexical":
                    vtxt = _fmt_num(r.get("dist_r_lex"), 3)
                elif fw == "semantic":
                    vtxt = _fmt_num(r.get("dist_r_sem"), 3)
                else:
                    vtxt = _fmt_num(r.get("dist_r_merge"), 3)
                lines.append(f"{j:>2} | {ptxt} | {vtxt}")

            ax.text(
                0.50, 0.50,
                "\n".join(lines),
                transform=ax.transAxes,
                ha="center", va="center",
                fontsize=5.8,
                color="#2F2B28",
                linespacing=1.18,
            )

    fig.savefig(out_img, dpi=max(220, int(dpi)), bbox_inches="tight", pad_inches=0.04)
    plt.close(fig)
    return out_img


def build_framework_merge_summary_image(
    out_img: str,
    audit_rows_by_g: Dict[str, List[dict]],
    dpi: int = 450,
) -> Optional[str]:
    """Build one A4 portrait summary image for merge only (h1/h2/normal)."""
    if plt is None:
        return None
    _setup_plot_zh_font()
    os.makedirs(os.path.dirname(out_img), exist_ok=True)

    import matplotlib.patches as patches

    def _fmt_path_short(s: str, max_len: int = 20) -> str:
        t = str(s or "").replace("\n", " ").strip()
        if len(t) <= max_len:
            return t
        return t[: max(1, max_len - 3)] + "..."

    def _fmt_num(v: Any, d: int = 3) -> str:
        try:
            return f"{float(v):.{d}f}"
        except Exception:
            return "n/a"

    def _pair_path_audit(r: dict, max_side: int = 12) -> str:
        pa = str(r.get("path1") or "")
        pb = str(r.get("path2") or "")
        return f"{_fmt_path_short(pa, max_side)}<->{_fmt_path_short(pb, max_side)}"

    fig = plt.figure(figsize=(8.27, 11.69))
    fig.patch.set_facecolor("#FFF7ED")
    gs = fig.add_gridspec(
        3, 1,
        left=0.035, right=0.965, top=0.945, bottom=0.03,
        wspace=0.02, hspace=0.05,
    )

    fig.suptitle("融合榜 × H1/H2/Normal 交叉审计总览（Top10）", fontsize=18, fontweight="bold", y=0.975)
    fig.text(
        0.5,
        0.955,
        "每格包含：融合统一U Top10 + 融合交叉审计 Top10（rank | path | score）",
        ha="center",
        va="center",
        fontsize=11.0,
        color="#5A4A3A",
    )

    for ri, g in enumerate(["h1", "h2", "normal"]):
        ax = fig.add_subplot(gs[ri, 0])
        ax.axis("off")

        box = patches.FancyBboxPatch(
            (0.005, 0.005), 0.99, 0.99,
            boxstyle="round,pad=0.008,rounding_size=0.015",
            transform=ax.transAxes,
            facecolor="#FFFEFC",
            edgecolor="#CFAE8E",
            linewidth=1.0,
        )
        ax.add_patch(box)

        ax.text(
            0.02, 0.975,
            f"融合 | {g.upper()}",
            transform=ax.transAxes,
            ha="left", va="top",
            fontsize=12.8, fontweight="bold",
            color="#7A4A2C",
        )

        audit_rows = list(audit_rows_by_g.get(g) or [])
        rows_total = _audit_sorted_for_framework(audit_rows, "merge")[:10]
        rows_audit = _audit_anomaly_sorted_for_framework(audit_rows, "merge")[:10]

        lines: List[str] = []
        lines.append("融合统一U Top10")
        lines.append("rk | path | score")
        for j, r in enumerate(rows_total, start=1):
            ptxt = _pair_path_audit(r, max_side=12)
            vtxt = _fmt_num(_audit_score_for_framework(r, "merge"), 3)
            lines.append(f"{j:>2} | {ptxt} | {vtxt}")

        lines.append("")
        lines.append("融合交叉审计 Top10")
        lines.append("rk | path | score")
        for j, r in enumerate(rows_audit, start=1):
            ptxt = _pair_path_audit(r, max_side=12)
            vtxt = _fmt_num(r.get("dist_r_merge"), 3)
            lines.append(f"{j:>2} | {ptxt} | {vtxt}")

        ax.text(
            0.50, 0.50,
            "\n".join(lines),
            transform=ax.transAxes,
            ha="center", va="center",
            fontsize=7.2,
            color="#2F2B28",
            linespacing=1.20,
        )

    fig.savefig(out_img, dpi=max(220, int(dpi)), bbox_inches="tight", pad_inches=0.04)
    plt.close(fig)
    return out_img


def write_word_report_framework_all(
    path: str,
    results_by_fw: Dict[str, Dict[str, Dict[str, Any]]],
    title: str,
    top_rows_per_mode: int = 50,
    comparison_plot_paths: Optional[List[str]] = None,
) -> None:
    """One combined Word for framework=all."""
    doc = Document()
    _apply_bnu_word_styles(doc)
    doc.add_paragraph(title, style="Heading 1")
    doc.add_paragraph(f"Generated at: {time.strftime('%Y-%m-%d %H:%M:%S')}", style="Normal")

    # 比较图放在最前
    for p in (comparison_plot_paths or []):
        if p and os.path.exists(p):
            try:
                doc.add_picture(p, width=Inches(7.8))
                doc.add_paragraph(style="Normal")
            except Exception:
                pass

    # 字数统计最先（正文段落中）
    ref_g = (results_by_fw.get("merge") or results_by_fw.get("lexical") or results_by_fw.get("semantic") or {})
    write_wordcount_section_all(doc, ref_g)

    for g in ["h1", "h2", "normal"]:
        doc.add_paragraph(f"{g.upper()} Top50（字面 / 语义 / 融合）", style="Heading 2")
        for fw, fw_cn in FRAMEWORK_ORDER:
            item = (results_by_fw.get(fw) or {}).get(g) or {}
            units = item.get("units") or []
            rows = (item.get("rows") or [])[: max(0, int(top_rows_per_mode))]
            doc.add_paragraph(f"{fw_cn} Top {len(rows)}", style="Heading 3")
            for i, r in enumerate(rows, start=1):
                try:
                    id1 = int(r["id1"])
                    id2 = int(r["id2"])
                except Exception:
                    continue
                if id1 < 0 or id2 < 0 or id1 >= len(units) or id2 >= len(units):
                    continue
                u1 = units[id1]
                u2 = units[id2]
                score = float(r.get("score_final") or 0.0)
                semq = float(r.get("semantic_q") or 0.0)
                rer = r.get("score_reranker_raw")
                rer_txt = f"{float(rer):.4f}" if rer is not None else "n/a"
                doc.add_paragraph(
                    f"#{i} score={score:.4f} sem_q={semq:.4f} reranker_raw={rer_txt} ids={id1}-{id2}",
                    style="Normal",
                )
                pm = doc.add_paragraph(style="Normal")
                pm.add_run("A: ")
                pm.add_run(str(r.get("path1") or unit_path(u1)))
                pm.add_run("\nB: ")
                pm.add_run(str(r.get("path2") or unit_path(u2)))
                if g == "normal":
                    doc.add_paragraph("A 正文（预处理：仅汉字）", style="Normal")
                    doc.add_paragraph(preprocess_keep_han(u1.text or ""), style="Normal")
                    doc.add_paragraph("B 正文（预处理：仅汉字）", style="Normal")
                    doc.add_paragraph(preprocess_keep_han(u2.text or ""), style="Normal")
                doc.add_paragraph(style="Normal")

    # Display-only collation section: merge + normal Top-N.
    merge_normal = ((results_by_fw.get("merge") or {}).get("normal") or {})
    if merge_normal:
        write_normal_collation_section(
            doc,
            units=merge_normal.get("units") or [],
            rows=merge_normal.get("rows") or [],
            topn=int(NORMAL_COLLATION_DEFAULTS["topn_merge_normal"]),
            detail_norm_edit_max=float(NORMAL_COLLATION_DEFAULTS["detail_norm_edit_max"]),
            skip_han_exact=bool(NORMAL_COLLATION_DEFAULTS["skip_han_exact"]),
            max_variant_items_per_pair=int(NORMAL_COLLATION_DEFAULTS["max_variant_items_per_pair"]),
        )

    doc.save(path)


def write_word_report_all(
    path: str,
    results_by_g: Dict[str, Dict[str, Any]],
    title: str,
    top_rows_per_g: int = 30,
    h2_max_chars: int = 1000,
    plot_paths: Optional[Dict[str, str]] = None,
    framework: str = "",
):
    """Combined Word report for h1/h2/normal top rows."""
    doc = Document()
    _apply_bnu_word_styles(doc)
    doc.add_paragraph(title, style="Heading 1")
    doc.add_paragraph(f"Generated at: {time.strftime('%Y-%m-%d %H:%M:%S')}", style="Normal")

    p0 = doc.add_paragraph(style="Normal")
    p0.add_run("This all-report summarizes top pairs for h1 / h2 / normal in one document. ")
    p0.add_run("For h1/h2 sections,正文 is hidden; normal shows full Han text.")

    # Put wordcount stats first for all granularity levels.
    write_wordcount_section_all(doc, results_by_g)

    for g in ["h1", "h2", "normal"]:
        item = results_by_g.get(g)
        if not item:
            continue
        units = item.get("units") or []
        rows = item.get("rows") or []
        show = rows[: max(0, int(top_rows_per_g))]

        doc.add_paragraph(f"{g.upper()} Top {len(show)}", style="Heading 2")
        doc.add_paragraph(
            f"Total pairs={len(rows)}; shown={len(show)}; score_mode={show[0].get('score_mode') if show else 'n/a'}",
            style="Normal",
        )
        pp = (plot_paths or {}).get(g)
        if pp and os.path.exists(pp):
            try:
                doc.add_picture(pp, width=Inches(7.8))
                doc.add_paragraph(style="Normal")
            except Exception:
                pass

        for i, r in enumerate(show, start=1):
            u1 = units[int(r["id1"])]
            u2 = units[int(r["id2"])]
            han1 = preprocess_keep_han(u1.text or "")
            han2 = preprocess_keep_han(u2.text or "")
            len1 = len(han1)
            len2 = len(han2)
            score = float(r.get("score_final") if r.get("score_final") is not None else (r.get("cosine_tf") or 0.0))
            cos_raw = float(r.get("cosine_tf_raw") or 0.0)
            sem_q = float(r.get("semantic_q") or 0.0)
            rer_raw = r.get("score_reranker_raw")
            rer_txt = f"{float(rer_raw):.4f}" if rer_raw is not None else "n/a"
            jac = float(r.get("jaccard") or 0.0)
            show_cov_meta = (normalize_style(g) in ("h1", "h2"))
            cov_f1 = (r.get("coverage_f1") if show_cov_meta else None)
            cov_info = ""
            if cov_f1 is not None:
                cov_info = f" covF1={float(cov_f1):.4f}"

            doc.add_paragraph(
                f"#{i} score={score:.4f} sem_q={sem_q:.4f} reranker_raw={rer_txt} cos_raw={cos_raw:.4f}{cov_info} jaccard={jac:.4f} ids={r['id1']}-{r['id2']} HanChars(A-B)={len1}-{len2}",
                style="Heading 3",
            )

            pm = doc.add_paragraph(style="Normal")
            pm.add_run("A: ")
            pm.add_run(str(r.get("path1") or unit_path(u1)))
            pm.add_run("\nB: ")
            pm.add_run(str(r.get("path2") or unit_path(u2)))
            hits = (str(r.get("coverage_hits") or "").strip() if show_cov_meta else "")
            if hits:
                pm.add_run("\ncoverage_hits: ")
                pm.add_run(hits)

            show1 = han1
            show2 = han2
            if g == "normal":
                doc.add_paragraph("A 正文（预处理：仅汉字）", style="Normal")
                doc.add_paragraph(show1, style="Normal")
                doc.add_paragraph("B 正文（预处理：仅汉字）", style="Normal")
                doc.add_paragraph(show2, style="Normal")
            doc.add_paragraph(style="Normal")

    if str(framework or "").strip().lower() == "merge":
        item_n = results_by_g.get("normal") or {}
        write_normal_collation_section(
            doc,
            units=item_n.get("units") or [],
            rows=item_n.get("rows") or [],
            topn=int(NORMAL_COLLATION_DEFAULTS["topn_merge_normal"]),
            detail_norm_edit_max=float(NORMAL_COLLATION_DEFAULTS["detail_norm_edit_max"]),
            skip_han_exact=bool(NORMAL_COLLATION_DEFAULTS["skip_han_exact"]),
            max_variant_items_per_pair=int(NORMAL_COLLATION_DEFAULTS["max_variant_items_per_pair"]),
        )

    doc.save(path)


def write_excel_report_all(
    path: str,
    results_by_g: Dict[str, Dict[str, Any]],
    top_rows_per_g: Optional[int] = None,
    weights_h: Optional[Dict[str, float]] = None,
    weights_n: Optional[Dict[str, float]] = None,
    framework: str = "merge",
    semantic_model_label: Optional[str] = None,
    reranker_model_label: Optional[str] = None,
):
    """Combined Excel with 3 sheets: h1/h2/normal.

    top_rows_per_g=None means output all rows for each sheet.
    """
    wb = openpyxl.Workbook()
    if wb.active is not None:
        wb.remove(wb.active)

    for g in ["h1", "h2", "normal"]:
        item = results_by_g.get(g)
        if not item:
            continue
        units = item.get("units") or []
        rows = (item.get("rows") or [])
        if top_rows_per_g is not None:
            rows = rows[: max(0, int(top_rows_per_g))]
        fw_cn = _framework_cn(framework)
        if framework == "lexical":
            model_tag = "Lexical"
        elif framework == "semantic":
            model_tag = semantic_model_label or "SemanticModel"
        else:
            if semantic_model_label and reranker_model_label:
                model_tag = f"{semantic_model_label}+{reranker_model_label}"
            else:
                model_tag = semantic_model_label or reranker_model_label or "MergeModel"
        ws = wb.create_sheet(title=_excel_safe_sheet_name(f"{fw_cn}_{g}_{model_tag}", fallback=f"{fw_cn}_{g}"))
        _write_excel_all_like_sheet(
            ws,
            framework=framework,
            granularity=g,
            units=units,
            rows=rows,
            weights_h=weights_h,
            weights_n=weights_n,
            semantic_model_label=semantic_model_label,
            reranker_model_label=reranker_model_label,
        )

    wb.save(path)
# =========================
# Integrated wordcount section (in the SAME Word report)
# =========================

def _unit_len_chars(u: Unit) -> int:
    """Count chars for a unit (Han-only after preprocessing)."""
    return _eff_len_chars(u.text or "")


def _summarize_lengths_by_key(units: List[Unit], key_fn) -> List[Tuple[Tuple, int, float, int]]:
    """Return list of (key, total_chars, avg_chars_per_unit, unit_count) sorted by total desc."""
    groups: Dict[Tuple, List[int]] = {}
    for u in units:
        L = _unit_len_chars(u)
        if L <= 0:
            continue
        groups.setdefault(key_fn(u), []).append(L)

    out: List[Tuple[Tuple, int, float, int]] = []
    for k, vals in groups.items():
        tot = int(sum(vals))
        cnt = int(len(vals))
        out.append((k, tot, float(tot / max(1, cnt)), cnt))
    out.sort(key=lambda x: x[1], reverse=True)
    return out


def write_wordcount_section(doc: Document, normal_units: List[Unit], show_top: int = 20) -> None:
    """Append wordcount statistics section.

    Requirements:
      - Show wordcount summary for three partitions: h1 / h2 / normal
      - Report total chars + average chars per unit
    Char-count definition: keep Han characters only
    (remove punctuation/special symbols/whitespace), then count.
    """
    doc.add_paragraph("字数统计", style="Heading 1")

    lens = [_unit_len_chars(u) for u in normal_units]
    lens_nonzero = [x for x in lens if x > 0]
    if not lens_nonzero:
        doc.add_paragraph("未检测到可统计的正文内容（全部为空或仅空白）。", style="Normal")
        return

    total = int(sum(lens_nonzero))
    mean = float(np.mean(lens_nonzero))
    p = doc.add_paragraph(style="Normal")
    p.add_run(f"normal units: {len(lens_nonzero)}；")
    p.add_run(f"Total chars: {total}；")
    p.add_run(f"Avg chars / normal-unit: {mean:.2f}。")

    # ---- h1 summary ----
    doc.add_paragraph("按 h1（书名）汇总", style="Heading 2")
    s_h1 = _summarize_lengths_by_key(normal_units, lambda u: (u.h1.strip() or "(empty)",))
    p1 = doc.add_paragraph(style="Normal")
    p1.add_run(f"h1 units: {len(s_h1)}；")
    p1.add_run(f"Total chars: {sum(x[1] for x in s_h1)}；")
    p1.add_run(f"Avg chars / h1-unit: {(sum(x[1] for x in s_h1)/max(1,len(s_h1))):.2f}")

    t1 = doc.add_table(rows=1, cols=5)
    h = t1.rows[0].cells
    h[0].text = "rank"
    h[1].text = "h1"
    h[2].text = "normal_units"
    h[3].text = "total_chars"
    h[4].text = "avg_chars"
    for i, (k, tot, avg, cnt) in enumerate(s_h1[: min(200, len(s_h1))]):
        r = t1.add_row().cells
        r[0].text = str(i + 1)
        r[1].text = k[0]
        r[2].text = str(cnt)
        r[3].text = str(tot)
        r[4].text = f"{avg:.2f}"

    doc.add_paragraph(style="Normal")

    # ---- h2 summary ----
    doc.add_paragraph("按 h2（书名+篇名）汇总", style="Heading 2")
    s_h2 = _summarize_lengths_by_key(normal_units, lambda u: (u.h1.strip() or "(empty)", u.h2.strip() or "(empty)"))
    p2 = doc.add_paragraph(style="Normal")
    p2.add_run(f"h2 units: {len(s_h2)}；")
    p2.add_run(f"Total chars: {sum(x[1] for x in s_h2)}；")
    p2.add_run(f"Avg chars / h2-unit: {(sum(x[1] for x in s_h2)/max(1,len(s_h2))):.2f}")

    t2 = doc.add_table(rows=1, cols=6)
    h = t2.rows[0].cells
    h[0].text = "rank"
    h[1].text = "h1"
    h[2].text = "h2"
    h[3].text = "normal_units"
    h[4].text = "total_chars"
    h[5].text = "avg_chars"
    for i, (k, tot, avg, cnt) in enumerate(s_h2[: min(200, len(s_h2))]):
        r = t2.add_row().cells
        r[0].text = str(i + 1)
        r[1].text = k[0]
        r[2].text = k[1]
        r[3].text = str(cnt)
        r[4].text = str(tot)
        r[5].text = f"{avg:.2f}"

    doc.add_paragraph(style="Normal")

    # ---- Top longest normal units ----
    doc.add_paragraph(f"Top {min(show_top, len(lens_nonzero))} 最长 normal 单位（按字数降序）", style="Heading 2")
    pairs = [(i, lens[i]) for i in range(len(normal_units)) if lens[i] > 0]
    pairs.sort(key=lambda x: x[1], reverse=True)

    t3 = doc.add_table(rows=1, cols=4)
    hdr = t3.rows[0].cells
    hdr[0].text = "rank"
    hdr[1].text = "row_id"
    hdr[2].text = "chars"
    hdr[3].text = "path"

    m = min(show_top, len(pairs))
    for i in range(m):
        idx, L = pairs[i]
        u = normal_units[idx]
        rr = t3.add_row().cells
        rr[0].text = str(i + 1)
        rr[1].text = str(idx)
        rr[2].text = str(int(L))
        rr[3].text = unit_path(u)

    doc.add_paragraph(style="Normal")


def export_units_jsonl(path: str, units: List[Unit]):
    with open(path, "w", encoding="utf-8") as f:
        for u in units:
            obj = {
                "uid": u.uid,
                "level": u.level,
                "h1": u.h1,
                "h2": u.h2,
                "h3": u.h3,
                "text": u.text,
                "para_ids": u.para_ids or [],
            }
            f.write(json.dumps(obj, ensure_ascii=False) + "\n")


def _json_safe_value(v: Any) -> Any:
    if v is None:
        return None
    if isinstance(v, (str, bool, int)):
        return v
    if isinstance(v, float):
        return v if math.isfinite(v) else None
    if isinstance(v, np.integer):
        return int(v)
    if isinstance(v, np.floating):
        fv = float(v)
        return fv if math.isfinite(fv) else None
    if isinstance(v, np.bool_):
        return bool(v)
    if isinstance(v, (list, tuple)):
        return [_json_safe_value(x) for x in v]
    if isinstance(v, dict):
        return {str(k): _json_safe_value(val) for k, val in v.items()}
    if hasattr(v, "isoformat"):
        try:
            return v.isoformat()
        except Exception:
            pass
    return str(v)


def _dedup_headers(headers: List[Any]) -> List[str]:
    used: Dict[str, int] = {}
    out: List[str] = []
    for i, h in enumerate(headers):
        base = str(h).strip() if h is not None else ""
        if not base:
            base = f"col_{i + 1}"
        cnt = int(used.get(base, 0)) + 1
        used[base] = cnt
        out.append(base if cnt == 1 else f"{base}__{cnt}")
    return out


def export_excel_machine_json(excel_path: str, json_path: str) -> Dict[str, int]:
    """Export full Excel workbook content into JSON for machine reading."""
    wb = openpyxl.load_workbook(excel_path, read_only=True, data_only=True)
    payload = {
        "source_excel": os.path.abspath(str(excel_path)),
        "generated_at": time.strftime("%Y-%m-%d %H:%M:%S"),
        "sheets": [],
    }
    total_rows = 0
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        it = ws.iter_rows(values_only=True)
        head = next(it, None)
        if head is None:
            payload["sheets"].append({
                "name": str(sheet_name),
                "columns": [],
                "row_count": 0,
                "rows": [],
            })
            continue

        cols = _dedup_headers(list(head))
        rows: List[Dict[str, Any]] = []
        for rr in it:
            if rr is None:
                continue
            rec: Dict[str, Any] = {}
            non_empty = False
            for i, k in enumerate(cols):
                v = rr[i] if i < len(rr) else None
                sv = _json_safe_value(v)
                rec[k] = sv
                if sv is not None and not (isinstance(sv, str) and sv == ""):
                    non_empty = True
            if non_empty:
                rows.append(rec)

        payload["sheets"].append({
            "name": str(sheet_name),
            "columns": cols,
            "row_count": len(rows),
            "rows": rows,
        })
        total_rows += len(rows)
    wb.close()

    os.makedirs(os.path.dirname(json_path), exist_ok=True)
    with open(json_path, "w", encoding="utf-8") as f:
        json.dump(payload, f, ensure_ascii=False)
    return {"sheet_count": len(payload["sheets"]), "row_count": int(total_rows)}


def _unit_export_key(u: Unit) -> Tuple[str, Any, str, str, str]:
    lvl = normalize_style(getattr(u, "level", "normal"))
    try:
        uid = int(getattr(u, "uid", 0))
    except Exception:
        uid = str(getattr(u, "uid", ""))
    return (
        lvl,
        uid,
        str(getattr(u, "h1", "") or ""),
        str(getattr(u, "h2", "") or ""),
        str(getattr(u, "h3", "") or ""),
    )


def _merge_units_unique(unit_groups: List[List[Unit]]) -> List[Unit]:
    seen = set()
    out: List[Unit] = []
    for group in unit_groups:
        for u in (group or []):
            k = _unit_export_key(u)
            if k in seen:
                continue
            seen.add(k)
            out.append(u)
    return out


def _collect_units_all_from_results_by_g(results_by_g: Dict[str, Dict[str, Any]]) -> List[Unit]:
    groups: List[List[Unit]] = []
    for g in ("h1", "h2", "normal"):
        item = results_by_g.get(g) or {}
        groups.append(list(item.get("units") or []))
    return _merge_units_unique(groups)


def _collect_units_all_from_results_by_fw(results_by_fw: Dict[str, Dict[str, Dict[str, Any]]]) -> List[Unit]:
    groups: List[List[Unit]] = []
    for g in ("h1", "h2", "normal"):
        groups.append(list(_pick_units_for_g(results_by_fw, g) or []))
    return _merge_units_unique(groups)


# =========================
# Evaluation (P0)
# =========================
def _norm_col(s: str) -> str:
    return str(s or "").strip().lower().replace(" ", "").replace("_", "")


def _read_table_records(path: str, sheet_name: Optional[str] = None) -> Tuple[List[str], List[Dict[str, Any]]]:
    """Read csv/xlsx records and return (normalized_headers, rows)."""
    ext = os.path.splitext(path)[1].lower()
    if ext == ".csv":
        rows: List[Dict[str, Any]] = []
        with open(path, "r", encoding="utf-8-sig", newline="") as f:
            reader = csv.DictReader(f)
            headers_raw = list(reader.fieldnames or [])
            headers_norm = [_norm_col(h) for h in headers_raw]
            for rec in reader:
                row: Dict[str, Any] = {}
                for h_raw, h_norm in zip(headers_raw, headers_norm):
                    row[h_norm] = rec.get(h_raw)
                rows.append(row)
        return headers_norm, rows

    wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    ws = wb[sheet_name] if (sheet_name and sheet_name in wb.sheetnames) else wb.active
    it = ws.iter_rows(values_only=True)
    try:
        header_row = next(it)
    except StopIteration:
        wb.close()
        return [], []

    headers_norm = [_norm_col(h) for h in (header_row or [])]
    rows = []
    for rr in it:
        if rr is None:
            continue
        row: Dict[str, Any] = {}
        all_empty = True
        for i, hv in enumerate(headers_norm):
            if not hv:
                continue
            v = rr[i] if i < len(rr) else None
            row[hv] = v
            if v is not None and str(v).strip() != "":
                all_empty = False
        if not all_empty:
            rows.append(row)
    wb.close()
    return headers_norm, rows


def _read_pair_scores_from_file(
    pred_path: str,
    score_col: str = "score_final",
    id1_col: str = "id1",
    id2_col: str = "id2",
    sheet_name: Optional[str] = None,
) -> Dict[int, Dict[int, float]]:
    """Return directed score map: query_id -> {candidate_id: score}."""
    _, rows = _read_table_records(pred_path, sheet_name=sheet_name)
    sid1 = _norm_col(id1_col)
    sid2 = _norm_col(id2_col)
    sscore = _norm_col(score_col)
    srank = _norm_col("rank_final")

    out: Dict[int, Dict[int, float]] = {}
    for r in rows:
        try:
            a = int(r.get(sid1))
            b = int(r.get(sid2))
        except Exception:
            continue
        if a == b:
            continue

        score = r.get(sscore)
        if score is None:
            # Fallback: rank_final -> descending surrogate score.
            try:
                rank_v = float(r.get(srank))
                score = 1.0 / max(1.0, rank_v)
            except Exception:
                score = None
        try:
            s = float(score)
        except Exception:
            continue

        out.setdefault(a, {})[b] = max(float(out.get(a, {}).get(b, -1e30)), s)
        out.setdefault(b, {})[a] = max(float(out.get(b, {}).get(a, -1e30)), s)
    return out


def _read_pair_labels_from_file(
    label_path: str,
    label_col: str = "label",
    id1_col: str = "id1",
    id2_col: str = "id2",
    sheet_name: Optional[str] = None,
    positive_threshold: float = 0.5,
) -> Dict[int, Dict[int, int]]:
    """Return directed binary labels: query_id -> {candidate_id: 0/1}."""
    _, rows = _read_table_records(label_path, sheet_name=sheet_name)
    sid1 = _norm_col(id1_col)
    sid2 = _norm_col(id2_col)
    slabel = _norm_col(label_col)

    out: Dict[int, Dict[int, int]] = {}
    thr = float(positive_threshold)
    for r in rows:
        try:
            a = int(r.get(sid1))
            b = int(r.get(sid2))
            lv = float(r.get(slabel))
        except Exception:
            continue
        if a == b:
            continue
        y = 1 if lv >= thr else 0
        out.setdefault(a, {})[b] = max(int(out.get(a, {}).get(b, 0)), y)
        out.setdefault(b, {})[a] = max(int(out.get(b, {}).get(a, 0)), y)
    return out


def evaluate_pair_ranking_metrics(
    pred_scores: Dict[int, Dict[int, float]],
    labels: Dict[int, Dict[int, int]],
    ks: List[int],
) -> Dict[str, Any]:
    """Compute Recall@K / nDCG@K / MRR on directed query-candidate ranking."""
    ks_clean = sorted({int(k) for k in ks if int(k) > 0})
    if not ks_clean:
        ks_clean = [1, 5, 10]

    qids_all = sorted(labels.keys())
    qids_pos = [q for q in qids_all if sum(int(v > 0) for v in labels.get(q, {}).values()) > 0]

    if not qids_pos:
        return {
            "queries_total": len(qids_all),
            "queries_with_positive": 0,
            "queries_with_predictions": 0,
            "mrr": 0.0,
            "recall": {f"@{k}": 0.0 for k in ks_clean},
            "ndcg": {f"@{k}": 0.0 for k in ks_clean},
        }

    recall_sum = {k: 0.0 for k in ks_clean}
    ndcg_sum = {k: 0.0 for k in ks_clean}
    mrr_sum = 0.0
    q_with_pred = 0

    for q in qids_pos:
        lab = labels.get(q, {})
        pos_total = int(sum(int(v > 0) for v in lab.values()))
        if pos_total <= 0:
            continue

        pred_map = pred_scores.get(q, {})
        if pred_map:
            q_with_pred += 1
            ranked = sorted(pred_map.items(), key=lambda x: (-float(x[1]), int(x[0])))
        else:
            ranked = []

        rels = [1 if int(lab.get(cid, 0)) > 0 else 0 for cid, _s in ranked]
        # MRR
        rr = 0.0
        for pos, rel in enumerate(rels, start=1):
            if rel > 0:
                rr = 1.0 / float(pos)
                break
        mrr_sum += rr

        for k in ks_clean:
            rel_k = rels[:k]
            hit_k = int(sum(rel_k))
            recall_sum[k] += float(hit_k / max(1, pos_total))

            if rel_k:
                dcg = 0.0
                for i, rel in enumerate(rel_k, start=1):
                    if rel > 0:
                        dcg += 1.0 / np.log2(float(i) + 1.0)
            else:
                dcg = 0.0
            ideal_hits = min(pos_total, k)
            idcg = sum(1.0 / np.log2(float(i) + 1.0) for i in range(1, ideal_hits + 1))
            ndcg_sum[k] += (0.0 if idcg <= 0 else float(dcg / idcg))

    qn = float(len(qids_pos))
    return {
        "queries_total": len(qids_all),
        "queries_with_positive": int(len(qids_pos)),
        "queries_with_predictions": int(q_with_pred),
        "mrr": float(mrr_sum / qn),
        "recall": {f"@{k}": float(recall_sum[k] / qn) for k in ks_clean},
        "ndcg": {f"@{k}": float(ndcg_sum[k] / qn) for k in ks_clean},
    }


def run_eval_interface(
    pred_path: str,
    label_path: str,
    *,
    pred_sheet: Optional[str] = None,
    label_sheet: Optional[str] = None,
    score_col: str = "score_final",
    label_col: str = "label",
    id1_col: str = "id1",
    id2_col: str = "id2",
    ks: Optional[List[int]] = None,
    positive_threshold: float = 0.5,
) -> Dict[str, Any]:
    """P0 evaluation interface: read output Excel + label file and compute ranking metrics."""
    ks_eff = ks or [1, 5, 10, 20, 50]
    pred_scores = _read_pair_scores_from_file(
        pred_path,
        score_col=score_col,
        id1_col=id1_col,
        id2_col=id2_col,
        sheet_name=pred_sheet,
    )
    labels = _read_pair_labels_from_file(
        label_path,
        label_col=label_col,
        id1_col=id1_col,
        id2_col=id2_col,
        sheet_name=label_sheet,
        positive_threshold=positive_threshold,
    )
    out = evaluate_pair_ranking_metrics(pred_scores, labels, ks_eff)
    out["pred_path"] = pred_path
    out["label_path"] = label_path
    out["score_col"] = score_col
    out["label_col"] = label_col
    out["ks"] = [int(k) for k in sorted({int(x) for x in ks_eff if int(x) > 0})]
    out["positive_threshold"] = float(positive_threshold)
    return out


def _parse_ks(s: str) -> List[int]:
    out: List[int] = []
    for p in str(s or "").split(","):
        p = p.strip()
        if not p:
            continue
        try:
            v = int(p)
        except Exception:
            continue
        if v > 0:
            out.append(v)
    return out or [1, 5, 10, 20, 50]


def _run_eval_cli(argv: Optional[List[str]] = None) -> int:
    parser = argparse.ArgumentParser(description="Evaluate pairwise ranking outputs.")
    parser.add_argument("--pred", required=True, help="Prediction file (xlsx/csv) with id1,id2,score_final.")
    parser.add_argument("--labels", required=True, help="Label file (xlsx/csv) with id1,id2,label.")
    parser.add_argument("--pred_sheet", default=None, help="Prediction sheet name (xlsx only).")
    parser.add_argument("--label_sheet", default=None, help="Label sheet name (xlsx only).")
    parser.add_argument("--score_col", default="score_final", help="Prediction score column name.")
    parser.add_argument("--label_col", default="label", help="Label column name.")
    parser.add_argument("--id1_col", default="id1", help="Pair column: id1")
    parser.add_argument("--id2_col", default="id2", help="Pair column: id2")
    parser.add_argument("--ks", default="1,5,10,20,50", help="Comma-separated K list, e.g. 1,5,10")
    parser.add_argument("--positive_threshold", type=float, default=0.5, help="Label >= threshold is positive.")
    parser.add_argument("--out_json", default=None, help="Optional output json path.")
    args = parser.parse_args(argv)

    metrics = run_eval_interface(
        pred_path=args.pred,
        label_path=args.labels,
        pred_sheet=args.pred_sheet,
        label_sheet=args.label_sheet,
        score_col=args.score_col,
        label_col=args.label_col,
        id1_col=args.id1_col,
        id2_col=args.id2_col,
        ks=_parse_ks(args.ks),
        positive_threshold=float(args.positive_threshold),
    )

    text = json.dumps(metrics, ensure_ascii=False, indent=2)
    print(text)
    if args.out_json:
        with open(args.out_json, "w", encoding="utf-8") as f:
            f.write(text + "\n")
    return 0


# =========================
# GUI App
# =========================
class App(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("Text Similarity Pipeline (Excel/Word)")
        self.geometry("1280x980")
        self.minsize(1220, 900)

        self.file_path = tk.StringVar()
        self.out_dir = tk.StringVar(value=desktop_dir())

        self.granularity = tk.StringVar(value="all")
        self.framework_mode = tk.StringVar(value="all")
        self.ngram_n = tk.IntVar(value=3)
        self.use_tfidf = tk.BooleanVar(value=True)

        # pairing mode by granularity (user-tunable)
        # defaults: h1=all; h2=topk(20); normal=topk(20)
        self.pair_mode_h1 = tk.StringVar(value="all")
        self.pair_topk_h1 = tk.IntVar(value=20)
        self.pair_mode_h2 = tk.StringVar(value="topk")
        self.pair_topk_h2 = tk.IntVar(value=20)
        self.pair_mode_n = tk.StringVar(value="topk")
        self.pair_topk_n = tk.IntVar(value=20)
        self.report_top_percent = tk.DoubleVar(value=1.0)  # default 1%

        self.export_jsonl = tk.BooleanVar(value=True)

        # composite Q-weight settings (user tunable)
        # h1/h2 defaults
        self.w_h_raw = tk.DoubleVar(value=0.35)
        self.w_h_cov = tk.DoubleVar(value=0.35)
        self.w_h_tfidf = tk.DoubleVar(value=0.20)
        self.w_h_jacc = tk.DoubleVar(value=0.10)
        # normal defaults
        self.w_n_raw = tk.DoubleVar(value=0.60)
        self.w_n_tfidf = tk.DoubleVar(value=0.20)
        self.w_n_jacc = tk.DoubleVar(value=0.20)

        # semantic similarity (embedding-only score)
        self.use_semantic = tk.BooleanVar(value=True)
        self.semantic_model_size = tk.StringVar(value="8B")
        self.semantic_topk = tk.IntVar(value=20)
        self.sem_embed_max_length = tk.IntVar(value=8192)
        self.sem_embed_overlong_stride_pct = tk.IntVar(value=75)  # stride = max_length * pct / 100
        self.semantic_cache_reuse = tk.BooleanVar(value=True)
        self.semantic_cache_write = tk.BooleanVar(value=True)
        self.use_reranker = tk.BooleanVar(value=True)
        self.reranker_model_size = tk.StringVar(value="8B")
        self.rerank_topn = tk.IntVar(value=int(RERANK_DEFAULT_TOPN_CAP))
        self.rerank_top_percent_non_h1 = tk.DoubleVar(value=float(RERANK_DEFAULT_TOP_PERCENT_NON_H1))

        # merge baseline
        self.merge_topk = tk.IntVar(value=int(MERGE_BASELINE["topk_total"]))
        self.merge_w_lex = tk.DoubleVar(value=float(MERGE_BASELINE["weight_lex"]))
        self.merge_w_sem = tk.DoubleVar(value=float(MERGE_BASELINE["weight_sem"]))

        self.status = tk.StringVar(value="Ready.")
        self.pair_progress = tk.StringVar(value="Pairs: 0/0")
        self._progress_total = 0
        self._run_started_at: Optional[float] = None
        self._build()

    def _build(self):
        pad = {"padx": 10, "pady": 5}
        style = ttk.Style()
        style.configure("Merge.TLabelframe.Label", foreground="#D1495B", font=("TkDefaultFont", 11, "bold"))
        style.configure("Lex.TLabelframe.Label", foreground="#C07A00", font=("TkDefaultFont", 10, "bold"))
        style.configure("Sem.TLabelframe.Label", foreground="#2F6EA8", font=("TkDefaultFont", 10, "bold"))

        frm = ttk.Frame(self)
        frm.pack(fill="both", expand=True)

        file_box = ttk.LabelFrame(frm, text="1) Input")
        file_box.pack(fill="x", **pad)
        ttk.Entry(file_box, textvariable=self.file_path).pack(side="left", fill="x", expand=True, padx=8, pady=8)
        ttk.Button(file_box, text="Browse", command=self.browse_file).pack(side="left", padx=8, pady=8)

        seg_box = ttk.LabelFrame(frm, text="2) Scope")
        seg_box.pack(fill="x", **pad)
        ttk.Label(seg_box, text="Granularity (default=all):").pack(side="left", padx=8, pady=8)
        for g in ["h1", "h2", "h3", "normal", "all"]:
            ttk.Radiobutton(seg_box, text=g, value=g, variable=self.granularity).pack(side="left", padx=6)

        fw_box = ttk.LabelFrame(frm, text="2b) Framework (default=all)")
        fw_box.pack(fill="x", **pad)
        ttk.Label(fw_box, text="Run mode:").pack(side="left", padx=8, pady=8)
        ttk.Radiobutton(fw_box, text="字面", value="lexical", variable=self.framework_mode).pack(side="left", padx=6)
        ttk.Radiobutton(fw_box, text="语义", value="semantic", variable=self.framework_mode).pack(side="left", padx=6)
        ttk.Radiobutton(fw_box, text="融合★重点", value="merge", variable=self.framework_mode).pack(side="left", padx=6)
        ttk.Radiobutton(fw_box, text="all", value="all", variable=self.framework_mode).pack(side="left", padx=6)
        ttk.Label(
            fw_box,
            text="字面",
            foreground="#C07A00",
            font=("TkDefaultFont", 9, "bold"),
        ).pack(side="right", padx=6)
        ttk.Label(
            fw_box,
            text="语义",
            foreground="#2F6EA8",
            font=("TkDefaultFont", 9, "bold"),
        ).pack(side="right", padx=6)
        ttk.Label(
            fw_box,
            text="融合",
            foreground="#D1495B",
            font=("TkDefaultFont", 9, "bold"),
        ).pack(side="right", padx=6)

        # Merge panel first (most important).
        merge_box = ttk.LabelFrame(frm, text="3) 融合参数（最重要）", style="Merge.TLabelframe")
        merge_box.pack(fill="x", **pad)
        ttk.Label(merge_box, text="Merge TopK:").grid(row=0, column=0, sticky="e", padx=6, pady=4)
        ttk.Entry(merge_box, textvariable=self.merge_topk, width=8).grid(row=0, column=1, sticky="w", padx=4, pady=4)
        ttk.Label(merge_box, text="Score weights (lex, sem):").grid(row=0, column=2, sticky="e", padx=10, pady=4)
        ttk.Entry(merge_box, textvariable=self.merge_w_lex, width=8).grid(row=0, column=3, sticky="w", padx=4, pady=4)
        ttk.Entry(merge_box, textvariable=self.merge_w_sem, width=8).grid(row=0, column=4, sticky="w", padx=4, pady=4)
        ttk.Label(merge_box, text="Baseline: lexical 0.7 + semantic 0.3", foreground="#D1495B").grid(
            row=0, column=5, sticky="w", padx=10, pady=4
        )

        # Lexical parameters in one panel.
        lex_box = ttk.LabelFrame(frm, text="4) 字面相似度参数（合并）", style="Lex.TLabelframe")
        lex_box.pack(fill="x", **pad)
        ttk.Label(lex_box, text="Char n-gram n:").grid(row=0, column=0, sticky="e", padx=8, pady=4)
        ttk.Spinbox(lex_box, from_=1, to=8, width=5, textvariable=self.ngram_n).grid(row=0, column=1, sticky="w", padx=6, pady=4)
        ttk.Checkbutton(lex_box, text="Compute TF-IDF", variable=self.use_tfidf).grid(row=0, column=2, sticky="w", padx=12, pady=4)
        if MinHash is None:
            ttk.Label(lex_box, text="(datasketch 未安装，MinHash 关闭)", foreground="gray").grid(row=0, column=3, sticky="w", padx=10, pady=4)

        ttk.Separator(lex_box, orient="horizontal").grid(row=1, column=0, columnspan=8, sticky="ew", padx=8, pady=4)
        ttk.Label(lex_box, text="Pairing (h1/h2/normal)").grid(row=2, column=0, sticky="w", padx=8, pady=2)
        ttk.Label(lex_box, text="Mode").grid(row=2, column=1, sticky="w", padx=8, pady=2)
        ttk.Label(lex_box, text="TopK").grid(row=2, column=2, sticky="w", padx=8, pady=2)
        ttk.Label(lex_box, text="Report Top%").grid(row=2, column=4, sticky="w", padx=8, pady=2)

        ttk.Label(lex_box, text="h1").grid(row=3, column=0, sticky="w", padx=8, pady=2)
        ttk.Combobox(lex_box, textvariable=self.pair_mode_h1, values=["all", "topk"], width=8, state="readonly").grid(row=3, column=1, sticky="w", padx=8, pady=2)
        ttk.Spinbox(lex_box, from_=1, to=500, width=8, textvariable=self.pair_topk_h1).grid(row=3, column=2, sticky="w", padx=8, pady=2)
        ttk.Entry(lex_box, textvariable=self.report_top_percent, width=8).grid(row=3, column=4, sticky="w", padx=8, pady=2)

        ttk.Label(lex_box, text="h2").grid(row=4, column=0, sticky="w", padx=8, pady=2)
        ttk.Combobox(lex_box, textvariable=self.pair_mode_h2, values=["all", "topk"], width=8, state="readonly").grid(row=4, column=1, sticky="w", padx=8, pady=2)
        ttk.Spinbox(lex_box, from_=1, to=500, width=8, textvariable=self.pair_topk_h2).grid(row=4, column=2, sticky="w", padx=8, pady=2)

        ttk.Label(lex_box, text="normal").grid(row=5, column=0, sticky="w", padx=8, pady=2)
        ttk.Combobox(lex_box, textvariable=self.pair_mode_n, values=["all", "topk"], width=8, state="readonly").grid(row=5, column=1, sticky="w", padx=8, pady=2)
        ttk.Spinbox(lex_box, from_=1, to=500, width=8, textvariable=self.pair_topk_n).grid(row=5, column=2, sticky="w", padx=8, pady=2)
        ttk.Label(lex_box, text="TopK 在 mode=all 时忽略；h3 使用 normal 配置。", foreground="gray").grid(
            row=6, column=0, columnspan=8, sticky="w", padx=8, pady=2
        )

        ttk.Separator(lex_box, orient="horizontal").grid(row=7, column=0, columnspan=8, sticky="ew", padx=8, pady=4)
        ttk.Label(lex_box, text="Composite weights: H1/H2 (raw, cov, tfidf, jacc)").grid(row=8, column=0, sticky="e", padx=8, pady=2)
        ttk.Entry(lex_box, textvariable=self.w_h_raw, width=7).grid(row=8, column=1, sticky="w", padx=3, pady=2)
        ttk.Entry(lex_box, textvariable=self.w_h_cov, width=7).grid(row=8, column=2, sticky="w", padx=3, pady=2)
        ttk.Entry(lex_box, textvariable=self.w_h_tfidf, width=7).grid(row=8, column=3, sticky="w", padx=3, pady=2)
        ttk.Entry(lex_box, textvariable=self.w_h_jacc, width=7).grid(row=8, column=4, sticky="w", padx=3, pady=2)

        ttk.Label(lex_box, text="Composite weights: Normal (raw, tfidf, jacc)").grid(row=9, column=0, sticky="e", padx=8, pady=2)
        ttk.Entry(lex_box, textvariable=self.w_n_raw, width=7).grid(row=9, column=1, sticky="w", padx=3, pady=2)
        ttk.Entry(lex_box, textvariable=self.w_n_tfidf, width=7).grid(row=9, column=2, sticky="w", padx=3, pady=2)
        ttk.Entry(lex_box, textvariable=self.w_n_jacc, width=7).grid(row=9, column=3, sticky="w", padx=3, pady=2)
        ttk.Label(lex_box, text="每层权重和需为 1.0。", foreground="gray").grid(row=8, column=5, rowspan=2, sticky="w", padx=10, pady=2)

        # Semantic parameters in one panel.
        sem_box = ttk.LabelFrame(frm, text="5) 语义相似度参数（合并）", style="Sem.TLabelframe")
        sem_box.pack(fill="x", **pad)
        ttk.Checkbutton(sem_box, text="Enable semantic similarity", variable=self.use_semantic).grid(row=0, column=0, sticky="w", padx=10, pady=4)
        ttk.Label(sem_box, text="Embedding model:").grid(row=0, column=1, sticky="e", padx=6, pady=4)
        ttk.Combobox(sem_box, textvariable=self.semantic_model_size, values=["0.6B", "4B", "8B", "Qwen2-7B", "BGE-M3"], width=12, state="readonly").grid(row=0, column=2, sticky="w", padx=4, pady=4)
        ttk.Label(sem_box, text="Sem TopK:").grid(row=0, column=3, sticky="e", padx=6, pady=4)
        ttk.Entry(sem_box, textvariable=self.semantic_topk, width=8).grid(row=0, column=4, sticky="w", padx=4, pady=4)
        ttk.Checkbutton(sem_box, text="Reuse semantic cache", variable=self.semantic_cache_reuse).grid(row=0, column=5, sticky="w", padx=10, pady=4)
        ttk.Checkbutton(sem_box, text="Write semantic cache", variable=self.semantic_cache_write).grid(row=0, column=6, sticky="w", padx=10, pady=4)

        ttk.Label(sem_box, text="Embed max_length:").grid(row=1, column=0, sticky="e", padx=6, pady=4)
        ttk.Combobox(sem_box, textvariable=self.sem_embed_max_length,
                      values=[512, 1024, 2048, 4096, 8192], width=7).grid(row=1, column=1, sticky="w", padx=4, pady=4)
        ttk.Label(sem_box, text="Overlong stride%:").grid(row=1, column=2, sticky="e", padx=6, pady=4)
        ttk.Spinbox(sem_box, from_=50, to=95, width=5, textvariable=self.sem_embed_overlong_stride_pct).grid(row=1, column=3, sticky="w", padx=4, pady=4)
        ttk.Label(sem_box, text="超过 max_length 的文本自动滑窗; stride = max_length × stride%", foreground="gray").grid(
            row=1, column=4, columnspan=3, sticky="w", padx=6, pady=4)

        ttk.Checkbutton(sem_box, text="Enable reranker (active top-candidates)", variable=self.use_reranker).grid(row=2, column=0, sticky="w", padx=10, pady=4)
        ttk.Label(sem_box, text="Reranker model:").grid(row=2, column=1, sticky="e", padx=6, pady=4)
        ttk.Combobox(sem_box, textvariable=self.reranker_model_size, values=["0.6B", "4B", "8B", "BGE-reranker"], width=14, state="readonly").grid(row=2, column=2, sticky="w", padx=4, pady=4)
        ttk.Label(sem_box, text=f"TopN cap (default={int(RERANK_DEFAULT_TOPN_CAP)}; 0=auto):").grid(row=2, column=3, sticky="e", padx=6, pady=4)
        ttk.Entry(sem_box, textvariable=self.rerank_topn, width=8).grid(row=2, column=4, sticky="w", padx=4, pady=4)
        ttk.Label(sem_box, text="Top% for h2/normal:").grid(row=2, column=5, sticky="e", padx=6, pady=4)
        ttk.Entry(sem_box, textvariable=self.rerank_top_percent_non_h1, width=8).grid(row=2, column=6, sticky="w", padx=4, pady=4)
        ttk.Label(
            sem_box,
            text="语义通道：H1/H2 用 MaxSim 精排（H2 Top3%并带500~1500护栏）；Normal 用 embedding cosine；reranker 仅在融合模式对 Top 候选重排。"
                 " 缓存 meta 含代码版本/脚本文件/输入文件/生成时间。",
            foreground="gray",
        ).grid(
            row=3, column=0, columnspan=8, sticky="w", padx=10, pady=4
        )

        out_box = ttk.LabelFrame(frm, text="6) Output")
        out_box.pack(fill="x", **pad)
        ttk.Label(out_box, text="Output dir:").pack(side="left", padx=8, pady=8)
        ttk.Entry(out_box, textvariable=self.out_dir).pack(side="left", fill="x", expand=True, padx=8, pady=8)
        ttk.Button(out_box, text="Change", command=self.browse_outdir).pack(side="left", padx=8, pady=8)
        ttk.Checkbutton(out_box, text="Export machine JSON + units_all.jsonl", variable=self.export_jsonl).pack(side="left", padx=10)

        run_box = ttk.Frame(frm)
        run_box.pack(fill="x", **pad)
        ttk.Button(run_box, text="Run", command=self.run).pack(side="left", padx=8)
        self.prog = ttk.Progressbar(run_box, mode="determinate", length=280)
        self.prog.pack(side="left", padx=10)
        ttk.Label(run_box, textvariable=self.pair_progress).pack(side="left", padx=10)
        ttk.Label(run_box, textvariable=self.status).pack(side="left", padx=12)

        log_box = ttk.LabelFrame(frm, text="Log")
        log_box.pack(fill="both", expand=True, **pad)
        self.log = ScrolledText(log_box, height=11)
        self.log.pack(fill="both", expand=True, padx=10, pady=8)
        self.log.configure(state="disabled")

        tips = ttk.LabelFrame(frm, text="Tips")
        tips.pack(fill="x", **pad)
        ttk.Label(
            tips,
            justify="left",
            text=(
                "默认 granularity=all + framework=all。\n"
                "framework=all 会自动跑 字面/语义/融合，并输出到同一套 all 报告。"
            ),
            foreground="gray",
        ).pack(anchor="w", padx=10, pady=6)

    def browse_file(self):
        path = filedialog.askopenfilename(
            title="Select Excel file",
            filetypes=[("Excel files", "*.xlsx *.xlsm *.xls"), ("All files", "*.*")]
        )
        if path:
            self.file_path.set(path)

    def browse_outdir(self):
        d = filedialog.askdirectory(title="Select output directory")
        if d:
            self.out_dir.set(d)

    def _on_ui(self, fn, *args, **kwargs):
        """Run UI operation on Tk main thread."""
        try:
            if threading.current_thread() is threading.main_thread():
                fn(*args, **kwargs)
            else:
                self.after(0, lambda: fn(*args, **kwargs))
        except Exception:
            pass

    def _status_set(self, msg: str):
        self._on_ui(self.status.set, str(msg))

    def _progress_stop(self):
        self._on_ui(self.prog.stop)

    def _safe_showinfo(self, title: str, text: str):
        self._on_ui(messagebox.showinfo, title, text)

    def _safe_showerror(self, title: str, text: str):
        self._on_ui(messagebox.showerror, title, text)

    def _append_log(self, msg: str):
        def _do():
            now_tag = time.strftime("%H:%M:%S")
            if self._run_started_at is None:
                prefix = f"[{now_tag}] "
            else:
                elapsed = _fmt_elapsed(time.perf_counter() - float(self._run_started_at))
                prefix = f"[{now_tag} +{elapsed}] "
            self.log.configure(state="normal")
            self.log.insert("end", prefix + str(msg) + "\n")
            self.log.see("end")
            self.log.configure(state="disabled")
        self._on_ui(_do)

    def _composite_weight_config(self) -> Tuple[Dict[str, float], Dict[str, float]]:
        w_h = {
            "cosine_tf_raw": float(self.w_h_raw.get()),
            "coverage": float(self.w_h_cov.get()),
            "cosine_tfidf": float(self.w_h_tfidf.get()),
            "jaccard": float(self.w_h_jacc.get()),
        }
        w_n = {
            "cosine_tf_raw": float(self.w_n_raw.get()),
            "coverage": 0.0,
            "cosine_tfidf": float(self.w_n_tfidf.get()),
            "jaccard": float(self.w_n_jacc.get()),
        }
        return w_h, w_n

    def _validate_composite_weights(self) -> bool:
        try:
            w_h, w_n = self._composite_weight_config()
        except Exception:
            messagebox.showerror("Error", "Composite weights must be numeric.")
            return False

        def _check(group_name: str, w: Dict[str, float]) -> Optional[str]:
            vals = [float(v) for v in w.values()]
            if any(v < 0 for v in vals):
                return f"{group_name} weights must be >= 0."
            s = float(sum(vals))
            if abs(s - 1.0) > 1e-6:
                return f"{group_name} 权重和必须等于 1.0（当前={s:.6f}）"
            return None

        err_h = _check("H1/H2", w_h)
        err_n = _check("Normal", w_n)
        if err_h or err_n:
            msg = "\n".join([x for x in [err_h, err_n] if x])
            messagebox.showerror("Error", msg)
            return False
        return True

    def _pairing_config_for_g(self, g: str) -> Tuple[str, int]:
        gg = normalize_style(g)
        if gg == "h1":
            mode = str(self.pair_mode_h1.get() or "all").strip().lower()
            k = int(self.pair_topk_h1.get())
        elif gg == "h2":
            mode = str(self.pair_mode_h2.get() or "topk").strip().lower()
            k = int(self.pair_topk_h2.get())
        else:
            # normal/h3 share one config row by design.
            mode = str(self.pair_mode_n.get() or "topk").strip().lower()
            k = int(self.pair_topk_n.get())
        if mode not in ("all", "topk"):
            mode = "topk"
        return mode, max(1, int(k))

    def _validate_pairing_config(self) -> bool:
        try:
            for g in ("h1", "h2", "normal"):
                mode, k = self._pairing_config_for_g(g)
                if mode not in ("all", "topk"):
                    raise ValueError(f"{g} mode must be all/topk")
                if int(k) < 1:
                    raise ValueError(f"{g} TopK must be >= 1")
        except Exception as e:
            messagebox.showerror("Error", f"Invalid pairing config: {e}")
            return False
        return True

    def _validate_semantic_config(self) -> bool:
        if not bool(self.use_semantic.get()):
            return True
        try:
            m = _normalize_embed_size_label(self.semantic_model_size.get() or "")
            if not _is_bge_model(m) and m not in QWEN_EMBED_MODEL_DIRS:
                raise ValueError(f"Unsupported model: {m}")
            self.semantic_model_size.set(m)
            if int(self.semantic_topk.get()) < 1:
                raise ValueError("Sem TopK must be >= 1")
        except Exception as e:
            messagebox.showerror("Error", f"Invalid semantic config: {e}")
            return False
        return True

    def _validate_reranker_config(self) -> bool:
        if not bool(self.use_reranker.get()):
            return True
        try:
            m = _normalize_reranker_size_label(self.reranker_model_size.get() or "")
            if not _is_bge_model(m) and m not in QWEN_RERANK_MODEL_DIRS:
                raise ValueError(f"Unsupported reranker model: {m}")
            self.reranker_model_size.set(m)
            if int(self.rerank_topn.get()) < 0:
                raise ValueError("rerank TopN must be >= 0")
            p = float(self.rerank_top_percent_non_h1.get())
            if p <= 0 or p > 100:
                raise ValueError("Top% for h2/normal must be in (0,100].")
        except Exception as e:
            messagebox.showerror("Error", f"Invalid reranker config: {e}")
            return False
        return True

    def _validate_merge_config(self) -> bool:
        try:
            k = int(self.merge_topk.get())
            if k < 2:
                raise ValueError("Merge TopK must be >= 2")
            wl = float(self.merge_w_lex.get())
            ws = float(self.merge_w_sem.get())
            if wl < 0 or ws < 0:
                raise ValueError("Merge weights must be >= 0")
            if abs((wl + ws) - 1.0) > 1e-6:
                raise ValueError(f"Merge weights must sum to 1.0 (current={wl+ws:.6f})")
        except Exception as e:
            messagebox.showerror("Error", f"Invalid merge config: {e}")
            return False
        return True

    def _reset_progress(self):
        self._progress_total = 0
        self.pair_progress.set("Pairs: 0/0")
        try:
            self.prog["value"] = 0
        except Exception:
            pass

    def _ui_progress(self, done: int, total: int, msg: str):
        def _do():
            self._progress_total = max(self._progress_total, int(total))
            self.pair_progress.set(f"Pairs: {done}/{total}")
            try:
                if total > 0:
                    self.prog["maximum"] = total
                    self.prog["value"] = done
            except Exception:
                pass
            if msg:
                self.status.set(msg)
        self._on_ui(_do)

    def run(self):
        if not self.file_path.get():
            messagebox.showerror("Error", "Please select an input file.")
            return
        outd = self.out_dir.get().strip()
        g = (self.granularity.get() or "").strip().lower()
        fw = str(self.framework_mode.get() or "merge").strip().lower()
        if g != "all" and fw != "all" and (not outd or not os.path.isdir(outd)):
            messagebox.showerror("Error", "Output directory is invalid.")
            return
        if not self._validate_pairing_config():
            return
        if fw not in VALID_FRAMEWORKS:
            messagebox.showerror("Error", f"Unsupported framework: {fw}")
            return
        if fw in ("lexical", "merge", "all"):
            if not self._validate_composite_weights():
                return
        if fw in ("semantic", "merge", "all"):
            if not bool(self.use_semantic.get()):
                messagebox.showerror("Error", "当前模式需要语义相似度，请启用 semantic similarity。")
                return
            if not self._validate_semantic_config():
                return
        if fw in ("merge", "all"):
            if not self._validate_merge_config():
                return
        if fw in ("merge", "all"):
            if not self._validate_reranker_config():
                return
        pair_cfg = {
            "h1": self._pairing_config_for_g("h1"),
            "h2": self._pairing_config_for_g("h2"),
            "normal": self._pairing_config_for_g("normal"),
        }
        if fw in ("lexical", "merge", "all"):
            w_h, w_n = self._composite_weight_config()
        else:
            w_h, w_n = None, None
        cfg = {
            "in_path": self.file_path.get().strip(),
            "out_dir": self.out_dir.get().strip(),
            "granularity": (self.granularity.get() or "").strip().lower(),
            "framework": str(self.framework_mode.get() or "merge").strip().lower(),
            "ngram_n": int(self.ngram_n.get()),
            "use_tfidf": bool(self.use_tfidf.get()),
            "pair_cfg": pair_cfg,
            "report_top_percent": float(self.report_top_percent.get()),
            "weights_h": w_h,
            "weights_n": w_n,
            "use_semantic": bool(self.use_semantic.get()),
            "semantic_model": _normalize_embed_size_label(self.semantic_model_size.get() or "8B"),
            "semantic_topk": int(self.semantic_topk.get()),
            "sem_embed_max_length": int(self.sem_embed_max_length.get()),
            "sem_embed_overlong_stride_pct": int(self.sem_embed_overlong_stride_pct.get()),
            "semantic_cache_reuse": bool(self.semantic_cache_reuse.get()),
            "semantic_cache_write": bool(self.semantic_cache_write.get()),
            "use_reranker": bool(self.use_reranker.get()),
            "reranker_model": _normalize_reranker_size_label(self.reranker_model_size.get() or "8B"),
            "rerank_topn": int(self.rerank_topn.get()),
            "rerank_top_percent_non_h1": float(self.rerank_top_percent_non_h1.get()),
            "merge_topk": int(self.merge_topk.get()),
            "merge_w_lex": float(self.merge_w_lex.get()),
            "merge_w_sem": float(self.merge_w_sem.get()),
            "export_jsonl": bool(self.export_jsonl.get()),
        }
        self._run_started_at = time.perf_counter()
        self._reset_progress()
        self._append_log("Run started.")
        threading.Thread(target=self._run_impl, args=(cfg,), daemon=True).start()

    def _run_impl(self, cfg: Optional[Dict[str, Any]] = None):
        try:
            run_t0 = time.perf_counter()
            run_stamp = time.strftime("%Y%m%d%H%M")
            # ── 在時間戳後附加腳本名 (e.g. "202602181030_pipeline17") ──
            import re as _re_scriptname
            _script_base = os.path.splitext(os.path.basename(__file__))[0]
            # 去掉括號及其內容: "pipeline17（merge 分数回填)" → "pipeline17"
            _script_base = _re_scriptname.sub(r'[（(].+?[）)]', '', _script_base).strip()
            if _script_base:
                run_stamp = f"{run_stamp}_{_script_base}"
            self._status_set("Reading file...")
            self._append_log("Reading input file...")
            self._append_log(f"Output file timestamp prefix: {run_stamp}")

            cfg = cfg or {}
            in_path = str(cfg.get("in_path") or "").strip()
            if not in_path:
                raise RuntimeError("Input file path is empty.")
            base = os.path.splitext(os.path.basename(in_path))[0]
            g_raw = str(cfg.get("granularity") or "all").strip().lower()
            fw_selected = str(cfg.get("framework") or "merge").strip().lower()
            outd_cfg = str(cfg.get("out_dir") or "").strip()
            ngram_n = int(max(1, cfg.get("ngram_n", 3)))
            use_tfidf = bool(cfg.get("use_tfidf", True))
            pair_cfg = cfg.get("pair_cfg") if isinstance(cfg.get("pair_cfg"), dict) else {}
            report_top_percent = float(cfg.get("report_top_percent", 1.0))
            w_h_cfg = cfg.get("weights_h")
            w_n_cfg = cfg.get("weights_n")

            use_semantic = bool(cfg.get("use_semantic", True))
            sem_model = _normalize_embed_size_label(cfg.get("semantic_model") or "8B")
            sem_topk_default = int(max(1, cfg.get("semantic_topk", 20)))
            sem_cache_reuse = bool(cfg.get("semantic_cache_reuse", True))
            sem_cache_write = bool(cfg.get("semantic_cache_write", True))
            sem_ml = int(cfg.get("sem_embed_max_length", 8192))
            sem_stride_pct = int(cfg.get("sem_embed_overlong_stride_pct", 75))

            rr_enabled = bool(cfg.get("use_reranker", True))
            rr_model = _normalize_reranker_size_label(cfg.get("reranker_model") or "8B")
            rr_topn_raw = int(cfg.get("rerank_topn", int(RERANK_DEFAULT_TOPN_CAP)))
            rr_topn = rr_topn_raw if rr_topn_raw > 0 else None
            rr_top_pct = float(cfg.get("rerank_top_percent_non_h1", float(RERANK_DEFAULT_TOP_PERCENT_NON_H1)))

            merge_k = int(max(2, cfg.get("merge_topk", int(MERGE_BASELINE["topk_total"]))))
            merge_w_lex = float(cfg.get("merge_w_lex", float(MERGE_BASELINE["weight_lex"])))
            merge_w_sem = float(cfg.get("merge_w_sem", float(MERGE_BASELINE["weight_sem"])))
            export_jsonl = bool(cfg.get("export_jsonl", True))
            self._append_log(
                "Semantic cache policy: "
                f"reuse={'on' if sem_cache_reuse else 'off'}, "
                f"write={'on' if sem_cache_write else 'off'}."
            )

            def pairing_for_g(gname: str) -> Tuple[str, int]:
                gg = normalize_style(gname)
                key = "h1" if gg == "h1" else ("h2" if gg == "h2" else "normal")
                raw = pair_cfg.get(key)
                if isinstance(raw, (list, tuple)) and len(raw) >= 2:
                    mode = str(raw[0] or "topk").strip().lower()
                    try:
                        topk_v = int(raw[1])
                    except Exception:
                        topk_v = 20
                    if mode not in ("all", "topk"):
                        mode = "topk"
                    return mode, max(1, topk_v)
                return self._pairing_config_for_g(gname)

            framework_all_mode = (fw_selected == "all")
            frameworks_to_run = ["lexical", "semantic", "merge"] if framework_all_mode else [fw_selected]

            run_all_gran = (g_raw == "all")
            dual_all_mode = bool(framework_all_mode and run_all_gran)
            if framework_all_mode and g_raw != "all":
                self._append_log("framework=all: granularity is forced to h1/h2/normal.")
            run_grans = ["h1", "h2", "normal"] if (run_all_gran or framework_all_mode) else [normalize_style(g_raw)]

            sem_label = _embed_model_label(sem_model)
            rer_label = _reranker_model_label(rr_model)
            if framework_all_mode:
                outd_user = outd_cfg
                base_out = outd_user if (outd_user and os.path.isdir(outd_user)) else desktop_dir()
                if dual_all_mode:
                    out_tag = _make_dual_all_output_name_tag(
                        upload_base=base,
                        semantic_model=sem_model,
                        reranker_model=rr_model,
                        reranker_enabled=rr_enabled,
                        time_prefix=run_stamp,
                    )
                else:
                    out_tag = _make_output_name_tag(
                        framework="all",
                        granularity="h1-h2-normal",
                        upload_base=base,
                        semantic_model=sem_model,
                        reranker_model=rr_model,
                        reranker_enabled=rr_enabled,
                        time_prefix=run_stamp,
                    )
                out_dir = os.path.join(base_out, out_tag)
                os.makedirs(out_dir, exist_ok=True)
                self._append_log(
                    f"framework=all detected. Output folder: {out_dir} "
                    f"(base={'user-selected' if (outd_user and os.path.isdir(outd_user)) else 'desktop-default'})"
                )
            elif run_all_gran:
                outd_user = outd_cfg
                base_out = outd_user if (outd_user and os.path.isdir(outd_user)) else desktop_dir()
                out_tag = _make_output_name_tag(
                    framework=fw_selected,
                    granularity="h1-h2-normal",
                    upload_base=base,
                    semantic_model=sem_model,
                    reranker_model=rr_model,
                    reranker_enabled=(rr_enabled if fw_selected == "merge" else False),
                    time_prefix=run_stamp,
                )
                out_dir = os.path.join(base_out, out_tag)
                os.makedirs(out_dir, exist_ok=True)
                self._append_log(
                    f"granularity=all detected. Output folder: {out_dir} "
                    f"(base={'user-selected' if (outd_user and os.path.isdir(outd_user)) else 'desktop-default'})"
                )
            else:
                out_tag = _make_output_name_tag(
                    framework=fw_selected,
                    granularity=normalize_style(g_raw),
                    upload_base=base,
                    semantic_model=sem_model,
                    reranker_model=rr_model,
                    reranker_enabled=(rr_enabled if fw_selected == "merge" else False),
                    time_prefix=run_stamp,
                )
                out_dir = os.path.join(outd_cfg, out_tag)
                os.makedirs(out_dir, exist_ok=True)

            # 1) Try explicit units-table (must include `level` + `text/正文`)
            units_from_excel = read_excel_as_units_table(in_path)
            if units_from_excel is not None:
                self._append_log("Detected Units table format (header includes 'level' + text/正文).")
                normal_units_all = select_units_from_units_table(units_from_excel, "normal")
                if not normal_units_all:
                    normal_units_all = units_from_excel

                def build_units_by_g(g: str) -> List[Unit]:
                    return select_units_from_units_table(units_from_excel, g)

            else:
                # 2) Try user's source-table (书名/篇名/正文)
                source_rows = read_excel_as_source_table(in_path)
                if source_rows is not None:
                    self._append_log("Detected Source table format (书名/篇名/正文).")
                    normal_units_all = build_units_from_source_rows(source_rows, "normal")

                    def build_units_by_g(g: str) -> List[Unit]:
                        return build_units_from_source_rows(source_rows, g)

                else:
                    # 3) Fallback: Paragraphs table (A=text, B=style)
                    self._append_log("Detected Paragraphs table format (A=text, B=style).")
                    paras = read_excel_as_paragraphs_table(in_path)
                    self._status_set(f"Loaded paragraphs: {len(paras)}. Aggregating...")
                    normal_units_all = segment_units_from_paragraphs(paras, "normal")

                    def build_units_by_g(g: str) -> List[Unit]:
                        return segment_units_from_paragraphs(paras, g)

            need_composite_weights = any(x in ("lexical", "merge") for x in frameworks_to_run)
            if need_composite_weights:
                w_h, w_n = dict(w_h_cfg or {}), dict(w_n_cfg or {})
                if not w_h or not w_n:
                    w_h, w_n = self._composite_weight_config()
            else:
                w_h, w_n = None, None
            merge_k_lex = int(max(1, round(merge_k * float(MERGE_BASELINE["recall_lex_ratio"]))))
            merge_k_sem = int(max(1, merge_k - merge_k_lex))
            results_by_fw: Dict[str, Dict[str, Dict[str, Any]]] = {}
            results_by_g: Dict[str, Dict[str, Any]] = {}
            # P0 cache: reuse units/representations/semantic embeddings across frameworks.
            units_cache: Dict[str, List[Unit]] = {}
            reps_cache: Dict[str, dict] = {}
            sem_emb_cache: Dict[str, np.ndarray] = {}
            sem_neighbors_cache: Dict[Tuple, Dict[int, List[int]]] = {}
            sem_maxsim_ctx_cache: Dict[str, dict] = {}

            # ── cache key 编入 model，防止切换模型后命中旧缓存 ──
            def _sek(g_: str) -> str:
                """Semantic Embedding cache Key: granularity + model."""
                return f"{g_}:{sem_model}"

            for fi, fw in enumerate(frameworks_to_run, start=1):
                results_curr: Dict[str, Dict[str, Any]] = {}
                for gi, g in enumerate(run_grans, start=1):
                    if g not in units_cache:
                        units_cache[g] = build_units_by_g(g)
                    units = units_cache[g]
                    self._append_log(f"[{fw}] Units loaded: {len(units)} (granularity={g}).")
                    if len(units) < 2:
                        msg = f"Not enough units for {g} (units={len(units)})."
                        if run_all_gran or framework_all_mode:
                            self._append_log("WARN: " + msg + " skipped.")
                            continue
                        raise RuntimeError(msg + " Check your granularity or Excel format.")

                    if g not in reps_cache:
                        self._status_set(f"[F{fi}/{len(frameworks_to_run)} G{gi}/{len(run_grans)}] {fw}/{g}: building representations...")
                        t_rep = time.perf_counter()
                        reps_cache[g] = build_representations(units, ngram_n, use_tfidf)
                        self._append_log(
                            f"[{fw}/{g}] representations built and cached in {_fmt_elapsed(time.perf_counter() - t_rep)}."
                        )
                    else:
                        self._append_log(f"[{fw}/{g}] representations cache hit.")
                    reps = reps_cache[g]

                    mode_run, topk_run = pairing_for_g(g)

                    sem_emb = None
                    sem_neighbors: Dict[int, List[int]] = {}
                    sem_maxsim_ctx = None
                    need_sem_channel = fw in ("semantic", "merge")
                    if need_sem_channel and use_semantic:
                        sem_topk_build = sem_topk_default
                        if fw == "merge" and mode_run == "topk":
                            sem_topk_build = merge_k_sem
                        try:
                            if normalize_style(g) in ("h1", "h2") and _semantic_maxsim_enabled(g):
                                if _sek(g) not in sem_maxsim_ctx_cache:
                                    self._status_set(f"[F{fi}/{len(frameworks_to_run)} G{gi}/{len(run_grans)}] {fw}/{g}: building semantic chunk context (MaxSim)...")
                                    t_ctx = time.perf_counter()
                                    sem_maxsim_ctx_cache[_sek(g)] = build_semantic_maxsim_context(
                                        units,
                                        model_size=sem_model,
                                        chunk_size=int(SEMANTIC_MAXSIM_CHUNK_SIZE),
                                        chunk_stride=int(SEMANTIC_MAXSIM_CHUNK_STRIDE),
                                        embed_max_length=sem_ml,
                                        embed_stride_pct=sem_stride_pct,
                                    )
                                    sem_emb_cache[_sek(g)] = np.asarray(sem_maxsim_ctx_cache[_sek(g)].get("mean_embeddings"), dtype=np.float32)
                                    self._append_log(
                                        f"[{fw}/{g}] semantic chunk context built in {_fmt_elapsed(time.perf_counter() - t_ctx)}: "
                                        f"chunks={int(sem_maxsim_ctx_cache[_sek(g)].get('chunk_total') or 0)}, "
                                        f"chunk_size={SEMANTIC_MAXSIM_CHUNK_SIZE}, stride={SEMANTIC_MAXSIM_CHUNK_STRIDE}."
                                    )
                                else:
                                    self._append_log(f"[{fw}/{g}] semantic chunk context cache hit.")
                                sem_maxsim_ctx = sem_maxsim_ctx_cache[_sek(g)]
                                sem_emb = sem_emb_cache[_sek(g)]
                            else:
                                if _sek(g) not in sem_emb_cache:
                                    self._status_set(f"[F{fi}/{len(frameworks_to_run)} G{gi}/{len(run_grans)}] {fw}/{g}: building semantic embeddings...")
                                    t_emb = time.perf_counter()
                                    texts = [_embed_text_for_semantic(u.text or "") for u in units]
                                    cache_path = _semantic_emb_cache_path(
                                        out_dir=out_dir,
                                        granularity=g,
                                        model_size=sem_model,
                                        texts=texts,
                                        max_length=sem_ml,
                                    )
                                    cached = None
                                    if sem_cache_reuse:
                                        cached = _load_semantic_emb_cache(cache_path, expected_n=len(texts))
                                    elif os.path.isfile(cache_path):
                                        self._append_log(
                                            f"[{fw}/{g}] semantic cache reuse disabled; existing cache ignored: {cache_path}"
                                        )
                                    if cached is not None:
                                        sem_emb_cache[_sek(g)] = cached
                                        meta = _load_semantic_emb_meta(cache_path)
                                        m_ver = str(meta.get("code_version") or "-")
                                        m_src = str(meta.get("script_name") or "-")
                                        m_in = str(meta.get("input_name") or "-")
                                        m_t = str(meta.get("generated_at") or "-")
                                        self._append_log(
                                            f"[{fw}/{g}] semantic embeddings disk cache hit in {_fmt_elapsed(time.perf_counter() - t_emb)}: "
                                            f"model={sem_model}, file={cache_path}, "
                                            f"meta(code={m_ver}, script={m_src}, input={m_in}, gen={m_t})."
                                        )
                                    else:
                                        sem_emb_cache[_sek(g)] = _encode_semantic_embeddings(texts, model_size=sem_model, max_length=sem_ml, stride_pct=sem_stride_pct)
                                        save_ok = False
                                        meta_ok = False
                                        if sem_cache_write:
                                            save_ok = _save_semantic_emb_cache(cache_path, sem_emb_cache[_sek(g)])
                                            if save_ok:
                                                meta_obj = _build_semantic_emb_cache_meta(
                                                    cache_path=cache_path,
                                                    input_file=in_path,
                                                    granularity=g,
                                                    model_size=sem_model,
                                                    max_length=sem_ml,
                                                    text_count=len(texts),
                                                    embs=sem_emb_cache[_sek(g)],
                                                )
                                                meta_ok = _save_semantic_emb_meta(cache_path, meta_obj)
                                        self._append_log(
                                            f"[{fw}/{g}] semantic embeddings built in {_fmt_elapsed(time.perf_counter() - t_emb)}: "
                                            f"model={sem_model}, disk_cache={'saved' if save_ok else ('disabled' if not sem_cache_write else 'skip')}, "
                                            f"file={cache_path}"
                                            f"{', meta=' + _semantic_emb_meta_path(cache_path) if (save_ok and meta_ok) else ''}."
                                        )
                                else:
                                    self._append_log(f"[{fw}/{g}] semantic embeddings cache hit.")
                                sem_emb = sem_emb_cache[_sek(g)]

                            sem_key = (g, sem_model, int(sem_topk_build))
                            if sem_key not in sem_neighbors_cache:
                                self._status_set(f"[F{fi}/{len(frameworks_to_run)} G{gi}/{len(run_grans)}] {fw}/{g}: building semantic neighbors(topk={sem_topk_build})...")
                                t_nei = time.perf_counter()
                                sem_neighbors_cache[sem_key] = build_semantic_neighbors(sem_emb, topk_neighbors=sem_topk_build)
                                self._append_log(
                                    f"[{fw}/{g}] semantic neighbors built and cached in {_fmt_elapsed(time.perf_counter() - t_nei)}: "
                                    f"sem_topk={sem_topk_build}."
                                )
                            else:
                                self._append_log(f"[{fw}/{g}] semantic neighbors cache hit: sem_topk={sem_topk_build}.")
                            sem_neighbors = sem_neighbors_cache.get(sem_key, {})
                        except Exception as e:
                            self._append_log(f"WARN [{fw}/{g}]: semantic channel failed: {e}")
                            sem_emb, sem_neighbors = None, {}
                    if need_sem_channel and sem_emb is None:
                        raise RuntimeError(f"{fw} framework requires semantic embeddings, but semantic channel is unavailable for {g}.")

                    if fw == "merge":
                        lex_topk_eff = int(merge_k_lex) if mode_run == "topk" else int(topk_run)
                        sem_topk_eff = int(merge_k_sem) if mode_run == "topk" else int(sem_topk_default)

                        self._status_set(f"[F{fi}/{len(frameworks_to_run)} G{gi}/{len(run_grans)}] {fw}/{g}: lexical channel...")
                        t_lex = time.perf_counter()
                        rows_lex = compute_similarities(
                            units,
                            reps,
                            mode=mode_run,
                            topk=lex_topk_eff,
                            include_tfidf=use_tfidf,
                            granularity=g,
                            normal_units=normal_units_all,
                            ngram_n=ngram_n,
                            coverage_seg_topk=int(COVERAGE_DEFAULTS["seg_topk"]),
                            coverage_min_cos=float(COVERAGE_DEFAULTS["min_cos"]),
                            coverage_min_jaccard=float(COVERAGE_DEFAULTS["min_jaccard"]),
                            coverage_hits_max=int(COVERAGE_DEFAULTS["hits_max"]),
                            semantic_embeddings=None,
                            semantic_neighbors={},
                            candidate_source="lexical",
                            framework="lexical",
                            progress_cb=self._ui_progress,
                        )
                        annotate_lexical_only_scores(rows_lex, g, weights_h=w_h, weights_n=w_n)
                        self._append_log(
                            f"[{fw}/{g}] lexical channel done in {_fmt_elapsed(time.perf_counter() - t_lex)}: "
                            f"pairs={len(rows_lex)}, topk={lex_topk_eff} (mode={mode_run})."
                        )

                        self._status_set(f"[F{fi}/{len(frameworks_to_run)} G{gi}/{len(run_grans)}] {fw}/{g}: semantic channel...")
                        t_sem = time.perf_counter()
                        rows_sem = compute_similarities(
                            units,
                            reps,
                            mode=mode_run,
                            topk=sem_topk_eff,
                            include_tfidf=use_tfidf,
                            granularity=g,
                            normal_units=normal_units_all,
                            ngram_n=ngram_n,
                            coverage_seg_topk=int(COVERAGE_DEFAULTS["seg_topk"]),
                            coverage_min_cos=float(COVERAGE_DEFAULTS["min_cos"]),
                            coverage_min_jaccard=float(COVERAGE_DEFAULTS["min_jaccard"]),
                            coverage_hits_max=int(COVERAGE_DEFAULTS["hits_max"]),
                            semantic_embeddings=sem_emb,
                            semantic_neighbors=sem_neighbors,
                            candidate_source="semantic",
                            framework="semantic",
                            progress_cb=self._ui_progress,
                        )
                        if sem_maxsim_ctx is not None:
                            t_ms = time.perf_counter()
                            ms_stat = apply_semantic_maxsim_refine(rows_sem, sem_maxsim_ctx, granularity=g)
                            self._append_log(
                                f"[{fw}/{g}] semantic MaxSim refine: applied={int(ms_stat.get('applied') or 0)}/"
                                f"{int(ms_stat.get('selected') or 0)} selected "
                                f"(total={len(rows_sem)}, lambda={float(ms_stat.get('lambda') or 0.0):.2f}, "
                                f"policy={ms_stat.get('policy')}) in {_fmt_elapsed(time.perf_counter() - t_ms)}."
                            )
                        annotate_semantic_only_scores(rows_sem, granularity=g)
                        self._append_log(
                            f"[{fw}/{g}] semantic channel done in {_fmt_elapsed(time.perf_counter() - t_sem)}: "
                            f"pairs={len(rows_sem)}, topk={sem_topk_eff} (mode={mode_run})."
                        )

                        rows = join_channel_rows_for_merge(rows_lex, rows_sem)
                        self._append_log(
                            f"[{fw}/{g}] join channels: lexical={len(rows_lex)} + semantic={len(rows_sem)} => union={len(rows)}."
                        )

                        # ── unified dual-channel completion ──
                        comp = complete_merge_dual_scores(
                            rows, reps, sem_emb, units, g,
                            normal_units=normal_units_all,
                            ngram_n=ngram_n,
                            coverage_cfg=COVERAGE_DEFAULTS,
                        )
                        self._append_log(
                            f"[{fw}/{g}] completion: filled_lex={comp['filled_lex']}, "
                            f"filled_sem={comp['filled_sem']}, "
                            f"missing_lex={comp['missing_lex']}, missing_sem={comp['missing_sem']}, "
                            f"missing_lex_fields={comp['missing_lex_fields']}, "
                            f"elapsed={comp.get('elapsed_s', '?')}s."
                        )
                        if comp["missing_lex"] or comp["missing_sem"] or comp["missing_lex_fields"]:
                            raise RuntimeError(
                                f"completion failed: missing_lex={comp['missing_lex']}, "
                                f"missing_sem={comp['missing_sem']}, "
                                f"missing_lex_fields={comp['missing_lex_fields']}"
                            )

                        annotate_merge_scores(
                            rows,
                            g,
                            weights_h=w_h,
                            weights_n=w_n,
                            w_lex=merge_w_lex,
                            w_sem=merge_w_sem,
                        )
                        self._append_log(
                            f"[{fw}/{g}] scoring=merge(join+completion, lex={merge_w_lex:.2f}, sem={merge_w_sem:.2f}), "
                            f"recall=lex{merge_k_lex}+sem{merge_k_sem} (k={merge_k})."
                        )
                    else:
                        if fw == "lexical":
                            candidate_source = "lexical"
                            topk_eff = int(topk_run)
                        else:
                            candidate_source = "semantic"
                            topk_eff = int(sem_topk_default) if mode_run == "topk" else int(topk_run)

                        self._status_set(f"[F{fi}/{len(frameworks_to_run)} G{gi}/{len(run_grans)}] {fw}/{g}: computing similarities...")
                        t_sim = time.perf_counter()
                        rows = compute_similarities(
                            units,
                            reps,
                            mode=mode_run,
                            topk=topk_eff,
                            include_tfidf=use_tfidf,
                            granularity=g,
                            normal_units=normal_units_all,
                            ngram_n=ngram_n,
                            coverage_seg_topk=int(COVERAGE_DEFAULTS["seg_topk"]),
                            coverage_min_cos=float(COVERAGE_DEFAULTS["min_cos"]),
                            coverage_min_jaccard=float(COVERAGE_DEFAULTS["min_jaccard"]),
                            coverage_hits_max=int(COVERAGE_DEFAULTS["hits_max"]),
                            semantic_embeddings=sem_emb,
                            semantic_neighbors=(sem_neighbors if candidate_source == "semantic" else {}),
                            candidate_source=candidate_source,
                            framework=fw,
                            progress_cb=self._ui_progress,
                        )
                        self._append_log(
                            f"[{fw}/{g}] candidate mode={candidate_source}, pairs={len(rows)}, "
                            f"compute_time={_fmt_elapsed(time.perf_counter() - t_sim)}."
                        )

                        if fw == "lexical":
                            annotate_lexical_only_scores(rows, g, weights_h=w_h, weights_n=w_n)
                            self._append_log(f"[{fw}/{g}] scoring=lexical_only.")
                        else:
                            if sem_maxsim_ctx is not None:
                                t_ms = time.perf_counter()
                                ms_stat = apply_semantic_maxsim_refine(rows, sem_maxsim_ctx, granularity=g)
                                self._append_log(
                                    f"[{fw}/{g}] semantic MaxSim refine: applied={int(ms_stat.get('applied') or 0)}/"
                                    f"{int(ms_stat.get('selected') or 0)} selected "
                                    f"(total={len(rows)}, lambda={float(ms_stat.get('lambda') or 0.0):.2f}, "
                                    f"policy={ms_stat.get('policy')}) in {_fmt_elapsed(time.perf_counter() - t_ms)}."
                                )
                            annotate_semantic_only_scores(rows, granularity=g)
                            self._append_log(f"[{fw}/{g}] scoring=semantic_only.")

                    if rr_enabled and fw != "merge":
                        self._append_log(f"[{fw}/{g}] reranker skipped (only enabled for merge framework).")

                    if rr_enabled and fw == "merge":
                        try:
                            rr_plan = _estimate_reranker_candidate_count(
                                total_rows=len(rows),
                                granularity=g,
                                topn=rr_topn,
                                top_percent_non_h1=rr_top_pct,
                            )
                            self._append_log(
                                f"[{fw}/{g}] reranker plan: candidates={rr_plan}/{len(rows)} "
                                f"(topn_cap={rr_topn_raw}, top_pct_non_h1={rr_top_pct})."
                            )
                            self._status_set(f"[F{fi}/{len(frameworks_to_run)} G{gi}/{len(run_grans)}] {fw}/{g}: reranker(sidecar)...")
                            t_rer = time.perf_counter()
                            annotate_reranker_sidecar(
                                rows,
                                units,
                                enabled=True,
                                model_size=rr_model,
                                topn=rr_topn,
                                granularity=g,
                                top_percent_non_h1=rr_top_pct,
                                prompt=RERANK_PROMPT_DEFAULT,
                            )
                            applied_n = 0
                            for _r in rows:
                                try:
                                    ap = int(_r.get("rerank_applied") or 0)
                                    if ap > 0:
                                        applied_n += 1
                                except Exception:
                                    pass
                            self._append_log(
                                f"[{fw}/{g}] reranker sidecar done: model={rr_model}, "
                                f"applied={applied_n}/{len(rows)} (no score/rank overwrite), "
                                f"topn_cap={rr_topn_raw}, top_pct_non_h1={rr_top_pct}, "
                                f"time={_fmt_elapsed(time.perf_counter() - t_rer)}."
                            )
                        except Exception as e:
                            self._append_log(f"WARN [{fw}/{g}]: reranker sidecar failed: {e}")

                    sem_label_row = sem_label if fw in ("semantic", "merge") else None
                    rer_label_row = rer_label if (fw == "merge" and rr_enabled) else None
                    for _r in rows:
                        _r["semantic_model_label"] = sem_label_row
                        _r["reranker_model_label"] = rer_label_row

                    rows.sort(key=lambda r: float(r.get("score_final") or 0.0), reverse=True)

                    results_curr[g] = {"units": units, "rows": rows}
    
                    if (not run_all_gran) and (not framework_all_mode):
                        file_tag = _make_output_name_tag(
                            framework=fw,
                            granularity=g,
                            upload_base=base,
                            semantic_model=sem_model,
                            reranker_model=rr_model,
                            reranker_enabled=(rr_enabled if fw == "merge" else False),
                            time_prefix=run_stamp,
                        )
                        out_excel = os.path.join(out_dir, f"{file_tag}.xlsx")
                        out_word = os.path.join(out_dir, f"{file_tag}.docx")
                        self._append_log("Writing Excel report (full text columns)...")
                        self._status_set("Writing Excel report (full text)...")
                        write_excel_report_fulltext(
                            out_excel,
                            units,
                            rows,
                            framework=fw,
                            granularity=g,
                            semantic_model_label=(sem_label if fw in ("semantic", "merge") else None),
                            reranker_model_label=(rer_label if (fw == "merge" and rr_enabled) else None),
                        )
                        self._append_log("Writing Word report (top rows)...")
                        self._status_set("Writing Word report (top rows)...")
                        write_word_report(
                            out_word,
                            units,
                            normal_units_all,
                            rows,
                            title=f"Similarity Report ({_framework_cn(fw)}): {file_tag}",
                            max_rows=100,
                            top_percent=report_top_percent,
                            framework=fw,
                            granularity=g,
                        )
                        out_jsonl = ""
                        out_json_excel = ""
                        if export_jsonl:
                            out_jsonl = os.path.join(out_dir, f"{run_stamp}_units_all.jsonl")
                            self._append_log("Exporting units_all.jsonl ...")
                            self._status_set("Exporting units_all.jsonl ...")
                            export_units_jsonl(out_jsonl, units)
                            out_json_excel = os.path.splitext(out_excel)[0] + ".json"
                            try:
                                self._append_log("Exporting machine JSON from Excel ...")
                                self._status_set("Exporting machine JSON ...")
                                st = export_excel_machine_json(out_excel, out_json_excel)
                                self._append_log(
                                    f"Machine JSON exported: sheets={int(st.get('sheet_count') or 0)}, "
                                    f"rows={int(st.get('row_count') or 0)}, file={out_json_excel}"
                                )
                            except Exception as e:
                                out_json_excel = ""
                                self._append_log(f"WARN: machine JSON export failed: {e}")
                        total_elapsed = _fmt_elapsed(time.perf_counter() - run_t0)
                        self._status_set("Done ✅")
                        self._append_log(f"Done ✅ total_elapsed={total_elapsed}")
                        self._progress_stop()
                        self._run_started_at = None
                        info_text = f"Outputs:\n{out_excel}\n{out_word}"
                        if out_jsonl:
                            info_text += f"\n{out_jsonl}"
                        if out_json_excel:
                            info_text += f"\n{out_json_excel}"
                        self._safe_showinfo("Done", info_text)
                        return

                if framework_all_mode:
                    if results_curr:
                        results_by_fw[fw] = results_curr
                else:
                    results_by_g = results_curr

            if framework_all_mode:
                if not results_by_fw:
                    raise RuntimeError("No valid results were produced for framework=all.")
                if dual_all_mode:
                    out_tag_all = _make_dual_all_output_name_tag(
                        upload_base=base,
                        semantic_model=sem_model,
                        reranker_model=rr_model,
                        reranker_enabled=rr_enabled,
                        time_prefix=run_stamp,
                    )
                else:
                    out_tag_all = _make_output_name_tag(
                        framework="all",
                        granularity="h1-h2-normal",
                        upload_base=base,
                        semantic_model=sem_model,
                        reranker_model=rr_model,
                        reranker_enabled=rr_enabled,
                        time_prefix=run_stamp,
                    )
                out_excel_total = os.path.join(out_dir, f"{out_tag_all}_total.xlsx")
                out_excel_audit = os.path.join(out_dir, f"{out_tag_all}_audit.xlsx")
                out_word_all = os.path.join(out_dir, f"{out_tag_all}.docx")
                out_top10_pdf = os.path.join(out_dir, f"{out_tag_all}_top10_tables.pdf")
                out_top10_img = os.path.join(out_dir, f"{out_tag_all}_top10_matrix_a4.png")
                out_merge_top10_img = os.path.join(out_dir, f"{out_tag_all}_merge_top10_matrix_a4.png")

                self._append_log("Enriching cross-channel reference scores (no extra workbook)...")
                self._status_set("Cross-channel enrichment...")
                try:
                    enrich_cross_channel_scores(
                        results_by_fw,
                        reps_cache=reps_cache,
                        sem_emb_cache=sem_emb_cache,
                        granularities=run_grans,
                        sem_model=sem_model,
                    )
                except Exception as e:
                    self._append_log(f"WARN: cross score enrichment failed: {e}")

                self._append_log("Building SHADOW audit rows (unified U by pair ids)...")
                self._status_set("Building shadow audit rows...")
                t_shadow = time.perf_counter()
                audit_rows_by_g = build_shadow_audit_rows(
                    results_by_fw,
                    topk_by_g=SHADOW_AUDIT_TOPK_DEFAULT,
                    reranker_enabled=rr_enabled,
                    reranker_model=rr_model,
                    reranker_prompt=RERANK_PROMPT_DEFAULT,
                    log_cb=self._append_log,
                )
                self._append_log(f"Shadow audit rows built in {_fmt_elapsed(time.perf_counter() - t_shadow)}.")

                self._append_log("Writing framework=all TOTAL Excel (9 sheets)...")
                self._status_set("Writing total Excel...")
                write_excel_report_framework_all(
                    out_excel_total,
                    results_by_fw,
                    top_rows_per_g=None,
                    weights_h=w_h,
                    weights_n=w_n,
                    semantic_model_label=sem_label,
                    reranker_model_label=(rer_label if rr_enabled else None),
                )

                self._append_log("Writing framework=all SHADOW AUDIT Excel (3 sheets)...")
                self._status_set("Writing shadow audit Excel...")
                write_excel_shadow_audit(
                    out_excel_audit,
                    audit_rows_by_g,
                    semantic_model_label=sem_label,
                    reranker_model_label=(rer_label if rr_enabled else None),
                )

                self._append_log("Writing framework=all Word report...")
                self._status_set("Writing framework=all Word report...")
                write_word_report_framework_all(
                    out_word_all,
                    results_by_fw,
                    title=f"Similarity ALL Report (字面/语义/融合): {out_tag_all}",
                    top_rows_per_mode=50,
                    comparison_plot_paths=[],
                )

                try:
                    self._append_log("Writing framework=all Top10 PDF (shadow-audit logic)...")
                    self._status_set("Writing Top10 PDF (shadow-audit)...")
                    got_pdf = build_framework_all_top10_pdf(
                        out_top10_pdf,
                        audit_rows_by_g,
                        weights_h=w_h,
                        weights_n=w_n,
                        merge_w_lex=merge_w_lex,
                        merge_w_sem=merge_w_sem,
                    )
                    if not got_pdf:
                        out_top10_pdf = ""
                        self._append_log("WARN: Top10 PDF generation skipped/failed.")
                except Exception as e:
                    out_top10_pdf = ""
                    self._append_log(f"WARN: Top10 PDF generation failed: {e}")

                try:
                    self._append_log("Writing framework=all A4 summary image (shadow-audit logic)...")
                    self._status_set("Writing A4 summary image (shadow-audit)...")
                    got_img = build_framework_all_summary_image(
                        out_top10_img,
                        audit_rows_by_g,
                        dpi=450,
                    )
                    if not got_img:
                        out_top10_img = ""
                        self._append_log("WARN: A4 summary image generation skipped/failed.")
                except Exception as e:
                    out_top10_img = ""
                    self._append_log(f"WARN: A4 summary image generation failed: {e}")

                try:
                    self._append_log("Writing framework=all MERGE-only A4 summary image...")
                    self._status_set("Writing merge-only A4 summary image...")
                    got_img_merge = build_framework_merge_summary_image(
                        out_merge_top10_img,
                        audit_rows_by_g,
                        dpi=450,
                    )
                    if not got_img_merge:
                        out_merge_top10_img = ""
                        self._append_log("WARN: Merge-only A4 summary image generation skipped/failed.")
                except Exception as e:
                    out_merge_top10_img = ""
                    self._append_log(f"WARN: Merge-only A4 summary image generation failed: {e}")

                out_jsonl_all = ""
                out_json_total = ""
                out_json_audit = ""
                if export_jsonl:
                    out_jsonl_all = os.path.join(out_dir, f"{run_stamp}_units_all.jsonl")
                    units_all = _collect_units_all_from_results_by_fw(results_by_fw)
                    self._append_log(f"Exporting units_all.jsonl ... records={len(units_all)}")
                    export_units_jsonl(out_jsonl_all, units_all)
                    out_json_total = os.path.splitext(out_excel_total)[0] + ".json"
                    out_json_audit = os.path.splitext(out_excel_audit)[0] + ".json"
                    try:
                        self._append_log("Exporting machine JSON from TOTAL Excel ...")
                        self._status_set("Exporting machine JSON (total) ...")
                        st_total = export_excel_machine_json(out_excel_total, out_json_total)
                        self._append_log(
                            f"Machine JSON(total) exported: sheets={int(st_total.get('sheet_count') or 0)}, "
                            f"rows={int(st_total.get('row_count') or 0)}, file={out_json_total}"
                        )
                    except Exception as e:
                        out_json_total = ""
                        self._append_log(f"WARN: total machine JSON export failed: {e}")
                    try:
                        self._append_log("Exporting machine JSON from AUDIT Excel ...")
                        self._status_set("Exporting machine JSON (audit) ...")
                        st_audit = export_excel_machine_json(out_excel_audit, out_json_audit)
                        self._append_log(
                            f"Machine JSON(audit) exported: sheets={int(st_audit.get('sheet_count') or 0)}, "
                            f"rows={int(st_audit.get('row_count') or 0)}, file={out_json_audit}"
                        )
                    except Exception as e:
                        out_json_audit = ""
                        self._append_log(f"WARN: audit machine JSON export failed: {e}")

                total_elapsed = _fmt_elapsed(time.perf_counter() - run_t0)
                self._status_set("Done ✅")
                self._append_log(f"Done ✅ total_elapsed={total_elapsed}")
                self._progress_stop()
                self._run_started_at = None
                self._safe_showinfo(
                    "Done",
                    f"Output folder:\n{out_dir}\n\nALL outputs:\n{out_excel_total}\n{out_excel_audit}\n{out_word_all}\n{out_top10_pdf or 'Top10 PDF skipped'}\n{out_top10_img or 'A4 summary image skipped'}\n{out_merge_top10_img or 'Merge-only A4 summary image skipped'}\n{out_jsonl_all or 'units_all.jsonl skipped'}\n{out_json_total or 'total JSON skipped'}\n{out_json_audit or 'audit JSON skipped'}"
                )
                return

            if not results_by_g:
                raise RuntimeError("No valid granularity results were produced in all-mode.")

            fw_label = FRAMEWORK_LABEL_MAP.get(fw_selected, "merge")
            out_tag_fw_all = _make_output_name_tag(
                framework=fw_selected,
                granularity="h1-h2-normal",
                upload_base=base,
                semantic_model=sem_model,
                reranker_model=rr_model,
                reranker_enabled=(rr_enabled if fw_selected == "merge" else False),
                time_prefix=run_stamp,
            )
            out_excel_all = os.path.join(out_dir, f"{out_tag_fw_all}.xlsx")
            out_word_all = os.path.join(out_dir, f"{out_tag_fw_all}.docx")
            out_plot_dir = os.path.join(out_dir, f"{out_tag_fw_all}_plots")
            if plt is None:
                self._append_log("WARN: matplotlib not available, plot images will be skipped.")
            plot_paths = build_top10_plot_bundle(
                results_by_g,
                out_dir=out_plot_dir,
                framework_key=fw_selected,
                file_stem=f"{out_tag_fw_all}_{fw_label}",
                weights_h=w_h,
                weights_n=w_n,
                merge_w_lex=merge_w_lex,
                merge_w_sem=merge_w_sem,
                rrf_k_h1=30,
            )

            self._append_log("Writing ALL Excel report (3 sheets, full rows each)...")
            self._status_set("Writing ALL Excel report...")
            write_excel_report_all(
                out_excel_all,
                results_by_g,
                top_rows_per_g=None,
                weights_h=w_h,
                weights_n=w_n,
                framework=fw_selected,
                semantic_model_label=(sem_label if fw_selected in ("semantic", "merge") else None),
                reranker_model_label=(rer_label if (fw_selected == "merge" and rr_enabled) else None),
            )

            self._append_log("Writing ALL Word report (h1/h2/normal top50)...")
            self._status_set("Writing ALL Word report...")
            write_word_report_all(
                out_word_all,
                results_by_g,
                title=f"Similarity ALL Report ({_framework_cn(fw_selected)}): {out_tag_fw_all}",
                top_rows_per_g=50,
                h2_max_chars=1000,
                plot_paths=plot_paths,
                framework=fw_selected,
            )

            out_jsonl_all = ""
            out_json_all = ""
            if export_jsonl:
                out_jsonl_all = os.path.join(out_dir, f"{run_stamp}_units_all.jsonl")
                units_all = _collect_units_all_from_results_by_g(results_by_g)
                self._append_log(f"Exporting units_all.jsonl ... records={len(units_all)}")
                export_units_jsonl(out_jsonl_all, units_all)
                out_json_all = os.path.splitext(out_excel_all)[0] + ".json"
                try:
                    self._append_log("Exporting machine JSON from ALL Excel ...")
                    self._status_set("Exporting machine JSON ...")
                    st = export_excel_machine_json(out_excel_all, out_json_all)
                    self._append_log(
                        f"Machine JSON exported: sheets={int(st.get('sheet_count') or 0)}, "
                        f"rows={int(st.get('row_count') or 0)}, file={out_json_all}"
                    )
                except Exception as e:
                    out_json_all = ""
                    self._append_log(f"WARN: ALL machine JSON export failed: {e}")

            total_elapsed = _fmt_elapsed(time.perf_counter() - run_t0)
            self._status_set("Done ✅")
            self._append_log(f"Done ✅ total_elapsed={total_elapsed}")
            self._progress_stop()
            self._run_started_at = None
            self._safe_showinfo(
                "Done",
                f"Output folder:\n{out_dir}\n\nALL outputs:\n{out_excel_all}\n{out_word_all}\n{out_jsonl_all or 'units_all.jsonl skipped'}\n{out_json_all or 'ALL JSON skipped'}\n\nPlots:\n{out_plot_dir}"
            )

        except Exception as e:
            self._status_set("Error.")
            self._append_log("ERROR: " + str(e))
            self._run_started_at = None
            self._safe_showerror("Error", str(e))


if __name__ == "__main__":
    # CLI evaluation interface:
    #   python pipeline.py eval --pred xxx.xlsx --labels yyy.xlsx
    #   python pipeline.py --eval --pred xxx.xlsx --labels yyy.xlsx
    if len(sys.argv) >= 2:
        cmd = str(sys.argv[1]).strip().lower()
        if cmd in ("eval", "--eval"):
            raise SystemExit(_run_eval_cli(sys.argv[2:]))
    app = App()
    try:
        def _sigint_handler(_signum, _frame):
            try:
                app.after(0, app.destroy)
            except Exception:
                pass
        signal.signal(signal.SIGINT, _sigint_handler)
    except Exception:
        pass
    try:
        app.mainloop()
    except KeyboardInterrupt:
        pass
